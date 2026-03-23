# -*- coding: utf-8 -*-
"""
Module tạo báo cáo Word tự động cho KPI/BSC NVKT (v2)

Refactored:
- Import scoring từ kpi_scoring.py (single source of truth)
- Import team mapping từ team_config.py (thay TEAM_SHORT_NAMES hardcode)
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os
import re

from kpi_scoring import (
    tinh_diem_C11_TP1, tinh_diem_C11_TP2,
    tinh_diem_C12_TP1, tinh_diem_C12_TP2,
    tinh_diem_C14, tinh_diem_C15,
    chuan_hoa_ty_le,
)
from team_config import BRCD_TEAMS

# Thiết lập matplotlib để hỗ trợ tiếng Việt
matplotlib.rcParams['font.family'] = 'DejaVu Sans'
matplotlib.use('Agg')  # Use non-interactive backend

# =============================================================================
# CẤU HÌNH
# =============================================================================
DEFAULT_KPI_FOLDER = "downloads/KPI"
DEFAULT_OUTPUT_FOLDER = "downloads/reports"

# Mapping tên đơn vị ngắn gọn - từ team_config (single source of truth)
TEAM_SHORT_NAMES = {}
for _team in BRCD_TEAMS:
    TEAM_SHORT_NAMES[f"Tổ Kỹ thuật Địa bàn {_team.short_name}"] = _team.short_name
    TEAM_SHORT_NAMES[f"Tổ Kỹ thuật địa bàn {_team.short_name}"] = _team.short_name
    # Thêm variant với short_name viết thường để xử lý case không khớp từ Excel
    _lower = _team.short_name.lower()
    if _lower != _team.short_name:
        TEAM_SHORT_NAMES[f"Tổ Kỹ thuật Địa bàn {_lower}"] = _team.short_name
        TEAM_SHORT_NAMES[f"Tổ Kỹ thuật địa bàn {_lower}"] = _team.short_name

# Thứ tự đơn vị BRCD (sorted by order)
TEAM_ORDER = [t.short_name
              for t in sorted(BRCD_TEAMS, key=lambda t: t.order) if t.active]


def _get_short_name(name):
    """Lấy tên ngắn từ tên đầy đủ đơn vị, xử lý case-insensitive."""
    # Kiểm tra trong dict trước
    if name in TEAM_SHORT_NAMES:
        return TEAM_SHORT_NAMES[name]
    # Thử strip prefix "Tổ Kỹ thuật" + "Địa bàn"/"địa bàn"
    m = re.match(r'(?i)Tổ\s+Kỹ\s+thuật\s+[Đđ]ịa\s+bàn\s+(.+)', str(name))
    if m:
        return m.group(1).strip().title()
    return name


def _match_unit(df, col, unit_name):
    """Match unit name flexibly - checks both short name and full name variants.

    Handles cases where data may contain full names like 'Tổ Kỹ thuật Địa bàn Suối hai'
    while unit_name is just 'Suối Hai', or vice versa.
    """
    short = _get_short_name(unit_name)
    mask = df[col].apply(lambda x: _get_short_name(str(x)) == short if pd.notna(x) else False)
    return df[mask]


# Màu sắc cho biểu đồ
CHART_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6B5B95']
BAR_COLORS = ['#4CAF50', '#2196F3', '#FF9800', '#E91E63', '#9C27B0']  # Xanh lá, Xanh dương, Cam, Hồng, Tím


# =============================================================================
# HÀM TIỆN ÍCH
# =============================================================================
def set_cell_shading(cell, color):
    """Đặt màu nền cho ô trong bảng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_border(table):
    """Đặt viền cho bảng"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def format_number(value, decimal_places=2):
    """Format số với xử lý NaN"""
    if pd.isna(value):
        return "N/A"
    return f"{value:.{decimal_places}f}"


def format_percent(value, decimal_places=2):
    """Format tỷ lệ phần trăm"""
    if pd.isna(value):
        return "N/A"
    return f"{value:.{decimal_places}f}%"


# =============================================================================
# HÀM ĐỌC DỮ LIỆU
# =============================================================================
def load_kpi_data(kpi_folder=DEFAULT_KPI_FOLDER):
    """
    Đọc dữ liệu KPI từ các file Excel
    
    Returns:
        tuple: (df_summary, df_detail) - DataFrame tóm tắt và chi tiết
    """
    kpi_path = Path(kpi_folder)
    
    # Đọc file tóm tắt
    summary_file = kpi_path / "KPI_NVKT_TomTat.xlsx"
    df_summary = pd.read_excel(summary_file)
    
    # Đọc file chi tiết
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    return df_summary, df_detail


def load_c1x_reports(data_folder="downloads/baocao_hanoi"):
    """
    Đọc dữ liệu chi tiết từ các file báo cáo C1.x
    
    Returns:
        dict: Dictionary chứa các DataFrame từ các sheet tổng hợp
    """
    data_path = Path(data_folder)
    reports = {}
    
    # C1.1 Report
    try:
        c11_file = data_path / "c1.1 report.xlsx"
        if c11_file.exists():
            reports['c11'] = pd.read_excel(c11_file, sheet_name='TH_C1.1')
            print("   ✅ Đọc C1.1 report thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc C1.1 report: {e}")
    
    # C1.2 Report
    try:
        c12_file = data_path / "c1.2 report.xlsx"
        if c12_file.exists():
            reports['c12'] = pd.read_excel(c12_file, sheet_name='TH_C1.2')
            print("   ✅ Đọc C1.2 report thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc C1.2 report: {e}")
    
    # C1.3 Report
    try:
        c13_file = data_path / "c1.3 report.xlsx"
        if c13_file.exists():
            reports['c13'] = pd.read_excel(c13_file, sheet_name='TH_C1.3')
            print("   ✅ Đọc C1.3 report thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc C1.3 report: {e}")
    
    # C1.4 Report
    try:
        c14_file = data_path / "c1.4 report.xlsx"
        if c14_file.exists():
            reports['c14'] = pd.read_excel(c14_file, sheet_name='TH_C1.4')
            print("   ✅ Đọc C1.4 report thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc C1.4 report: {e}")
    
    # C1.5 Chi tiết Report - Sheet TH_TTVTST
    try:
        c15_file = data_path / "c1.5_chitiet_report.xlsx"
        if c15_file.exists():
            reports['c15_ttvtst'] = pd.read_excel(c15_file, sheet_name='TH_TTVTST')
            print("   ✅ Đọc C1.5 report (TH_TTVTST) thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc C1.5 report: {e}")
    
    return reports


def load_exclusion_comparison_data(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc dữ liệu so sánh trước/sau giảm trừ từ các file Excel
    
    Args:
        exclusion_folder: Thư mục chứa các file kết quả sau giảm trừ
        
    Returns:
        dict: Dictionary chứa DataFrames cho từng chỉ tiêu
              - 'c11_sm4': So sánh C1.1 SM4 (Sửa chữa báo hỏng)
              - 'c11_sm2': So sánh C1.1 SM2 (Sửa chữa chủ động)
              - 'c12_sm1': So sánh C1.2 SM1 (Hỏng lặp lại)
              - 'c12_sm4': So sánh C1.2 SM4 (Tỷ lệ báo hỏng BRCĐ)
              - 'c14': So sánh C1.4 (Độ hài lòng)
              - 'c15': So sánh C1.5 (Thiết lập dịch vụ BRCĐ)
              - 'tong_hop': Tổng hợp tất cả chỉ tiêu
    """
    data_path = Path(exclusion_folder)
    comparison_data = {}
    
    if not data_path.exists():
        print(f"   ⚠️ Không tìm thấy thư mục giảm trừ: {exclusion_folder}")
        return comparison_data
    
    # C1.1 SM4 - Sửa chữa báo hỏng đúng quy định
    try:
        c11_sm4_file = data_path / "So_sanh_C11_SM4.xlsx"
        if c11_sm4_file.exists():
            comparison_data['c11_sm4'] = {
                'chi_tiet': pd.read_excel(c11_sm4_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c11_sm4_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc So_sanh_C11_SM4.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc So_sanh_C11_SM4.xlsx: {e}")
    
    # C1.1 SM2 - Sửa chữa chủ động
    try:
        c11_sm2_file = data_path / "So_sanh_C11_SM2.xlsx"
        if c11_sm2_file.exists():
            comparison_data['c11_sm2'] = {
                'chi_tiet': pd.read_excel(c11_sm2_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c11_sm2_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc So_sanh_C11_SM2.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc So_sanh_C11_SM2.xlsx: {e}")
    
    # C1.2 SM1 - Hỏng lặp lại
    try:
        c12_sm1_file = data_path / "So_sanh_C12_SM1.xlsx"
        if c12_sm1_file.exists():
            comparison_data['c12_sm1'] = {
                'chi_tiet': pd.read_excel(c12_sm1_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c12_sm1_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc So_sanh_C12_SM1.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc So_sanh_C12_SM1.xlsx: {e}")
    
    # C1.2 SM4 - Tỷ lệ báo hỏng BRCĐ
    try:
        c12_sm4_file = data_path / "SM4-C12-ti-le-su-co-dv-brcd.xlsx"
        if c12_sm4_file.exists():
            comparison_data['c12_sm4'] = {
                'chi_tiet': pd.read_excel(c12_sm4_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c12_sm4_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc SM4-C12-ti-le-su-co-dv-brcd.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc SM4-C12: {e}")
    
    # C1.4 - Độ hài lòng khách hàng
    try:
        c14_file = data_path / "So_sanh_C14.xlsx"
        if c14_file.exists():
            comparison_data['c14'] = {
                'chi_tiet': pd.read_excel(c14_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c14_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc So_sanh_C14.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc So_sanh_C14.xlsx: {e}")
    
    
    # C1.5 - Tỷ lệ thiết lập dịch vụ BRCĐ đạt thời gian quy định
    try:
        c15_file = data_path / "So_sanh_C15.xlsx"
        if c15_file.exists():
            comparison_data['c15'] = {
                'chi_tiet': pd.read_excel(c15_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c15_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ✅ Đọc So_sanh_C15.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc So_sanh_C15.xlsx: {e}")
    
    # Tổng hợp giảm trừ
    try:
        tong_hop_file = data_path / "Tong_hop_giam_tru.xlsx"
        if tong_hop_file.exists():
            comparison_data['tong_hop'] = pd.read_excel(tong_hop_file)
            print("   ✅ Đọc Tong_hop_giam_tru.xlsx thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc Tong_hop_giam_tru.xlsx: {e}")
    
    return comparison_data


def load_unit_level_exclusion_data(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc dữ liệu thống kê theo đơn vị (Tổ) từ sheet Thong_ke_theo_don_vi
    
    Returns:
        dict: Dictionary chứa DataFrames thống kê theo đơn vị cho từng chỉ tiêu
    """
    data_path = Path(exclusion_folder)
    unit_data = {}
    
    if not data_path.exists():
        print(f"   ⚠️ Không tìm thấy thư mục giảm trừ: {exclusion_folder}")
        return unit_data
    
    # C1.1 SM4
    try:
        c11_sm4_file = data_path / "So_sanh_C11_SM4.xlsx"
        if c11_sm4_file.exists():
            unit_data['c11_sm4'] = pd.read_excel(c11_sm4_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.1 SM4 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.1 SM4: {e}")
    
    # C1.1 SM2
    try:
        c11_sm2_file = data_path / "So_sanh_C11_SM2.xlsx"
        if c11_sm2_file.exists():
            unit_data['c11_sm2'] = pd.read_excel(c11_sm2_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.1 SM2 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.1 SM2: {e}")
    
    # C1.2 SM1
    try:
        c12_sm1_file = data_path / "So_sanh_C12_SM1.xlsx"
        if c12_sm1_file.exists():
            unit_data['c12_sm1'] = pd.read_excel(c12_sm1_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.2 SM1 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.2 SM1: {e}")
    
    
    # C1.2 SM4 - Tỷ lệ sự cố BRCĐ
    try:
        c12_sm4_file = data_path / "SM4-C12-ti-le-su-co-dv-brcd.xlsx"
        if c12_sm4_file.exists():
            unit_data['c12_sm4'] = pd.read_excel(c12_sm4_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.2 SM4 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.2 SM4: {e}")
    
    # C1.4
    try:
        c14_file = data_path / "So_sanh_C14.xlsx"
        if c14_file.exists():
            unit_data['c14'] = pd.read_excel(c14_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.4 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.4: {e}")
    
    # C1.5
    try:
        c15_file = data_path / "So_sanh_C15.xlsx"
        if c15_file.exists():
            unit_data['c15'] = pd.read_excel(c15_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ✅ Đọc unit stats C1.5 thành công")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc unit stats C1.5: {e}")

    # Normalize 'Đơn vị' column in all loaded DataFrames to short names
    for key, df in unit_data.items():
        if 'Đơn vị' in df.columns:
            unit_data[key] = df.copy()
            unit_data[key]['Đơn vị'] = df['Đơn vị'].apply(
                lambda x: _get_short_name(str(x)) if pd.notna(x) else x
            )

    return unit_data


def load_bsc_unit_scores_from_comparison(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc điểm BSC đã tính sẵn từ file Tong_hop_Diem_BSC_Don_Vi.xlsx
    File này chứa cả điểm Trước và Sau giảm trừ cho từng đơn vị và TTVT Sơn Tây
    
    Returns:
        dict: Dictionary với cấu trúc:
            {
                'units': DataFrame chứa điểm các đơn vị (sheet Tong_hop_Don_vi),
                'individuals': DataFrame chứa điểm cá nhân (sheet Chi_tiet_Ca_nhan)
            }
    """
    file_path = Path(exclusion_folder) / "Tong_hop_Diem_BSC_Don_Vi.xlsx"
    result = {'units': None, 'individuals': None}
    
    if not file_path.exists():
        print(f"   ⚠️ Không tìm thấy file: {file_path}")
        return result
    
    try:
        result['units'] = pd.read_excel(file_path, sheet_name='Tong_hop_Don_vi')
        print(f"   ✅ Đọc điểm BSC đơn vị từ Tong_hop_Diem_BSC_Don_Vi.xlsx: {len(result['units'])} dòng")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc sheet Tong_hop_Don_vi: {e}")
    
    try:
        result['individuals'] = pd.read_excel(file_path, sheet_name='Chi_tiet_Ca_nhan')
        print(f"   ✅ Đọc điểm BSC cá nhân từ Tong_hop_Diem_BSC_Don_Vi.xlsx: {len(result['individuals'])} dòng")
    except Exception as e:
        print(f"   ⚠️ Không thể đọc sheet Chi_tiet_Ca_nhan: {e}")
    
    return result


def load_nvkt_exclusion_summary(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc dữ liệu KPI sau giảm trừ theo NVKT từ các file so sánh thành phần
    (Sử dụng lại logic của load_nvkt_exclusion_detail để đảm bảo nhất quán)
    """
    # Vì load_nvkt_exclusion_detail đã tổng hợp từ các file gốc và tính điểm lại,
    # nên ta có thể dùng lại kết quả của nó.
    return load_nvkt_exclusion_detail(exclusion_folder)


def add_kpi_summary_table_after_exclusion(doc, df_exclusion, team_name):
    """
    Thêm bảng tổng hợp KPI sau giảm trừ vào document cho 1 tổ
    
    Args:
        doc: Document Word
        df_exclusion: DataFrame dữ liệu sau giảm trừ
        team_name: Tên tổ cần lọc
    """
    if df_exclusion is None or df_exclusion.empty:
        doc.add_paragraph("(Không có dữ liệu sau giảm trừ)")
        return
    
    # Lọc theo tổ
    df = _match_unit(df_exclusion, 'don_vi', team_name).copy()
    if df.empty:
        doc.add_paragraph("(Không có dữ liệu sau giảm trừ cho tổ này)")
        return
    
    # Sắp xếp
    df = df.sort_values('nvkt')

    # Tạo bảng - bao gồm C1.1, C1.2, C1.4, C1.5 sau giảm trừ
    headers = ['STT', 'NVKT', 'C1.1', 'C1.2', 'C1.4', 'C1.5']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], '2E7D32')  # Màu xanh lá để phân biệt với bảng thô
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells

        data = [
            str(idx),
            row.get('nvkt', ''),
            format_number(row.get('Diem_C1.1', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = str(value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            
            # Màu nền xen kẽ
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')  # Xanh lá nhạt
            
            # Tô màu điểm theo mức (chỉ các cột điểm)
            if i >= 2:
                try:
                    val = float(value) if value and value != 'N/A' else None
                    if val is not None:
                        if val >= 4.5:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                            run.font.bold = True
                        elif val < 3:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.bold = True
                except (ValueError, TypeError):
                    pass


def chuan_hoa_ten_nvkt(name):
    """
    Chuẩn hóa tên NVKT về dạng Title Case để tránh trùng lặp do viết hoa/thường khác nhau
    Ví dụ: "Bùi văn Cường" -> "Bùi Văn Cường"
    """
    if pd.isna(name) or name is None or str(name).strip() == '' or str(name) == 'nan':
        return None
    return str(name).strip().title()


def load_nvkt_raw_detail(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc dữ liệu KPI chi tiết TRƯỚC giảm trừ (Thô) theo NVKT từ các file so sánh
    Sử dụng các cột có hậu tố (Thô) thay vì (Sau GT)
    """
    data_path = Path(exclusion_folder)
    if not data_path.exists():
        return None

    print("   🔄 Đang tổng hợp dữ liệu chi tiết NVKT (Thô) từ các file so sánh...")
    
    files = {
        'c11_sm2': ('So_sanh_C11_SM2.xlsx', 'So_sanh_chi_tiet'),
        'c11_sm4': ('So_sanh_C11_SM4.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm1': ('So_sanh_C12_SM1.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm4': ('SM4-C12-ti-le-su-co-dv-brcd.xlsx', 'So_sanh_chi_tiet'),
        'c14': ('So_sanh_C14.xlsx', 'So_sanh_chi_tiet'),
        'c15': ('So_sanh_C15.xlsx', 'So_sanh_chi_tiet')
    }
    
    dfs = {}
    for key, (filename, sheet) in files.items():
        try:
            f_path = data_path / filename
            if f_path.exists():
                df = pd.read_excel(f_path, sheet_name=sheet)
                if 'NVKT' in df.columns:
                    # Chuẩn hóa tên NVKT về Title Case để tránh trùng lặp
                    df['NVKT'] = df['NVKT'].apply(chuan_hoa_ten_nvkt)
                    df = df[df['NVKT'].notna()]  # Loại bỏ các dòng có NVKT null
                    dfs[key] = df
        except Exception as e:
            print(f"   ⚠️ Lỗi đọc file {filename}: {e}")

    if not dfs:
        return None

    # Collect all NVKTs (đã được chuẩn hóa)
    all_nvkt = set()
    nvkt_info = {}
    for df in dfs.values():
        if 'NVKT' in df.columns:
            for _, row in df.iterrows():
                nvkt = row['NVKT']
                if nvkt:
                    all_nvkt.add(nvkt)
                    if nvkt not in nvkt_info:
                        don_vi = row.get('TEN_DOI', '') or row.get('Đơn vị', '')
                        nvkt_info[nvkt] = {'don_vi': don_vi}

    if not all_nvkt:
        return None

    summary_data = []
    for nvkt in all_nvkt:
        info = nvkt_info.get(nvkt, {})
        row_data = {
            'nvkt': nvkt,
            'don_vi': info.get('don_vi', ''),
            'c11_tp1_tong_phieu': 0, 'c11_tp1_phieu_dat': 0, 'c11_tp1_ty_le': 0, 'diem_c11_tp1': 5,
            'c11_tp2_tong_phieu': 0, 'c11_tp2_phieu_dat': 0, 'c11_tp2_ty_le': 0, 'diem_c11_tp2': 5,
            'Diem_C1.1': 0,
            'c12_tp1_phieu_hll': 0, 'c12_tp1_phieu_bh': 0, 'c12_tp1_ty_le': 0, 'diem_c12_tp1': 5,
            'c12_tp2_tong_tb': 0, 'c12_tp2_phieu_bh': 0, 'c12_tp2_ty_le': 0, 'diem_c12_tp2': 5,
            'Diem_C1.2': 0,
            # C1.4 - Mặc định np.nan nếu không có dữ liệu (hiển thị N/A)
            'c14_phieu_ks': np.nan, 'c14_phieu_khl': np.nan, 'c14_ty_le': np.nan, 'diem_c14': np.nan, 'Diem_C1.4': np.nan,
            # C1.5 - Mặc định np.nan nếu không có dữ liệu (hiển thị N/A)
            'c15_tong_phieu': np.nan, 'c15_phieu_dat': np.nan, 'c15_phieu_khong_dat': np.nan, 'c15_ty_le': np.nan, 'diem_c15': np.nan, 'Diem_C1.5': np.nan
        }

        # C1.1 SM2 (TP1) - Thô
        if 'c11_sm2' in dfs:
            r = dfs['c11_sm2'][dfs['c11_sm2']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp1_tong_phieu'] = r.get('Tổng phiếu (Thô)', 0)
                row_data['c11_tp1_phieu_dat'] = r.get('Số phiếu đạt (Thô)', 0)
                row_data['c11_tp1_ty_le'] = r.get('Tỷ lệ % (Thô)', 0)
                row_data['diem_c11_tp1'] = r.get('Điểm BSC (Thô)', 5)
        
        # C1.1 SM4 (TP2) - Thô
        if 'c11_sm4' in dfs:
            r = dfs['c11_sm4'][dfs['c11_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp2_tong_phieu'] = r.get('Tổng phiếu (Thô)', 0)
                row_data['c11_tp2_phieu_dat'] = r.get('Số phiếu đạt (Thô)', 0)
                row_data['c11_tp2_ty_le'] = r.get('Tỷ lệ % (Thô)', 0)
                row_data['diem_c11_tp2'] = r.get('Điểm BSC (Thô)', 5)
        
        row_data['Diem_C1.1'] = 0.3 * row_data['diem_c11_tp1'] + 0.7 * row_data['diem_c11_tp2']

        # C1.2 SM1 (TP1) - Thô
        if 'c12_sm1' in dfs:
            r = dfs['c12_sm1'][dfs['c12_sm1']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp1_phieu_hll'] = r.get('Số phiếu HLL (Thô)', 0)
                row_data['c12_tp1_phieu_bh'] = r.get('Số phiếu báo hỏng (Thô)', 0)
                row_data['c12_tp1_ty_le'] = r.get('Tỷ lệ HLL % (Thô)', 0)
                row_data['diem_c12_tp1'] = r.get('Điểm BSC (Thô)', 5)
        
        # C1.2 SM4 (TP2) - Thô
        if 'c12_sm4' in dfs:
            r = dfs['c12_sm4'][dfs['c12_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp2_tong_tb'] = r.get('Tổng TB (Thô)', 0)
                row_data['c12_tp2_phieu_bh'] = r.get('Số phiếu báo hỏng (Thô)', 0)
                row_data['c12_tp2_ty_le'] = r.get('Tỷ lệ báo hỏng (%) (Thô)', 0)
                row_data['diem_c12_tp2'] = r.get('Điểm BSC (Thô)', 5)

        row_data['Diem_C1.2'] = 0.5 * row_data['diem_c12_tp1'] + 0.5 * row_data['diem_c12_tp2']

        # C1.4 - Thô
        if 'c14' in dfs:
            r = dfs['c14'][dfs['c14']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c14_phieu_ks'] = r.get('Tổng phiếu KS (Thô)', 0)
                row_data['c14_phieu_khl'] = r.get('Số phiếu KHL (Thô)', 0)
                row_data['c14_ty_le'] = r.get('Tỷ lệ HL (%) (Thô)', 0)
                row_data['diem_c14'] = r.get('Điểm BSC (Thô)', 5)
                row_data['Diem_C1.4'] = row_data['diem_c14']

        # C1.5 - Thô
        if 'c15' in dfs:
            r = dfs['c15'][dfs['c15']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c15_tong_phieu'] = r.get('Tổng Hoàn công (Thô)', 0)
                row_data['c15_phieu_dat'] = r.get('Phiếu đạt (Thô)', 0)
                row_data['c15_phieu_khong_dat'] = r.get('Phiếu không đạt (Thô)', 0)
                row_data['c15_ty_le'] = r.get('Tỷ lệ đạt (%) (Thô)', 0)
                row_data['diem_c15'] = r.get('Điểm BSC (Thô)', 5)
                row_data['Diem_C1.5'] = row_data['diem_c15']

        summary_data.append(row_data)

    df_result = pd.DataFrame(summary_data)
    print(f"   ✅ Tổng hợp xong dữ liệu chi tiết NVKT Thô ({len(df_result)} nhân viên)")
    return df_result



def load_nvkt_exclusion_detail(exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Đọc dữ liệu KPI chi tiết sau giảm trừ theo NVKT từ các file so sánh thành phần
    """
    data_path = Path(exclusion_folder)
    if not data_path.exists():
        return None

    print("   🔄 Đang tổng hợp dữ liệu chi tiết NVKT từ các file so sánh...")
    
    # Danh sách các file cần đọc
    files = {
        'c11_sm2': ('So_sanh_C11_SM2.xlsx', 'So_sanh_chi_tiet'),
        'c11_sm4': ('So_sanh_C11_SM4.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm1': ('So_sanh_C12_SM1.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm4': ('SM4-C12-ti-le-su-co-dv-brcd.xlsx', 'So_sanh_chi_tiet'),
        'c14': ('So_sanh_C14.xlsx', 'So_sanh_chi_tiet'),
        'c15': ('So_sanh_C15.xlsx', 'So_sanh_chi_tiet')
    }
    
    dfs = {}
    
    for key, (filename, sheet) in files.items():
        try:
            # Đọc file, bỏ qua dòng tiêu đề phụ nếu có (thường header=0 là đủ nếu cột nằm ở dòng 1)
            f_path = data_path / filename
            if f_path.exists():
                df = pd.read_excel(f_path, sheet_name=sheet)
                # Chuẩn hóa tên cột NVKT và TEN_DOI
                if 'Mã nhân viên' in df.columns:
                    df.rename(columns={'Mã nhân viên': 'NVKT'}, inplace=True)
                if 'Tên nhân viên' in df.columns:
                    df.rename(columns={'Tên nhân viên': 'TEN_NV'}, inplace=True)
                
                # Đảm bảo có cột NVKT để merge
                if 'NVKT' in df.columns:
                    # Chuẩn hóa tên NVKT về Title Case để tránh trùng lặp
                    df['NVKT'] = df['NVKT'].apply(chuan_hoa_ten_nvkt)
                    df = df[df['NVKT'].notna()]  # Loại bỏ các dòng có NVKT null
                    dfs[key] = df
                else:
                    print(f"   ⚠️ File {filename}: Không tìm thấy cột NVKT")
            else:
                print(f"   ⚠️ Không tìm thấy file {filename}")
        except Exception as e:
             print(f"   ⚠️ Lỗi đọc file {filename}: {e}")

    if not dfs:
        return None

    # Lấy danh sách tất cả NVKT từ các file (đã được chuẩn hóa Title Case)
    all_nvkt = set()
    nvkt_info = {} # Lưu thông tin NVKT (Tên, Tổ)

    for df in dfs.values():
        if 'NVKT' in df.columns:
            for _, row in df.iterrows():
                nvkt = row['NVKT']
                if nvkt:
                    all_nvkt.add(nvkt)
                    # Lưu thông tin bổ sung nếu chưa có
                    if nvkt not in nvkt_info:
                        don_vi = row.get('TEN_DOI', '') or row.get('Đơn vị', '')
                        # Logic: Giữ nguyên tên đơn vị từ file, sau này lọc theo tên đó
                        nvkt_info[nvkt] = {
                            'don_vi': don_vi,
                            'ten_nv': row.get('TEN_NV', '') or row.get('Tên nhân viên', '')
                        }

    if not all_nvkt:
        return None

    # Tạo DataFrame tổng hợp
    summary_data = []

    for nvkt in all_nvkt:
        info = nvkt_info.get(nvkt, {})
        row_data = {
            'nvkt': nvkt,
            'don_vi': info.get('don_vi', ''),
            'ten_nv': info.get('ten_nv', ''),
            # C1.1
            'c11_tp1_tong_phieu': 0, 'c11_tp1_phieu_dat': 0, 'c11_tp1_ty_le': 0, 'diem_c11_tp1': 5,
            'c11_tp2_tong_phieu': 0, 'c11_tp2_phieu_dat': 0, 'c11_tp2_ty_le': 0, 'diem_c11_tp2': 5,
            'Diem_C1.1': 0,
            # C1.2 - column names fixed to match table renderer
            'c12_tp1_phieu_hll': 0, 'c12_tp1_phieu_bh': 0, 'c12_tp1_ty_le': 0, 'diem_c12_tp1': 5,
            'c12_tp2_tong_tb': 0, 'c12_tp2_phieu_bh': 0, 'c12_tp2_ty_le': 0, 'diem_c12_tp2': 5,
            'Diem_C1.2': 0,
            # C1.4 - Mặc định np.nan nếu không có dữ liệu (hiển thị N/A)
            'c14_phieu_ks': np.nan, 'c14_phieu_khl': np.nan, 'c14_ty_le': np.nan, 'diem_c14': np.nan, 'Diem_C1.4': np.nan,
            # C1.5 - Mặc định np.nan nếu không có dữ liệu (hiển thị N/A)
            'c15_tong_phieu': np.nan, 'c15_phieu_dat': np.nan, 'c15_phieu_khong_dat': np.nan, 'c15_ty_le': np.nan, 'diem_c15': np.nan, 'Diem_C1.5': np.nan
        }
        
        # Fill C1.1 SM2 (TP1)
        if 'c11_sm2' in dfs:
            r = dfs['c11_sm2'][dfs['c11_sm2']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp1_tong_phieu'] = r.get('Tổng phiếu (Sau GT)', 0)
                row_data['c11_tp1_phieu_dat'] = r.get('Số phiếu đạt (Sau GT)', 0)
                row_data['c11_tp1_ty_le'] = r.get('Tỷ lệ % (Sau GT)', 0)
                row_data['diem_c11_tp1'] = r.get('Điểm BSC (Sau GT)', 5)
        
        # Fill C1.1 SM4 (TP2)
        if 'c11_sm4' in dfs:
            r = dfs['c11_sm4'][dfs['c11_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp2_tong_phieu'] = r.get('Tổng phiếu (Sau GT)', 0)
                row_data['c11_tp2_phieu_dat'] = r.get('Số phiếu đạt (Sau GT)', 0)
                row_data['c11_tp2_ty_le'] = r.get('Tỷ lệ % (Sau GT)', 0)
                row_data['diem_c11_tp2'] = r.get('Điểm BSC (Sau GT)', 5)
        
        # Calculate C1.1 Score
        row_data['Diem_C1.1'] = 0.3 * row_data['diem_c11_tp1'] + 0.7 * row_data['diem_c11_tp2']

        # Fill C1.2 SM1 (TP1) - Hỏng lặp lại
        if 'c12_sm1' in dfs:
            r = dfs['c12_sm1'][dfs['c12_sm1']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp1_phieu_hll'] = r.get('Số phiếu HLL (Sau GT)', 0)
                row_data['c12_tp1_phieu_bh'] = r.get('Số phiếu báo hỏng (Sau GT)', 0)
                row_data['c12_tp1_ty_le'] = r.get('Tỷ lệ HLL % (Sau GT)', 0)
                row_data['diem_c12_tp1'] = r.get('Điểm BSC (Sau GT)', 5)
        
        # Fill C1.2 SM4 (TP2) - Tỷ lệ sự cố
        if 'c12_sm4' in dfs:
            r = dfs['c12_sm4'][dfs['c12_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp2_tong_tb'] = r.get('Tổng TB (Thô)', 0)
                row_data['c12_tp2_phieu_bh'] = r.get('Số phiếu báo hỏng (Sau GT)', 0)
                row_data['c12_tp2_ty_le'] = r.get('Tỷ lệ báo hỏng (%) (Sau GT)', 0)
                row_data['diem_c12_tp2'] = r.get('Điểm BSC (Sau GT)', 5)

        # Calculate C1.2 Score
        row_data['Diem_C1.2'] = 0.5 * row_data['diem_c12_tp1'] + 0.5 * row_data['diem_c12_tp2']

        # Fill C1.4 - Độ hài lòng khách hàng
        if 'c14' in dfs:
            r = dfs['c14'][dfs['c14']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c14_phieu_ks'] = r.get('Tổng phiếu KS (Sau GT)', 0)
                row_data['c14_phieu_khl'] = r.get('Số phiếu KHL (Sau GT)', 0)
                row_data['c14_ty_le'] = r.get('Tỷ lệ HL (%) (Sau GT)', 0)
                row_data['diem_c14'] = r.get('Điểm BSC (Sau GT)', 5)
                row_data['Diem_C1.4'] = row_data['diem_c14']

        # Fill C1.5 - Thiết lập dịch vụ BRCĐ
        if 'c15' in dfs:
            r = dfs['c15'][dfs['c15']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c15_tong_phieu'] = r.get('Tổng Hoàn công (Sau GT)', 0)
                row_data['c15_phieu_dat'] = r.get('Phiếu đạt (Sau GT)', 0)
                row_data['c15_phieu_khong_dat'] = r.get('Phiếu không đạt (Sau GT)', 0)
                row_data['c15_ty_le'] = r.get('Tỷ lệ đạt (%) (Sau GT)', 0)
                row_data['diem_c15'] = r.get('Điểm BSC (Sau GT)', 5)
                row_data['Diem_C1.5'] = row_data['diem_c15']

        summary_data.append(row_data)

    df_result = pd.DataFrame(summary_data)
    print(f"   ✅ Tổng hợp xong dữ liệu chi tiết NVKT ({len(df_result)} nhân viên)")
    return df_result


def add_c11_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Thêm bảng chi tiết C1.1 sau giảm trừ
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.1 sau giảm trừ)")
        return
    
    df = _match_unit(df_exclusion_detail, 'don_vi', team_name).copy()
    if df.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.1 sau giảm trừ cho tổ này)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi tiết chỉ tiêu C1.1 - Chất lượng sửa chữa thuê bao BRCĐ (sau giảm trừ)', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('TP1 = Sửa chữa chủ động (SCCD ≤72h) | TP2 = Sửa chữa theo báo hỏng (SC BH) | Sau GT = Sau giảm trừ')
    
    headers = ['STT', 'NVKT', 'Tổng SCCD', 'Đạt ≤72h', 'TL(%)', 'Điểm TP1',
               'Tổng SC BH', 'Đúng hạn', 'TL(%)', 'Điểm TP2', 'Điểm C1.1']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - màu xanh lá đậm hơn để phân biệt
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '1B5E20')  # Xanh lá đậm
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c11_tp1_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp1_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp1', np.nan)),
            format_number(row.get('c11_tp2_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp2_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp2', np.nan)),
            format_number(row.get('Diem_C1.1', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'C8E6C9')  # Xanh lá nhạt hơn


def add_c12_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Thêm bảng chi tiết C1.2 sau giảm trừ
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.2 sau giảm trừ)")
        return
    
    df = _match_unit(df_exclusion_detail, 'don_vi', team_name).copy()
    if df.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.2 sau giảm trừ cho tổ này)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi tiết chỉ tiêu C1.2 - Tỷ lệ thuê bao báo hỏng (sau giảm trừ)', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('TP1 = Hỏng lặp (≥2 lần/7 ngày) | TP2 = Tỷ lệ BH/TB quản lý | Sau GT = Sau giảm trừ')
    
    headers = ['STT', 'NVKT', 'Hỏng lặp', 'Tổng BH', 'TL(%)', 'Điểm TP1',
               'Phiếu BH', 'TB QL', 'TL(‰)', 'Điểm TP2', 'Điểm C1.2']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - màu xanh dương đậm hơn
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '0D47A1')  # Xanh dương đậm
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c12_tp1_phieu_hll', np.nan), 0),
            format_number(row.get('c12_tp1_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp1', np.nan)),
            format_number(row.get('c12_tp2_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp2_tong_tb', np.nan), 0),
            format_number(row.get('c12_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp2', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'BBDEFB')  # Xanh dương nhạt


def add_c14_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Thêm bảng chi tiết C1.4 sau giảm trừ - Độ hài lòng khách hàng
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.4 sau giảm trừ)")
        return
    
    df = _match_unit(df_exclusion_detail, 'don_vi', team_name).copy()
    if df.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.4 sau giảm trừ cho tổ này)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi tiết chỉ tiêu C1.4 - Độ hài lòng khách hàng (sau giảm trừ)', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('KS = Khảo sát | Không HL = Không hài lòng | Sau GT = Sau giảm trừ')
    
    headers = ['STT', 'NVKT', 'Tổng KS', 'Không HL', 'Tỷ lệ HL (%)', 'Điểm C1.4']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - màu cam đậm hơn
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'E65100')  # Cam đậm
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c14_phieu_ks', np.nan), 0),
            format_number(row.get('c14_phieu_khl', np.nan), 0),
            format_number(row.get('c14_ty_le', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFE0B2')  # Cam nhạt


def add_c15_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Thêm bảng chi tiết C1.5 sau giảm trừ - Thiết lập dịch vụ BRCĐ
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.5 sau giảm trừ)")
        return

    df = _match_unit(df_exclusion_detail, 'don_vi', team_name).copy()
    if df.empty:
        doc.add_paragraph("(Không có dữ liệu chi tiết C1.5 sau giảm trừ cho tổ này)")
        return

    df = df.sort_values('nvkt')

    doc.add_heading('Chi tiết chỉ tiêu C1.5 - Thiết lập dịch vụ BRCĐ đạt thời gian quy định (sau giảm trừ)', level=3)

    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('Đạt TG = Hoàn thành đúng thời gian | Sau GT = Sau giảm trừ (loại bỏ phiếu loại trừ)')

    headers = ['STT', 'NVKT', 'Đạt TG', 'Không đạt', 'Tổng phiếu', 'Tỷ lệ (%)', 'Điểm C1.5']

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header - màu tím đậm hơn để phân biệt với bảng trước
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '4A148C')  # Tím đậm
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c15_phieu_dat', np.nan), 0),
            format_number(row.get('c15_phieu_khong_dat', np.nan), 0),
            format_number(row.get('c15_tong_phieu', np.nan), 0),
            format_number(row.get('c15_ty_le', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)

            if idx % 2 == 0:
                set_cell_shading(cells[i], 'EDE7F6')  # Tím nhạt


def create_nvkt_bar_chart_after_exclusion(df_exclusion, team_name, output_path=None):
    """
    Tạo biểu đồ cột so sánh điểm KPI sau giảm trừ theo NVKT trong 1 tổ
    
    Args:
        df_exclusion: DataFrame chứa dữ liệu KPI sau giảm trừ
        team_name: Tên tổ cần tạo biểu đồ
        output_path: Đường dẫn lưu file (None = trả về bytes)
    
    Returns:
        bytes hoặc str, None nếu không có dữ liệu
    """
    if df_exclusion is None or df_exclusion.empty:
        return None
    
    # Lọc theo tổ
    df = _match_unit(df_exclusion, 'don_vi', team_name).copy()
    if df.empty or len(df) == 0:
        return None
    
    # Sắp xếp theo tên
    df = df.sort_values('nvkt')
    
    # Lấy tên ngắn của tổ
    short_name = _get_short_name(team_name)
    
    # Chuẩn bị dữ liệu
    nvkts = df['nvkt'].tolist()
    c11 = df['Diem_C1.1'].fillna(0).tolist()
    c12 = df['Diem_C1.2'].fillna(0).tolist()
    c14 = df['Diem_C1.4'].fillna(0).tolist()
    c15 = df['Diem_C1.5'].fillna(0).tolist()

    # Tạo biểu đồ
    fig, ax = plt.subplots(figsize=(12, 6))

    x = np.arange(len(nvkts))
    width = 0.2

    # Các cột - bao gồm C1.1, C1.2, C1.4, C1.5 sau giảm trừ
    bars1 = ax.bar(x - 1.5*width, c11, width, label='C1.1', color='#66BB6A')  # Xanh lá
    bars2 = ax.bar(x - 0.5*width, c12, width, label='C1.2', color='#42A5F5')  # Xanh dương
    bars3 = ax.bar(x + 0.5*width, c14, width, label='C1.4', color='#FFA726')  # Cam
    bars4 = ax.bar(x + 1.5*width, c15, width, label='C1.5', color='#AB47BC')  # Tím

    # Thêm giá trị lên cột
    for bars in [bars1, bars2, bars3, bars4]:
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax.annotate(f'{height:.1f}',
                           xy=(bar.get_x() + bar.get_width() / 2, height),
                           xytext=(0, 3),
                           textcoords="offset points",
                           ha='center', va='bottom', fontsize=7)
    
    ax.set_xlabel('Nhân viên kỹ thuật', fontsize=11)
    ax.set_ylabel('Điểm KPI', fontsize=11)
    ax.set_title(f'SO SÁNH ĐIỂM KPI SAU GIẢM TRỪ - {short_name.upper()}', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(nvkts, rotation=45, ha='right', fontsize=9)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_unit_comparison_chart(unit_data, chi_tieu='c11_sm4', output_path=None):
    """
    Tạo biểu đồ grouped bar so sánh tỷ lệ trước/sau GT theo đơn vị
    
    Args:
        unit_data: Dictionary từ load_unit_level_exclusion_data()
        chi_tieu: 'c11_sm4', 'c11_sm2', 'c12_sm1', 'c14'
        output_path: Đường dẫn lưu file
    """
    if not unit_data or chi_tieu not in unit_data:
        return None
    
    df = unit_data[chi_tieu]
    
    # Lấy cột tỷ lệ
    tyle_tho_col = None
    tyle_sau_col = None
    for col in df.columns:
        if 'Tỷ lệ' in col and 'Thô' in col:
            tyle_tho_col = col
        elif 'Tỷ lệ' in col and 'Sau GT' in col:
            tyle_sau_col = col
    
    if not tyle_tho_col or not tyle_sau_col:
        return None
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    width = 0.35
    
    don_vi = df['Đơn vị'].values
    tyle_tho = df[tyle_tho_col].fillna(0).values
    tyle_sau = df[tyle_sau_col].fillna(0).values
    
    bars1 = ax.bar(x - width/2, tyle_tho, width, label='Trước giảm trừ', color='#EF5350', alpha=0.8)
    bars2 = ax.bar(x + width/2, tyle_sau, width, label='Sau giảm trừ', color='#66BB6A', alpha=0.8)
    
    # Thêm giá trị lên cột
    for bar, val in zip(bars1, tyle_tho):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
    for bar, val in zip(bars2, tyle_sau):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
    
    # Tiêu đề theo chỉ tiêu
    titles = {
        'c11_sm4': 'C1.1 SM4 - Sửa chữa báo hỏng',
        'c11_sm2': 'C1.1 SM2 - Sửa chữa chủ động', 
        'c12_sm1': 'C1.2 SM1 - Hỏng lặp lại',
        'c14': 'C1.4 - Độ hài lòng khách hàng'
    }
    
    ax.set_xlabel('Đơn vị', fontsize=11)
    ax.set_ylabel('Tỷ lệ (%)', fontsize=11)
    ax.set_title(f'SO SÁNH TRƯỚC/SAU GIẢM TRỪ THEO ĐƠN VỊ\n{titles.get(chi_tieu, chi_tieu)}', 
                fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(don_vi, rotation=15, ha='right', fontsize=10)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_unit_exclusion_table(doc, unit_data, chi_tieu='c11_sm4'):
    """
    Thêm bảng thống kê theo đơn vị vào document
    
    Args:
        doc: Document Word
        unit_data: Dictionary từ load_unit_level_exclusion_data()
        chi_tieu: 'c11_sm4', 'c11_sm2', 'c12_sm1', 'c14'
    """
    if not unit_data or chi_tieu not in unit_data:
        return
    
    df = unit_data[chi_tieu]
    
    # Định nghĩa tiêu đề và màu theo chỉ tiêu
    config = {
        'c11_sm4': {'title': 'C1.1 SM4 - Sửa chữa báo hỏng theo đơn vị (Sau GT)', 'color': 'C62828'},
        'c11_sm2': {'title': 'C1.1 SM2 - Sửa chữa chủ động theo đơn vị (Sau GT)', 'color': 'AD1457'},
        'c12_sm1': {'title': 'C1.2 SM1 - Hỏng lặp lại theo đơn vị (Sau GT)', 'color': '0D47A1'},
        'c14': {'title': 'C1.4 - Độ hài lòng theo đơn vị (Sau GT)', 'color': 'E65100'}
    }
    
    cfg = config.get(chi_tieu, {'title': chi_tieu, 'color': '333333'})
    
    doc.add_heading(cfg['title'], level=4)
    
    # Lấy các cột cần hiển thị
    display_cols = ['Đơn vị', 'Tổng phiếu (Thô)', 'Phiếu loại trừ', 'Tổng phiếu (Sau GT)']
    tyle_cols = [c for c in df.columns if 'Tỷ lệ' in c]
    thay_doi_cols = [c for c in df.columns if 'Thay đổi' in c]
    
    headers = display_cols + tyle_cols[:2] + thay_doi_cols[:1]
    headers = [h for h in headers if h in df.columns]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        # Rút gọn tên header
        h = headers[i]
        short_h = h.replace('(Thô)', '(T)').replace('(Sau GT)', '(S)').replace('Tổng phiếu', 'Tổng')
        header.text = short_h
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, cfg['color'])
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            val = row.get(h, '')
            if pd.isna(val):
                val = ''
            elif isinstance(val, (int, float)):
                if 'Tỷ lệ' in h or 'Thay đổi' in h:
                    val = f"{val:.2f}%"
                else:
                    val = str(int(val))
            cells[i].text = str(val)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            # Tô màu cho dòng TTVT
            if 'TTVT' in str(row.get('Đơn vị', '')):
                run.font.bold = True
                set_cell_shading(cells[i], 'E0E0E0')
            elif idx % 2 == 0:
                set_cell_shading(cells[i], 'F5F5F5')
    
    doc.add_paragraph()


def add_unit_level_exclusion_section(doc, unit_data, c1x_reports=None):
    """
    Thêm phần thống kê giảm trừ theo đơn vị vào document
    Bao gồm biểu đồ BSC sau giảm trừ, bảng và biểu đồ chi tiết cho từng chỉ tiêu
    """
    if not unit_data:
        return
    
    doc.add_heading('Thống kê giảm trừ theo đơn vị (Tổ)', level=3)
    
    p = doc.add_paragraph()
    p.add_run('📊 Số liệu dưới đây thể hiện kết quả các chỉ tiêu BSC trước và sau giảm trừ, ')
    p.add_run('được tổng hợp theo từng Tổ kỹ thuật và toàn TTVT Sơn Tây.')
    doc.add_paragraph()
    
    # (Biểu đồ BSC sau giảm trừ đã được đặt ở section 1.1.b - trước phần này)
    
    # C1.1 SM4
    if 'c11_sm4' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c11_sm4')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c11_sm4')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Không thể tạo biểu đồ C1.1 SM4: {e}")
        doc.add_paragraph()
    
    # C1.1 SM2
    if 'c11_sm2' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c11_sm2')
        doc.add_paragraph()
    
    # C1.2 SM1
    if 'c12_sm1' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c12_sm1')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c12_sm1')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Không thể tạo biểu đồ C1.2 SM1: {e}")
        doc.add_paragraph()
    
    # C1.4
    if 'c14' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c14')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c14')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Không thể tạo biểu đồ C1.4: {e}")
        doc.add_paragraph()


def create_comparison_bar_chart(comparison_data, output_path=None):
    """
    Tạo biểu đồ grouped bar so sánh tỷ lệ trước/sau giảm trừ
    
    Args:
        comparison_data: Dictionary từ load_exclusion_comparison_data()
        output_path: Đường dẫn lưu file (None = trả về bytes)
    
    Returns:
        bytes hoặc str: Dữ liệu ảnh
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        return None
    
    df = comparison_data['tong_hop']
    
    # Sắp xếp theo thứ tự
    chi_tieu_order = ['C1.1 SM4', 'C1.1 SM2', 'C1.2', 'C1.2 Tỷ lệ BRCĐ báo hỏng', 'C1.4 Độ hài lòng KH']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    width = 0.35
    
    tyle_tho = df['Tỷ lệ % (Thô)'].fillna(0).values
    tyle_sau = df['Tỷ lệ % (Sau GT)'].fillna(0).values
    chi_tieu = df['Chỉ tiêu'].values
    
    bars1 = ax.bar(x - width/2, tyle_tho, width, label='Trước giảm trừ', color='#E57373', alpha=0.8)
    bars2 = ax.bar(x + width/2, tyle_sau, width, label='Sau giảm trừ', color='#81C784', alpha=0.8)
    
    # Thêm giá trị lên cột
    for bar, val in zip(bars1, tyle_tho):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                   f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
    for bar, val in zip(bars2, tyle_sau):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                   f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Chỉ tiêu', fontsize=11)
    ax.set_ylabel('Tỷ lệ (%)', fontsize=11)
    ax.set_title('SO SÁNH TỶ LỆ TRƯỚC/SAU GIẢM TRỪ', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(chi_tieu, rotation=15, ha='right', fontsize=9)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_exclusion_summary_table(doc, comparison_data):
    """
    Thêm bảng tổng hợp so sánh trước/sau giảm trừ vào document
    
    Args:
        doc: Document Word
        comparison_data: Dictionary từ load_exclusion_comparison_data()
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        doc.add_paragraph("⚠️ Không có dữ liệu giảm trừ")
        return
    
    df = comparison_data['tong_hop']
    
    doc.add_heading('BẢNG TỔNG HỢP SO SÁNH TRƯỚC/SAU GIẢM TRỪ', level=3)
    
    headers = ['Chỉ tiêu', 'Tổng phiếu (Thô)', 'Loại trừ', 'Tổng phiếu (Sau GT)', 
               'Tỷ lệ % (Thô)', 'Tỷ lệ % (Sau GT)', 'Thay đổi %']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'D32F2F')  # Đỏ đậm
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        
        thay_doi = row.get('Thay đổi %', 0)
        if pd.isna(thay_doi):
            thay_doi = 0
        
        data = [
            str(row.get('Chỉ tiêu', '')),
            str(int(row.get('Tổng phiếu (Thô)', 0))) if pd.notna(row.get('Tổng phiếu (Thô)')) else 'N/A',
            str(int(row.get('Phiếu loại trừ', 0))) if pd.notna(row.get('Phiếu loại trừ')) else 'N/A',
            str(int(row.get('Tổng phiếu (Sau GT)', 0))) if pd.notna(row.get('Tổng phiếu (Sau GT)')) else 'N/A',
            format_number(row.get('Tỷ lệ % (Thô)', 0)) + '%',
            format_number(row.get('Tỷ lệ % (Sau GT)', 0)) + '%',
            f"{thay_doi:+.2f}%"
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            # Tô màu chênh lệch
            if i == 6:  # Cột thay đổi
                if thay_doi > 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Xanh lá (tăng)
                elif thay_doi < 0:
                    run.font.color.rgb = RGBColor(200, 0, 0)  # Đỏ (giảm)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFEBEE')  # Hồng nhạt
    
    doc.add_paragraph()


def add_c1x_detail_with_exclusion(doc, c1x_reports, comparison_data, chi_tieu='c11'):
    """
    Thêm bảng chi tiết cho 1 chỉ tiêu với cả dữ liệu trước và sau giảm trừ
    
    Args:
        doc: Document Word
        c1x_reports: Dictionary từ load_c1x_reports()
        comparison_data: Dictionary từ load_exclusion_comparison_data()
        chi_tieu: 'c11', 'c12', hoặc 'c14'
    """
    if chi_tieu == 'c11':
        title = 'C1.1 - Tỷ lệ sửa chữa phiếu chất lượng & báo hỏng'
        # Hiển thị bảng gốc
        if 'c11' in c1x_reports:
            doc.add_heading(f'{title} (DỮ LIỆU THÔ)', level=3)
            df = c1x_reports['c11']
            _add_c11_table(doc, df)
        
        # Hiển thị bảng sau giảm trừ nếu có
        if 'c11_sm4' in comparison_data:
            doc.add_heading(f'{title} (SAU GIẢM TRỪ)', level=3)
            df_sau = comparison_data['c11_sm4']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.1 SM4')
        
        if 'c11_sm2' in comparison_data:
            df_sau = comparison_data['c11_sm2']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.1 SM2')
    
    elif chi_tieu == 'c12':
        title = 'C1.2 - Tỷ lệ báo hỏng lặp lại & Tỷ lệ sự cố dịch vụ'
        if 'c12' in c1x_reports:
            doc.add_heading(f'{title} (DỮ LIỆU THÔ)', level=3)
            df = c1x_reports['c12']
            _add_c12_table(doc, df)
        
        if 'c12_sm1' in comparison_data:
            doc.add_heading(f'{title} (SAU GIẢM TRỪ)', level=3)
            df_sau = comparison_data['c12_sm1']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.2 SM1')
    
    elif chi_tieu == 'c14':
        title = 'C1.4 - Độ hài lòng khách hàng sau sửa chữa'
        if 'c14' in c1x_reports:
            doc.add_heading(f'{title} (DỮ LIỆU THÔ)', level=3)
            df = c1x_reports['c14']
            _add_c14_table(doc, df)
        
        if 'c14' in comparison_data:
            doc.add_heading(f'{title} (SAU GIẢM TRỪ)', level=3)
            df_sau = comparison_data['c14']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.4')


def _add_exclusion_summary_mini(doc, df_tong_hop, label):
    """Helper: Thêm mini summary table cho 1 chỉ tiêu sau giảm trừ"""
    p = doc.add_paragraph()
    p.add_run(f'📊 {label}: ').bold = True
    
    if df_tong_hop is not None and len(df_tong_hop) > 0:
        row = df_tong_hop.iloc[0]
        tyle_tho = row.get('Tỷ lệ % (Thô)', row.get('Tỷ lệ HLL % (Thô)', 0))
        tyle_sau = row.get('Tỷ lệ % (Sau GT)', row.get('Tỷ lệ HLL % (Sau GT)', 0))
        thay_doi = row.get('Thay đổi %', 0)
        
        if pd.isna(tyle_tho):
            tyle_tho = 0
        if pd.isna(tyle_sau):
            tyle_sau = 0
        if pd.isna(thay_doi):
            thay_doi = 0
        
        p.add_run(f'Trước GT: {tyle_tho:.2f}% → Sau GT: {tyle_sau:.2f}% ')
        
        thay_doi_run = p.add_run(f'(Δ: {thay_doi:+.2f}%)')
        if thay_doi > 0:
            thay_doi_run.font.color.rgb = RGBColor(0, 128, 0)
        elif thay_doi < 0:
            thay_doi_run.font.color.rgb = RGBColor(200, 0, 0)


def _add_c11_table(doc, df):
    """Helper: Thêm bảng C1.1 gốc"""
    headers = ['Đơn vị', 'SC Chủ động (SM1)', 'Đạt (SM2)', 'TL SC CĐ (%)', 
               'Báo hỏng (SM3)', 'Đạt ĐH (SM4)', 'TL SCBH (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = _get_short_name(row['Đơn vị'])
        if short_name == 'Tổng':
            short_name = 'TTVT Sơn Tây'
        data = [
            short_name,
            str(int(row.get('SM1', 0))),
            str(int(row.get('SM2', 0))),
            format_number(row.get('Tỷ lệ sửa chữa phiếu chất lượng chủ động dịch vụ FiberVNN, MyTV đạt yêu cầu', 0)),
            str(int(row.get('SM3', 0))),
            str(int(row.get('SM4', 0))),
            format_number(row.get('Tỷ lệ phiếu sửa chữa báo hỏng dịch vụ BRCD đúng quy định không tính hẹn', 0)),
            format_number(row.get('Chỉ tiêu BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')
            if short_name == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'C8E6C9')
    doc.add_paragraph()


def _add_c12_table(doc, df):
    """Helper: Thêm bảng C1.2 gốc"""
    headers = ['Đơn vị', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)', 
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = _get_short_name(row['Đơn vị'])
        if short_name == 'Tổng':
            short_name = 'TTVT Sơn Tây'
        data = [
            short_name,
            str(int(row.get('SM1', 0))),
            str(int(row.get('SM2', 0))),
            format_number(row.get('Tỷ lệ thuê bao báo hỏng dịch vụ BRCĐ lặp lại', 0)),
            str(int(row.get('SM3', 0))),
            str(int(row.get('SM4', 0))),
            format_number(row.get('Tỷ lệ sự cố dịch vụ BRCĐ', 0)),
            format_number(row.get('Chỉ tiêu BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')
            if short_name == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
    doc.add_paragraph()


def _add_c14_table(doc, df):
    """Helper: Thêm bảng C1.4 gốc"""
    headers = ['Đơn vị', 'Tổng phiếu', 'Đã KS', 'KS TC', 'KH HL', 
               'KHL KT PV', 'TL HL PV (%)', 'TL KH HL (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'F57C00')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = _get_short_name(row['Đơn vị'])
        if short_name == 'Tổng':
            short_name = 'TTVT Sơn Tây'
        data = [
            short_name,
            str(int(row.get('Tổng phiếu', 0))),
            str(int(row.get('SL đã KS', 0))),
            str(int(row.get('SL KS thành công', 0))),
            str(int(row.get('SL KH hài lòng', 0))),
            str(int(row.get('Không HL KT phục vụ', 0))),
            format_number(row.get('Tỷ lệ HL KT phục vụ', 0)),
            format_number(row.get('Tỷ lệ KH hài lòng', 0)),
            format_number(row.get('Điểm BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')
            if short_name == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'FFE0B2')
def add_c11_unit_level_exclusion_table(doc, unit_data, c1x_reports=None):
    """
    Thêm bảng C1.1 tổng hợp theo đơn vị (cấp tổ) sau giảm trừ
    Tương tự bảng C1.1 gốc nhưng với số liệu sau giảm trừ

    Args:
        doc: Document Word
        unit_data: Dictionary từ load_unit_level_exclusion_data()
        c1x_reports: Dictionary chứa báo cáo C1.x gốc (để lấy số liệu SM1, SM3)
    """
    if not unit_data:
        return

    # Kiểm tra có dữ liệu C1.1 không
    if 'c11_sm2' not in unit_data or 'c11_sm4' not in unit_data:
        return

    doc.add_heading('C1.1 - Tỷ lệ sửa chữa phiếu chất lượng & báo hỏng (sau giảm trừ)', level=3)

    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 GHI CHÚ: ').bold = True
    p.add_run('Bảng này hiển thị số liệu C1.1 tổng hợp theo đơn vị sau khi loại bỏ các phiếu thuộc diện giảm trừ. ')
    p.add_run('SM1, SM3 là số liệu thô (không áp dụng giảm trừ). SM2, SM4 là số liệu sau giảm trừ.')
    doc.add_paragraph()

    df_sm2 = unit_data['c11_sm2']
    df_sm4 = unit_data['c11_sm4']

    # Lấy danh sách đơn vị
    team_order = TEAM_ORDER

    # Tạo bảng
    headers = ['Đơn vị', 'SC Chủ động (SM1)', 'Đạt (SM2)', 'TL SC CĐ (%)',
               'Báo hỏng (SM3)', 'Đạt ĐH (SM4)', 'TL SCBH (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header - sử dụng màu xanh lá đậm hơn để phân biệt với bảng gốc
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1B5E20')  # Xanh lá đậm
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # Xử lý từng đơn vị
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = _get_short_name(don_vi)

        # Lấy dữ liệu SM2 (sửa chữa chủ động)
        sm2_row = _match_unit(df_sm2, 'Đơn vị', don_vi)
        if sm2_row.empty:
            sm1 = 0
            sm2 = 0
            tyle_sm2 = 0
        else:
            sm2_row = sm2_row.iloc[0]
            sm1 = sm2_row.get('Tổng phiếu (Sau GT)', 0)  # SAU GIẢM TRỪ
            sm2 = sm2_row.get('Phiếu đạt (Sau GT)', 0)
            tyle_sm2 = sm2_row.get('Tỷ lệ % (Sau GT)', 0)
            if pd.notna(tyle_sm2) and tyle_sm2 > 1:
                tyle_sm2 = tyle_sm2 / 100

        # Lấy dữ liệu SM4 (sửa chữa báo hỏng)
        sm4_row = _match_unit(df_sm4, 'Đơn vị', don_vi)
        if sm4_row.empty:
            sm3 = 0
            sm4 = 0
            tyle_sm4 = 0
        else:
            sm4_row = sm4_row.iloc[0]
            sm3 = sm4_row.get('Tổng phiếu (Sau GT)', 0)  # SAU GIẢM TRỪ
            sm4 = sm4_row.get('Phiếu đạt (Sau GT)', 0)
            tyle_sm4 = sm4_row.get('Tỷ lệ % (Sau GT)', 0)
            if pd.notna(tyle_sm4) and tyle_sm4 > 1:
                tyle_sm4 = tyle_sm4 / 100

        # Tính điểm BSC
        diem_tp1 = tinh_diem_C11_TP1(tyle_sm2)
        diem_tp2 = tinh_diem_C11_TP2(tyle_sm4)
        diem_bsc = 0.30 * diem_tp1 + 0.70 * diem_tp2

        data = [
            short_name,
            str(int(sm1)) if pd.notna(sm1) else '0',
            str(int(sm2)) if pd.notna(sm2) else '0',
            format_number(tyle_sm2 * 100 if pd.notna(tyle_sm2) else 0),
            str(int(sm3)) if pd.notna(sm3) else '0',
            str(int(sm4)) if pd.notna(sm4) else '0',
            format_number(tyle_sm4 * 100 if pd.notna(tyle_sm4) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')

    # Thêm dòng tổng (TTVT Sơn Tây)
    cells = table.add_row().cells

    # Lấy dữ liệu tổng từ SM2
    sm2_tong = _match_unit(df_sm2, 'Đơn vị', 'TTVT Sơn Tây')
    if sm2_tong.empty:
        sm1_tong = 0
        sm2_tong_dat = 0
        tyle_sm2_tong = 0
    else:
        sm2_tong = sm2_tong.iloc[0]
        sm1_tong = sm2_tong.get('Tổng phiếu (Sau GT)', 0)  # SAU GIẢM TRỪ
        sm2_tong_dat = sm2_tong.get('Phiếu đạt (Sau GT)', 0)
        tyle_sm2_tong = sm2_tong.get('Tỷ lệ % (Sau GT)', 0)
        if pd.notna(tyle_sm2_tong) and tyle_sm2_tong > 1:
            tyle_sm2_tong = tyle_sm2_tong / 100

    # Lấy dữ liệu tổng từ SM4
    sm4_tong = _match_unit(df_sm4, 'Đơn vị', 'TTVT Sơn Tây')
    if sm4_tong.empty:
        sm3_tong = 0
        sm4_tong_dat = 0
        tyle_sm4_tong = 0
    else:
        sm4_tong = sm4_tong.iloc[0]
        sm3_tong = sm4_tong.get('Tổng phiếu (Sau GT)', 0)  # SAU GIẢM TRỪ
        sm4_tong_dat = sm4_tong.get('Phiếu đạt (Sau GT)', 0)
        tyle_sm4_tong = sm4_tong.get('Tỷ lệ % (Sau GT)', 0)
        if pd.notna(tyle_sm4_tong) and tyle_sm4_tong > 1:
            tyle_sm4_tong = tyle_sm4_tong / 100

    # Tính điểm BSC tổng
    diem_tp1_tong = tinh_diem_C11_TP1(tyle_sm2_tong)
    diem_tp2_tong = tinh_diem_C11_TP2(tyle_sm4_tong)
    diem_bsc_tong = 0.30 * diem_tp1_tong + 0.70 * diem_tp2_tong

    data_tong = [
        'TTVT Sơn Tây',
        str(int(sm1_tong)) if pd.notna(sm1_tong) else '0',
        str(int(sm2_tong_dat)) if pd.notna(sm2_tong_dat) else '0',
        format_number(tyle_sm2_tong * 100 if pd.notna(tyle_sm2_tong) else 0),
        str(int(sm3_tong)) if pd.notna(sm3_tong) else '0',
        str(int(sm4_tong_dat)) if pd.notna(sm4_tong_dat) else '0',
        format_number(tyle_sm4_tong * 100 if pd.notna(tyle_sm4_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], 'A5D6A7')  # Xanh lá đậm hơn cho dòng tổng

    doc.add_paragraph()


def add_c12_unit_level_exclusion_table(doc, unit_data, c1x_reports=None):
    """
    Thêm bảng C1.2 tổng hợp theo đơn vị (cấp tổ) sau giảm trừ
    Dữ liệu từ: So_sanh_C12_SM1.xlsx (TP1) và SM4-C12-ti-le-su-co-dv-brcd.xlsx (TP2)
    """
    if not unit_data or 'c12_sm1' not in unit_data or 'c12_sm4' not in unit_data:
        return

    doc.add_heading('C1.2 - Tỷ lệ báo hỏng lặp lại & Tỷ lệ sự cố dịch vụ (sau giảm trừ)', level=3)

    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 GHI CHÚ: ').bold = True
    p.add_run('Tất cả số liệu đã được giảm trừ, lấy từ các file so sánh.')
    doc.add_paragraph()

    df_sm1 = unit_data['c12_sm1']  # So_sanh_C12_SM1.xlsx
    df_sm4 = unit_data['c12_sm4']  # SM4-C12-ti-le-su-co-dv-brcd.xlsx

    team_order = TEAM_ORDER + ['TTVT Sơn Tây']

    # Tạo bảng
    headers = ['Đơn vị', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)',
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0D47A1')  # Xanh dương đậm
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # Xử lý từng đơn vị
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = _get_short_name(don_vi)

        # Lấy dữ liệu SM1, SM2 (hỏng lặp lại sau giảm trừ) từ So_sanh_C12_SM1.xlsx
        sm1_row = _match_unit(df_sm1, 'Đơn vị', don_vi)
        if sm1_row.empty:
            sm1 = 0
            sm2 = 0  # BH (SM2) = Phiếu báo hỏng (Sau GT)
            tyle_hll = 0
        else:
            sm1_row = sm1_row.iloc[0]
            sm1 = sm1_row.get('Phiếu HLL (Sau GT)', 0)
            sm2 = sm1_row.get('Phiếu báo hỏng (Sau GT)', 0)  # SAU GIẢM TRỪ
            tyle_hll = sm1_row.get('Tỷ lệ HLL % (Sau GT)', 0)
            if pd.notna(tyle_hll) and tyle_hll > 1:
                tyle_hll = tyle_hll / 100

        # Lấy dữ liệu TP2 (SM3, SM4) từ SM4-C12-ti-le-su-co-dv-brcd.xlsx (cột Sau GT)
        sm4_row = _match_unit(df_sm4, 'Đơn vị', don_vi)
        if sm4_row.empty:
            sm3 = 0
            sm4 = 0
            tyle_sc = 0
        else:
            sm4_row = sm4_row.iloc[0]
            sm3 = sm4_row.get('Phiếu báo hỏng (Sau GT)', 0)       # BH SC
            sm4 = sm4_row.get('Tổng TB (Sau GT)', 0)              # TB
            tyle_sc = sm4_row.get('Tỷ lệ báo hỏng % (Sau GT)', 0) # TL SC (%)
            # Chuyển đổi sang thập phân nếu cần
            if pd.notna(tyle_sc) and tyle_sc > 1:
                tyle_sc = tyle_sc / 100

        # Tính điểm BSC
        diem_tp1 = tinh_diem_C12_TP1(tyle_hll)
        diem_tp2 = tinh_diem_C12_TP2(tyle_sc)
        diem_bsc = 0.50 * diem_tp1 + 0.50 * diem_tp2

        data = [
            short_name,
            str(int(sm1)) if pd.notna(sm1) else '0',
            str(int(sm2)) if pd.notna(sm2) else '0',
            format_number(tyle_hll * 100 if pd.notna(tyle_hll) else 0),
            str(int(sm3)) if pd.notna(sm3) else '0',
            str(int(sm4)) if pd.notna(sm4) else '0',
            format_number(tyle_sc * 100 if pd.notna(tyle_sc) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            # Tô đậm và nền xanh cho dòng TTVT
            if don_vi == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
            elif idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')

    doc.add_paragraph()


def add_c14_unit_level_exclusion_table(doc, unit_data):
    """
    Thêm bảng C1.4 tổng hợp theo đơn vị (cấp tổ) sau giảm trừ

    Args:
        doc: Document Word
        unit_data: Dictionary từ load_unit_level_exclusion_data()
    """
    if not unit_data or 'c14' not in unit_data:
        return

    doc.add_heading('C1.4 - Độ hài lòng khách hàng sau sửa chữa (sau giảm trừ)', level=3)

    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 GHI CHÚ: ').bold = True
    p.add_run('Số liệu sau khi loại bỏ các phiếu khảo sát thuộc diện giảm trừ.')
    doc.add_paragraph()

    df_c14 = unit_data['c14']

    team_order = TEAM_ORDER

    # Tạo bảng
    headers = ['Đơn vị', 'Tổng phiếu', 'KH hài lòng', 'KH không HL', 'TL HL (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E65100')  # Cam đậm
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # Xử lý từng đơn vị
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = _get_short_name(don_vi)

        # Lấy dữ liệu C1.4
        c14_row = _match_unit(df_c14, 'Đơn vị', don_vi)
        if c14_row.empty:
            tong_phieu = 0
            phieu_khl = 0
            tyle_hl = 0
        else:
            c14_row = c14_row.iloc[0]
            tong_phieu = c14_row.get('Tổng phiếu (Sau GT)', 0)
            phieu_khl = c14_row.get('Phiếu KHL (Sau GT)', 0)
            tyle_hl = c14_row.get('Tỷ lệ HL % (Sau GT)', 0)
            if pd.notna(tyle_hl) and tyle_hl > 1:
                tyle_hl = tyle_hl / 100

        phieu_hl = tong_phieu - phieu_khl if pd.notna(tong_phieu) and pd.notna(phieu_khl) else 0
        diem_bsc = tinh_diem_C14(tyle_hl)

        data = [
            short_name,
            str(int(tong_phieu)) if pd.notna(tong_phieu) else '0',
            str(int(phieu_hl)) if pd.notna(phieu_hl) else '0',
            str(int(phieu_khl)) if pd.notna(phieu_khl) else '0',
            format_number(tyle_hl * 100 if pd.notna(tyle_hl) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFE0B2')

    # Thêm dòng tổng (TTVT Sơn Tây)
    cells = table.add_row().cells

    c14_tong_row = _match_unit(df_c14, 'Đơn vị', 'TTVT Sơn Tây')
    if c14_tong_row.empty:
        tong_phieu_tong = 0
        phieu_khl_tong = 0
        tyle_hl_tong = 0
    else:
        c14_tong_row = c14_tong_row.iloc[0]
        tong_phieu_tong = c14_tong_row.get('Tổng phiếu (Sau GT)', 0)
        phieu_khl_tong = c14_tong_row.get('Phiếu KHL (Sau GT)', 0)
        tyle_hl_tong = c14_tong_row.get('Tỷ lệ HL % (Sau GT)', 0)
        if pd.notna(tyle_hl_tong) and tyle_hl_tong > 1:
            tyle_hl_tong = tyle_hl_tong / 100

    phieu_hl_tong = tong_phieu_tong - phieu_khl_tong if pd.notna(tong_phieu_tong) and pd.notna(phieu_khl_tong) else 0
    diem_bsc_tong = tinh_diem_C14(tyle_hl_tong)

    data_tong = [
        'TTVT Sơn Tây',
        str(int(tong_phieu_tong)) if pd.notna(tong_phieu_tong) else '0',
        str(int(phieu_hl_tong)) if pd.notna(phieu_hl_tong) else '0',
        str(int(phieu_khl_tong)) if pd.notna(phieu_khl_tong) else '0',
        format_number(tyle_hl_tong * 100 if pd.notna(tyle_hl_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], 'FFCC80')

    doc.add_paragraph()


def add_c15_unit_level_exclusion_table(doc, unit_data):
    """
    Thêm bảng C1.5 tổng hợp theo đơn vị (cấp tổ) sau giảm trừ

    Args:
        doc: Document Word
        unit_data: Dictionary từ load_unit_level_exclusion_data()
    """
    if not unit_data or 'c15' not in unit_data:
        return

    doc.add_heading('C1.5 - Tỷ lệ thiết lập dịch vụ đạt thời gian quy định (sau giảm trừ)', level=3)

    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 GHI CHÚ: ').bold = True
    p.add_run('Số liệu sau khi loại bỏ các phiếu lắp đặt thuộc diện giảm trừ.')
    doc.add_paragraph()

    df_c15 = unit_data['c15']

    team_order = TEAM_ORDER

    # Tạo bảng
    headers = ['Đơn vị', 'Phiếu đạt', 'Phiếu không đạt', 'Tổng phiếu', 'Tỷ lệ đạt (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '00695C')  # Xanh ngọc đậm
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # Xử lý từng đơn vị
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = _get_short_name(don_vi)

        # Lấy dữ liệu C1.5
        c15_row = _match_unit(df_c15, 'Đơn vị', don_vi)
        if c15_row.empty:
            phieu_dat = 0
            tong_phieu = 0
            tyle_dat = 0
        else:
            c15_row = c15_row.iloc[0]
            phieu_dat = c15_row.get('Phiếu đạt (Sau GT)', 0)
            tong_phieu = c15_row.get('Tổng phiếu (Sau GT)', 0)
            tyle_dat = c15_row.get('Tỷ lệ đạt % (Sau GT)', 0)
            if pd.notna(tyle_dat) and tyle_dat > 1:
                tyle_dat = tyle_dat / 100

        phieu_ko_dat = tong_phieu - phieu_dat if pd.notna(tong_phieu) and pd.notna(phieu_dat) else 0
        diem_bsc = tinh_diem_C15(tyle_dat)

        data = [
            short_name,
            str(int(phieu_dat)) if pd.notna(phieu_dat) else '0',
            str(int(phieu_ko_dat)) if pd.notna(phieu_ko_dat) else '0',
            str(int(tong_phieu)) if pd.notna(tong_phieu) else '0',
            format_number(tyle_dat * 100 if pd.notna(tyle_dat) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'B2DFDB')

    # Thêm dòng tổng (TTVT Sơn Tây)
    cells = table.add_row().cells

    c15_tong_row = _match_unit(df_c15, 'Đơn vị', 'TTVT Sơn Tây')
    if c15_tong_row.empty:
        phieu_dat_tong = 0
        tong_phieu_tong = 0
        tyle_dat_tong = 0
    else:
        c15_tong_row = c15_tong_row.iloc[0]
        phieu_dat_tong = c15_tong_row.get('Phiếu đạt (Sau GT)', 0)
        tong_phieu_tong = c15_tong_row.get('Tổng phiếu (Sau GT)', 0)
        tyle_dat_tong = c15_tong_row.get('Tỷ lệ đạt % (Sau GT)', 0)
        if pd.notna(tyle_dat_tong) and tyle_dat_tong > 1:
            tyle_dat_tong = tyle_dat_tong / 100

    phieu_ko_dat_tong = tong_phieu_tong - phieu_dat_tong if pd.notna(tong_phieu_tong) and pd.notna(phieu_dat_tong) else 0
    diem_bsc_tong = tinh_diem_C15(tyle_dat_tong)

    data_tong = [
        'TTVT Sơn Tây',
        str(int(phieu_dat_tong)) if pd.notna(phieu_dat_tong) else '0',
        str(int(phieu_ko_dat_tong)) if pd.notna(phieu_ko_dat_tong) else '0',
        str(int(tong_phieu_tong)) if pd.notna(tong_phieu_tong) else '0',
        format_number(tyle_dat_tong * 100 if pd.notna(tyle_dat_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], '80CBC4')

    doc.add_paragraph()


def add_c11_exclusion_table(doc, comparison_data):
    """
    Thêm bảng C1.1 sau giảm trừ chi tiết theo NVKT (riêng biệt)
    """
    if not comparison_data:
        return

    has_data = False
    
    # C1.1 SM4 - Sửa chữa báo hỏng
    if 'c11_sm4' in comparison_data:
        has_data = True
        doc.add_heading('C1.1 - SAU GIẢM TRỪ (SM4 - Sửa chữa báo hỏng)', level=4)
        df = comparison_data['c11_sm4']['chi_tiet']
        
        headers = ['NVKT', 'Tổng phiếu (Thô)', 'Tổng phiếu (Sau GT)', 
                   'Số phiếu đạt (Thô)', 'Số phiếu đạt (Sau GT)',
                   'Tỷ lệ % (Thô)', 'Tỷ lệ % (Sau GT)', 'Chênh lệch %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'C62828')  # Đỏ đậm
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Chênh lệch %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('Tổng phiếu (Thô)', 0))) if pd.notna(row.get('Tổng phiếu (Thô)')) else '0',
                str(int(row.get('Tổng phiếu (Sau GT)', 0))) if pd.notna(row.get('Tổng phiếu (Sau GT)')) else '0',
                str(int(row.get('Số phiếu đạt (Thô)', 0))) if pd.notna(row.get('Số phiếu đạt (Thô)')) else '0',
                str(int(row.get('Số phiếu đạt (Sau GT)', 0))) if pd.notna(row.get('Số phiếu đạt (Sau GT)')) else '0',
                format_number(row.get('Tỷ lệ % (Thô)', 0)),
                format_number(row.get('Tỷ lệ % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:  # Cột chênh lệch
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FFEBEE')
        
        # Thêm tổng hợp
        if 'tong_hop' in comparison_data['c11_sm4']:
            df_th = comparison_data['c11_sm4']['tong_hop']
            if len(df_th) > 0:
                row_th = df_th.iloc[0]
                p = doc.add_paragraph()
                p.add_run('📊 Tổng hợp C1.1 SM4: ').bold = True
                tyle_tho = row_th.get('Tỷ lệ % (Thô)', 0)
                tyle_sau = row_th.get('Tỷ lệ % (Sau GT)', 0)
                thay_doi = row_th.get('Thay đổi %', 0)
                if pd.isna(thay_doi): thay_doi = 0
                p.add_run(f'Trước: {tyle_tho:.2f}% → Sau: {tyle_sau:.2f}% (Δ: {thay_doi:+.2f}%)')
        
        doc.add_paragraph()
    
    # C1.1 SM2 - Sửa chữa chủ động
    if 'c11_sm2' in comparison_data:
        has_data = True
        doc.add_heading('C1.1 - SAU GIẢM TRỪ (SM2 - Sửa chữa chủ động)', level=4)
        df = comparison_data['c11_sm2']['chi_tiet']
        
        headers = ['NVKT', 'Tổng phiếu (Thô)', 'Tổng phiếu (Sau GT)', 
                   'Số phiếu đạt (Thô)', 'Số phiếu đạt (Sau GT)',
                   'Tỷ lệ % (Thô)', 'Tỷ lệ % (Sau GT)', 'Chênh lệch %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'AD1457')  # Hồng đậm
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Chênh lệch %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('Tổng phiếu (Thô)', 0))) if pd.notna(row.get('Tổng phiếu (Thô)')) else '0',
                str(int(row.get('Tổng phiếu (Sau GT)', 0))) if pd.notna(row.get('Tổng phiếu (Sau GT)')) else '0',
                str(int(row.get('Số phiếu đạt (Thô)', 0))) if pd.notna(row.get('Số phiếu đạt (Thô)')) else '0',
                str(int(row.get('Số phiếu đạt (Sau GT)', 0))) if pd.notna(row.get('Số phiếu đạt (Sau GT)')) else '0',
                format_number(row.get('Tỷ lệ % (Thô)', 0)),
                format_number(row.get('Tỷ lệ % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FCE4EC')
        
        doc.add_paragraph()


def add_c12_exclusion_table(doc, comparison_data):
    """
    Thêm bảng C1.2 sau giảm trừ (riêng biệt)
    """
    if not comparison_data:
        return
    
    # C1.2 SM1 - Hỏng lặp lại
    if 'c12_sm1' in comparison_data:
        doc.add_heading('C1.2 - SAU GIẢM TRỪ (SM1 - Hỏng lặp lại)', level=4)
        df = comparison_data['c12_sm1']['chi_tiet']
        
        headers = ['NVKT', 'Phiếu HLL (Thô)', 'Phiếu HLL (Sau GT)', 
                   'Phiếu BH (Thô)', 'Phiếu BH (Sau GT)',
                   'Tỷ lệ HLL % (Thô)', 'Tỷ lệ HLL % (Sau GT)', 'Chênh lệch %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '0D47A1')  # Xanh dương đậm
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Chênh lệch %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('Số phiếu HLL (Thô)', 0))) if pd.notna(row.get('Số phiếu HLL (Thô)')) else '0',
                str(int(row.get('Số phiếu HLL (Sau GT)', 0))) if pd.notna(row.get('Số phiếu HLL (Sau GT)')) else '0',
                str(int(row.get('Số phiếu báo hỏng (Thô)', 0))) if pd.notna(row.get('Số phiếu báo hỏng (Thô)')) else '0',
                str(int(row.get('Số phiếu báo hỏng (Sau GT)', 0))) if pd.notna(row.get('Số phiếu báo hỏng (Sau GT)')) else '0',
                format_number(row.get('Tỷ lệ HLL % (Thô)', 0)),
                format_number(row.get('Tỷ lệ HLL % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'E3F2FD')
        
        doc.add_paragraph()


def add_c14_exclusion_table(doc, comparison_data):
    """
    Thêm bảng C1.4 sau giảm trừ (riêng biệt)
    """
    if not comparison_data or 'c14' not in comparison_data:
        return
    
    doc.add_heading('C1.4 - SAU GIẢM TRỪ (Độ hài lòng khách hàng)', level=4)
    df = comparison_data['c14']['chi_tiet']
    
    headers = ['NVKT', 'Tổng KS (Thô)', 'KHL (Thô)', 'Tỷ lệ HL % (Thô)',
               'Tổng KS (Sau GT)', 'KHL (Sau GT)', 'Tỷ lệ HL % (Sau GT)', 'Chênh lệch %']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E65100')  # Cam đậm
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        chenh_lech = row.get('Chênh lệch %', 0)
        if pd.isna(chenh_lech):
            chenh_lech = 0
        data = [
            str(row.get('NVKT', '')),
            str(int(row.get('Tổng phiếu KS (Thô)', 0))) if pd.notna(row.get('Tổng phiếu KS (Thô)')) else '0',
            str(int(row.get('Số phiếu KHL (Thô)', 0))) if pd.notna(row.get('Số phiếu KHL (Thô)')) else '0',
            format_number(row.get('Tỷ lệ HL (%) (Thô)', 0)),
            str(int(row.get('Tổng phiếu KS (Sau GT)', 0))) if pd.notna(row.get('Tổng phiếu KS (Sau GT)')) else '0',
            str(int(row.get('Số phiếu KHL (Sau GT)', 0))) if pd.notna(row.get('Số phiếu KHL (Sau GT)')) else '0',
            format_number(row.get('Tỷ lệ HL (%) (Sau GT)', 0)),
            f"{chenh_lech:+.2f}%"
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if i == 7:
                if chenh_lech > 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif chenh_lech < 0:
                    run.font.color.rgb = RGBColor(200, 0, 0)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')
    
    doc.add_paragraph()


def create_exclusion_bar_chart(comparison_data, output_path=None):
    """
    Tạo biểu đồ bar riêng cho dữ liệu sau giảm trừ
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        return None
    
    df = comparison_data['tong_hop']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    
    tyle_sau = df['Tỷ lệ % (Sau GT)'].fillna(0).values
    chi_tieu = df['Chỉ tiêu'].values
    
    # Màu sắc theo mức độ tốt/xấu
    colors = []
    for val in tyle_sau:
        if val >= 95:
            colors.append('#4CAF50')  # Xanh lá - tốt
        elif val >= 90:
            colors.append('#FFC107')  # Vàng - trung bình
        else:
            colors.append('#F44336')  # Đỏ - cần cải thiện
    
    bars = ax.bar(x, tyle_sau, color=colors, alpha=0.8, edgecolor='black', linewidth=0.5)
    
    # Thêm giá trị lên cột
    for bar, val in zip(bars, tyle_sau):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=10, fontweight='bold')
    
    ax.set_xlabel('Chỉ tiêu', fontsize=12)
    ax.set_ylabel('Tỷ lệ (%)', fontsize=12)
    ax.set_title('TỶ LỆ CÁC CHỈ TIÊU SAU GIẢM TRỪ', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(chi_tieu, rotation=15, ha='right', fontsize=10)
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    ax.set_ylim(0, max(tyle_sau) * 1.15 if len(tyle_sau) > 0 else 100)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_c1x_overview_table(doc, c1x_reports, comparison_data=None, unit_data=None, exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Thêm bảng tổng quan chi tiết từ các báo cáo C1.x vào document
    Nếu có comparison_data, sẽ thêm bảng số liệu sau giảm trừ ngay sau bảng thô

    Args:
        doc: Document Word
        c1x_reports: Dictionary chứa các DataFrame từ load_c1x_reports()
        comparison_data: Dictionary chứa dữ liệu so sánh từ load_exclusion_comparison_data()
        unit_data: Dictionary chứa dữ liệu thống kê theo đơn vị từ load_unit_level_exclusion_data()
        exclusion_folder: Thư mục chứa dữ liệu giảm trừ
    """
    doc.add_heading('1.3. Số liệu chi tiết các chỉ tiêu BSC theo Đội/TTVT', level=2)

    # =========================================================================
    # Bảng C1.1 - Tỷ lệ sửa chữa
    # =========================================================================
    # DÙNG DỮ LIỆU TỪ So_sanh_C11_SM2.xlsx (TP1) và So_sanh_C11_SM4.xlsx (TP2)
    doc.add_heading('C1.1 - Tỷ lệ sửa chữa phiếu chất lượng & báo hỏng', level=3)
    
    headers = ['Đơn vị', 'SC Chủ động (SM1)', 'Đạt (SM2)', 'TL SC CĐ (%)', 
               'Báo hỏng (SM3)', 'Đạt ĐH (SM4)', 'TL SCBH (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Đọc dữ liệu từ file so sánh
    teams_order = [t.short_name for t in sorted(BRCD_TEAMS, key=lambda t: t.order) if t.active] + ['TTVT Sơn Tây']

    def get_short_name_c11(don_vi):
        if not don_vi: return None
        result = _get_short_name(str(don_vi))
        if result != str(don_vi):
            return result
        for orig, short in TEAM_SHORT_NAMES.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # Lấy dữ liệu TP1 (SC Chủ động) từ So_sanh_C11_SM2.xlsx
    tp1_data = {}  # short_name -> {sm1, sm2, tyle, diem}
    if unit_data and 'c11_sm2' in unit_data:
        for _, row in unit_data['c11_sm2'].iterrows():
            short_name = get_short_name_c11(row.get('Đơn vị', ''))
            if short_name:
                tp1_data[short_name] = {
                    'sm1': row.get('Tổng phiếu (Thô)', 0),
                    'sm2': row.get('Phiếu đạt (Thô)', 0),
                    'tyle': row.get('Tỷ lệ % (Thô)', 0),
                }
    
    # Lấy dữ liệu TP2 (Báo hỏng) từ So_sanh_C11_SM4.xlsx
    tp2_data = {}  # short_name -> {sm3, sm4, tyle, diem}
    if unit_data and 'c11_sm4' in unit_data:
        for _, row in unit_data['c11_sm4'].iterrows():
            short_name = get_short_name_c11(row.get('Đơn vị', ''))
            if short_name:
                tp2_data[short_name] = {
                    'sm3': row.get('Tổng phiếu (Thô)', 0),   # Báo hỏng = Tổng phiếu
                    'sm4': row.get('Phiếu đạt (Thô)', 0),     # Đạt ĐH = Phiếu đạt
                    'tyle': row.get('Tỷ lệ % (Thô)', 0),
                    'diem': row.get('Điểm BSC (Thô)', 0),
                }
    
    # Lấy điểm BSC tổng hợp từ Tong_hop_Diem_BSC_Don_Vi.xlsx
    bsc_scores_c11 = {}
    bsc_data_c11 = load_bsc_unit_scores_from_comparison(exclusion_folder)
    if bsc_data_c11 and bsc_data_c11.get('units') is not None:
        for _, row in bsc_data_c11['units'].iterrows():
            short_name = get_short_name_c11(row.get('don_vi', ''))
            if short_name:
                bsc_scores_c11[short_name] = row.get('Diem_C1.1 (Trước)', 0)
    
    # Tạo dữ liệu bảng
    for idx, team in enumerate(teams_order, 1):
        cells = table.add_row().cells
        tp1 = tp1_data.get(team, {})
        tp2 = tp2_data.get(team, {})
        bsc_score = bsc_scores_c11.get(team, 0)
        
        data = [
            team,
            str(int(tp1.get('sm1', 0))),
            str(int(tp1.get('sm2', 0))),
            format_number(tp1.get('tyle', 0)),
            str(int(tp2.get('sm3', 0))),
            str(int(tp2.get('sm4', 0))),
            format_number(tp2.get('tyle', 0)),
            format_number(bsc_score)
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')
            # Tô đậm dòng tổng
            if team == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'C8E6C9')
    
    doc.add_paragraph()

    # Thêm bảng C1.1 tổng hợp theo đơn vị (tổ) sau giảm trừ nếu có
    if unit_data:
        add_c11_unit_level_exclusion_table(doc, unit_data, c1x_reports)

    # Thêm bảng C1.1 chi tiết theo NVKT sau giảm trừ nếu có
    # COMMENT: Bỏ bảng chi tiết từng NVKT này vì đã có trong PHẦN 2
    # if comparison_data:
    #     add_c11_exclusion_table(doc, comparison_data)
    # =========================================================================
    # Bảng C1.2 - Tỷ lệ báo hỏng lặp lại & sự cố
    # =========================================================================
    # DÙNG DỮ LIỆU TỪ So_sanh_C12_SM1.xlsx (TP1) và SM4-C12-ti-le-su-co-dv-brcd.xlsx (TP2)
    doc.add_heading('C1.2 - Tỷ lệ báo hỏng lặp lại & Tỷ lệ sự cố dịch vụ', level=3)
    
    headers = ['Đơn vị', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)', 
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Đọc dữ liệu từ file so sánh
    teams_order_c12 = [t.short_name for t in sorted(BRCD_TEAMS, key=lambda t: t.order) if t.active] + ['TTVT Sơn Tây']

    def get_short_name_c12(don_vi):
        if not don_vi: return None
        result = _get_short_name(str(don_vi))
        if result != str(don_vi):
            return result
        for orig, short in TEAM_SHORT_NAMES.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # Lấy dữ liệu TP1 (HLL) từ So_sanh_C12_SM1.xlsx
    tp1_c12_data = {}  # short_name -> {sm1, sm2, tyle}
    if unit_data and 'c12_sm1' in unit_data:
        for _, row in unit_data['c12_sm1'].iterrows():
            short_name = get_short_name_c12(row.get('Đơn vị', ''))
            if short_name:
                tp1_c12_data[short_name] = {
                    'sm1': row.get('Phiếu HLL (Thô)', 0),           # HLL (SM1)
                    'sm2': row.get('Phiếu báo hỏng (Thô)', 0),      # BH (SM2)
                    'tyle': row.get('Tỷ lệ HLL % (Thô)', 0),        # TL HLL (%)
                }
    
    # Lấy dữ liệu TP2 (Sự cố) từ SM4-C12-ti-le-su-co-dv-brcd.xlsx
    tp2_c12_data = {}  # short_name -> {sm3, sm4, tyle}
    if unit_data and 'c12_sm4' in unit_data:
        for _, row in unit_data['c12_sm4'].iterrows():
            short_name = get_short_name_c12(row.get('Đơn vị', ''))
            if short_name:
                tp2_c12_data[short_name] = {
                    'sm3': row.get('Phiếu báo hỏng (Thô)', 0),        # BH SC (SM3)
                    'sm4': row.get('Tổng TB (Thô)', 0),               # TB (SM4)
                    'tyle': row.get('Tỷ lệ báo hỏng % (Thô)', 0),     # TL SC (%)
                }
    
    # Lấy điểm BSC tổng hợp từ Tong_hop_Diem_BSC_Don_Vi.xlsx
    bsc_scores_c12 = {}
    bsc_data_c12 = load_bsc_unit_scores_from_comparison(exclusion_folder)
    if bsc_data_c12 and bsc_data_c12.get('units') is not None:
        for _, row in bsc_data_c12['units'].iterrows():
            short_name = get_short_name_c12(row.get('don_vi', ''))
            if short_name:
                bsc_scores_c12[short_name] = row.get('Diem_C1.2 (Trước)', 0)
    
    # Tạo dữ liệu bảng
    for idx, team in enumerate(teams_order_c12, 1):
        cells = table.add_row().cells
        tp1 = tp1_c12_data.get(team, {})
        tp2 = tp2_c12_data.get(team, {})
        bsc_score = bsc_scores_c12.get(team, 0)
        
        data = [
            team,
            str(int(tp1.get('sm1', 0) or 0)),
            str(int(tp1.get('sm2', 0) or 0)),
            format_number(tp1.get('tyle', 0) or 0),
            str(int(tp2.get('sm3', 0) or 0)),
            str(int(tp2.get('sm4', 0) or 0)),
            format_number(tp2.get('tyle', 0) or 0),
            format_number(bsc_score or 0)
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')
            if team == 'TTVT Sơn Tây':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
    
    doc.add_paragraph()

    # Thêm bảng C1.2 tổng hợp theo đơn vị (tổ) sau giảm trừ nếu có
    if unit_data:
        add_c12_unit_level_exclusion_table(doc, unit_data, c1x_reports)

    # Thêm bảng C1.2 chi tiết theo NVKT sau giảm trừ nếu có
    # COMMENT: Bỏ bảng chi tiết từng NVKT này vì đã có trong PHẦN 2
    # if comparison_data:
    #     add_c12_exclusion_table(doc, comparison_data)
    
    # =========================================================================
    # Bảng C1.3 - Kênh TSL
    # =========================================================================
    if 'c13' in c1x_reports:
        doc.add_heading('C1.3 - Chỉ tiêu kênh thuê leased line (TSL)', level=3)
        df = c1x_reports['c13']
        
        headers = ['Đơn vị', 'SC TSL (SM1)', 'Đạt (SM2)', 'TL SC (%)', 
                   'HLL (SM3)', 'BH (SM4)', 'TL HLL (%)', 'Số TB (SM6)', 'TL SC (%)', 'Điểm BSC']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '6A1B9A')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            short_name = _get_short_name(row['Đơn vị'])
            if short_name == 'Tổng':
                short_name = 'TTVT Sơn Tây'
            data = [
                short_name,
                str(int(row.get('SM1', 0))),
                str(int(row.get('SM2', 0))),
                format_number(row.get('Tỷ lệ sửa chữa dịch vụ kênh TSL hoàn thành đúng thời gian quy định', 0)),
                str(int(row.get('SM3', 0))),
                str(int(row.get('SM4', 0))),
                format_number(row.get('Tỷ lệ thuê bao báo hỏng dịch vụ kênh TSL lặp lại', 0)),
                str(int(row.get('SM6', 0))),
                format_number(row.get('Tỷ lệ sự cố dịch vụ kênh TSL', 0)),
                format_number(row.get('Chỉ tiêu BSC', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'F3E5F5')
                if short_name == 'TTVT Sơn Tây':
                    run.font.bold = True
                    set_cell_shading(cells[i], 'E1BEE7')
        
        doc.add_paragraph()
    
    # =========================================================================
    # Bảng C1.4 - Hài lòng khách hàng
    # =========================================================================
    if 'c14' in c1x_reports:
        doc.add_heading('C1.4 - Độ hài lòng khách hàng sau sửa chữa', level=3)
        df = c1x_reports['c14']
        
        headers = ['Đơn vị', 'Tổng phiếu', 'Đã KS', 'KS TC', 'KH HL', 
                   'KHL KT PV', 'TL HL PV (%)', 'TL KH HL (%)', 'Điểm BSC']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'F57C00')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            short_name = _get_short_name(row['Đơn vị'])
            if short_name == 'Tổng':
                short_name = 'TTVT Sơn Tây'
            data = [
                short_name,
                str(int(row.get('Tổng phiếu', 0))),
                str(int(row.get('SL đã KS', 0))),
                str(int(row.get('SL KS thành công', 0))),
                str(int(row.get('SL KH hài lòng', 0))),
                str(int(row.get('Không HL KT phục vụ', 0))),
                format_number(row.get('Tỷ lệ HL KT phục vụ', 0)),
                format_number(row.get('Tỷ lệ KH hài lòng', 0)),
                format_number(row.get('Điểm BSC', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FFF3E0')
                if short_name == 'TTVT Sơn Tây':
                    run.font.bold = True
                    set_cell_shading(cells[i], 'FFE0B2')
        
        doc.add_paragraph()

    # Thêm bảng C1.4 tổng hợp theo đơn vị (tổ) sau giảm trừ nếu có
    if unit_data:
        add_c14_unit_level_exclusion_table(doc, unit_data)

    # Thêm bảng C1.4 chi tiết theo NVKT sau giảm trừ nếu có
    # COMMENT: Bỏ bảng chi tiết từng NVKT này vì đã có trong PHẦN 2
    # if comparison_data:
    #     add_c14_exclusion_table(doc, comparison_data)
    
    # =========================================================================
    # Bảng C1.5 - Tỷ lệ thiết lập dịch vụ đạt
    # =========================================================================
    if 'c15_ttvtst' in c1x_reports:
        doc.add_heading('C1.5 - Tỷ lệ thiết lập dịch vụ đạt thời gian quy định', level=3)
        df = c1x_reports['c15_ttvtst']
        
        headers = ['Đơn vị', 'Phiếu đạt', 'Phiếu không đạt', 'Tổng HC', 'Tỉ lệ đạt (%)']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '00796B')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            don_vi = row.get('DOIVT', '')
            short_name = _get_short_name(don_vi)
            data = [
                short_name,
                str(int(row.get('Phiếu đạt', 0))),
                str(int(row.get('Phiếu không đạt', 0))),
                str(int(row.get('Tổng Hoàn công', 0))),
                format_number(row.get('Tỉ lệ đạt (%)', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(10)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'E0F2F1')
                # Tô đậm dòng TTVT Sơn Tây 
                if 'TTVT' in don_vi:
                    run.font.bold = True
                    set_cell_shading(cells[i], 'B2DFDB')

        doc.add_paragraph()

    # Thêm bảng C1.5 tổng hợp theo đơn vị (tổ) sau giảm trừ nếu có
    if unit_data:
        add_c15_unit_level_exclusion_table(doc, unit_data)


# =============================================================================
# HÀM TẠO BIỂU ĐỒ
# =============================================================================
def create_team_comparison_chart(c1x_reports, output_path=None, bsc_data=None):
    """
    Tạo biểu đồ so sánh điểm BSC thực tế giữa 4 tổ
    
    Args:
        c1x_reports: Dictionary chứa các DataFrame từ load_c1x_reports()
        output_path: Đường dẫn lưu file ảnh (None = trả về bytes)
        bsc_data: Dictionary từ load_bsc_unit_scores_from_comparison() (ưu tiên sử dụng)
    
    Returns:
        bytes hoặc str: Dữ liệu ảnh hoặc đường dẫn file
    """
    # Chuẩn bị dữ liệu từ các báo cáo C1.x
    teams_order = ['Phúc Thọ', 'Quảng Oai', 'Suối Hai', 'Sơn Tây']
    
    # Khởi tạo dict chứa điểm BSC (bao gồm C1.5)
    bsc_scores = {team: {'C1.1': 0, 'C1.2': 0, 'C1.3': 0, 'C1.4': 0, 'C1.5': 0} for team in teams_order}
    
    # Map tên đơn vị
    def get_short_name(don_vi):
        if not don_vi: return None
        result = _get_short_name(str(don_vi))
        if result != str(don_vi):
            return result
        for orig, short in TEAM_SHORT_NAMES.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return _get_short_name(don_vi)

    # ƯU TIÊN: Lấy điểm từ bsc_data (Tong_hop_Diem_BSC_Don_Vi.xlsx) - cột (Trước)
    if bsc_data and bsc_data.get('units') is not None and not bsc_data['units'].empty:
        print("  📊 Biểu đồ: Sử dụng điểm từ Tong_hop_Diem_BSC_Don_Vi.xlsx (cột Trước)")
        for _, row in bsc_data['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.1'] = row.get('Diem_C1.1 (Trước)', 0) or 0
                bsc_scores[short_name]['C1.2'] = row.get('Diem_C1.2 (Trước)', 0) or 0
                bsc_scores[short_name]['C1.4'] = row.get('Diem_C1.4 (Trước)', 0) or 0
                bsc_scores[short_name]['C1.5'] = row.get('Diem_C1.5 (Trước)', 0) or 0
        # C1.3 vẫn lấy từ c1x_reports (không có trong comparison)
        if c1x_reports and 'c13' in c1x_reports:
            for _, row in c1x_reports['c13'].iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Chỉ tiêu BSC', 0)
    else:
        # FALLBACK: Lấy điểm BSC từ c1x_reports (số liệu gốc)
        if 'c11' in c1x_reports:
            df = c1x_reports['c11']
            for _, row in df.iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = _get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.1'] = row.get('Chỉ tiêu BSC', 0)
        
        if 'c12' in c1x_reports:
            df = c1x_reports['c12']
            for _, row in df.iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = _get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.2'] = row.get('Chỉ tiêu BSC', 0)
        
        if 'c13' in c1x_reports:
            df = c1x_reports['c13']
            for _, row in df.iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = _get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Chỉ tiêu BSC', 0)
        
        if 'c14' in c1x_reports:
            df = c1x_reports['c14']
            for _, row in df.iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = _get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.4'] = row.get('Điểm BSC', 0)
        
        if 'c15_ttvtst' in c1x_reports:
            df = c1x_reports['c15_ttvtst']
            for _, row in df.iterrows():
                don_vi = row.get('DOIVT', '')
                short_name = _get_short_name(don_vi)
                if short_name in teams_order:
                    ty_le = row.get('Tỉ lệ đạt (%)', 0)
                    if ty_le >= 99.5:
                        diem_bsc = 5.0
                    elif ty_le <= 89.5:
                        diem_bsc = 1.0
                    else:
                        diem_bsc = 1 + 4 * (ty_le - 89.5) / 10
                    bsc_scores[short_name]['C1.5'] = round(diem_bsc, 2)
    
    # Tạo DataFrame từ dữ liệu
    chart_data = pd.DataFrame(bsc_scores).T
    chart_data = chart_data.reindex(teams_order)  # Đảm bảo thứ tự
    
    # Tạo biểu đồ
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(teams_order))
    width = 0.15  # Thu hẹp để có chỗ cho 5 cột
    
    metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    
    for i, metric in enumerate(metrics):
        values = chart_data[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
        # Thêm giá trị lên cột
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                       f'{val:.2f}', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Tổ Kỹ thuật', fontsize=12)
    ax.set_ylabel('Điểm BSC', fontsize=12)
    ax.set_title('SO SÁNH ĐIỂM BSC THỰC TẾ GIỮA CÁC TỔ', fontsize=14, fontweight='bold')
    ax.set_xticks(x + width * 2)  # Điều chỉnh vị trí label
    ax.set_xticklabels(teams_order, fontsize=11)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    # Lưu hoặc trả về bytes
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_team_bsc_after_exclusion_chart(unit_data, c1x_reports=None, output_path=None, bsc_data=None):
    """
    Tạo biểu đồ so sánh điểm BSC SAU GIẢM TRỪ giữa 4 tổ
    
    Args:
        unit_data: Dữ liệu unit_level từ load_unit_level_exclusion_data()
        c1x_reports: Dữ liệu c1x_reports (fallback)
        output_path: Đường dẫn lưu file
        bsc_data: Dictionary từ load_bsc_unit_scores_from_comparison() (ưu tiên sử dụng)
    """
    teams_order = ['Phúc Thọ', 'Quảng Oai', 'Suối Hai', 'Sơn Tây']
    
    # Khởi tạo dict chứa điểm BSC
    bsc_scores = {team: {'C1.1': 0, 'C1.2': 0, 'C1.3': 0, 'C1.4': 0, 'C1.5': 0} for team in teams_order}
    
    # Mapping tên đội trong Excel -> tên ngắn
    def get_short_name(don_vi):
        if not don_vi: return None
        result = _get_short_name(str(don_vi))
        if result != str(don_vi):
            return result
        for orig, short in TEAM_SHORT_NAMES.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # ================================================================
    # ƯU TIÊN: Lấy điểm từ bsc_data (Tong_hop_Diem_BSC_Don_Vi.xlsx) - cột (Sau)
    # ================================================================
    if bsc_data and bsc_data.get('units') is not None and not bsc_data['units'].empty:
        print("  📊 Biểu đồ sau GT: Sử dụng điểm từ Tong_hop_Diem_BSC_Don_Vi.xlsx (cột Sau)")
        for _, row in bsc_data['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.1'] = row.get('Diem_C1.1 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.2'] = row.get('Diem_C1.2 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.4'] = row.get('Diem_C1.4 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.5'] = row.get('Diem_C1.5 (Sau)', 0) or 0
        # C1.3 vẫn lấy từ c1x_reports (không có giảm trừ)
        if c1x_reports and 'c13' in c1x_reports:
            for _, row in c1x_reports['c13'].iterrows():
                don_vi = row.get('Đơn vị', '')
                short_name = get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Chỉ tiêu BSC', 0)
        
        # Tạo DataFrame và biểu đồ rồi return (không chạy tiếp phần tính từ unit_data)
        chart_data = pd.DataFrame(bsc_scores).T
        chart_data = chart_data.reindex(teams_order)
        
        fig, ax = plt.subplots(figsize=(14, 6))
        x = np.arange(len(teams_order))
        width = 0.15
        metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
        
        for i, metric in enumerate(metrics):
            values = chart_data[metric].fillna(0).values
            bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
            for bar, val in zip(bars, values):
                if val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                           f'{val:.2f}', ha='center', va='bottom', fontsize=8)
        
        ax.set_xlabel('Tổ Kỹ thuật', fontsize=12)
        ax.set_ylabel('Điểm BSC', fontsize=12)
        ax.set_title('ĐIỂM BSC SAU GIẢM TRỪ GIỮA CÁC TỔ', fontsize=14, fontweight='bold')
        ax.set_xticks(x + width * 2)
        ax.set_xticklabels(teams_order, fontsize=11)
        ax.set_ylim(0, 6)
        ax.legend(loc='upper right')
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        
        if output_path:
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            return output_path
        else:
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            buf.seek(0)
            return buf
    
    # FALLBACK: Tính từ unit_data (cách cũ) - chỉ chạy nếu không có bsc_data
    # Sử dụng scoring functions từ kpi_scoring (imported ở top level)

    # ================================================================
    # Tính C1.1 = 0.30*TP1 + 0.70*TP2
    # ================================================================
    # TP1 từ c11_sm2 (Sửa chữa chủ động), TP2 từ c11_sm4 (Sửa chữa BH)
    c11_tp1 = {}  # team -> tỷ lệ thập phân
    c11_tp2 = {}  # team -> tỷ lệ thập phân
    
    if unit_data and 'c11_sm2' in unit_data:
        df = unit_data['c11_sm2']
        for _, row in df.iterrows():
            short = get_short_name(row.get('Đơn vị', ''))
            if short and short in teams_order:
                tyle = row.get('Tỷ lệ % (Sau GT)', 0) or 0
                c11_tp1[short] = tyle / 100 if tyle > 1 else tyle  # Chuyển về thập phân
    
    if unit_data and 'c11_sm4' in unit_data:
        df = unit_data['c11_sm4']
        for _, row in df.iterrows():
            short = get_short_name(row.get('Đơn vị', ''))
            if short and short in teams_order:
                tyle = row.get('Tỷ lệ % (Sau GT)', 0) or 0
                c11_tp2[short] = tyle / 100 if tyle > 1 else tyle
    
    for team in teams_order:
        tp1 = c11_tp1.get(team)
        tp2 = c11_tp2.get(team)
        diem_tp1 = tinh_diem_C11_TP1(tp1)
        diem_tp2 = tinh_diem_C11_TP2(tp2)
        bsc_scores[team]['C1.1'] = round(0.30 * diem_tp1 + 0.70 * diem_tp2, 2)
    
    # ================================================================
    # Tính C1.2 = 0.50*TP1 + 0.50*TP2
    # ================================================================
    # TP1 từ c12_sm1 (HLL), TP2 từ báo cáo gốc (không có trong unit_data chưa)
    c12_tp1 = {}  # team -> tỷ lệ HLL thập phân
    
    if unit_data and 'c12_sm1' in unit_data:
        df = unit_data['c12_sm1']
        for _, row in df.iterrows():
            short = get_short_name(row.get('Đơn vị', ''))
            if short and short in teams_order:
                tyle_col = [c for c in df.columns if 'Tỷ lệ' in c and 'Sau GT' in c]
                if tyle_col:
                    tyle = row.get(tyle_col[0], 0) or 0
                    c12_tp1[short] = tyle / 100 if tyle > 1 else tyle
    
    
    # C1.2 TP2 - Tỷ lệ sự cố BRCĐ từ unit_data['c12_sm4'] (SAU GIẢM TRỪ)
    c12_tp2 = {}
    
    # Hàm tính điểm TP2 từ tỷ lệ sự cố (‰)
    def tinh_diem_tp2_from_percentage(tyle_percent):
        """Tính điểm TP2 từ tỷ lệ % - chuyển sang ‰ trước"""
        if pd.isna(tyle_percent) or tyle_percent is None:
            return 5
        # Chuyển % sang ‰: 1.76% = 17.6‰
        tyle_permil = tyle_percent * 10
        if tyle_permil <= 15:  # ≤1.5%
            return 5
        elif tyle_permil < 25:  # <2.5%
            return 5 - 4 * (tyle_permil - 15) / 10
        else:
            return 1
    
    if unit_data and 'c12_sm4' in unit_data:
        # Đọc từ file SM4-C12-ti-le-su-co-dv-brcd.xlsx
        df = unit_data['c12_sm4']
        for _, row in df.iterrows():
            short = get_short_name(row.get('TEN_DOI', ''))
            if short and short in teams_order:
                tyle = row.get('Tỷ lệ báo hỏng (%) (Sau GT)', 0) or 0
                # Tính điểm từ tỷ lệ
                diem_tp2 = tinh_diem_tp2_from_percentage(tyle)
                c12_tp2[short] = diem_tp2
    elif c1x_reports and 'c12' in c1x_reports:
        # Fallback: dùng dữ liệu gốc nếu không có dữ liệu sau giảm trừ
        df = c1x_reports['c12']
        for _, row in df.iterrows():
            don_vi = row.get('Đơn vị', '')
            short_name = _get_short_name(don_vi)
            if short_name in teams_order:
                diem_tp2 = row.get('Điểm C1.2 TP2', 5)
                c12_tp2[short_name] = diem_tp2
    
    for team in teams_order:
        tp1 = c12_tp1.get(team)
        diem_tp1 = tinh_diem_C12_TP1(tp1)
        diem_tp2 = c12_tp2.get(team, 5)  # Default 5 nếu không có dữ liệu
        bsc_scores[team]['C1.2'] = round(0.50 * diem_tp1 + 0.50 * diem_tp2, 2)
    
    # ================================================================
    # C1.3 - giữ nguyên từ c1x_reports (không có giảm trừ)
    # ================================================================
    if c1x_reports and 'c13' in c1x_reports:
        df = c1x_reports['c13']
        for _, row in df.iterrows():
            don_vi = row.get('Đơn vị', '')
            short_name = _get_short_name(don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.3'] = row.get('Chỉ tiêu BSC', 0) or 0
    
    # ================================================================
    # C1.4 - sau giảm trừ
    # ================================================================
    if unit_data and 'c14' in unit_data:
        df = unit_data['c14']
        for _, row in df.iterrows():
            short = get_short_name(row.get('Đơn vị', ''))
            if short and short in teams_order:
                tyle_col = [c for c in df.columns if 'Tỷ lệ HL' in c and 'Sau GT' in c]
                if tyle_col:
                    tyle = row.get(tyle_col[0], 0) or 0
                    tyle_dec = tyle / 100 if tyle > 1 else tyle
                    bsc_scores[short]['C1.4'] = round(tinh_diem_C14(tyle_dec), 2)
    
    # ================================================================
    # C1.5 - SỬ DỤNG DỮ LIỆU SAU GIẢM TRỪ từ unit_data
    # ================================================================
    if unit_data and 'c15' in unit_data:
        # Sử dụng dữ liệu SAU GIẢM TRỪ từ file So_sanh_C15.xlsx
        df = unit_data['c15']
        for _, row in df.iterrows():
            don_vi = row.get('Đơn vị', '')
            short_name = _get_short_name(don_vi)
            if short_name in teams_order:
                ty_le = row.get('Tỷ lệ đạt % (Sau GT)', 0) or 0
                ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                bsc_scores[short_name]['C1.5'] = round(tinh_diem_C15(ty_le_dec), 2)
    elif c1x_reports and 'c15_ttvtst' in c1x_reports:
        # Fallback: nếu không có dữ liệu sau giảm trừ, dùng dữ liệu gốc
        df = c1x_reports['c15_ttvtst']
        for _, row in df.iterrows():
            don_vi = row.get('DOIVT', '')
            short_name = _get_short_name(don_vi)
            if short_name in teams_order:
                ty_le = row.get('Tỉ lệ đạt (%)', 0) or 0
                ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                bsc_scores[short_name]['C1.5'] = round(tinh_diem_C15(ty_le_dec), 2)
    
    # Tạo DataFrame từ dữ liệu
    chart_data = pd.DataFrame(bsc_scores).T
    chart_data = chart_data.reindex(teams_order)
    
    # Tạo biểu đồ
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(teams_order))
    width = 0.15
    
    metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    
    for i, metric in enumerate(metrics):
        values = chart_data[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                       f'{val:.2f}', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Tổ Kỹ thuật', fontsize=12)
    ax.set_ylabel('Điểm BSC', fontsize=12)
    ax.set_title('ĐIỂM BSC SAU GIẢM TRỪ GIỮA CÁC TỔ', fontsize=14, fontweight='bold')
    ax.set_xticks(x + width * 2)
    ax.set_xticklabels(teams_order, fontsize=11)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf

def create_nvkt_bar_chart(df_summary, team_name, output_path=None):
    """
    Tạo biểu đồ cột so sánh điểm KPI theo NVKT trong 1 tổ
    
    Args:
        df_summary: DataFrame chứa dữ liệu KPI
        team_name: Tên tổ cần tạo biểu đồ
        output_path: Đường dẫn lưu file (None = trả về bytes)
    
    Returns:
        bytes hoặc str
    """
    # Lọc dữ liệu theo tổ
    df_team = _match_unit(df_summary, 'don_vi', team_name).copy()
    
    if df_team.empty:
        return None
    
    # Sắp xếp theo tên NVKT
    df_team = df_team.sort_values('nvkt')
    
    # Tạo biểu đồ
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(df_team))
    width = 0.2
    
    metrics = ['Diem_C1.1', 'Diem_C1.2', 'Diem_C1.4', 'Diem_C1.5']
    labels = ['C1.1', 'C1.2', 'C1.4', 'C1.5']
    
    for i, (metric, label) in enumerate(zip(metrics, labels)):
        values = df_team[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=label, color=BAR_COLORS[i])
        # Thêm giá trị lên cột
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.03,
                       f'{val:.1f}', ha='center', va='bottom', fontsize=7, rotation=90)
    
    short_name = _get_short_name(team_name)
    ax.set_xlabel('NVKT', fontsize=11)
    ax.set_ylabel('Điểm KPI', fontsize=11)
    ax.set_title(f'ĐIỂM KPI THEO NVKT - {short_name.upper()}', fontsize=13, fontweight='bold')
    ax.set_xticks(x + width * 1.5)
    ax.set_xticklabels(df_team['nvkt'].values, fontsize=8, rotation=45, ha='right')
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


# =============================================================================
# HÀM TẠO BẢNG TRONG WORD
# =============================================================================
def add_kpi_summary_table(doc, df_summary, team_name=None):
    """
    Thêm bảng tổng hợp KPI vào document
    
    Args:
        doc: Document Word
        df_summary: DataFrame dữ liệu
        team_name: Lọc theo tổ (None = tất cả)
    """
    if team_name:
        df = _match_unit(df_summary, 'don_vi', team_name).copy()
    else:
        df = df_summary.copy()
    
    # Sắp xếp
    df = df.sort_values(['don_vi', 'nvkt'])
    
    # Tạo bảng
    headers = ['STT', 'Đơn vị', 'NVKT', 'C1.1', 'C1.2', 'C1.4', 'C1.5']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], '1F4E79')
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        
        short_name = _get_short_name(row['don_vi'])
        
        data = [
            str(idx),
            short_name,
            row['nvkt'],
            format_number(row.get('Diem_C1.1', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            # Tô màu dòng xen kẽ
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F4FD')


def add_c11_detail_table(doc, df_detail, team_name=None):
    """
    Thêm bảng chi tiết C1.1 (Thành phần 1 + Thành phần 2)
    """
    if team_name:
        df = _match_unit(df_detail, 'don_vi', team_name).copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    # Tiêu đề
    doc.add_heading('Chi tiết chỉ tiêu C1.1 - Chất lượng sửa chữa thuê bao BRCĐ', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('TP1 = Sửa chữa chủ động (SCCD ≤72h) | TP2 = Sửa chữa theo báo hỏng (SC BH)')
    
    headers = ['STT', 'NVKT', 'Tổng SCCD', 'Đạt ≤72h', 'TL(%)', 'Điểm TP1',
               'Tổng SC BH', 'Đúng hạn', 'TL(%)', 'Điểm TP2', 'Điểm C1.1']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '2E7D32')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c11_tp1_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp1_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp1', np.nan)),
            format_number(row.get('c11_tp2_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp2_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp2', np.nan)),
            format_number(row.get('Diem_C1.1', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')


def add_c12_detail_table(doc, df_detail, team_name=None):
    """
    Thêm bảng chi tiết C1.2 (Thành phần 1 + Thành phần 2)
    """
    if team_name:
        df = _match_unit(df_detail, 'don_vi', team_name).copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi tiết chỉ tiêu C1.2 - Tỷ lệ thuê bao báo hỏng', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('TP1 = Hỏng lặp (≥2 lần/7 ngày) | TP2 = Tỷ lệ BH/TB quản lý | BH = Báo hỏng | TB QL = Thuê bao quản lý')
    
    headers = ['STT', 'NVKT', 'Hỏng lặp', 'Tổng BH', 'TL(%)', 'Điểm TP1',
               'Phiếu BH', 'TB QL', 'TL(‰)', 'Điểm TP2', 'Điểm C1.2']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '1565C0')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c12_tp1_phieu_hll', np.nan), 0),
            format_number(row.get('c12_tp1_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp1', np.nan)),
            format_number(row.get('c12_tp2_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp2_tong_tb', np.nan), 0),
            format_number(row.get('c12_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp2', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')


def add_c14_detail_table(doc, df_detail, team_name=None):
    """
    Thêm bảng chi tiết C1.4 - Độ hài lòng khách hàng
    """
    if team_name:
        df = _match_unit(df_detail, 'don_vi', team_name).copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi tiết chỉ tiêu C1.4 - Độ hài lòng khách hàng', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('KS = Khảo sát | Không HL = Không hài lòng | HL = Hài lòng')
    
    headers = ['STT', 'NVKT', 'Tổng KS', 'Không HL', 'Tỷ lệ HL (%)', 'Điểm C1.4']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'F57C00')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c14_phieu_ks', np.nan), 0),
            format_number(row.get('c14_phieu_khl', np.nan), 0),
            format_number(row.get('c14_ty_le', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')


def add_c15_detail_table(doc, df_detail, team_name=None):
    """
    Thêm bảng chi tiết C1.5 - Tỷ lệ thiết lập dịch vụ
    """
    if team_name:
        df = _match_unit(df_detail, 'don_vi', team_name).copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi tiết chỉ tiêu C1.5 - Thiết lập dịch vụ BRCĐ đạt thời gian quy định', level=3)
    
    # Chú thích
    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('Đạt TG = Hoàn thành đúng thời gian (ngoài CCCO: ≤24h, trong CCCO: theo quy định)')
    
    headers = ['STT', 'NVKT', 'Đạt TG', 'Không đạt', 'Tổng phiếu', 'Tỷ lệ (%)', 'Điểm C1.5']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '7B1FA2')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Dữ liệu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c15_phieu_dat', np.nan), 0),
            format_number(row.get('c15_phieu_khong_dat', np.nan), 0),
            format_number(row.get('c15_tong_phieu', np.nan), 0),
            format_number(row.get('c15_ty_le', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'F3E5F5')


# =============================================================================
# HÀM SHC CHO BÁO CÁO TỔNG HỢP
# =============================================================================

def load_shc_summary_by_unit(data_folder="downloads/baocao_hanoi"):
    """
    Đọc dữ liệu SHC tổng hợp theo đơn vị từ sheet Xu_huong_theo_don_vi
    
    Returns:
        dict: {'units': {...}, 'dates': [...]} hoặc None
    """
    import glob
    
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    latest_file = max(files, key=os.path.getmtime)
    
    try:
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_don_vi')
        
        # Cột đầu tiên là Đơn vị, các cột còn lại là ngày
        date_columns = [col for col in df.columns if col != 'Đơn vị']
        
        result = {'units': {}, 'dates': date_columns}
        
        for _, row in df.iterrows():
            unit_name = row['Đơn vị']
            values = [int(row[col]) if pd.notna(row[col]) else 0 for col in date_columns]
            result['units'][unit_name] = values
        
        return result
    except Exception as e:
        print(f"   ⚠️ Lỗi đọc SHC summary: {e}")
        return None


def load_shc_by_nvkt_for_unit(unit_name, data_folder="downloads/baocao_hanoi"):
    """
    Đọc dữ liệu SHC theo từng NVKT cho 1 đơn vị từ sheet Xu_huong_theo_NVKT
    
    Returns:
        dict: {'nvkt_list': [...], 'dates': [...], 'data': {nvkt: [values]}} hoặc None
    """
    import glob
    
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    latest_file = max(files, key=os.path.getmtime)
    
    try:
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_NVKT')
        
        # Lọc theo đơn vị
        df_unit = _match_unit(df, 'Đơn vị', unit_name)
        
        if df_unit.empty:
            return None
        
        date_columns = [col for col in df.columns if col not in ['Đơn vị', 'NVKT']]
        
        result = {
            'nvkt_list': [],
            'dates': [str(d) for d in date_columns],
            'data': {}
        }
        
        for _, row in df_unit.iterrows():
            nvkt = row['NVKT']
            values = [int(row[col]) if pd.notna(row[col]) else 0 for col in date_columns]
            result['nvkt_list'].append(nvkt)
            result['data'][nvkt] = values
        
        return result
    except Exception as e:
        print(f"   ⚠️ Lỗi đọc SHC by NVKT: {e}")
        return None


def create_nvkt_shc_grouped_chart(nvkt_data, unit_name, output_path=None):
    """
    Tạo biểu đồ nhóm cột SHC theo NVKT, mỗi ngày 1 màu khác nhau
    """
    if not nvkt_data or not nvkt_data['data']:
        return None
    
    nvkt_list = nvkt_data['nvkt_list']
    dates = nvkt_data['dates']
    data = nvkt_data['data']
    
    # Sử dụng họ tên đầy đủ
    nvkt_labels = nvkt_list
    
    # Setup figure
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(nvkt_list))
    n_dates = len(dates)
    width = 0.8 / n_dates  # Chiều rộng mỗi cột
    
    # Màu sắc cho từng ngày
    colors = plt.cm.tab10(np.linspace(0, 1, n_dates))
    
    # Vẽ từng ngày
    for i, date in enumerate(dates):
        values = [data[nvkt][i] for nvkt in nvkt_list]
        offset = (i - n_dates/2 + 0.5) * width
        bars = ax.bar(x + offset, values, width, label=date, color=colors[i])
        
        # Thêm giá trị lên cột (chỉ nếu > 0)
        for bar, val in zip(bars, values):
            if val > 0:
                ax.annotate(f'{val}',
                           xy=(bar.get_x() + bar.get_width() / 2, bar.get_height()),
                           xytext=(0, 1),
                           textcoords="offset points",
                           ha='center', va='bottom',
                           fontsize=7, fontweight='bold')
    
    short_name = _get_short_name(unit_name)
    ax.set_xlabel('NVKT', fontsize=11)
    ax.set_ylabel('Số TB suy hao cao', fontsize=11)
    ax.set_title(f'KẾT QUẢ XỬ LÝ SUY HAO CAO - {short_name}', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(nvkt_labels, rotation=45, ha='right', fontsize=9)
    ax.legend(title='Ngày', loc='upper right', fontsize=8, ncol=2)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_shc_overview_chart(shc_data, output_path=None):
    """
    Tạo biểu đồ tổng hợp SHC theo ngày cho tất cả đơn vị (stacked bar)
    """
    if not shc_data:
        return None
    
    dates = [str(d) for d in shc_data['dates']]
    units = shc_data['units']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(dates))
    width = 0.2
    colors = ['#2196F3', '#4CAF50', '#FF9800', '#E91E63']
    
    unit_names = list(units.keys())
    for i, unit_name in enumerate(unit_names):
        short_name = _get_short_name(unit_name)
        values = units[unit_name]
        bars = ax.bar(x + i * width, values, width, label=short_name, color=colors[i % len(colors)])
        
        # Thêm giá trị lên cột
        for bar, val in zip(bars, values):
            if val > 0:
                ax.annotate(f'{val}', xy=(bar.get_x() + bar.get_width() / 2, bar.get_height()),
                           xytext=(0, 2), textcoords="offset points",
                           ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Ngày', fontsize=11)
    ax.set_ylabel('Số TB suy hao cao', fontsize=11)
    ax.set_title('XU HƯỚNG SUY HAO CAO THEO ĐƠN VỊ', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x + width * (len(unit_names) - 1) / 2)
    ax.set_xticklabels(dates, rotation=45, ha='right', fontsize=9)
    ax.legend(loc='upper right')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_shc_overview_section(doc, data_folder="downloads/baocao_hanoi"):
    """
    Thêm phần tổng quan SHC vào báo cáo (PHẦN 1)
    """
    shc_data = load_shc_summary_by_unit(data_folder)
    
    if not shc_data:
        return
    
    doc.add_heading('1.4. Tổng quan Suy Hao Cao', level=2)
    
    dates = shc_data['dates']
    units = shc_data['units']
    
    # Bảng tổng hợp
    table = doc.add_table(rows=1, cols=len(dates) + 2)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    headers = ['Đơn vị'] + [str(d) for d in dates] + ['Tổng']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Data rows
    grand_total = 0
    for unit_name, values in units.items():
        cells = table.add_row().cells
        short_name = _get_short_name(unit_name)
        cells[0].text = short_name
        
        total = sum(values)
        grand_total += total
        
        for j, val in enumerate(values):
            cells[j + 1].text = str(val)
            cells[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[j + 1].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if val > 5:
                run.font.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
        
        cells[-1].text = str(total)
        cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[-1].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
    
    # Dòng tổng
    cells = table.add_row().cells
    cells[0].text = 'TỔNG CỘNG'
    cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(cells[0], 'E3F2FD')
    
    totals_by_date = [sum(units[u][i] for u in units) for i in range(len(dates))]
    for j, total in enumerate(totals_by_date):
        cells[j + 1].text = str(total)
        cells[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[j + 1].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cells[j + 1], 'E3F2FD')
    
    cells[-1].text = str(grand_total)
    cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[-1].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(cells[-1], 'E3F2FD')
    
    doc.add_paragraph()
    
    # Biểu đồ
    try:
        chart = create_shc_overview_chart(shc_data)
        if chart:
            doc.add_picture(chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ⚠️ Không thể tạo biểu đồ SHC: {e}")
    
    doc.add_paragraph()


def add_shc_unit_section(doc, unit_name, data_folder="downloads/baocao_hanoi"):
    """
    Thêm phần SHC chi tiết cho 1 đơn vị (trong PHẦN 2)
    """
    shc_data = load_shc_summary_by_unit(data_folder)
    
    if not shc_data or unit_name not in shc_data['units']:
        return
    
    dates = shc_data['dates']
    values = shc_data['units'][unit_name]
    
    short_name = _get_short_name(unit_name)
    doc.add_heading(f'Số liệu Suy Hao Cao - {short_name}', level=3)
    
    # Bảng dữ liệu - cải thiện format
    table = doc.add_table(rows=2, cols=len(dates) + 1)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header row - bao gồm cột Chỉ tiêu
    headers = ['Ngày'] + [str(d) for d in dates]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Data row
    table.rows[1].cells[0].text = 'Số TB SHC'
    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_shading(table.rows[1].cells[0], 'E3F2FD')
    
    for j, val in enumerate(values):
        cell = table.rows[1].cells[j + 1]
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Tô màu theo mức độ (dựa trên ngưỡng phù hợp cho đơn vị)
        avg_val = sum(values) / len(values) if values else 0
        if val <= avg_val * 0.5:
            set_cell_shading(cell, 'C8E6C9')  # Xanh nhạt - tốt
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif val <= avg_val * 1.2:
            set_cell_shading(cell, 'FFF9C4')  # Vàng nhạt - trung bình
        else:
            set_cell_shading(cell, 'FFCDD2')  # Đỏ nhạt - cao
            run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    
    # Thống kê chi tiết
    total = sum(values)
    avg = total / len(values) if values else 0
    max_val = max(values) if values else 0
    min_val = min(values) if values else 0
    max_day = dates[values.index(max_val)] if values else 'N/A'
    min_day = dates[values.index(min_val)] if values else 'N/A'
    
    p = doc.add_paragraph()
    p.add_run(f'📊 Tổng: {total} | Trung bình: {avg:.1f}/ngày | ').bold = True
    p.add_run(f'Cao nhất: {max_val} ({max_day}) | Thấp nhất: {min_val} ({min_day})')
    
    doc.add_paragraph()
    
    # Biểu đồ nhóm theo NVKT (như hình mẫu)
    try:
        nvkt_data = load_shc_by_nvkt_for_unit(unit_name, data_folder)
        if nvkt_data:
            chart = create_nvkt_shc_grouped_chart(nvkt_data, unit_name)
            if chart:
                doc.add_picture(chart, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ⚠️ Không thể tạo biểu đồ SHC cho {short_name}: {e}")
    
    doc.add_paragraph()


# =============================================================================
# HÀM TẠO BÁO CÁO WORD HOÀN CHỈNH
# =============================================================================

def generate_kpi_report(kpi_folder=DEFAULT_KPI_FOLDER, output_folder=DEFAULT_OUTPUT_FOLDER, 
                         report_month=None, report_title=None, include_exclusion=True,
                         exclusion_folder="downloads/kq_sau_giam_tru_hni"):
    """
    Tạo báo cáo Word hoàn chỉnh với bảng biểu và biểu đồ KPI
    
    Args:
        kpi_folder: Thư mục chứa file KPI Excel
        output_folder: Thư mục xuất báo cáo Word
        report_month: Tháng báo cáo (vd: "01/2026"), mặc định là tháng hiện tại
        report_title: Tiêu đề tùy chỉnh
        include_exclusion: Bao gồm dữ liệu sau giảm trừ (mặc định True)
        exclusion_folder: Thư mục chứa dữ liệu sau giảm trừ
        
    Returns:
        str: Đường dẫn file Word đã tạo
    """
    print("="*60)
    print("📝 BẮT ĐẦU TẠO BÁO CÁO WORD KPI")
    print("="*60)
    
    # Xác định tháng báo cáo
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # Tạo thư mục output nếu chưa có
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Đọc dữ liệu
    print("📊 Đọc dữ liệu KPI...")
    df_summary, df_detail = load_kpi_data(kpi_folder)
    
    # Lấy danh sách các tổ
    teams = df_summary['don_vi'].unique()
    print(f"   Tìm thấy {len(teams)} tổ kỹ thuật")
    
    # Tạo document
    doc = Document()
    
    # Thiết lập style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # =========================================================================
    # TIÊU ĐỀ
    # =========================================================================
    created_time = datetime.now().strftime('%d/%m/%Y %H:%M')
    title = doc.add_heading(level=0)
    title_run = title.add_run(report_title or f'BÁO CÁO KẾT QUẢ BSC/KPI\nTHÁNG {report_month}')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Ngày tạo: {created_time}")
    doc.add_paragraph()
    
    # =========================================================================
    # PHẦN 1: TỔNG QUAN
    # =========================================================================
    print("📈 Tạo phần Tổng quan...")
    doc.add_heading('PHẦN 1: TỔNG QUAN', level=1)
    
    # Đọc dữ liệu chi tiết từ các báo cáo C1.x (cần cho biểu đồ BSC)
    print("📊 Đọc dữ liệu chi tiết từ các báo cáo C1.x...")
    c1x_reports = load_c1x_reports()
    
    # Đọc dữ liệu giảm trừ nếu được yêu cầu
    comparison_data = None
    if include_exclusion:
        print("📊 Đọc dữ liệu so sánh trước/sau giảm trừ...")
        comparison_data = load_exclusion_comparison_data(exclusion_folder)
    
    # 1.1 Biểu đồ so sánh điểm BSC thực tế 4 tổ
    doc.add_heading('1.1. So sánh điểm BSC thực tế giữa các tổ', level=2)
    if c1x_reports:
        # Load điểm BSC từ Tong_hop_Diem_BSC_Don_Vi.xlsx nếu có
        bsc_scores_for_chart = load_bsc_unit_scores_from_comparison(exclusion_folder) if include_exclusion else None
        team_chart = create_team_comparison_chart(c1x_reports, bsc_data=bsc_scores_for_chart)
        doc.add_picture(team_chart, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("(Không có dữ liệu C1.x để tạo biểu đồ)")
    doc.add_paragraph()
    
    # 1.1.b Biểu đồ BSC SAU GIẢM TRỪ (ngay sau biểu đồ thô)
    if include_exclusion and comparison_data:
        doc.add_heading('1.1.b. So sánh điểm BSC sau giảm trừ giữa các tổ', level=2)
        try:
            unit_data = load_unit_level_exclusion_data(exclusion_folder)
            if unit_data:
                # Sử dụng bsc_scores_for_chart đã load ở trên
                bsc_after_chart = create_team_bsc_after_exclusion_chart(unit_data, c1x_reports, bsc_data=bsc_scores_for_chart)
                if bsc_after_chart:
                    doc.add_picture(bsc_after_chart, width=Inches(6.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Không thể tạo biểu đồ BSC sau giảm trừ: {e}")
        doc.add_paragraph()
    
    # 1.2 Thêm bảng thống kê tổng hợp theo tổ - sử dụng điểm BSC thực tế
    doc.add_heading('1.2. Thống kê điểm BSC theo đơn vị', level=2)
    
    # Chú thích giải thích các chỉ tiêu
    legend = doc.add_paragraph()
    legend.add_run('📋 CHÚ THÍCH CÁC CHỈ TIÊU BSC - VIỄN CẢNH KHÁCH HÀNG (C)').bold = True
    
    # C1.1
    p11 = doc.add_paragraph()
    run_title = p11.add_run('C1.1 - Chất lượng sửa chữa thuê bao BRCĐ: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p11.add_run('Gồm 2 thành phần:\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_tp1 = p11.add_run('   • TP1 (30%): Sửa chữa chủ động - Tỷ lệ phiếu SCCD hoàn thành ≤72h\n')
    run_tp1.italic = True
    run_tp1.font.size = Pt(10)
    
    run_tp2 = p11.add_run('   • TP2 (70%): Sửa chữa theo báo hỏng - Tỷ lệ phiếu BH hoàn thành đúng hạn')
    run_tp2.italic = True
    run_tp2.font.size = Pt(10)
    
    # C1.2
    p12 = doc.add_paragraph()
    run_title = p12.add_run('C1.2 - Tỷ lệ thuê bao báo hỏng: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p12.add_run('Gồm 2 thành phần:\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_tp1 = p12.add_run('   • TP1 (50%): Hỏng lặp lại - Tỷ lệ TB báo hỏng ≥2 lần/7 ngày\n')
    run_tp1.italic = True
    run_tp1.font.size = Pt(10)
    
    run_tp2 = p12.add_run('   • TP2 (50%): Tỷ lệ sự cố - Tỷ lệ phiếu BH / Tổng TB quản lý (‰)')
    run_tp2.italic = True
    run_tp2.font.size = Pt(10)
    
    # C1.3
    p13 = doc.add_paragraph()
    run_title = p13.add_run('C1.3 - Chất lượng sửa chữa kênh TSL (Leased Line): ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p13.add_run('Áp dụng cho các dịch vụ Internet trực tiếp, kênh thuê riêng, MegaWan, Metronet, Siptrunking')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    # C1.4
    p14 = doc.add_paragraph()
    run_title = p14.add_run('C1.4 - Độ hài lòng khách hàng: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p14.add_run('Tỷ lệ khách hàng hài lòng sau khi được sửa chữa (qua khảo sát)')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    # C1.5
    p15 = doc.add_paragraph()
    run_title = p15.add_run('C1.5 - Thiết lập dịch vụ BRCĐ đạt thời gian quy định: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p15.add_run('Tỷ lệ phiếu lắp đặt hoàn thành đúng hạn\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_detail = p15.add_run('   • Ngoài CCCO: ≤24h | Trong CCCO: Phiếu trước 17h xong trong ngày')
    run_detail.italic = True
    run_detail.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # =========================================================================
    # SỬ DỤNG DỮ LIỆU TỪ Tong_hop_Diem_BSC_Don_Vi.xlsx (có cả Trước và Sau GT)
    # =========================================================================
    teams_order = [t.short_name for t in sorted(BRCD_TEAMS, key=lambda t: t.order) if t.active]

    def get_short_name(don_vi):
        if not don_vi: return None
        result = _get_short_name(str(don_vi))
        if result != str(don_vi):
            return result
        for orig, short in TEAM_SHORT_NAMES.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return don_vi
    
    # Đọc dữ liệu từ Tong_hop_Diem_BSC_Don_Vi.xlsx
    print("📊 Đọc điểm BSC từ Tong_hop_Diem_BSC_Don_Vi.xlsx...")
    bsc_scores = load_bsc_unit_scores_from_comparison(exclusion_folder)
    
    # Khởi tạo dữ liệu BSC mặc định
    bsc_data = {team: {
        'C1.1_truoc': 0, 'C1.1_sau': 0,
        'C1.2_truoc': 0, 'C1.2_sau': 0,
        'C1.3': 0,  # C1.3 giữ nguyên, không có giảm trừ
        'C1.4_truoc': 0, 'C1.4_sau': 0,
        'C1.5_truoc': 0, 'C1.5_sau': 0
    } for team in teams_order + ['TTVT Sơn Tây']}
    
    # Đọc điểm từ Tong_hop_Diem_BSC_Don_Vi.xlsx
    if bsc_scores['units'] is not None and not bsc_scores['units'].empty:
        for _, row in bsc_scores['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in bsc_data:
                bsc_data[short_name]['C1.1_truoc'] = row.get('Diem_C1.1 (Trước)', 0) or 0
                bsc_data[short_name]['C1.1_sau'] = row.get('Diem_C1.1 (Sau)', 0) or 0
                bsc_data[short_name]['C1.2_truoc'] = row.get('Diem_C1.2 (Trước)', 0) or 0
                bsc_data[short_name]['C1.2_sau'] = row.get('Diem_C1.2 (Sau)', 0) or 0
                bsc_data[short_name]['C1.4_truoc'] = row.get('Diem_C1.4 (Trước)', 0) or 0
                bsc_data[short_name]['C1.4_sau'] = row.get('Diem_C1.4 (Sau)', 0) or 0
                bsc_data[short_name]['C1.5_truoc'] = row.get('Diem_C1.5 (Trước)', 0) or 0
                bsc_data[short_name]['C1.5_sau'] = row.get('Diem_C1.5 (Sau)', 0) or 0
    
    # Đọc C1.3 từ c1x_reports (giữ nguyên vì không có giảm trừ)
    if c1x_reports and 'c13' in c1x_reports:
        for _, row in c1x_reports['c13'].iterrows():
            don_vi = row.get('Đơn vị', '')
            short_name = get_short_name(don_vi)
            if short_name in bsc_data:
                bsc_data[short_name]['C1.3'] = row.get('Chỉ tiêu BSC', 0) or 0
            elif don_vi == 'Tổng':
                bsc_data['TTVT Sơn Tây']['C1.3'] = row.get('Chỉ tiêu BSC', 0) or 0
    
    # Tạo bảng với cấu trúc: Đơn vị | C1.1 (Trước/Sau) | C1.2 (Trước/Sau) | C1.3 | C1.4 (Trước/Sau) | C1.5 (Trước/Sau)
    headers = ['Đơn vị', 'C1.1', '', 'C1.2', '', 'C1.3', 'C1.4', '', 'C1.5', '']
    sub_headers = ['', 'Trước', 'Sau', 'Trước', 'Sau', '', 'Trước', 'Sau', 'Trước', 'Sau']
    
    table = doc.add_table(rows=2, cols=10)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header row 1 - Merge các ô
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Đơn vị'
    header_cells[0].merge(table.rows[1].cells[0])  # Merge với dòng dưới
    
    # C1.1 header với merge
    header_cells[1].text = 'C1.1'
    header_cells[1].merge(header_cells[2])
    
    # C1.2 header với merge
    header_cells[3].text = 'C1.2'
    header_cells[3].merge(header_cells[4])
    
    # C1.3 header (không có trước/sau)
    header_cells[5].text = 'C1.3'
    header_cells[5].merge(table.rows[1].cells[5])
    
    # C1.4 header với merge
    header_cells[6].text = 'C1.4'
    header_cells[6].merge(header_cells[7])
    
    # C1.5 header với merge
    header_cells[8].text = 'C1.5'
    header_cells[8].merge(header_cells[9])
    
    # Format header row 1
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    
    # Header row 2 (sub headers: Trước/Sau)
    sub_header_cells = table.rows[1].cells
    sub_header_values = ['', 'Trước', 'Sau', 'Trước', 'Sau', '', 'Trước', 'Sau', 'Trước', 'Sau']
    for i, val in enumerate(sub_header_values):
        if val:
            sub_header_cells[i].text = val
            sub_header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(sub_header_cells[i], '2E75B6')
            if sub_header_cells[i].paragraphs[0].runs:
                run = sub_header_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    
    # Data rows
    for idx, team in enumerate(teams_order, 1):
        cells = table.add_row().cells
        scores = bsc_data[team]
        row_data = [
            team,
            format_number(scores['C1.1_truoc']),
            format_number(scores['C1.1_sau']),
            format_number(scores['C1.2_truoc']),
            format_number(scores['C1.2_sau']),
            format_number(scores['C1.3']),
            format_number(scores['C1.4_truoc']),
            format_number(scores['C1.4_sau']),
            format_number(scores['C1.5_truoc']),
            format_number(scores['C1.5_sau'])
        ]
        for i, value in enumerate(row_data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cells[i].paragraphs[0].runs:
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F4FD')
    
    # Dòng TTVT Sơn Tây
    cells = table.add_row().cells
    ttvt_scores = bsc_data['TTVT Sơn Tây']
    ttvt_data = [
        'TTVT Sơn Tây',
        format_number(ttvt_scores['C1.1_truoc']),
        format_number(ttvt_scores['C1.1_sau']),
        format_number(ttvt_scores['C1.2_truoc']),
        format_number(ttvt_scores['C1.2_sau']),
        format_number(ttvt_scores['C1.3']),
        format_number(ttvt_scores['C1.4_truoc']),
        format_number(ttvt_scores['C1.4_sau']),
        format_number(ttvt_scores['C1.5_truoc']),
        format_number(ttvt_scores['C1.5_sau'])
    ]
    for i, value in enumerate(ttvt_data):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cells[i].paragraphs[0].runs:
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.bold = True
        set_cell_shading(cells[i], 'B2DFDB')
    
    doc.add_paragraph()
    
    # 1.4 Số liệu chi tiết các chỉ tiêu BSC theo Đội/TTVT (sử dụng c1x_reports đã load)
    # Nếu có comparison_data, sẽ thêm bảng sau giảm trừ ngay sau mỗi bảng thô
    if c1x_reports:
        # Load unit_data nếu cần cho bảng tổng hợp sau giảm trừ
        unit_data = None
        if include_exclusion:
            unit_data = load_unit_level_exclusion_data(exclusion_folder)
        add_c1x_overview_table(doc, c1x_reports, comparison_data, unit_data, exclusion_folder)
    
    # 1.5 Tổng quan Suy Hao Cao
    print("📊 Thêm phần Suy Hao Cao...")
    add_shc_overview_section(doc, data_folder="downloads/baocao_hanoi")
    
    # 1.5 SỐ LIỆU SAU GIẢM TRỪ - TỔNG HỢP (nếu có)
    if include_exclusion and comparison_data:
        print("📊 Thêm phần tổng hợp số liệu sau giảm trừ...")
        doc.add_heading('1.5. Tổng hợp số liệu sau giảm trừ', level=2)
        
        # Chú thích
        p_note = doc.add_paragraph()
        p_note.add_run('📋 GHI CHÚ: ').bold = True
        p_note.add_run('Số liệu sau giảm trừ được tính sau khi loại bỏ các phiếu báo hỏng thuộc diện loại trừ. Bảng chi tiết đã được hiển thị ngay sau mỗi bảng chỉ tiêu ở phần 1.3.')
        doc.add_paragraph()
        
        # Bảng tổng hợp so sánh
        add_exclusion_summary_table(doc, comparison_data)
        
        # Biểu đồ riêng cho dữ liệu sau giảm trừ
        doc.add_heading('Biểu đồ tỷ lệ sau giảm trừ', level=3)
        try:
            exclusion_chart = create_exclusion_bar_chart(comparison_data)
            if exclusion_chart:
                doc.add_picture(exclusion_chart, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Không thể tạo biểu đồ sau giảm trừ: {e}")
        
        doc.add_paragraph()
        
        # Phần thống kê theo đơn vị (Tổ và TTVT)
        print("📊 Thêm phần thống kê theo đơn vị...")
        unit_data = load_unit_level_exclusion_data(exclusion_folder)
        if unit_data:
            add_unit_level_exclusion_section(doc, unit_data, c1x_reports)
    
    doc.add_page_break()
    
    # =========================================================================
    # PHẦN 2: CHI TIẾT THEO TỔ
    # =========================================================================
    print("📋 Tạo phần Chi tiết theo tổ...")
    doc.add_heading('PHẦN 2: CHI TIẾT THEO TỪNG TỔ', level=1)
    
    # Đọc dữ liệu KPI sau giảm trừ theo NVKT
    df_exclusion_nvkt = None
    df_exclusion_detail = None
    df_raw_detail = None  # Dữ liệu Thô (trước giảm trừ) từ So_sanh_*.xlsx
    if include_exclusion:
        df_exclusion_nvkt = load_nvkt_exclusion_summary(exclusion_folder)
        df_exclusion_detail = load_nvkt_exclusion_detail(exclusion_folder)
        df_raw_detail = load_nvkt_raw_detail(exclusion_folder)  # Thô từ cùng file
    
    for team_idx, team_name in enumerate(teams, 1):
        short_name = _get_short_name(team_name)
        print(f"   📁 Tổ {team_idx}: {short_name}")
        
        doc.add_heading(f'2.{team_idx}. {short_name}', level=2)
        
        # Bảng KPI tổng hợp của tổ
        doc.add_heading(f'Bảng điểm KPI tổng hợp', level=3)
        add_kpi_summary_table(doc, df_summary, team_name)
        doc.add_paragraph()
        
        # Bảng KPI sau giảm trừ (nếu có)
        if df_exclusion_nvkt is not None:
            doc.add_heading('Bảng điểm KPI tổng hợp (sau giảm trừ)', level=3)
            add_kpi_summary_table_after_exclusion(doc, df_exclusion_nvkt, team_name)
            doc.add_paragraph()
        
        # Biểu đồ cột so sánh NVKT
        doc.add_heading(f'Biểu đồ so sánh điểm KPI theo NVKT', level=3)
        nvkt_chart = create_nvkt_bar_chart(df_summary, team_name)
        if nvkt_chart:
            doc.add_picture(nvkt_chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Biểu đồ sau giảm trừ (nếu có)
        if df_exclusion_nvkt is not None:
            doc.add_heading('Biểu đồ so sánh điểm KPI theo NVKT (sau giảm trừ)', level=3)
            nvkt_chart_gt = create_nvkt_bar_chart_after_exclusion(df_exclusion_nvkt, team_name)
            if nvkt_chart_gt:
                doc.add_picture(nvkt_chart_gt, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # Bảng chi tiết từng chỉ tiêu (Thô - trước giảm trừ)
        if df_raw_detail is not None:
            add_c11_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c11_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.1 chi tiết sau giảm trừ
        if df_exclusion_detail is not None:
            add_c11_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c12_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c12_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.2 chi tiết sau giảm trừ
        if df_exclusion_detail is not None:
            add_c12_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c14_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c14_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.4 chi tiết sau giảm trừ
        if df_exclusion_detail is not None:
            add_c14_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c15_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c15_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()

        # C1.5 chi tiết sau giảm trừ
        if df_exclusion_detail is not None:
            add_c15_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()

        # Số liệu Suy Hao Cao cho tổ
        add_shc_unit_section(doc, team_name, data_folder="downloads/baocao_hanoi")
        
        # Thêm page break sau mỗi tổ (trừ tổ cuối)
        if team_idx < len(teams):
            doc.add_page_break()
    
    # =========================================================================
    # PHẦN 3: KẾT LUẬN
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('PHẦN 3: KẾT LUẬN VÀ KIẾN NGHỊ', level=1)
    
    # Thống kê tổng quan
    total_nvkt = len(df_summary)
    avg_c11 = df_summary['Diem_C1.1'].mean()
    avg_c12 = df_summary['Diem_C1.2'].mean()
    avg_c14 = df_summary['Diem_C1.4'].mean()
    avg_c15 = df_summary['Diem_C1.5'].mean()
    
    doc.add_heading('3.1. Thống kê tổng quan', level=2)
    doc.add_paragraph(f"• Tổng số NVKT: {total_nvkt}")
    doc.add_paragraph(f"• Điểm trung bình C1.1: {format_number(avg_c11)}")
    doc.add_paragraph(f"• Điểm trung bình C1.2: {format_number(avg_c12)}")
    doc.add_paragraph(f"• Điểm trung bình C1.4: {format_number(avg_c14)}")
    doc.add_paragraph(f"• Điểm trung bình C1.5: {format_number(avg_c15)}")
    
    # Đánh giá
    doc.add_heading('3.2. Đánh giá chung', level=2)
    doc.add_paragraph("(Phần này cần bổ sung nội dung đánh giá theo thực tế)")
    
    doc.add_heading('3.3. Kiến nghị', level=2)
    doc.add_paragraph("(Phần này cần bổ sung nội dung kiến nghị theo thực tế)")
    
    # =========================================================================
    # PHỤ LỤC: BẢNG ĐIỂM KPI CHI TIẾT THEO NVKT
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('PHỤ LỤC: BẢNG ĐIỂM KPI CHI TIẾT THEO NVKT', level=1)
    
    # Bảng KPI tổng hợp theo NVKT
    doc.add_heading('Bảng điểm KPI tổng hợp theo NVKT', level=2)
    add_kpi_summary_table(doc, df_summary)
    doc.add_paragraph()
    
    # Tạo biểu đồ bar cho từng tổ
    doc.add_heading('Biểu đồ điểm KPI theo NVKT (theo từng tổ)', level=2)
    
    for team_name in teams:
        short_name = _get_short_name(team_name)
        doc.add_heading(f'{short_name}', level=3)
        
        nvkt_chart = create_nvkt_bar_chart(df_summary, team_name)
        if nvkt_chart:
            doc.add_picture(nvkt_chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # =========================================================================
    # LƯU FILE
    # =========================================================================
    output_file = output_path / f"Bao_cao_KPI_NVKT_{report_month.replace('/', '_')}.docx"
    doc.save(output_file)
    
    print("="*60)
    print(f"✅ ĐÃ TẠO BÁO CÁO WORD THÀNH CÔNG!")
    print(f"   📄 File: {output_file}")
    print("="*60)
    
    return str(output_file)


# =============================================================================
# HÀM TẠO BÁO CÁO CÁ NHÂN CHO TỪNG NVKT
# =============================================================================

def sanitize_filename(name):
    """
    Chuẩn hóa tên file - loại bỏ ký tự đặc biệt và thay khoảng trắng bằng _
    """
    import re
    # Thay khoảng trắng bằng _
    name = name.replace(' ', '_')
    # Loại bỏ ký tự đặc biệt (giữ lại chữ cái Unicode, số và _)
    name = re.sub(r'[^\w\s-]', '', name, flags=re.UNICODE)
    return name


def create_individual_radar_chart(nvkt_data, output_path=None):
    """
    Tạo biểu đồ radar so sánh điểm KPI của 1 NVKT
    
    Args:
        nvkt_data: Dictionary chứa điểm KPI của NVKT
        output_path: Đường dẫn lưu file (None = trả về bytes)
    
    Returns:
        bytes hoặc str
    """
    # Lấy điểm các chỉ tiêu
    metrics = ['Diem_C1.1', 'Diem_C1.2', 'Diem_C1.4', 'Diem_C1.5']
    labels = ['C1.1\nSửa chữa', 'C1.2\nBáo hỏng', 'C1.4\nHài lòng', 'C1.5\nThiết lập DV']
    
    values = []
    for m in metrics:
        val = nvkt_data.get(m, 0)
        if pd.isna(val):
            val = 0
        values.append(val)
    
    # Số lượng biến
    num_vars = len(labels)
    
    # Tính góc cho mỗi trục
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    
    # Đóng vòng radar
    values += values[:1]
    angles += angles[:1]
    
    # Tạo figure
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    
    # Vẽ radar
    ax.fill(angles, values, color='#2E86AB', alpha=0.25)
    ax.plot(angles, values, color='#2E86AB', linewidth=2, marker='o', markersize=8)
    
    # Vẽ đường chuẩn 5 điểm
    target_values = [5] * (num_vars + 1)
    ax.plot(angles, target_values, color='#C73E1D', linewidth=1.5, linestyle='--', alpha=0.7, label='Mục tiêu (5 điểm)')
    
    # Thiết lập các trục
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=11)
    ax.set_ylim(0, 5.5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(['1', '2', '3', '4', '5'], fontsize=9)
    
    # Thêm giá trị lên điểm
    for angle, value, label in zip(angles[:-1], values[:-1], labels):
        ax.annotate(f'{value:.2f}', 
                   xy=(angle, value), 
                   xytext=(angle, value + 0.3),
                   ha='center', va='bottom', fontsize=10, fontweight='bold')
    
    ax.legend(loc='upper right', bbox_to_anchor=(1.2, 1.1))
    ax.set_title('BIỂU ĐỒ ĐIỂM KPI', fontsize=14, fontweight='bold', pad=20)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_individual_summary_table(doc, nvkt_data):
    """
    Thêm bảng tổng hợp điểm KPI cho 1 NVKT vào document
    
    Args:
        doc: Document Word
        nvkt_data: Dictionary hoặc Series chứa dữ liệu NVKT
    """
    # Tạo bảng 2 cột: Chỉ tiêu - Điểm
    headers = ['Chỉ tiêu', 'Điểm']
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
    
    # Dữ liệu các chỉ tiêu
    kpi_info = [
        ('C1.1 - Tỷ lệ sửa chữa phiếu chất lượng & báo hỏng', 'Diem_C1.1'),
        ('C1.2 - Tỷ lệ báo hỏng lặp lại & sự cố dịch vụ', 'Diem_C1.2'),
        ('C1.4 - Độ hài lòng khách hàng', 'Diem_C1.4'),
        ('C1.5 - Tỷ lệ thiết lập dịch vụ đạt thời gian quy định', 'Diem_C1.5'),
    ]
    
    for idx, (label, col) in enumerate(kpi_info, 1):
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        value = nvkt_data.get(col, np.nan)
        cells[1].text = format_number(value)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, cell in enumerate(cells):
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            if idx % 2 == 0:
                set_cell_shading(cell, 'E8F4FD')
            
            # Tô màu điểm theo mức
            if i == 1 and not pd.isna(value):
                if value >= 4.5:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Xanh lá
                    run.font.bold = True
                elif value < 3:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Đỏ
                    run.font.bold = True


def add_individual_c11_detail(doc, nvkt_data):
    """
    Thêm chi tiết chỉ tiêu C1.1 cho 1 NVKT
    """
    doc.add_heading('2. CHI TIẾT CHỈ TIÊU C1.1 - CHẤT LƯỢNG SỬA CHỮA BRCĐ', level=2)
    
    # Thành phần 1
    doc.add_heading('2.1. Thành phần 1: Sửa chữa chủ động (SCCD ≤72h) - 30%', level=3)
    
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'
    set_table_border(table1)
    
    headers1 = ['Tổng SCCD', 'Đạt ≤72h', 'Tỷ lệ (%)', 'Điểm']
    for i, cell in enumerate(table1.rows[0].cells):
        cell.text = headers1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data1 = [
        format_number(nvkt_data.get('c11_tp1_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c11_tp1_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c11_tp1_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c11_tp1', np.nan))
    ]
    for i, cell in enumerate(table1.rows[1].cells):
        cell.text = data1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Thành phần 2
    doc.add_heading('2.2. Thành phần 2: Sửa chữa theo báo hỏng (đúng hạn) - 70%', level=3)
    
    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    set_table_border(table2)
    
    headers2 = ['Tổng SC BH', 'Đúng hạn', 'Tỷ lệ (%)', 'Điểm']
    for i, cell in enumerate(table2.rows[0].cells):
        cell.text = headers2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '388E3C')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data2 = [
        format_number(nvkt_data.get('c11_tp2_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c11_tp2_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c11_tp2_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c11_tp2', np.nan))
    ]
    for i, cell in enumerate(table2.rows[1].cells):
        cell.text = data2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Điểm tổng hợp
    p = doc.add_paragraph()
    p.add_run('➤ ĐIỂM TỔNG HỢP C1.1: ').bold = True
    diem_c11 = nvkt_data.get('Diem_C1.1', np.nan)
    run = p.add_run(format_number(diem_c11))
    run.bold = True
    run.font.size = Pt(12)
    if not pd.isna(diem_c11):
        if diem_c11 >= 4.5:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif diem_c11 < 3:
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    p.add_run(' (= TP1 × 30% + TP2 × 70%)')


def add_individual_c12_detail(doc, nvkt_data):
    """
    Thêm chi tiết chỉ tiêu C1.2 cho 1 NVKT
    """
    doc.add_heading('3. CHI TIẾT CHỈ TIÊU C1.2 - TỶ LỆ THUÊ BAO BÁO HỎNG', level=2)
    
    # Thành phần 1
    doc.add_heading('3.1. Thành phần 1: Hỏng lặp (≥2 lần/7 ngày) - 50%', level=3)
    
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'
    set_table_border(table1)
    
    headers1 = ['TB hỏng lặp', 'Tổng BH', 'Tỷ lệ (%)', 'Điểm']
    for i, cell in enumerate(table1.rows[0].cells):
        cell.text = headers1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data1 = [
        format_number(nvkt_data.get('c12_tp1_phieu_hll', np.nan), 0),
        format_number(nvkt_data.get('c12_tp1_phieu_bh', np.nan), 0),
        format_number(nvkt_data.get('c12_tp1_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c12_tp1', np.nan))
    ]
    for i, cell in enumerate(table1.rows[1].cells):
        cell.text = data1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Thành phần 2
    doc.add_heading('3.2. Thành phần 2: Tỷ lệ BH/TB quản lý (‰) - 50%', level=3)
    
    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    set_table_border(table2)
    
    headers2 = ['Phiếu BH', 'TB quản lý', 'Tỷ lệ (‰)', 'Điểm']
    for i, cell in enumerate(table2.rows[0].cells):
        cell.text = headers2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1976D2')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data2 = [
        format_number(nvkt_data.get('c12_tp2_phieu_bh', np.nan), 0),
        format_number(nvkt_data.get('c12_tp2_tong_tb', np.nan), 0),
        format_number(nvkt_data.get('c12_tp2_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c12_tp2', np.nan))
    ]
    for i, cell in enumerate(table2.rows[1].cells):
        cell.text = data2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Điểm tổng hợp
    p = doc.add_paragraph()
    p.add_run('➤ ĐIỂM TỔNG HỢP C1.2: ').bold = True
    diem_c12 = nvkt_data.get('Diem_C1.2', np.nan)
    run = p.add_run(format_number(diem_c12))
    run.bold = True
    run.font.size = Pt(12)
    if not pd.isna(diem_c12):
        if diem_c12 >= 4.5:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif diem_c12 < 3:
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    p.add_run(' (= TP1 × 50% + TP2 × 50%)')


def add_individual_c14_detail(doc, nvkt_data):
    """
    Thêm chi tiết chỉ tiêu C1.4 cho 1 NVKT
    """
    doc.add_heading('4. CHI TIẾT CHỈ TIÊU C1.4 - ĐỘ HÀI LÒNG KHÁCH HÀNG', level=2)
    
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    set_table_border(table)
    
    headers = ['Phiếu KS thành công', 'Phiếu KH không HL', 'Tỷ lệ HL (%)', 'Điểm']
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'F57C00')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data = [
        format_number(nvkt_data.get('c14_phieu_ks', np.nan), 0),
        format_number(nvkt_data.get('c14_phieu_khl', np.nan), 0),
        format_number(nvkt_data.get('c14_ty_le', np.nan)),
        format_number(nvkt_data.get('Diem_C1.4', np.nan))
    ]
    for i, cell in enumerate(table.rows[1].cells):
        cell.text = data[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Công thức tính điểm
    p = doc.add_paragraph()
    p.add_run('📌 Công thức tính điểm: ').bold = True
    p.add_run('≥ 99.5% = 5 điểm, > 95% = nội suy, ≤ 95% = 1 điểm')


def add_individual_c15_detail(doc, nvkt_data):
    """
    Thêm chi tiết chỉ tiêu C1.5 cho 1 NVKT
    """
    doc.add_heading('5. CHI TIẾT CHỈ TIÊU C1.5 - TỶ LỆ THIẾT LẬP DỊCH VỤ ĐẠT', level=2)
    
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    set_table_border(table)
    
    headers = ['Phiếu đạt', 'Phiếu không đạt', 'Tổng phiếu', 'Tỷ lệ đạt (%)', 'Điểm']
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '7B1FA2')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data = [
        format_number(nvkt_data.get('c15_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c15_phieu_khong_dat', np.nan), 0),
        format_number(nvkt_data.get('c15_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c15_ty_le', np.nan)),
        format_number(nvkt_data.get('Diem_C1.5', np.nan))
    ]
    for i, cell in enumerate(table.rows[1].cells):
        cell.text = data[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Công thức tính điểm
    p = doc.add_paragraph()
    p.add_run('📌 Công thức tính điểm: ').bold = True
    p.add_run('≥ 99.5% = 5 điểm, 89.5% < KQ < 99.5% = nội suy, ≤ 89.5% = 1 điểm')


def load_shc_trend_data(nvkt_name, data_folder="downloads/baocao_hanoi"):
    """
    Đọc dữ liệu xu hướng SHC cho 1 NVKT từ file Bao_cao_xu_huong_*.xlsx
    
    Args:
        nvkt_name: Tên NVKT cần tìm
        data_folder: Thư mục chứa file báo cáo
    
    Returns:
        dict: {'dates': [...], 'values': [...], 'don_vi': '...'} hoặc None
    """
    import glob
    
    # Tìm file Bao_cao_xu_huong mới nhất
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    # Sắp xếp theo thời gian chỉnh sửa file (mới nhất cuối cùng)
    latest_file = max(files, key=os.path.getmtime)
    print(f"   📊 Sử dụng file SHC: {os.path.basename(latest_file)}")
    
    try:
        # Đọc sheet Xu_huong_theo_NVKT
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_NVKT')
        
        # Tìm NVKT trong cột 'NVKT'
        nvkt_row = df[df['NVKT'] == nvkt_name]
        
        if nvkt_row.empty:
            return None
        
        row = nvkt_row.iloc[0]
        don_vi = row.get('Đơn vị', '')
        
        # Lấy các cột ngày (không phải 'Đơn vị' và 'NVKT')
        date_columns = [col for col in df.columns if col not in ['Đơn vị', 'NVKT']]
        
        dates = []
        values = []
        for col in date_columns:
            dates.append(str(col))
            val = row[col]
            values.append(int(val) if pd.notna(val) else 0)
        
        return {
            'dates': dates,
            'values': values,
            'don_vi': don_vi
        }
    except Exception as e:
        print(f"   ⚠️ Lỗi đọc dữ liệu SHC: {e}")
        return None


def create_shc_trend_bar_chart(shc_data, nvkt_name, output_path=None):
    """
    Tạo biểu đồ cột thể hiện xu hướng số TB suy hao cao theo ngày
    
    Args:
        shc_data: dict với keys 'dates' và 'values'
        nvkt_name: Tên NVKT
        output_path: Đường dẫn lưu file (None = trả về bytes)
    
    Returns:
        bytes hoặc str: Chart image
    """
    if not shc_data or not shc_data.get('dates') or not shc_data.get('values'):
        return None
    
    dates = shc_data['dates']
    values = shc_data['values']
    
    # Tạo figure
    fig, ax = plt.subplots(figsize=(10, 5))
    
    # Vẽ biểu đồ cột
    x_pos = range(len(dates))
    bars = ax.bar(x_pos, values, color='#2196F3', edgecolor='#1565C0', linewidth=1)
    
    # Thêm giá trị lên cột
    for bar, val in zip(bars, values):
        height = bar.get_height()
        ax.annotate(f'{val}',
                   xy=(bar.get_x() + bar.get_width() / 2, height),
                   xytext=(0, 3),
                   textcoords="offset points",
                   ha='center', va='bottom',
                   fontsize=9, fontweight='bold')
    
    # Thiết lập trục
    ax.set_xticks(x_pos)
    ax.set_xticklabels(dates, rotation=45, ha='right', fontsize=9)
    ax.set_ylabel('Số TB suy hao cao', fontsize=11)
    ax.set_xlabel('Ngày', fontsize=11)
    ax.set_title(f'XU HƯỚNG SỐ TB SUY HAO CAO - {nvkt_name}', fontsize=12, fontweight='bold', pad=15)
    
    # Grid và layout
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    # Đường xu hướng trung bình
    if len(values) > 1:
        avg = sum(values) / len(values)
        ax.axhline(y=avg, color='#E91E63', linestyle='--', linewidth=1.5, alpha=0.7, 
                   label=f'TB: {avg:.1f}')
        ax.legend(loc='upper right')
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi"):
    """
    Thêm phần Số liệu Suy Hao Cao vào báo cáo cá nhân
    Bao gồm: Bảng dữ liệu + Biểu đồ bar
    
    Args:
        doc: Document Word
        nvkt_name: Tên NVKT
        data_folder: Thư mục chứa file báo cáo
    """
    # Load dữ liệu SHC
    shc_data = load_shc_trend_data(nvkt_name, data_folder)
    
    if not shc_data:
        return  # Không có dữ liệu SHC
    
    doc.add_page_break()
    doc.add_heading('6. SỐ LIỆU SUY HAO CAO', level=2)
    
    dates = shc_data['dates']
    values = shc_data['values']
    
    # Mô tả
    p = doc.add_paragraph()
    p.add_run('📊 Xu hướng số thuê bao suy hao cao theo ngày:').bold = True
    
    doc.add_paragraph()
    
    # Tạo bảng dữ liệu (chia thành các nhóm nếu nhiều ngày)
    max_cols = 10  # Số cột tối đa mỗi bảng
    
    for i in range(0, len(dates), max_cols):
        chunk_dates = dates[i:i+max_cols]
        chunk_values = values[i:i+max_cols]
        
        table = doc.add_table(rows=2, cols=len(chunk_dates))
        table.style = 'Table Grid'
        set_table_border(table)
        
        # Header row - Ngày
        for j, date in enumerate(chunk_dates):
            cell = table.rows[0].cells[j]
            cell.text = str(date)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1E88E5')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        
        # Data row - Số lượng
        for j, val in enumerate(chunk_values):
            cell = table.rows[1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.bold = True
            
            # Tô màu theo mức độ
            if val == 0:
                set_cell_shading(cell, 'C8E6C9')  # Xanh lá nhạt
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif val <= 3:
                set_cell_shading(cell, 'FFF9C4')  # Vàng nhạt
            else:
                set_cell_shading(cell, 'FFCDD2')  # Đỏ nhạt
                run.font.color.rgb = RGBColor(200, 0, 0)
        
        doc.add_paragraph()
    
    # Thống kê tổng quan
    total = sum(values)
    avg = total / len(values) if values else 0
    max_val = max(values) if values else 0
    min_val = min(values) if values else 0
    
    p = doc.add_paragraph()
    p.add_run(f'📈 Tổng số TB SHC trong kỳ: ').bold = True
    p.add_run(f'{total}')
    
    p = doc.add_paragraph()
    p.add_run(f'📊 Trung bình/ngày: ').bold = True
    p.add_run(f'{avg:.1f}')
    
    p = doc.add_paragraph()
    p.add_run(f'⬆️ Cao nhất: ').bold = True
    p.add_run(f'{max_val}')
    p.add_run(f'  |  ')
    p.add_run(f'⬇️ Thấp nhất: ').bold = True
    p.add_run(f'{min_val}')
    
    doc.add_paragraph()
    
    # Biểu đồ bar
    try:
        chart = create_shc_trend_bar_chart(shc_data, nvkt_name)
        if chart:
            doc.add_picture(chart, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ⚠️ Không thể tạo biểu đồ SHC: {e}")


def generate_individual_kpi_report(nvkt_name, don_vi, kpi_folder=DEFAULT_KPI_FOLDER, 
                                    output_folder=DEFAULT_OUTPUT_FOLDER, report_month=None):
    """
    Tạo báo cáo Word cho 1 NVKT cụ thể
    
    Args:
        nvkt_name: Tên NVKT (vd: "Bùi Văn Duẩn")
        don_vi: Tên đơn vị/tổ (vd: "Tổ Kỹ thuật Địa bàn Phúc Thọ")
        kpi_folder: Thư mục chứa file KPI
        output_folder: Thư mục xuất báo cáo
        report_month: Tháng báo cáo (vd: "01/2026")
    
    Returns:
        str: Đường dẫn file Word đã tạo
    """
    # Xác định tháng báo cáo
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # Đọc dữ liệu KPI chi tiết
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    # Lọc dữ liệu cho NVKT cụ thể
    nvkt_df = _match_unit(df_detail[df_detail['nvkt'] == nvkt_name], 'don_vi', don_vi)
    
    if nvkt_df.empty:
        print(f"⚠️ Không tìm thấy dữ liệu cho NVKT: {nvkt_name} - {don_vi}")
        return None
    
    nvkt_data = nvkt_df.iloc[0].to_dict()
    
    # Tạo thư mục output theo tổ
    short_name = _get_short_name(don_vi)
    folder_name = sanitize_filename(short_name)
    output_path = Path(output_folder) / "individual_reports" / folder_name
    output_path.mkdir(parents=True, exist_ok=True)
    
    # Tạo document
    doc = Document()
    
    # Thiết lập style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # =========================================================================
    # TIÊU ĐỀ
    # =========================================================================
    created_time = datetime.now().strftime('%d/%m/%Y %H:%M')
    
    title = doc.add_heading(level=0)
    title_run = title.add_run('BÁO CÁO KẾT QUẢ BSC/KPI CÁ NHÂN')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(level=1)
    subtitle_run = subtitle.add_run(f'THÁNG {report_month}')
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Thông tin cá nhân
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Họ và tên:', nvkt_name),
        ('Đơn vị:', short_name),
        ('Ngày tạo báo cáo:', created_time)
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # PHẦN 1: TỔNG QUAN
    # =========================================================================
    doc.add_heading('1. TỔNG QUAN ĐIỂM KPI', level=2)
    
    # Bảng tổng hợp
    add_individual_summary_table(doc, nvkt_data)
    doc.add_paragraph()
    
    # Biểu đồ radar
    try:
        radar_chart = create_individual_radar_chart(nvkt_data)
        doc.add_picture(radar_chart, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ⚠️ Không thể tạo biểu đồ radar: {e}")
    
    doc.add_page_break()
    
    # =========================================================================
    # CHI TIẾT TỪNG CHỈ TIÊU
    # =========================================================================
    add_individual_c11_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c12_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c14_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c15_detail(doc, nvkt_data)
    
    # =========================================================================
    # PHẦN 6: SỐ LIỆU SUY HAO CAO
    # =========================================================================
    add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi")
    
    # =========================================================================
    # LƯU FILE
    # =========================================================================
    safe_name = sanitize_filename(nvkt_name)
    output_file = output_path / f"Bao_cao_KPI_{safe_name}_{report_month.replace('/', '_')}.docx"
    doc.save(output_file)
    
    return str(output_file)



def generate_all_individual_reports(kpi_folder=DEFAULT_KPI_FOLDER, output_folder=DEFAULT_OUTPUT_FOLDER, 
                                     report_month=None):
    """
    Tạo báo cáo cho TẤT CẢ NVKT
    
    Args:
        kpi_folder: Thư mục chứa file KPI
        output_folder: Thư mục xuất báo cáo
        report_month: Tháng báo cáo
    
    Returns:
        list: Danh sách đường dẫn các file đã tạo
    """
    print("="*60)
    print("📝 BẮT ĐẦU TẠO BÁO CÁO KPI CÁ NHÂN CHO TẤT CẢ NVKT")
    print("="*60)
    
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # Đọc dữ liệu KPI
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    # Lấy danh sách NVKT
    nvkt_list = df_detail[['don_vi', 'nvkt']].drop_duplicates()
    total = len(nvkt_list)
    
    print(f"📊 Tìm thấy {total} NVKT")
    print()
    
    success_files = []
    failed_count = 0
    
    for idx, row in nvkt_list.iterrows():
        don_vi = row['don_vi']
        nvkt_name = row['nvkt']
        short_name = _get_short_name(don_vi)
        
        current = len(success_files) + failed_count + 1
        print(f"   [{current}/{total}] {nvkt_name} ({short_name})...", end=" ")
        
        try:
            result = generate_individual_kpi_report(
                nvkt_name=nvkt_name,
                don_vi=don_vi,
                kpi_folder=kpi_folder,
                output_folder=output_folder,
                report_month=report_month
            )
            if result:
                success_files.append(result)
                print("✅")
            else:
                failed_count += 1
                print("❌ (không có dữ liệu)")
        except Exception as e:
            failed_count += 1
            print(f"❌ ({str(e)[:30]})")
    
    print()
    print("="*60)
    print(f"✅ HOÀN THÀNH!")
    print(f"   📄 Đã tạo: {len(success_files)} báo cáo")
    print(f"   ❌ Thất bại: {failed_count}")
    print(f"   📁 Thư mục: {Path(output_folder) / 'individual_reports'}")
    print("="*60)
    
    return success_files


def generate_all_individual_reports_after_exclusion(kpi_folder, output_root, report_month=None):
    """
    Tạo báo cáo cá nhân sau giảm trừ, phân loại theo thư mục Tổ kỹ thuật
    Lưu tại: {output_root}/ca_nhan/{tên tổ kỹ thuật}/
    """
    print("="*60)
    print("📝 BẮT ĐẦU TẠO BÁO CÁO KPI CÁ NHÂN SAU GIẢM TRỪ")
    print("="*60)
    
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
        
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_SauGiamTru_ChiTiet.xlsx"
    summary_file = kpi_path / "KPI_NVKT_SauGiamTru_TomTat.xlsx"
    
    if not detail_file.exists():
        print(f"❌ Không tìm thấy file: {detail_file}")
        return 0
        
    df_detail = pd.read_excel(detail_file)
    nvkt_list = df_detail[['don_vi', 'nvkt']].drop_duplicates()
    total = len(nvkt_list)
    
    print(f"📊 Tìm thấy {total} NVKT sau giảm trừ")
    
    success_count = 0
    for idx, row in nvkt_list.iterrows():
        don_vi = row['don_vi']
        nvkt_name = row['nvkt']
        
        # Đảm bảo don_vi là chuỗi
        don_vi_str = str(don_vi) if pd.notna(don_vi) else "Unknown"
        
        # Tạo thư mục cho từng Đội (Tổ)
        team_folder_name = sanitize_filename(don_vi_str)
        team_output_path = Path(output_root) / "ca_nhan" / team_folder_name
        team_output_path.mkdir(parents=True, exist_ok=True)
        
        current = success_count + 1
        print(f"   [{current}/{total}] {nvkt_name} ({don_vi})...", end=" ")
        
        try:
            # Tạo document mới
            doc = Document()
            
            # Thiết lập style mặc định cho doc
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Lấy data NVKT
            # Lấy data NVKT - sử dụng logic lọc an toàn với NaN
            if pd.isna(don_vi):
                mask = (df_detail['nvkt'] == nvkt_name) & (df_detail['don_vi'].isna())
            else:
                nvkt_filtered = df_detail[df_detail['nvkt'] == nvkt_name]
                nvkt_df_match_temp = _match_unit(nvkt_filtered, 'don_vi', don_vi)
                mask = df_detail.index.isin(nvkt_df_match_temp.index)
            
            nvkt_df_match = df_detail[mask]
            if nvkt_df_match.empty:
                print(f"❌ (Không tìm thấy data)")
                continue
                
            nvkt_data = nvkt_df_match.iloc[0].to_dict()
            
            short_name = _get_short_name(don_vi_str)
            created_time = datetime.now().strftime("%d/%m/%Y %H:%M")
            
            # Header
            header = doc.sections[0].header
            p = header.paragraphs[0]
            p.text = f"BÁO CÁO KẾT QUẢ KPI CÁ NHÂN - THÁNG {report_month} (SAU GIẢM TRỪ)"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Title
            title = doc.add_heading(level=0)
            title_run = title.add_run('BÁO CÁO KẾT QUẢ BSC/KPI CÁ NHÂN (SAU GIẢM TRỪ)')
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(level=1)
            subtitle_run = subtitle.add_run(f'THÁNG {report_month}')
            subtitle_run.font.size = Pt(16)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Thông tin cá nhân
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            info_data = [
                ('Họ và tên:', nvkt_name),
                ('Đơn vị:', short_name),
                ('Ngày tạo báo cáo:', created_time)
            ]
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                info_table.rows[i].cells[1].text = value
            
            doc.add_paragraph()
            
            # Phần 1: Tổng quan
            doc.add_heading('1. TỔNG QUAN ĐIỂM KPI', level=2)
            add_individual_summary_table(doc, nvkt_data)
            
            # Biểu đồ radar
            try:
                radar_chart = create_individual_radar_chart(nvkt_data)
                doc.add_picture(radar_chart, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
                
            doc.add_page_break()
            
            # Chi tiết từng chỉ tiêu
            add_individual_c11_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c12_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c14_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c15_detail(doc, nvkt_data)
            
            # Suy hao cao - truyền data_folder mặc định
            add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi")
            
            # Lưu file
            safe_name = sanitize_filename(nvkt_name)
            current_date = datetime.now().strftime("%d_%m_%Y")
            filename = f"{safe_name}_Bao_cao_KPI_{current_date}.docx"
            output_file = team_output_path / filename
            doc.save(output_file)
            
            success_count += 1
            print("✅")
        except Exception as e:
            print(f"❌ (Lỗi: {e})")
            
    print(f"\n✅ Hoàn thành: Đã tạo {success_count}/{total} báo cáo cá nhân sau giảm trừ.")
    print(f"📁 Thư mục xuất: {output_root}/ca_nhan/")
    
    return success_count


# =============================================================================
# MAIN - Chạy trực tiếp module
# =============================================================================
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='Tạo báo cáo KPI/BSC cho NVKT')
    parser.add_argument('--individual', action='store_true', 
                       help='Tạo báo cáo cá nhân cho từng NVKT')
    parser.add_argument('--all', action='store_true',
                       help='Tạo báo cáo cho tất cả NVKT (dùng với --individual)')
    parser.add_argument('--nvkt', type=str, default=None,
                       help='Tên NVKT cụ thể (dùng với --individual)')
    parser.add_argument('--donvi', type=str, default=None,
                       help='Tên đơn vị của NVKT (dùng với --individual --nvkt)')
    parser.add_argument('--month', type=str, default="01/2026",
                       help='Tháng báo cáo (vd: 01/2026)')
    parser.add_argument('--kpi-folder', type=str, default="downloads/KPI",
                       help='Thư mục chứa file KPI')
    parser.add_argument('--output-folder', type=str, default="downloads/reports",
                       help='Thư mục xuất báo cáo')
    
    args = parser.parse_args()
    
    if args.individual:
        if args.all:
            # Tạo báo cáo cho tất cả NVKT
            generate_all_individual_reports(
                kpi_folder=args.kpi_folder,
                output_folder=args.output_folder,
                report_month=args.month
            )
        elif args.nvkt and args.donvi:
            # Tạo báo cáo cho 1 NVKT cụ thể
            result = generate_individual_kpi_report(
                nvkt_name=args.nvkt,
                don_vi=args.donvi,
                kpi_folder=args.kpi_folder,
                output_folder=args.output_folder,
                report_month=args.month
            )
            if result:
                print(f"✅ Đã tạo báo cáo: {result}")
            else:
                print("❌ Không thể tạo báo cáo")
        else:
            print("❌ Lỗi: Cần chỉ định --all hoặc --nvkt và --donvi")
            print("   Ví dụ: python report_generator.py --individual --all")
            print("   Hoặc:  python report_generator.py --individual --nvkt 'Bùi Văn Duẩn' --donvi 'Tổ Kỹ thuật Địa bàn Phúc Thọ'")
    else:
        # Tạo cả báo cáo tổng hợp và báo cáo cá nhân (mặc định)
        print("=" * 60)
        print("📊 TẠO BÁO CÁO TỔNG HỢP")
        print("=" * 60)
        report_path = generate_kpi_report(
            kpi_folder=args.kpi_folder,
            output_folder=args.output_folder,
            report_month=args.month
        )
        print(f"\n📁 Đường dẫn báo cáo tổng hợp: {report_path}")
        
        print("\n")
        print("=" * 60)
        print("📝 TẠO BÁO CÁO CÁ NHÂN CHO TẤT CẢ NVKT")
        print("=" * 60)
        generate_all_individual_reports(
            kpi_folder=args.kpi_folder,
            output_folder=args.output_folder,
            report_month=args.month
        )

