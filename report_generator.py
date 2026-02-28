# -*- coding: utf-8 -*-
"""
Module t·∫°o b√°o c√°o Word t·ª± ƒë·ªông cho KPI/BSC NVKT
T·∫°o file Word v·ªõi b·∫£ng bi·ªÉu v√† bi·ªÉu ƒë·ªì

Author: Auto-generated
Date: 2026-01-08
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os

# Thi·∫øt l·∫≠p matplotlib ƒë·ªÉ h·ªó tr·ª£ ti·∫øng Vi·ªát
matplotlib.rcParams['font.family'] = 'DejaVu Sans'
matplotlib.use('Agg')  # Use non-interactive backend

# =============================================================================
# C·∫§U H√åNH
# =============================================================================
DEFAULT_KPI_FOLDER = "downloads/KPI"
DEFAULT_OUTPUT_FOLDER = "downloads/reports"

# Mapping t√™n ƒë∆°n v·ªã ng·∫Øn g·ªçn
TEAM_SHORT_NAMES = {
    "T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç": "Ph√∫c Th·ªç",
    "T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai": "Qu·∫£ng Oai", 
    "T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai": "Su·ªëi Hai",
    "T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y": "S∆°n T√¢y"
}

# M√†u s·∫Øc cho bi·ªÉu ƒë·ªì
CHART_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6B5B95']
BAR_COLORS = ['#4CAF50', '#2196F3', '#FF9800', '#E91E63', '#9C27B0']  # Xanh l√°, Xanh d∆∞∆°ng, Cam, H·ªìng, T√≠m


# =============================================================================
# H√ÄM TI·ªÜN √çCH
# =============================================================================
def set_cell_shading(cell, color):
    """ƒê·∫∑t m√†u n·ªÅn cho √¥ trong b·∫£ng"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_border(table):
    """ƒê·∫∑t vi·ªÅn cho b·∫£ng"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def format_number(value, decimal_places=2):
    """Format s·ªë v·ªõi x·ª≠ l√Ω NaN"""
    if pd.isna(value):
        return "N/A"
    return f"{value:.{decimal_places}f}"


def format_percent(value, decimal_places=2):
    """Format t·ª∑ l·ªá ph·∫ßn trƒÉm"""
    if pd.isna(value):
        return "N/A"
    return f"{value:.{decimal_places}f}%"


# =============================================================================
# H√ÄM ƒê·ªåC D·ªÆ LI·ªÜU
# =============================================================================
def load_kpi_data(kpi_folder=DEFAULT_KPI_FOLDER):
    """
    ƒê·ªçc d·ªØ li·ªáu KPI t·ª´ c√°c file Excel
    
    Returns:
        tuple: (df_summary, df_detail) - DataFrame t√≥m t·∫Øt v√† chi ti·∫øt
    """
    kpi_path = Path(kpi_folder)
    
    # ƒê·ªçc file t√≥m t·∫Øt
    summary_file = kpi_path / "KPI_NVKT_TomTat.xlsx"
    df_summary = pd.read_excel(summary_file)
    
    # ƒê·ªçc file chi ti·∫øt
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    return df_summary, df_detail


def load_c1x_reports(data_folder="downloads/baocao_hanoi"):
    """
    ƒê·ªçc d·ªØ li·ªáu chi ti·∫øt t·ª´ c√°c file b√°o c√°o C1.x
    
    Returns:
        dict: Dictionary ch·ª©a c√°c DataFrame t·ª´ c√°c sheet t·ªïng h·ª£p
    """
    data_path = Path(data_folder)
    reports = {}
    
    # C1.1 Report
    try:
        c11_file = data_path / "c1.1 report.xlsx"
        if c11_file.exists():
            reports['c11'] = pd.read_excel(c11_file, sheet_name='TH_C1.1')
            print("   ‚úÖ ƒê·ªçc C1.1 report th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc C1.1 report: {e}")
    
    # C1.2 Report
    try:
        c12_file = data_path / "c1.2 report.xlsx"
        if c12_file.exists():
            reports['c12'] = pd.read_excel(c12_file, sheet_name='TH_C1.2')
            print("   ‚úÖ ƒê·ªçc C1.2 report th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc C1.2 report: {e}")
    
    # C1.3 Report
    try:
        c13_file = data_path / "c1.3 report.xlsx"
        if c13_file.exists():
            reports['c13'] = pd.read_excel(c13_file, sheet_name='TH_C1.3')
            print("   ‚úÖ ƒê·ªçc C1.3 report th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc C1.3 report: {e}")
    
    # C1.4 Report
    try:
        c14_file = data_path / "c1.4 report.xlsx"
        if c14_file.exists():
            reports['c14'] = pd.read_excel(c14_file, sheet_name='TH_C1.4')
            print("   ‚úÖ ƒê·ªçc C1.4 report th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc C1.4 report: {e}")
    
    # C1.5 Chi ti·∫øt Report - Sheet TH_TTVTST
    try:
        c15_file = data_path / "c1.5_chitiet_report.xlsx"
        if c15_file.exists():
            reports['c15_ttvtst'] = pd.read_excel(c15_file, sheet_name='TH_TTVTST')
            print("   ‚úÖ ƒê·ªçc C1.5 report (TH_TTVTST) th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc C1.5 report: {e}")
    
    return reports


def load_exclusion_comparison_data(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc d·ªØ li·ªáu so s√°nh tr∆∞·ªõc/sau gi·∫£m tr·ª´ t·ª´ c√°c file Excel
    
    Args:
        exclusion_folder: Th∆∞ m·ª•c ch·ª©a c√°c file k·∫øt qu·∫£ sau gi·∫£m tr·ª´
        
    Returns:
        dict: Dictionary ch·ª©a DataFrames cho t·ª´ng ch·ªâ ti√™u
              - 'c11_sm4': So s√°nh C1.1 SM4 (S·ª≠a ch·ªØa b√°o h·ªèng)
              - 'c11_sm2': So s√°nh C1.1 SM2 (S·ª≠a ch·ªØa ch·ªß ƒë·ªông)
              - 'c12_sm1': So s√°nh C1.2 SM1 (H·ªèng l·∫∑p l·∫°i)
              - 'c12_sm4': So s√°nh C1.2 SM4 (T·ª∑ l·ªá b√°o h·ªèng BRCƒê)
              - 'c14': So s√°nh C1.4 (ƒê·ªô h√†i l√≤ng)
              - 'c15': So s√°nh C1.5 (Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê)
              - 'tong_hop': T·ªïng h·ª£p t·∫•t c·∫£ ch·ªâ ti√™u
    """
    data_path = Path(exclusion_folder)
    comparison_data = {}
    
    if not data_path.exists():
        print(f"   ‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y th∆∞ m·ª•c gi·∫£m tr·ª´: {exclusion_folder}")
        return comparison_data
    
    # C1.1 SM4 - S·ª≠a ch·ªØa b√°o h·ªèng ƒë√∫ng quy ƒë·ªãnh
    try:
        c11_sm4_file = data_path / "So_sanh_C11_SM4.xlsx"
        if c11_sm4_file.exists():
            comparison_data['c11_sm4'] = {
                'chi_tiet': pd.read_excel(c11_sm4_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c11_sm4_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc So_sanh_C11_SM4.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc So_sanh_C11_SM4.xlsx: {e}")
    
    # C1.1 SM2 - S·ª≠a ch·ªØa ch·ªß ƒë·ªông
    try:
        c11_sm2_file = data_path / "So_sanh_C11_SM2.xlsx"
        if c11_sm2_file.exists():
            comparison_data['c11_sm2'] = {
                'chi_tiet': pd.read_excel(c11_sm2_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c11_sm2_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc So_sanh_C11_SM2.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc So_sanh_C11_SM2.xlsx: {e}")
    
    # C1.2 SM1 - H·ªèng l·∫∑p l·∫°i
    try:
        c12_sm1_file = data_path / "So_sanh_C12_SM1.xlsx"
        if c12_sm1_file.exists():
            comparison_data['c12_sm1'] = {
                'chi_tiet': pd.read_excel(c12_sm1_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c12_sm1_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc So_sanh_C12_SM1.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc So_sanh_C12_SM1.xlsx: {e}")
    
    # C1.2 SM4 - T·ª∑ l·ªá b√°o h·ªèng BRCƒê
    try:
        c12_sm4_file = data_path / "SM4-C12-ti-le-su-co-dv-brcd.xlsx"
        if c12_sm4_file.exists():
            comparison_data['c12_sm4'] = {
                'chi_tiet': pd.read_excel(c12_sm4_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c12_sm4_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc SM4-C12-ti-le-su-co-dv-brcd.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc SM4-C12: {e}")
    
    # C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng
    try:
        c14_file = data_path / "So_sanh_C14.xlsx"
        if c14_file.exists():
            comparison_data['c14'] = {
                'chi_tiet': pd.read_excel(c14_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c14_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc So_sanh_C14.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc So_sanh_C14.xlsx: {e}")
    
    
    # C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê ƒë·∫°t th·ªùi gian quy ƒë·ªãnh
    try:
        c15_file = data_path / "So_sanh_C15.xlsx"
        if c15_file.exists():
            comparison_data['c15'] = {
                'chi_tiet': pd.read_excel(c15_file, sheet_name='So_sanh_chi_tiet'),
                'tong_hop': pd.read_excel(c15_file, sheet_name='Thong_ke_tong_hop')
            }
            print("   ‚úÖ ƒê·ªçc So_sanh_C15.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc So_sanh_C15.xlsx: {e}")
    
    # T·ªïng h·ª£p gi·∫£m tr·ª´
    try:
        tong_hop_file = data_path / "Tong_hop_giam_tru.xlsx"
        if tong_hop_file.exists():
            comparison_data['tong_hop'] = pd.read_excel(tong_hop_file)
            print("   ‚úÖ ƒê·ªçc Tong_hop_giam_tru.xlsx th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc Tong_hop_giam_tru.xlsx: {e}")
    
    return comparison_data


def load_unit_level_exclusion_data(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc d·ªØ li·ªáu th·ªëng k√™ theo ƒë∆°n v·ªã (T·ªï) t·ª´ sheet Thong_ke_theo_don_vi
    
    Returns:
        dict: Dictionary ch·ª©a DataFrames th·ªëng k√™ theo ƒë∆°n v·ªã cho t·ª´ng ch·ªâ ti√™u
    """
    data_path = Path(exclusion_folder)
    unit_data = {}
    
    if not data_path.exists():
        print(f"   ‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y th∆∞ m·ª•c gi·∫£m tr·ª´: {exclusion_folder}")
        return unit_data
    
    # C1.1 SM4
    try:
        c11_sm4_file = data_path / "So_sanh_C11_SM4.xlsx"
        if c11_sm4_file.exists():
            unit_data['c11_sm4'] = pd.read_excel(c11_sm4_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.1 SM4 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.1 SM4: {e}")
    
    # C1.1 SM2
    try:
        c11_sm2_file = data_path / "So_sanh_C11_SM2.xlsx"
        if c11_sm2_file.exists():
            unit_data['c11_sm2'] = pd.read_excel(c11_sm2_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.1 SM2 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.1 SM2: {e}")
    
    # C1.2 SM1
    try:
        c12_sm1_file = data_path / "So_sanh_C12_SM1.xlsx"
        if c12_sm1_file.exists():
            unit_data['c12_sm1'] = pd.read_excel(c12_sm1_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.2 SM1 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.2 SM1: {e}")
    
    
    # C1.2 SM4 - T·ª∑ l·ªá s·ª± c·ªë BRCƒê
    try:
        c12_sm4_file = data_path / "SM4-C12-ti-le-su-co-dv-brcd.xlsx"
        if c12_sm4_file.exists():
            unit_data['c12_sm4'] = pd.read_excel(c12_sm4_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.2 SM4 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.2 SM4: {e}")
    
    # C1.4
    try:
        c14_file = data_path / "So_sanh_C14.xlsx"
        if c14_file.exists():
            unit_data['c14'] = pd.read_excel(c14_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.4 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.4: {e}")
    
    # C1.5
    try:
        c15_file = data_path / "So_sanh_C15.xlsx"
        if c15_file.exists():
            unit_data['c15'] = pd.read_excel(c15_file, sheet_name='Thong_ke_theo_don_vi')
            print("   ‚úÖ ƒê·ªçc unit stats C1.5 th√†nh c√¥ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc unit stats C1.5: {e}")
    
    return unit_data


def load_bsc_unit_scores_from_comparison(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc ƒëi·ªÉm BSC ƒë√£ t√≠nh s·∫µn t·ª´ file Tong_hop_Diem_BSC_Don_Vi.xlsx
    File n√†y ch·ª©a c·∫£ ƒëi·ªÉm Tr∆∞·ªõc v√† Sau gi·∫£m tr·ª´ cho t·ª´ng ƒë∆°n v·ªã v√† TTVT S∆°n T√¢y
    
    Returns:
        dict: Dictionary v·ªõi c·∫•u tr√∫c:
            {
                'units': DataFrame ch·ª©a ƒëi·ªÉm c√°c ƒë∆°n v·ªã (sheet Tong_hop_Don_vi),
                'individuals': DataFrame ch·ª©a ƒëi·ªÉm c√° nh√¢n (sheet Chi_tiet_Ca_nhan)
            }
    """
    file_path = Path(exclusion_folder) / "Tong_hop_Diem_BSC_Don_Vi.xlsx"
    result = {'units': None, 'individuals': None}
    
    if not file_path.exists():
        print(f"   ‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y file: {file_path}")
        return result
    
    try:
        result['units'] = pd.read_excel(file_path, sheet_name='Tong_hop_Don_vi')
        print(f"   ‚úÖ ƒê·ªçc ƒëi·ªÉm BSC ƒë∆°n v·ªã t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx: {len(result['units'])} d√≤ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc sheet Tong_hop_Don_vi: {e}")
    
    try:
        result['individuals'] = pd.read_excel(file_path, sheet_name='Chi_tiet_Ca_nhan')
        print(f"   ‚úÖ ƒê·ªçc ƒëi·ªÉm BSC c√° nh√¢n t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx: {len(result['individuals'])} d√≤ng")
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ ƒë·ªçc sheet Chi_tiet_Ca_nhan: {e}")
    
    return result


def load_nvkt_exclusion_summary(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc d·ªØ li·ªáu KPI sau gi·∫£m tr·ª´ theo NVKT t·ª´ c√°c file so s√°nh th√†nh ph·∫ßn
    (S·ª≠ d·ª•ng l·∫°i logic c·ªßa load_nvkt_exclusion_detail ƒë·ªÉ ƒë·∫£m b·∫£o nh·∫•t qu√°n)
    """
    # V√¨ load_nvkt_exclusion_detail ƒë√£ t·ªïng h·ª£p t·ª´ c√°c file g·ªëc v√† t√≠nh ƒëi·ªÉm l·∫°i,
    # n√™n ta c√≥ th·ªÉ d√πng l·∫°i k·∫øt qu·∫£ c·ªßa n√≥.
    return load_nvkt_exclusion_detail(exclusion_folder)


def add_kpi_summary_table_after_exclusion(doc, df_exclusion, team_name):
    """
    Th√™m b·∫£ng t·ªïng h·ª£p KPI sau gi·∫£m tr·ª´ v√†o document cho 1 t·ªï
    
    Args:
        doc: Document Word
        df_exclusion: DataFrame d·ªØ li·ªáu sau gi·∫£m tr·ª´
        team_name: T√™n t·ªï c·∫ßn l·ªçc
    """
    if df_exclusion is None or df_exclusion.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu sau gi·∫£m tr·ª´)")
        return
    
    # L·ªçc theo t·ªï
    df = df_exclusion[df_exclusion['don_vi'] == team_name].copy()
    if df.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu sau gi·∫£m tr·ª´ cho t·ªï n√†y)")
        return
    
    # S·∫Øp x·∫øp
    df = df.sort_values('nvkt')

    # T·∫°o b·∫£ng - bao g·ªìm C1.1, C1.2, C1.4, C1.5 sau gi·∫£m tr·ª´
    headers = ['STT', 'NVKT', 'C1.1', 'C1.2', 'C1.4', 'C1.5']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], '2E7D32')  # M√†u xanh l√° ƒë·ªÉ ph√¢n bi·ªát v·ªõi b·∫£ng th√¥
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells

        data = [
            str(idx),
            row.get('nvkt', ''),
            format_number(row.get('Diem_C1.1', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = str(value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(10)
            
            # M√†u n·ªÅn xen k·∫Ω
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')  # Xanh l√° nh·∫°t
            
            # T√¥ m√†u ƒëi·ªÉm theo m·ª©c (ch·ªâ c√°c c·ªôt ƒëi·ªÉm)
            if i >= 2:
                try:
                    val = float(value) if value and value != 'N/A' else None
                    if val is not None:
                        if val >= 4.5:
                            run.font.color.rgb = RGBColor(0, 128, 0)
                            run.font.bold = True
                        elif val < 3:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                            run.font.bold = True
                except (ValueError, TypeError):
                    pass


def chuan_hoa_ten_nvkt(name):
    """
    Chu·∫©n h√≥a t√™n NVKT v·ªÅ d·∫°ng Title Case ƒë·ªÉ tr√°nh tr√πng l·∫∑p do vi·∫øt hoa/th∆∞·ªùng kh√°c nhau
    V√≠ d·ª•: "B√πi vƒÉn C∆∞·ªùng" -> "B√πi VƒÉn C∆∞·ªùng"
    """
    if pd.isna(name) or name is None or str(name).strip() == '' or str(name) == 'nan':
        return None
    return str(name).strip().title()


def load_nvkt_raw_detail(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc d·ªØ li·ªáu KPI chi ti·∫øt TR∆Ø·ªöC gi·∫£m tr·ª´ (Th√¥) theo NVKT t·ª´ c√°c file so s√°nh
    S·ª≠ d·ª•ng c√°c c·ªôt c√≥ h·∫≠u t·ªë (Th√¥) thay v√¨ (Sau GT)
    """
    data_path = Path(exclusion_folder)
    if not data_path.exists():
        return None

    print("   üîÑ ƒêang t·ªïng h·ª£p d·ªØ li·ªáu chi ti·∫øt NVKT (Th√¥) t·ª´ c√°c file so s√°nh...")
    
    files = {
        'c11_sm2': ('So_sanh_C11_SM2.xlsx', 'So_sanh_chi_tiet'),
        'c11_sm4': ('So_sanh_C11_SM4.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm1': ('So_sanh_C12_SM1.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm4': ('SM4-C12-ti-le-su-co-dv-brcd.xlsx', 'So_sanh_chi_tiet'),
        'c14': ('So_sanh_C14.xlsx', 'So_sanh_chi_tiet'),
        'c15': ('So_sanh_C15.xlsx', 'So_sanh_chi_tiet')
    }
    
    dfs = {}
    for key, (filename, sheet) in files.items():
        try:
            f_path = data_path / filename
            if f_path.exists():
                df = pd.read_excel(f_path, sheet_name=sheet)
                if 'NVKT' in df.columns:
                    # Chu·∫©n h√≥a t√™n NVKT v·ªÅ Title Case ƒë·ªÉ tr√°nh tr√πng l·∫∑p
                    df['NVKT'] = df['NVKT'].apply(chuan_hoa_ten_nvkt)
                    df = df[df['NVKT'].notna()]  # Lo·∫°i b·ªè c√°c d√≤ng c√≥ NVKT null
                    dfs[key] = df
        except Exception as e:
            print(f"   ‚ö†Ô∏è L·ªói ƒë·ªçc file {filename}: {e}")

    if not dfs:
        return None

    # Collect all NVKTs (ƒë√£ ƒë∆∞·ª£c chu·∫©n h√≥a)
    all_nvkt = set()
    nvkt_info = {}
    for df in dfs.values():
        if 'NVKT' in df.columns:
            for _, row in df.iterrows():
                nvkt = row['NVKT']
                if nvkt:
                    all_nvkt.add(nvkt)
                    if nvkt not in nvkt_info:
                        don_vi = row.get('TEN_DOI', '') or row.get('ƒê∆°n v·ªã', '')
                        nvkt_info[nvkt] = {'don_vi': don_vi}

    if not all_nvkt:
        return None

    summary_data = []
    for nvkt in all_nvkt:
        info = nvkt_info.get(nvkt, {})
        row_data = {
            'nvkt': nvkt,
            'don_vi': info.get('don_vi', ''),
            'c11_tp1_tong_phieu': 0, 'c11_tp1_phieu_dat': 0, 'c11_tp1_ty_le': 0, 'diem_c11_tp1': 5,
            'c11_tp2_tong_phieu': 0, 'c11_tp2_phieu_dat': 0, 'c11_tp2_ty_le': 0, 'diem_c11_tp2': 5,
            'Diem_C1.1': 0,
            'c12_tp1_phieu_hll': 0, 'c12_tp1_phieu_bh': 0, 'c12_tp1_ty_le': 0, 'diem_c12_tp1': 5,
            'c12_tp2_tong_tb': 0, 'c12_tp2_phieu_bh': 0, 'c12_tp2_ty_le': 0, 'diem_c12_tp2': 5,
            'Diem_C1.2': 0,
            # C1.4 - M·∫∑c ƒë·ªãnh np.nan n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu (hi·ªÉn th·ªã N/A)
            'c14_phieu_ks': np.nan, 'c14_phieu_khl': np.nan, 'c14_ty_le': np.nan, 'diem_c14': np.nan, 'Diem_C1.4': np.nan,
            # C1.5 - M·∫∑c ƒë·ªãnh np.nan n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu (hi·ªÉn th·ªã N/A)
            'c15_tong_phieu': np.nan, 'c15_phieu_dat': np.nan, 'c15_phieu_khong_dat': np.nan, 'c15_ty_le': np.nan, 'diem_c15': np.nan, 'Diem_C1.5': np.nan
        }

        # C1.1 SM2 (TP1) - Th√¥
        if 'c11_sm2' in dfs:
            r = dfs['c11_sm2'][dfs['c11_sm2']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp1_tong_phieu'] = r.get('T·ªïng phi·∫øu (Th√¥)', 0)
                row_data['c11_tp1_phieu_dat'] = r.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 0)
                row_data['c11_tp1_ty_le'] = r.get('T·ª∑ l·ªá % (Th√¥)', 0)
                row_data['diem_c11_tp1'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)
        
        # C1.1 SM4 (TP2) - Th√¥
        if 'c11_sm4' in dfs:
            r = dfs['c11_sm4'][dfs['c11_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp2_tong_phieu'] = r.get('T·ªïng phi·∫øu (Th√¥)', 0)
                row_data['c11_tp2_phieu_dat'] = r.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 0)
                row_data['c11_tp2_ty_le'] = r.get('T·ª∑ l·ªá % (Th√¥)', 0)
                row_data['diem_c11_tp2'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)
        
        row_data['Diem_C1.1'] = 0.3 * row_data['diem_c11_tp1'] + 0.7 * row_data['diem_c11_tp2']

        # C1.2 SM1 (TP1) - Th√¥
        if 'c12_sm1' in dfs:
            r = dfs['c12_sm1'][dfs['c12_sm1']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp1_phieu_hll'] = r.get('S·ªë phi·∫øu HLL (Th√¥)', 0)
                row_data['c12_tp1_phieu_bh'] = r.get('S·ªë phi·∫øu b√°o h·ªèng (Th√¥)', 0)
                row_data['c12_tp1_ty_le'] = r.get('T·ª∑ l·ªá HLL % (Th√¥)', 0)
                row_data['diem_c12_tp1'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)
        
        # C1.2 SM4 (TP2) - Th√¥
        if 'c12_sm4' in dfs:
            r = dfs['c12_sm4'][dfs['c12_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp2_tong_tb'] = r.get('T·ªïng TB (Th√¥)', 0)
                row_data['c12_tp2_phieu_bh'] = r.get('S·ªë phi·∫øu b√°o h·ªèng (Th√¥)', 0)
                row_data['c12_tp2_ty_le'] = r.get('T·ª∑ l·ªá b√°o h·ªèng (%) (Th√¥)', 0)
                row_data['diem_c12_tp2'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)

        row_data['Diem_C1.2'] = 0.5 * row_data['diem_c12_tp1'] + 0.5 * row_data['diem_c12_tp2']

        # C1.4 - Th√¥
        if 'c14' in dfs:
            r = dfs['c14'][dfs['c14']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c14_phieu_ks'] = r.get('T·ªïng phi·∫øu KS (Th√¥)', 0)
                row_data['c14_phieu_khl'] = r.get('S·ªë phi·∫øu KHL (Th√¥)', 0)
                row_data['c14_ty_le'] = r.get('T·ª∑ l·ªá HL (%) (Th√¥)', 0)
                row_data['diem_c14'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)
                row_data['Diem_C1.4'] = row_data['diem_c14']

        # C1.5 - Th√¥
        if 'c15' in dfs:
            r = dfs['c15'][dfs['c15']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c15_tong_phieu'] = r.get('T·ªïng Ho√†n c√¥ng (Th√¥)', 0)
                row_data['c15_phieu_dat'] = r.get('Phi·∫øu ƒë·∫°t (Th√¥)', 0)
                row_data['c15_phieu_khong_dat'] = r.get('Phi·∫øu kh√¥ng ƒë·∫°t (Th√¥)', 0)
                row_data['c15_ty_le'] = r.get('T·ª∑ l·ªá ƒë·∫°t (%) (Th√¥)', 0)
                row_data['diem_c15'] = r.get('ƒêi·ªÉm BSC (Th√¥)', 5)
                row_data['Diem_C1.5'] = row_data['diem_c15']

        summary_data.append(row_data)

    df_result = pd.DataFrame(summary_data)
    print(f"   ‚úÖ T·ªïng h·ª£p xong d·ªØ li·ªáu chi ti·∫øt NVKT Th√¥ ({len(df_result)} nh√¢n vi√™n)")
    return df_result



def load_nvkt_exclusion_detail(exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    ƒê·ªçc d·ªØ li·ªáu KPI chi ti·∫øt sau gi·∫£m tr·ª´ theo NVKT t·ª´ c√°c file so s√°nh th√†nh ph·∫ßn
    """
    data_path = Path(exclusion_folder)
    if not data_path.exists():
        return None

    print("   üîÑ ƒêang t·ªïng h·ª£p d·ªØ li·ªáu chi ti·∫øt NVKT t·ª´ c√°c file so s√°nh...")
    
    # Danh s√°ch c√°c file c·∫ßn ƒë·ªçc
    files = {
        'c11_sm2': ('So_sanh_C11_SM2.xlsx', 'So_sanh_chi_tiet'),
        'c11_sm4': ('So_sanh_C11_SM4.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm1': ('So_sanh_C12_SM1.xlsx', 'So_sanh_chi_tiet'),
        'c12_sm4': ('SM4-C12-ti-le-su-co-dv-brcd.xlsx', 'So_sanh_chi_tiet'),
        'c14': ('So_sanh_C14.xlsx', 'So_sanh_chi_tiet'),
        'c15': ('So_sanh_C15.xlsx', 'So_sanh_chi_tiet')
    }
    
    dfs = {}
    
    for key, (filename, sheet) in files.items():
        try:
            # ƒê·ªçc file, b·ªè qua d√≤ng ti√™u ƒë·ªÅ ph·ª• n·∫øu c√≥ (th∆∞·ªùng header=0 l√† ƒë·ªß n·∫øu c·ªôt n·∫±m ·ªü d√≤ng 1)
            f_path = data_path / filename
            if f_path.exists():
                df = pd.read_excel(f_path, sheet_name=sheet)
                # Chu·∫©n h√≥a t√™n c·ªôt NVKT v√† TEN_DOI
                if 'M√£ nh√¢n vi√™n' in df.columns:
                    df.rename(columns={'M√£ nh√¢n vi√™n': 'NVKT'}, inplace=True)
                if 'T√™n nh√¢n vi√™n' in df.columns:
                    df.rename(columns={'T√™n nh√¢n vi√™n': 'TEN_NV'}, inplace=True)
                
                # ƒê·∫£m b·∫£o c√≥ c·ªôt NVKT ƒë·ªÉ merge
                if 'NVKT' in df.columns:
                    # Chu·∫©n h√≥a t√™n NVKT v·ªÅ Title Case ƒë·ªÉ tr√°nh tr√πng l·∫∑p
                    df['NVKT'] = df['NVKT'].apply(chuan_hoa_ten_nvkt)
                    df = df[df['NVKT'].notna()]  # Lo·∫°i b·ªè c√°c d√≤ng c√≥ NVKT null
                    dfs[key] = df
                else:
                    print(f"   ‚ö†Ô∏è File {filename}: Kh√¥ng t√¨m th·∫•y c·ªôt NVKT")
            else:
                print(f"   ‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y file {filename}")
        except Exception as e:
             print(f"   ‚ö†Ô∏è L·ªói ƒë·ªçc file {filename}: {e}")

    if not dfs:
        return None

    # L·∫•y danh s√°ch t·∫•t c·∫£ NVKT t·ª´ c√°c file (ƒë√£ ƒë∆∞·ª£c chu·∫©n h√≥a Title Case)
    all_nvkt = set()
    nvkt_info = {} # L∆∞u th√¥ng tin NVKT (T√™n, T·ªï)

    for df in dfs.values():
        if 'NVKT' in df.columns:
            for _, row in df.iterrows():
                nvkt = row['NVKT']
                if nvkt:
                    all_nvkt.add(nvkt)
                    # L∆∞u th√¥ng tin b·ªï sung n·∫øu ch∆∞a c√≥
                    if nvkt not in nvkt_info:
                        don_vi = row.get('TEN_DOI', '') or row.get('ƒê∆°n v·ªã', '')
                        # Logic: Gi·ªØ nguy√™n t√™n ƒë∆°n v·ªã t·ª´ file, sau n√†y l·ªçc theo t√™n ƒë√≥
                        nvkt_info[nvkt] = {
                            'don_vi': don_vi,
                            'ten_nv': row.get('TEN_NV', '') or row.get('T√™n nh√¢n vi√™n', '')
                        }

    if not all_nvkt:
        return None

    # T·∫°o DataFrame t·ªïng h·ª£p
    summary_data = []

    for nvkt in all_nvkt:
        info = nvkt_info.get(nvkt, {})
        row_data = {
            'nvkt': nvkt,
            'don_vi': info.get('don_vi', ''),
            'ten_nv': info.get('ten_nv', ''),
            # C1.1
            'c11_tp1_tong_phieu': 0, 'c11_tp1_phieu_dat': 0, 'c11_tp1_ty_le': 0, 'diem_c11_tp1': 5,
            'c11_tp2_tong_phieu': 0, 'c11_tp2_phieu_dat': 0, 'c11_tp2_ty_le': 0, 'diem_c11_tp2': 5,
            'Diem_C1.1': 0,
            # C1.2 - column names fixed to match table renderer
            'c12_tp1_phieu_hll': 0, 'c12_tp1_phieu_bh': 0, 'c12_tp1_ty_le': 0, 'diem_c12_tp1': 5,
            'c12_tp2_tong_tb': 0, 'c12_tp2_phieu_bh': 0, 'c12_tp2_ty_le': 0, 'diem_c12_tp2': 5,
            'Diem_C1.2': 0,
            # C1.4 - M·∫∑c ƒë·ªãnh np.nan n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu (hi·ªÉn th·ªã N/A)
            'c14_phieu_ks': np.nan, 'c14_phieu_khl': np.nan, 'c14_ty_le': np.nan, 'diem_c14': np.nan, 'Diem_C1.4': np.nan,
            # C1.5 - M·∫∑c ƒë·ªãnh np.nan n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu (hi·ªÉn th·ªã N/A)
            'c15_tong_phieu': np.nan, 'c15_phieu_dat': np.nan, 'c15_phieu_khong_dat': np.nan, 'c15_ty_le': np.nan, 'diem_c15': np.nan, 'Diem_C1.5': np.nan
        }
        
        # Fill C1.1 SM2 (TP1)
        if 'c11_sm2' in dfs:
            r = dfs['c11_sm2'][dfs['c11_sm2']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp1_tong_phieu'] = r.get('T·ªïng phi·∫øu (Sau GT)', 0)
                row_data['c11_tp1_phieu_dat'] = r.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)', 0)
                row_data['c11_tp1_ty_le'] = r.get('T·ª∑ l·ªá % (Sau GT)', 0)
                row_data['diem_c11_tp1'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)
        
        # Fill C1.1 SM4 (TP2)
        if 'c11_sm4' in dfs:
            r = dfs['c11_sm4'][dfs['c11_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c11_tp2_tong_phieu'] = r.get('T·ªïng phi·∫øu (Sau GT)', 0)
                row_data['c11_tp2_phieu_dat'] = r.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)', 0)
                row_data['c11_tp2_ty_le'] = r.get('T·ª∑ l·ªá % (Sau GT)', 0)
                row_data['diem_c11_tp2'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)
        
        # Calculate C1.1 Score
        row_data['Diem_C1.1'] = 0.3 * row_data['diem_c11_tp1'] + 0.7 * row_data['diem_c11_tp2']

        # Fill C1.2 SM1 (TP1) - H·ªèng l·∫∑p l·∫°i
        if 'c12_sm1' in dfs:
            r = dfs['c12_sm1'][dfs['c12_sm1']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp1_phieu_hll'] = r.get('S·ªë phi·∫øu HLL (Sau GT)', 0)
                row_data['c12_tp1_phieu_bh'] = r.get('S·ªë phi·∫øu b√°o h·ªèng (Sau GT)', 0)
                row_data['c12_tp1_ty_le'] = r.get('T·ª∑ l·ªá HLL % (Sau GT)', 0)
                row_data['diem_c12_tp1'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)
        
        # Fill C1.2 SM4 (TP2) - T·ª∑ l·ªá s·ª± c·ªë
        if 'c12_sm4' in dfs:
            r = dfs['c12_sm4'][dfs['c12_sm4']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c12_tp2_tong_tb'] = r.get('T·ªïng TB (Th√¥)', 0)
                row_data['c12_tp2_phieu_bh'] = r.get('S·ªë phi·∫øu b√°o h·ªèng (Sau GT)', 0)
                row_data['c12_tp2_ty_le'] = r.get('T·ª∑ l·ªá b√°o h·ªèng (%) (Sau GT)', 0)
                row_data['diem_c12_tp2'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)

        # Calculate C1.2 Score
        row_data['Diem_C1.2'] = 0.5 * row_data['diem_c12_tp1'] + 0.5 * row_data['diem_c12_tp2']

        # Fill C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng
        if 'c14' in dfs:
            r = dfs['c14'][dfs['c14']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c14_phieu_ks'] = r.get('T·ªïng phi·∫øu KS (Sau GT)', 0)
                row_data['c14_phieu_khl'] = r.get('S·ªë phi·∫øu KHL (Sau GT)', 0)
                row_data['c14_ty_le'] = r.get('T·ª∑ l·ªá HL (%) (Sau GT)', 0)
                row_data['diem_c14'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)
                row_data['Diem_C1.4'] = row_data['diem_c14']

        # Fill C1.5 - Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê
        if 'c15' in dfs:
            r = dfs['c15'][dfs['c15']['NVKT'] == nvkt]
            if not r.empty:
                r = r.iloc[0]
                row_data['c15_tong_phieu'] = r.get('T·ªïng Ho√†n c√¥ng (Sau GT)', 0)
                row_data['c15_phieu_dat'] = r.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
                row_data['c15_phieu_khong_dat'] = r.get('Phi·∫øu kh√¥ng ƒë·∫°t (Sau GT)', 0)
                row_data['c15_ty_le'] = r.get('T·ª∑ l·ªá ƒë·∫°t (%) (Sau GT)', 0)
                row_data['diem_c15'] = r.get('ƒêi·ªÉm BSC (Sau GT)', 5)
                row_data['Diem_C1.5'] = row_data['diem_c15']

        summary_data.append(row_data)

    df_result = pd.DataFrame(summary_data)
    print(f"   ‚úÖ T·ªïng h·ª£p xong d·ªØ li·ªáu chi ti·∫øt NVKT ({len(df_result)} nh√¢n vi√™n)")
    return df_result


def add_c11_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.1 sau gi·∫£m tr·ª´
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.1 sau gi·∫£m tr·ª´)")
        return
    
    df = df_exclusion_detail[df_exclusion_detail['don_vi'] == team_name].copy()
    if df.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.1 sau gi·∫£m tr·ª´ cho t·ªï n√†y)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.1 - Ch·∫•t l∆∞·ª£ng s·ª≠a ch·ªØa thu√™ bao BRCƒê (sau gi·∫£m tr·ª´)', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('TP1 = S·ª≠a ch·ªØa ch·ªß ƒë·ªông (SCCD ‚â§72h) | TP2 = S·ª≠a ch·ªØa theo b√°o h·ªèng (SC BH) | Sau GT = Sau gi·∫£m tr·ª´')
    
    headers = ['STT', 'NVKT', 'T·ªïng SCCD', 'ƒê·∫°t ‚â§72h', 'TL(%)', 'ƒêi·ªÉm TP1',
               'T·ªïng SC BH', 'ƒê√∫ng h·∫°n', 'TL(%)', 'ƒêi·ªÉm TP2', 'ƒêi·ªÉm C1.1']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - m√†u xanh l√° ƒë·∫≠m h∆°n ƒë·ªÉ ph√¢n bi·ªát
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '1B5E20')  # Xanh l√° ƒë·∫≠m
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c11_tp1_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp1_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp1', np.nan)),
            format_number(row.get('c11_tp2_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp2_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp2', np.nan)),
            format_number(row.get('Diem_C1.1', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'C8E6C9')  # Xanh l√° nh·∫°t h∆°n


def add_c12_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.2 sau gi·∫£m tr·ª´
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.2 sau gi·∫£m tr·ª´)")
        return
    
    df = df_exclusion_detail[df_exclusion_detail['don_vi'] == team_name].copy()
    if df.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.2 sau gi·∫£m tr·ª´ cho t·ªï n√†y)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.2 - T·ª∑ l·ªá thu√™ bao b√°o h·ªèng (sau gi·∫£m tr·ª´)', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('TP1 = H·ªèng l·∫∑p (‚â•2 l·∫ßn/7 ng√†y) | TP2 = T·ª∑ l·ªá BH/TB qu·∫£n l√Ω | Sau GT = Sau gi·∫£m tr·ª´')
    
    headers = ['STT', 'NVKT', 'H·ªèng l·∫∑p', 'T·ªïng BH', 'TL(%)', 'ƒêi·ªÉm TP1',
               'Phi·∫øu BH', 'TB QL', 'TL(‚Ä∞)', 'ƒêi·ªÉm TP2', 'ƒêi·ªÉm C1.2']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - m√†u xanh d∆∞∆°ng ƒë·∫≠m h∆°n
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '0D47A1')  # Xanh d∆∞∆°ng ƒë·∫≠m
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c12_tp1_phieu_hll', np.nan), 0),
            format_number(row.get('c12_tp1_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp1', np.nan)),
            format_number(row.get('c12_tp2_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp2_tong_tb', np.nan), 0),
            format_number(row.get('c12_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp2', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'BBDEFB')  # Xanh d∆∞∆°ng nh·∫°t


def add_c14_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.4 sau gi·∫£m tr·ª´ - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.4 sau gi·∫£m tr·ª´)")
        return
    
    df = df_exclusion_detail[df_exclusion_detail['don_vi'] == team_name].copy()
    if df.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.4 sau gi·∫£m tr·ª´ cho t·ªï n√†y)")
        return
    
    df = df.sort_values('nvkt')
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng (sau gi·∫£m tr·ª´)', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('KS = Kh·∫£o s√°t | Kh√¥ng HL = Kh√¥ng h√†i l√≤ng | Sau GT = Sau gi·∫£m tr·ª´')
    
    headers = ['STT', 'NVKT', 'T·ªïng KS', 'Kh√¥ng HL', 'T·ª∑ l·ªá HL (%)', 'ƒêi·ªÉm C1.4']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header - m√†u cam ƒë·∫≠m h∆°n
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'E65100')  # Cam ƒë·∫≠m
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c14_phieu_ks', np.nan), 0),
            format_number(row.get('c14_phieu_khl', np.nan), 0),
            format_number(row.get('c14_ty_le', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFE0B2')  # Cam nh·∫°t


def add_c15_detail_table_after_exclusion(doc, df_exclusion_detail, team_name):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.5 sau gi·∫£m tr·ª´ - Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê
    """
    if df_exclusion_detail is None or df_exclusion_detail.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.5 sau gi·∫£m tr·ª´)")
        return

    df = df_exclusion_detail[df_exclusion_detail['don_vi'] == team_name].copy()
    if df.empty:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu chi ti·∫øt C1.5 sau gi·∫£m tr·ª´ cho t·ªï n√†y)")
        return

    df = df.sort_values('nvkt')

    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.5 - Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê ƒë·∫°t th·ªùi gian quy ƒë·ªãnh (sau gi·∫£m tr·ª´)', level=3)

    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('ƒê·∫°t TG = Ho√†n th√†nh ƒë√∫ng th·ªùi gian | Sau GT = Sau gi·∫£m tr·ª´ (lo·∫°i b·ªè phi·∫øu lo·∫°i tr·ª´)')

    headers = ['STT', 'NVKT', 'ƒê·∫°t TG', 'Kh√¥ng ƒë·∫°t', 'T·ªïng phi·∫øu', 'T·ª∑ l·ªá (%)', 'ƒêi·ªÉm C1.5']

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header - m√†u t√≠m ƒë·∫≠m h∆°n ƒë·ªÉ ph√¢n bi·ªát v·ªõi b·∫£ng tr∆∞·ªõc
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '4A148C')  # T√≠m ƒë·∫≠m
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c15_phieu_dat', np.nan), 0),
            format_number(row.get('c15_phieu_khong_dat', np.nan), 0),
            format_number(row.get('c15_tong_phieu', np.nan), 0),
            format_number(row.get('c15_ty_le', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)

            if idx % 2 == 0:
                set_cell_shading(cells[i], 'EDE7F6')  # T√≠m nh·∫°t


def create_nvkt_bar_chart_after_exclusion(df_exclusion, team_name, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì c·ªôt so s√°nh ƒëi·ªÉm KPI sau gi·∫£m tr·ª´ theo NVKT trong 1 t·ªï
    
    Args:
        df_exclusion: DataFrame ch·ª©a d·ªØ li·ªáu KPI sau gi·∫£m tr·ª´
        team_name: T√™n t·ªï c·∫ßn t·∫°o bi·ªÉu ƒë·ªì
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file (None = tr·∫£ v·ªÅ bytes)
    
    Returns:
        bytes ho·∫∑c str, None n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu
    """
    if df_exclusion is None or df_exclusion.empty:
        return None
    
    # L·ªçc theo t·ªï
    df = df_exclusion[df_exclusion['don_vi'] == team_name].copy()
    if df.empty or len(df) == 0:
        return None
    
    # S·∫Øp x·∫øp theo t√™n
    df = df.sort_values('nvkt')
    
    # L·∫•y t√™n ng·∫Øn c·ªßa t·ªï
    short_name = TEAM_SHORT_NAMES.get(team_name, team_name)
    
    # Chu·∫©n b·ªã d·ªØ li·ªáu
    nvkts = df['nvkt'].tolist()
    c11 = df['Diem_C1.1'].fillna(0).tolist()
    c12 = df['Diem_C1.2'].fillna(0).tolist()
    c14 = df['Diem_C1.4'].fillna(0).tolist()
    c15 = df['Diem_C1.5'].fillna(0).tolist()

    # T·∫°o bi·ªÉu ƒë·ªì
    fig, ax = plt.subplots(figsize=(12, 6))

    x = np.arange(len(nvkts))
    width = 0.2

    # C√°c c·ªôt - bao g·ªìm C1.1, C1.2, C1.4, C1.5 sau gi·∫£m tr·ª´
    bars1 = ax.bar(x - 1.5*width, c11, width, label='C1.1', color='#66BB6A')  # Xanh l√°
    bars2 = ax.bar(x - 0.5*width, c12, width, label='C1.2', color='#42A5F5')  # Xanh d∆∞∆°ng
    bars3 = ax.bar(x + 0.5*width, c14, width, label='C1.4', color='#FFA726')  # Cam
    bars4 = ax.bar(x + 1.5*width, c15, width, label='C1.5', color='#AB47BC')  # T√≠m

    # Th√™m gi√° tr·ªã l√™n c·ªôt
    for bars in [bars1, bars2, bars3, bars4]:
        for bar in bars:
            height = bar.get_height()
            if height > 0:
                ax.annotate(f'{height:.1f}',
                           xy=(bar.get_x() + bar.get_width() / 2, height),
                           xytext=(0, 3),
                           textcoords="offset points",
                           ha='center', va='bottom', fontsize=7)
    
    ax.set_xlabel('Nh√¢n vi√™n k·ªπ thu·∫≠t', fontsize=11)
    ax.set_ylabel('ƒêi·ªÉm KPI', fontsize=11)
    ax.set_title(f'SO S√ÅNH ƒêI·ªÇM KPI SAU GI·∫¢M TR·ª™ - {short_name.upper()}', fontsize=13, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(nvkts, rotation=45, ha='right', fontsize=9)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_unit_comparison_chart(unit_data, chi_tieu='c11_sm4', output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì grouped bar so s√°nh t·ª∑ l·ªá tr∆∞·ªõc/sau GT theo ƒë∆°n v·ªã
    
    Args:
        unit_data: Dictionary t·ª´ load_unit_level_exclusion_data()
        chi_tieu: 'c11_sm4', 'c11_sm2', 'c12_sm1', 'c14'
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file
    """
    if not unit_data or chi_tieu not in unit_data:
        return None
    
    df = unit_data[chi_tieu]
    
    # L·∫•y c·ªôt t·ª∑ l·ªá
    tyle_tho_col = None
    tyle_sau_col = None
    for col in df.columns:
        if 'T·ª∑ l·ªá' in col and 'Th√¥' in col:
            tyle_tho_col = col
        elif 'T·ª∑ l·ªá' in col and 'Sau GT' in col:
            tyle_sau_col = col
    
    if not tyle_tho_col or not tyle_sau_col:
        return None
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    width = 0.35
    
    don_vi = df['ƒê∆°n v·ªã'].values
    tyle_tho = df[tyle_tho_col].fillna(0).values
    tyle_sau = df[tyle_sau_col].fillna(0).values
    
    bars1 = ax.bar(x - width/2, tyle_tho, width, label='Tr∆∞·ªõc gi·∫£m tr·ª´', color='#EF5350', alpha=0.8)
    bars2 = ax.bar(x + width/2, tyle_sau, width, label='Sau gi·∫£m tr·ª´', color='#66BB6A', alpha=0.8)
    
    # Th√™m gi√° tr·ªã l√™n c·ªôt
    for bar, val in zip(bars1, tyle_tho):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
    for bar, val in zip(bars2, tyle_sau):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.3,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=9)
    
    # Ti√™u ƒë·ªÅ theo ch·ªâ ti√™u
    titles = {
        'c11_sm4': 'C1.1 SM4 - S·ª≠a ch·ªØa b√°o h·ªèng',
        'c11_sm2': 'C1.1 SM2 - S·ª≠a ch·ªØa ch·ªß ƒë·ªông', 
        'c12_sm1': 'C1.2 SM1 - H·ªèng l·∫∑p l·∫°i',
        'c14': 'C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng'
    }
    
    ax.set_xlabel('ƒê∆°n v·ªã', fontsize=11)
    ax.set_ylabel('T·ª∑ l·ªá (%)', fontsize=11)
    ax.set_title(f'SO S√ÅNH TR∆Ø·ªöC/SAU GI·∫¢M TR·ª™ THEO ƒê∆†N V·ªä\n{titles.get(chi_tieu, chi_tieu)}', 
                fontsize=12, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(don_vi, rotation=15, ha='right', fontsize=10)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_unit_exclusion_table(doc, unit_data, chi_tieu='c11_sm4'):
    """
    Th√™m b·∫£ng th·ªëng k√™ theo ƒë∆°n v·ªã v√†o document
    
    Args:
        doc: Document Word
        unit_data: Dictionary t·ª´ load_unit_level_exclusion_data()
        chi_tieu: 'c11_sm4', 'c11_sm2', 'c12_sm1', 'c14'
    """
    if not unit_data or chi_tieu not in unit_data:
        return
    
    df = unit_data[chi_tieu]
    
    # ƒê·ªãnh nghƒ©a ti√™u ƒë·ªÅ v√† m√†u theo ch·ªâ ti√™u
    config = {
        'c11_sm4': {'title': 'C1.1 SM4 - S·ª≠a ch·ªØa b√°o h·ªèng theo ƒë∆°n v·ªã (Sau GT)', 'color': 'C62828'},
        'c11_sm2': {'title': 'C1.1 SM2 - S·ª≠a ch·ªØa ch·ªß ƒë·ªông theo ƒë∆°n v·ªã (Sau GT)', 'color': 'AD1457'},
        'c12_sm1': {'title': 'C1.2 SM1 - H·ªèng l·∫∑p l·∫°i theo ƒë∆°n v·ªã (Sau GT)', 'color': '0D47A1'},
        'c14': {'title': 'C1.4 - ƒê·ªô h√†i l√≤ng theo ƒë∆°n v·ªã (Sau GT)', 'color': 'E65100'}
    }
    
    cfg = config.get(chi_tieu, {'title': chi_tieu, 'color': '333333'})
    
    doc.add_heading(cfg['title'], level=4)
    
    # L·∫•y c√°c c·ªôt c·∫ßn hi·ªÉn th·ªã
    display_cols = ['ƒê∆°n v·ªã', 'T·ªïng phi·∫øu (Th√¥)', 'Phi·∫øu lo·∫°i tr·ª´', 'T·ªïng phi·∫øu (Sau GT)']
    tyle_cols = [c for c in df.columns if 'T·ª∑ l·ªá' in c]
    thay_doi_cols = [c for c in df.columns if 'Thay ƒë·ªïi' in c]
    
    headers = display_cols + tyle_cols[:2] + thay_doi_cols[:1]
    headers = [h for h in headers if h in df.columns]
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        # R√∫t g·ªçn t√™n header
        h = headers[i]
        short_h = h.replace('(Th√¥)', '(T)').replace('(Sau GT)', '(S)').replace('T·ªïng phi·∫øu', 'T·ªïng')
        header.text = short_h
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, cfg['color'])
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            val = row.get(h, '')
            if pd.isna(val):
                val = ''
            elif isinstance(val, (int, float)):
                if 'T·ª∑ l·ªá' in h or 'Thay ƒë·ªïi' in h:
                    val = f"{val:.2f}%"
                else:
                    val = str(int(val))
            cells[i].text = str(val)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            # T√¥ m√†u cho d√≤ng TTVT
            if 'TTVT' in str(row.get('ƒê∆°n v·ªã', '')):
                run.font.bold = True
                set_cell_shading(cells[i], 'E0E0E0')
            elif idx % 2 == 0:
                set_cell_shading(cells[i], 'F5F5F5')
    
    doc.add_paragraph()


def add_unit_level_exclusion_section(doc, unit_data, c1x_reports=None):
    """
    Th√™m ph·∫ßn th·ªëng k√™ gi·∫£m tr·ª´ theo ƒë∆°n v·ªã v√†o document
    Bao g·ªìm bi·ªÉu ƒë·ªì BSC sau gi·∫£m tr·ª´, b·∫£ng v√† bi·ªÉu ƒë·ªì chi ti·∫øt cho t·ª´ng ch·ªâ ti√™u
    """
    if not unit_data:
        return
    
    doc.add_heading('Th·ªëng k√™ gi·∫£m tr·ª´ theo ƒë∆°n v·ªã (T·ªï)', level=3)
    
    p = doc.add_paragraph()
    p.add_run('üìä S·ªë li·ªáu d∆∞·ªõi ƒë√¢y th·ªÉ hi·ªán k·∫øt qu·∫£ c√°c ch·ªâ ti√™u BSC tr∆∞·ªõc v√† sau gi·∫£m tr·ª´, ')
    p.add_run('ƒë∆∞·ª£c t·ªïng h·ª£p theo t·ª´ng T·ªï k·ªπ thu·∫≠t v√† to√†n TTVT S∆°n T√¢y.')
    doc.add_paragraph()
    
    # (Bi·ªÉu ƒë·ªì BSC sau gi·∫£m tr·ª´ ƒë√£ ƒë∆∞·ª£c ƒë·∫∑t ·ªü section 1.1.b - tr∆∞·ªõc ph·∫ßn n√†y)
    
    # C1.1 SM4
    if 'c11_sm4' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c11_sm4')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c11_sm4')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì C1.1 SM4: {e}")
        doc.add_paragraph()
    
    # C1.1 SM2
    if 'c11_sm2' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c11_sm2')
        doc.add_paragraph()
    
    # C1.2 SM1
    if 'c12_sm1' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c12_sm1')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c12_sm1')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì C1.2 SM1: {e}")
        doc.add_paragraph()
    
    # C1.4
    if 'c14' in unit_data:
        add_unit_exclusion_table(doc, unit_data, 'c14')
        try:
            chart = create_unit_comparison_chart(unit_data, 'c14')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì C1.4: {e}")
        doc.add_paragraph()


def create_comparison_bar_chart(comparison_data, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì grouped bar so s√°nh t·ª∑ l·ªá tr∆∞·ªõc/sau gi·∫£m tr·ª´
    
    Args:
        comparison_data: Dictionary t·ª´ load_exclusion_comparison_data()
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file (None = tr·∫£ v·ªÅ bytes)
    
    Returns:
        bytes ho·∫∑c str: D·ªØ li·ªáu ·∫£nh
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        return None
    
    df = comparison_data['tong_hop']
    
    # S·∫Øp x·∫øp theo th·ª© t·ª±
    chi_tieu_order = ['C1.1 SM4', 'C1.1 SM2', 'C1.2', 'C1.2 T·ª∑ l·ªá BRCƒê b√°o h·ªèng', 'C1.4 ƒê·ªô h√†i l√≤ng KH']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    width = 0.35
    
    tyle_tho = df['T·ª∑ l·ªá % (Th√¥)'].fillna(0).values
    tyle_sau = df['T·ª∑ l·ªá % (Sau GT)'].fillna(0).values
    chi_tieu = df['Ch·ªâ ti√™u'].values
    
    bars1 = ax.bar(x - width/2, tyle_tho, width, label='Tr∆∞·ªõc gi·∫£m tr·ª´', color='#E57373', alpha=0.8)
    bars2 = ax.bar(x + width/2, tyle_sau, width, label='Sau gi·∫£m tr·ª´', color='#81C784', alpha=0.8)
    
    # Th√™m gi√° tr·ªã l√™n c·ªôt
    for bar, val in zip(bars1, tyle_tho):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                   f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
    for bar, val in zip(bars2, tyle_sau):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                   f'{val:.1f}%', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Ch·ªâ ti√™u', fontsize=11)
    ax.set_ylabel('T·ª∑ l·ªá (%)', fontsize=11)
    ax.set_title('SO S√ÅNH T·ª∂ L·ªÜ TR∆Ø·ªöC/SAU GI·∫¢M TR·ª™', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(chi_tieu, rotation=15, ha='right', fontsize=9)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_exclusion_summary_table(doc, comparison_data):
    """
    Th√™m b·∫£ng t·ªïng h·ª£p so s√°nh tr∆∞·ªõc/sau gi·∫£m tr·ª´ v√†o document
    
    Args:
        doc: Document Word
        comparison_data: Dictionary t·ª´ load_exclusion_comparison_data()
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        doc.add_paragraph("‚ö†Ô∏è Kh√¥ng c√≥ d·ªØ li·ªáu gi·∫£m tr·ª´")
        return
    
    df = comparison_data['tong_hop']
    
    doc.add_heading('B·∫¢NG T·ªîNG H·ª¢P SO S√ÅNH TR∆Ø·ªöC/SAU GI·∫¢M TR·ª™', level=3)
    
    headers = ['Ch·ªâ ti√™u', 'T·ªïng phi·∫øu (Th√¥)', 'Lo·∫°i tr·ª´', 'T·ªïng phi·∫øu (Sau GT)', 
               'T·ª∑ l·ªá % (Th√¥)', 'T·ª∑ l·ªá % (Sau GT)', 'Thay ƒë·ªïi %']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'D32F2F')  # ƒê·ªè ƒë·∫≠m
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        
        thay_doi = row.get('Thay ƒë·ªïi %', 0)
        if pd.isna(thay_doi):
            thay_doi = 0
        
        data = [
            str(row.get('Ch·ªâ ti√™u', '')),
            str(int(row.get('T·ªïng phi·∫øu (Th√¥)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Th√¥)')) else 'N/A',
            str(int(row.get('Phi·∫øu lo·∫°i tr·ª´', 0))) if pd.notna(row.get('Phi·∫øu lo·∫°i tr·ª´')) else 'N/A',
            str(int(row.get('T·ªïng phi·∫øu (Sau GT)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Sau GT)')) else 'N/A',
            format_number(row.get('T·ª∑ l·ªá % (Th√¥)', 0)) + '%',
            format_number(row.get('T·ª∑ l·ªá % (Sau GT)', 0)) + '%',
            f"{thay_doi:+.2f}%"
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            # T√¥ m√†u ch√™nh l·ªách
            if i == 6:  # C·ªôt thay ƒë·ªïi
                if thay_doi > 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Xanh l√° (tƒÉng)
                elif thay_doi < 0:
                    run.font.color.rgb = RGBColor(200, 0, 0)  # ƒê·ªè (gi·∫£m)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFEBEE')  # H·ªìng nh·∫°t
    
    doc.add_paragraph()


def add_c1x_detail_with_exclusion(doc, c1x_reports, comparison_data, chi_tieu='c11'):
    """
    Th√™m b·∫£ng chi ti·∫øt cho 1 ch·ªâ ti√™u v·ªõi c·∫£ d·ªØ li·ªáu tr∆∞·ªõc v√† sau gi·∫£m tr·ª´
    
    Args:
        doc: Document Word
        c1x_reports: Dictionary t·ª´ load_c1x_reports()
        comparison_data: Dictionary t·ª´ load_exclusion_comparison_data()
        chi_tieu: 'c11', 'c12', ho·∫∑c 'c14'
    """
    if chi_tieu == 'c11':
        title = 'C1.1 - T·ª∑ l·ªá s·ª≠a ch·ªØa phi·∫øu ch·∫•t l∆∞·ª£ng & b√°o h·ªèng'
        # Hi·ªÉn th·ªã b·∫£ng g·ªëc
        if 'c11' in c1x_reports:
            doc.add_heading(f'{title} (D·ªÆ LI·ªÜU TH√î)', level=3)
            df = c1x_reports['c11']
            _add_c11_table(doc, df)
        
        # Hi·ªÉn th·ªã b·∫£ng sau gi·∫£m tr·ª´ n·∫øu c√≥
        if 'c11_sm4' in comparison_data:
            doc.add_heading(f'{title} (SAU GI·∫¢M TR·ª™)', level=3)
            df_sau = comparison_data['c11_sm4']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.1 SM4')
        
        if 'c11_sm2' in comparison_data:
            df_sau = comparison_data['c11_sm2']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.1 SM2')
    
    elif chi_tieu == 'c12':
        title = 'C1.2 - T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i & T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª•'
        if 'c12' in c1x_reports:
            doc.add_heading(f'{title} (D·ªÆ LI·ªÜU TH√î)', level=3)
            df = c1x_reports['c12']
            _add_c12_table(doc, df)
        
        if 'c12_sm1' in comparison_data:
            doc.add_heading(f'{title} (SAU GI·∫¢M TR·ª™)', level=3)
            df_sau = comparison_data['c12_sm1']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.2 SM1')
    
    elif chi_tieu == 'c14':
        title = 'C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng sau s·ª≠a ch·ªØa'
        if 'c14' in c1x_reports:
            doc.add_heading(f'{title} (D·ªÆ LI·ªÜU TH√î)', level=3)
            df = c1x_reports['c14']
            _add_c14_table(doc, df)
        
        if 'c14' in comparison_data:
            doc.add_heading(f'{title} (SAU GI·∫¢M TR·ª™)', level=3)
            df_sau = comparison_data['c14']['tong_hop']
            _add_exclusion_summary_mini(doc, df_sau, 'C1.4')


def _add_exclusion_summary_mini(doc, df_tong_hop, label):
    """Helper: Th√™m mini summary table cho 1 ch·ªâ ti√™u sau gi·∫£m tr·ª´"""
    p = doc.add_paragraph()
    p.add_run(f'üìä {label}: ').bold = True
    
    if df_tong_hop is not None and len(df_tong_hop) > 0:
        row = df_tong_hop.iloc[0]
        tyle_tho = row.get('T·ª∑ l·ªá % (Th√¥)', row.get('T·ª∑ l·ªá HLL % (Th√¥)', 0))
        tyle_sau = row.get('T·ª∑ l·ªá % (Sau GT)', row.get('T·ª∑ l·ªá HLL % (Sau GT)', 0))
        thay_doi = row.get('Thay ƒë·ªïi %', 0)
        
        if pd.isna(tyle_tho):
            tyle_tho = 0
        if pd.isna(tyle_sau):
            tyle_sau = 0
        if pd.isna(thay_doi):
            thay_doi = 0
        
        p.add_run(f'Tr∆∞·ªõc GT: {tyle_tho:.2f}% ‚Üí Sau GT: {tyle_sau:.2f}% ')
        
        thay_doi_run = p.add_run(f'(Œî: {thay_doi:+.2f}%)')
        if thay_doi > 0:
            thay_doi_run.font.color.rgb = RGBColor(0, 128, 0)
        elif thay_doi < 0:
            thay_doi_run.font.color.rgb = RGBColor(200, 0, 0)


def _add_c11_table(doc, df):
    """Helper: Th√™m b·∫£ng C1.1 g·ªëc"""
    headers = ['ƒê∆°n v·ªã', 'SC Ch·ªß ƒë·ªông (SM1)', 'ƒê·∫°t (SM2)', 'TL SC Cƒê (%)', 
               'B√°o h·ªèng (SM3)', 'ƒê·∫°t ƒêH (SM4)', 'TL SCBH (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(row['ƒê∆°n v·ªã'], row['ƒê∆°n v·ªã'])
        if short_name == 'T·ªïng':
            short_name = 'TTVT S∆°n T√¢y'
        data = [
            short_name,
            str(int(row.get('SM1', 0))),
            str(int(row.get('SM2', 0))),
            format_number(row.get('T·ª∑ l·ªá s·ª≠a ch·ªØa phi·∫øu ch·∫•t l∆∞·ª£ng ch·ªß ƒë·ªông d·ªãch v·ª• FiberVNN, MyTV ƒë·∫°t y√™u c·∫ßu', 0)),
            str(int(row.get('SM3', 0))),
            str(int(row.get('SM4', 0))),
            format_number(row.get('T·ª∑ l·ªá phi·∫øu s·ª≠a ch·ªØa b√°o h·ªèng d·ªãch v·ª• BRCD ƒë√∫ng quy ƒë·ªãnh kh√¥ng t√≠nh h·∫πn', 0)),
            format_number(row.get('Ch·ªâ ti√™u BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')
            if short_name == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'C8E6C9')
    doc.add_paragraph()


def _add_c12_table(doc, df):
    """Helper: Th√™m b·∫£ng C1.2 g·ªëc"""
    headers = ['ƒê∆°n v·ªã', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)', 
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(row['ƒê∆°n v·ªã'], row['ƒê∆°n v·ªã'])
        if short_name == 'T·ªïng':
            short_name = 'TTVT S∆°n T√¢y'
        data = [
            short_name,
            str(int(row.get('SM1', 0))),
            str(int(row.get('SM2', 0))),
            format_number(row.get('T·ª∑ l·ªá thu√™ bao b√°o h·ªèng d·ªãch v·ª• BRCƒê l·∫∑p l·∫°i', 0)),
            str(int(row.get('SM3', 0))),
            str(int(row.get('SM4', 0))),
            format_number(row.get('T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª• BRCƒê', 0)),
            format_number(row.get('Ch·ªâ ti√™u BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')
            if short_name == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
    doc.add_paragraph()


def _add_c14_table(doc, df):
    """Helper: Th√™m b·∫£ng C1.4 g·ªëc"""
    headers = ['ƒê∆°n v·ªã', 'T·ªïng phi·∫øu', 'ƒê√£ KS', 'KS TC', 'KH HL', 
               'KHL KT PV', 'TL HL PV (%)', 'TL KH HL (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'F57C00')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(row['ƒê∆°n v·ªã'], row['ƒê∆°n v·ªã'])
        if short_name == 'T·ªïng':
            short_name = 'TTVT S∆°n T√¢y'
        data = [
            short_name,
            str(int(row.get('T·ªïng phi·∫øu', 0))),
            str(int(row.get('SL ƒë√£ KS', 0))),
            str(int(row.get('SL KS th√†nh c√¥ng', 0))),
            str(int(row.get('SL KH h√†i l√≤ng', 0))),
            str(int(row.get('Kh√¥ng HL KT ph·ª•c v·ª•', 0))),
            format_number(row.get('T·ª∑ l·ªá HL KT ph·ª•c v·ª•', 0)),
            format_number(row.get('T·ª∑ l·ªá KH h√†i l√≤ng', 0)),
            format_number(row.get('ƒêi·ªÉm BSC', 0))
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')
            if short_name == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'FFE0B2')
def add_c11_unit_level_exclusion_table(doc, unit_data, c1x_reports=None):
    """
    Th√™m b·∫£ng C1.1 t·ªïng h·ª£p theo ƒë∆°n v·ªã (c·∫•p t·ªï) sau gi·∫£m tr·ª´
    T∆∞∆°ng t·ª± b·∫£ng C1.1 g·ªëc nh∆∞ng v·ªõi s·ªë li·ªáu sau gi·∫£m tr·ª´

    Args:
        doc: Document Word
        unit_data: Dictionary t·ª´ load_unit_level_exclusion_data()
        c1x_reports: Dictionary ch·ª©a b√°o c√°o C1.x g·ªëc (ƒë·ªÉ l·∫•y s·ªë li·ªáu SM1, SM3)
    """
    if not unit_data:
        return

    # Ki·ªÉm tra c√≥ d·ªØ li·ªáu C1.1 kh√¥ng
    if 'c11_sm2' not in unit_data or 'c11_sm4' not in unit_data:
        return

    doc.add_heading('C1.1 - T·ª∑ l·ªá s·ª≠a ch·ªØa phi·∫øu ch·∫•t l∆∞·ª£ng & b√°o h·ªèng (sau gi·∫£m tr·ª´)', level=3)

    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã GHI CH√ö: ').bold = True
    p.add_run('B·∫£ng n√†y hi·ªÉn th·ªã s·ªë li·ªáu C1.1 t·ªïng h·ª£p theo ƒë∆°n v·ªã sau khi lo·∫°i b·ªè c√°c phi·∫øu thu·ªôc di·ªán gi·∫£m tr·ª´. ')
    p.add_run('SM1, SM3 l√† s·ªë li·ªáu th√¥ (kh√¥ng √°p d·ª•ng gi·∫£m tr·ª´). SM2, SM4 l√† s·ªë li·ªáu sau gi·∫£m tr·ª´.')
    doc.add_paragraph()

    df_sm2 = unit_data['c11_sm2']
    df_sm4 = unit_data['c11_sm4']

    # L·∫•y danh s√°ch ƒë∆°n v·ªã
    team_order = ['T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai',
                  'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y']

    # T·∫°o b·∫£ng
    headers = ['ƒê∆°n v·ªã', 'SC Ch·ªß ƒë·ªông (SM1)', 'ƒê·∫°t (SM2)', 'TL SC Cƒê (%)',
               'B√°o h·ªèng (SM3)', 'ƒê·∫°t ƒêH (SM4)', 'TL SCBH (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header - s·ª≠ d·ª•ng m√†u xanh l√° ƒë·∫≠m h∆°n ƒë·ªÉ ph√¢n bi·ªát v·ªõi b·∫£ng g·ªëc
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1B5E20')  # Xanh l√° ƒë·∫≠m
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # H√†m t√≠nh ƒëi·ªÉm BSC cho C1.1
    def tinh_diem_C11_TP1(tyle):
        """T√≠nh ƒëi·ªÉm TP1 (30%) - S·ª≠a ch·ªØa ch·ªß ƒë·ªông"""
        if pd.isna(tyle): return 5
        if tyle >= 0.99: return 5
        elif tyle > 0.96: return 1 + 4 * (tyle - 0.96) / 0.03
        else: return 1

    def tinh_diem_C11_TP2(tyle):
        """T√≠nh ƒëi·ªÉm TP2 (70%) - S·ª≠a ch·ªØa b√°o h·ªèng"""
        if pd.isna(tyle): return 5
        if tyle >= 0.85: return 5
        elif tyle >= 0.82: return 4 + (tyle - 0.82) / 0.03
        elif tyle >= 0.79: return 3 + (tyle - 0.79) / 0.03
        elif tyle >= 0.76: return 2
        else: return 1

    # X·ª≠ l√Ω t·ª´ng ƒë∆°n v·ªã
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)

        # L·∫•y d·ªØ li·ªáu SM2 (s·ª≠a ch·ªØa ch·ªß ƒë·ªông)
        sm2_row = df_sm2[df_sm2['ƒê∆°n v·ªã'] == don_vi]
        if sm2_row.empty:
            sm1 = 0
            sm2 = 0
            tyle_sm2 = 0
        else:
            sm2_row = sm2_row.iloc[0]
            sm1 = sm2_row.get('T·ªïng phi·∫øu (Sau GT)', 0)  # SAU GI·∫¢M TR·ª™
            sm2 = sm2_row.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
            tyle_sm2 = sm2_row.get('T·ª∑ l·ªá % (Sau GT)', 0)
            if pd.notna(tyle_sm2) and tyle_sm2 > 1:
                tyle_sm2 = tyle_sm2 / 100

        # L·∫•y d·ªØ li·ªáu SM4 (s·ª≠a ch·ªØa b√°o h·ªèng)
        sm4_row = df_sm4[df_sm4['ƒê∆°n v·ªã'] == don_vi]
        if sm4_row.empty:
            sm3 = 0
            sm4 = 0
            tyle_sm4 = 0
        else:
            sm4_row = sm4_row.iloc[0]
            sm3 = sm4_row.get('T·ªïng phi·∫øu (Sau GT)', 0)  # SAU GI·∫¢M TR·ª™
            sm4 = sm4_row.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
            tyle_sm4 = sm4_row.get('T·ª∑ l·ªá % (Sau GT)', 0)
            if pd.notna(tyle_sm4) and tyle_sm4 > 1:
                tyle_sm4 = tyle_sm4 / 100

        # T√≠nh ƒëi·ªÉm BSC
        diem_tp1 = tinh_diem_C11_TP1(tyle_sm2)
        diem_tp2 = tinh_diem_C11_TP2(tyle_sm4)
        diem_bsc = 0.30 * diem_tp1 + 0.70 * diem_tp2

        data = [
            short_name,
            str(int(sm1)) if pd.notna(sm1) else '0',
            str(int(sm2)) if pd.notna(sm2) else '0',
            format_number(tyle_sm2 * 100 if pd.notna(tyle_sm2) else 0),
            str(int(sm3)) if pd.notna(sm3) else '0',
            str(int(sm4)) if pd.notna(sm4) else '0',
            format_number(tyle_sm4 * 100 if pd.notna(tyle_sm4) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')

    # Th√™m d√≤ng t·ªïng (TTVT S∆°n T√¢y)
    cells = table.add_row().cells

    # L·∫•y d·ªØ li·ªáu t·ªïng t·ª´ SM2
    sm2_tong = df_sm2[df_sm2['ƒê∆°n v·ªã'] == 'TTVT S∆°n T√¢y']
    if sm2_tong.empty:
        sm1_tong = 0
        sm2_tong_dat = 0
        tyle_sm2_tong = 0
    else:
        sm2_tong = sm2_tong.iloc[0]
        sm1_tong = sm2_tong.get('T·ªïng phi·∫øu (Sau GT)', 0)  # SAU GI·∫¢M TR·ª™
        sm2_tong_dat = sm2_tong.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
        tyle_sm2_tong = sm2_tong.get('T·ª∑ l·ªá % (Sau GT)', 0)
        if pd.notna(tyle_sm2_tong) and tyle_sm2_tong > 1:
            tyle_sm2_tong = tyle_sm2_tong / 100

    # L·∫•y d·ªØ li·ªáu t·ªïng t·ª´ SM4
    sm4_tong = df_sm4[df_sm4['ƒê∆°n v·ªã'] == 'TTVT S∆°n T√¢y']
    if sm4_tong.empty:
        sm3_tong = 0
        sm4_tong_dat = 0
        tyle_sm4_tong = 0
    else:
        sm4_tong = sm4_tong.iloc[0]
        sm3_tong = sm4_tong.get('T·ªïng phi·∫øu (Sau GT)', 0)  # SAU GI·∫¢M TR·ª™
        sm4_tong_dat = sm4_tong.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
        tyle_sm4_tong = sm4_tong.get('T·ª∑ l·ªá % (Sau GT)', 0)
        if pd.notna(tyle_sm4_tong) and tyle_sm4_tong > 1:
            tyle_sm4_tong = tyle_sm4_tong / 100

    # T√≠nh ƒëi·ªÉm BSC t·ªïng
    diem_tp1_tong = tinh_diem_C11_TP1(tyle_sm2_tong)
    diem_tp2_tong = tinh_diem_C11_TP2(tyle_sm4_tong)
    diem_bsc_tong = 0.30 * diem_tp1_tong + 0.70 * diem_tp2_tong

    data_tong = [
        'TTVT S∆°n T√¢y',
        str(int(sm1_tong)) if pd.notna(sm1_tong) else '0',
        str(int(sm2_tong_dat)) if pd.notna(sm2_tong_dat) else '0',
        format_number(tyle_sm2_tong * 100 if pd.notna(tyle_sm2_tong) else 0),
        str(int(sm3_tong)) if pd.notna(sm3_tong) else '0',
        str(int(sm4_tong_dat)) if pd.notna(sm4_tong_dat) else '0',
        format_number(tyle_sm4_tong * 100 if pd.notna(tyle_sm4_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], 'A5D6A7')  # Xanh l√° ƒë·∫≠m h∆°n cho d√≤ng t·ªïng

    doc.add_paragraph()


def add_c12_unit_level_exclusion_table(doc, unit_data, c1x_reports=None):
    """
    Th√™m b·∫£ng C1.2 t·ªïng h·ª£p theo ƒë∆°n v·ªã (c·∫•p t·ªï) sau gi·∫£m tr·ª´
    D·ªØ li·ªáu t·ª´: So_sanh_C12_SM1.xlsx (TP1) v√† SM4-C12-ti-le-su-co-dv-brcd.xlsx (TP2)
    """
    if not unit_data or 'c12_sm1' not in unit_data or 'c12_sm4' not in unit_data:
        return

    doc.add_heading('C1.2 - T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i & T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª• (sau gi·∫£m tr·ª´)', level=3)

    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã GHI CH√ö: ').bold = True
    p.add_run('T·∫•t c·∫£ s·ªë li·ªáu ƒë√£ ƒë∆∞·ª£c gi·∫£m tr·ª´, l·∫•y t·ª´ c√°c file so s√°nh.')
    doc.add_paragraph()

    df_sm1 = unit_data['c12_sm1']  # So_sanh_C12_SM1.xlsx
    df_sm4 = unit_data['c12_sm4']  # SM4-C12-ti-le-su-co-dv-brcd.xlsx

    team_order = ['T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai',
                  'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y', 'TTVT S∆°n T√¢y']

    # T·∫°o b·∫£ng
    headers = ['ƒê∆°n v·ªã', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)',
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '0D47A1')  # Xanh d∆∞∆°ng ƒë·∫≠m
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # H√†m t√≠nh ƒëi·ªÉm BSC cho C1.2
    def tinh_diem_C12_TP1(tyle):
        """T√≠nh ƒëi·ªÉm TP1 (50%) - H·ªèng l·∫∑p l·∫°i"""
        if pd.isna(tyle): return 5
        if tyle <= 0.025: return 5
        elif tyle < 0.04: return 5 - 4 * (tyle - 0.025) / 0.015
        else: return 1

    def tinh_diem_C12_TP2(kq):
        """T√≠nh ƒëi·ªÉm TP2 (50%) - T·ª∑ l·ªá s·ª± c·ªë (%) - kq l√† th·∫≠p ph√¢n"""
        if pd.isna(kq): return 5
        if kq <= 0.02: return 5
        elif kq < 0.03: return 5 - 4 * (kq - 0.02) / 0.01
        else: return 1

    # X·ª≠ l√Ω t·ª´ng ƒë∆°n v·ªã
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)

        # L·∫•y d·ªØ li·ªáu SM1, SM2 (h·ªèng l·∫∑p l·∫°i sau gi·∫£m tr·ª´) t·ª´ So_sanh_C12_SM1.xlsx
        sm1_row = df_sm1[df_sm1['ƒê∆°n v·ªã'] == don_vi]
        if sm1_row.empty:
            sm1 = 0
            sm2 = 0  # BH (SM2) = Phi·∫øu b√°o h·ªèng (Sau GT)
            tyle_hll = 0
        else:
            sm1_row = sm1_row.iloc[0]
            sm1 = sm1_row.get('Phi·∫øu HLL (Sau GT)', 0)
            sm2 = sm1_row.get('Phi·∫øu b√°o h·ªèng (Sau GT)', 0)  # SAU GI·∫¢M TR·ª™
            tyle_hll = sm1_row.get('T·ª∑ l·ªá HLL % (Sau GT)', 0)
            if pd.notna(tyle_hll) and tyle_hll > 1:
                tyle_hll = tyle_hll / 100

        # L·∫•y d·ªØ li·ªáu TP2 (SM3, SM4) t·ª´ SM4-C12-ti-le-su-co-dv-brcd.xlsx (c·ªôt Sau GT)
        sm4_row = df_sm4[df_sm4['ƒê∆°n v·ªã'] == don_vi]
        if sm4_row.empty:
            sm3 = 0
            sm4 = 0
            tyle_sc = 0
        else:
            sm4_row = sm4_row.iloc[0]
            sm3 = sm4_row.get('Phi·∫øu b√°o h·ªèng (Sau GT)', 0)       # BH SC
            sm4 = sm4_row.get('T·ªïng TB (Sau GT)', 0)              # TB
            tyle_sc = sm4_row.get('T·ª∑ l·ªá b√°o h·ªèng % (Sau GT)', 0) # TL SC (%)
            # Chuy·ªÉn ƒë·ªïi sang th·∫≠p ph√¢n n·∫øu c·∫ßn
            if pd.notna(tyle_sc) and tyle_sc > 1:
                tyle_sc = tyle_sc / 100

        # T√≠nh ƒëi·ªÉm BSC
        diem_tp1 = tinh_diem_C12_TP1(tyle_hll)
        diem_tp2 = tinh_diem_C12_TP2(tyle_sc)
        diem_bsc = 0.50 * diem_tp1 + 0.50 * diem_tp2

        data = [
            short_name,
            str(int(sm1)) if pd.notna(sm1) else '0',
            str(int(sm2)) if pd.notna(sm2) else '0',
            format_number(tyle_hll * 100 if pd.notna(tyle_hll) else 0),
            str(int(sm3)) if pd.notna(sm3) else '0',
            str(int(sm4)) if pd.notna(sm4) else '0',
            format_number(tyle_sc * 100 if pd.notna(tyle_sc) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            # T√¥ ƒë·∫≠m v√† n·ªÅn xanh cho d√≤ng TTVT
            if don_vi == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
            elif idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')

    doc.add_paragraph()


def add_c14_unit_level_exclusion_table(doc, unit_data):
    """
    Th√™m b·∫£ng C1.4 t·ªïng h·ª£p theo ƒë∆°n v·ªã (c·∫•p t·ªï) sau gi·∫£m tr·ª´

    Args:
        doc: Document Word
        unit_data: Dictionary t·ª´ load_unit_level_exclusion_data()
    """
    if not unit_data or 'c14' not in unit_data:
        return

    doc.add_heading('C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng sau s·ª≠a ch·ªØa (sau gi·∫£m tr·ª´)', level=3)

    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã GHI CH√ö: ').bold = True
    p.add_run('S·ªë li·ªáu sau khi lo·∫°i b·ªè c√°c phi·∫øu kh·∫£o s√°t thu·ªôc di·ªán gi·∫£m tr·ª´.')
    doc.add_paragraph()

    df_c14 = unit_data['c14']

    team_order = ['T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai',
                  'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y']

    # T·∫°o b·∫£ng
    headers = ['ƒê∆°n v·ªã', 'T·ªïng phi·∫øu', 'KH h√†i l√≤ng', 'KH kh√¥ng HL', 'TL HL (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E65100')  # Cam ƒë·∫≠m
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # H√†m t√≠nh ƒëi·ªÉm BSC cho C1.4
    def tinh_diem_C14(tyle):
        """T√≠nh ƒëi·ªÉm C1.4 - ƒê·ªô h√†i l√≤ng"""
        if pd.isna(tyle): return 5
        if tyle >= 0.995: return 5
        elif tyle > 0.95: return 1 + 4 * (tyle - 0.95) / 0.045
        else: return 1

    # X·ª≠ l√Ω t·ª´ng ƒë∆°n v·ªã
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)

        # L·∫•y d·ªØ li·ªáu C1.4
        c14_row = df_c14[df_c14['ƒê∆°n v·ªã'] == don_vi]
        if c14_row.empty:
            tong_phieu = 0
            phieu_khl = 0
            tyle_hl = 0
        else:
            c14_row = c14_row.iloc[0]
            tong_phieu = c14_row.get('T·ªïng phi·∫øu (Sau GT)', 0)
            phieu_khl = c14_row.get('Phi·∫øu KHL (Sau GT)', 0)
            tyle_hl = c14_row.get('T·ª∑ l·ªá HL % (Sau GT)', 0)
            if pd.notna(tyle_hl) and tyle_hl > 1:
                tyle_hl = tyle_hl / 100

        phieu_hl = tong_phieu - phieu_khl if pd.notna(tong_phieu) and pd.notna(phieu_khl) else 0
        diem_bsc = tinh_diem_C14(tyle_hl)

        data = [
            short_name,
            str(int(tong_phieu)) if pd.notna(tong_phieu) else '0',
            str(int(phieu_hl)) if pd.notna(phieu_hl) else '0',
            str(int(phieu_khl)) if pd.notna(phieu_khl) else '0',
            format_number(tyle_hl * 100 if pd.notna(tyle_hl) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFE0B2')

    # Th√™m d√≤ng t·ªïng (TTVT S∆°n T√¢y)
    cells = table.add_row().cells

    c14_tong_row = df_c14[df_c14['ƒê∆°n v·ªã'] == 'TTVT S∆°n T√¢y']
    if c14_tong_row.empty:
        tong_phieu_tong = 0
        phieu_khl_tong = 0
        tyle_hl_tong = 0
    else:
        c14_tong_row = c14_tong_row.iloc[0]
        tong_phieu_tong = c14_tong_row.get('T·ªïng phi·∫øu (Sau GT)', 0)
        phieu_khl_tong = c14_tong_row.get('Phi·∫øu KHL (Sau GT)', 0)
        tyle_hl_tong = c14_tong_row.get('T·ª∑ l·ªá HL % (Sau GT)', 0)
        if pd.notna(tyle_hl_tong) and tyle_hl_tong > 1:
            tyle_hl_tong = tyle_hl_tong / 100

    phieu_hl_tong = tong_phieu_tong - phieu_khl_tong if pd.notna(tong_phieu_tong) and pd.notna(phieu_khl_tong) else 0
    diem_bsc_tong = tinh_diem_C14(tyle_hl_tong)

    data_tong = [
        'TTVT S∆°n T√¢y',
        str(int(tong_phieu_tong)) if pd.notna(tong_phieu_tong) else '0',
        str(int(phieu_hl_tong)) if pd.notna(phieu_hl_tong) else '0',
        str(int(phieu_khl_tong)) if pd.notna(phieu_khl_tong) else '0',
        format_number(tyle_hl_tong * 100 if pd.notna(tyle_hl_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], 'FFCC80')

    doc.add_paragraph()


def add_c15_unit_level_exclusion_table(doc, unit_data):
    """
    Th√™m b·∫£ng C1.5 t·ªïng h·ª£p theo ƒë∆°n v·ªã (c·∫•p t·ªï) sau gi·∫£m tr·ª´

    Args:
        doc: Document Word
        unit_data: Dictionary t·ª´ load_unit_level_exclusion_data()
    """
    if not unit_data or 'c15' not in unit_data:
        return

    doc.add_heading('C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• ƒë·∫°t th·ªùi gian quy ƒë·ªãnh (sau gi·∫£m tr·ª´)', level=3)

    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã GHI CH√ö: ').bold = True
    p.add_run('S·ªë li·ªáu sau khi lo·∫°i b·ªè c√°c phi·∫øu l·∫Øp ƒë·∫∑t thu·ªôc di·ªán gi·∫£m tr·ª´.')
    doc.add_paragraph()

    df_c15 = unit_data['c15']

    team_order = ['T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai',
                  'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai', 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y']

    # T·∫°o b·∫£ng
    headers = ['ƒê∆°n v·ªã', 'Phi·∫øu ƒë·∫°t', 'Phi·∫øu kh√¥ng ƒë·∫°t', 'T·ªïng phi·∫øu', 'T·ª∑ l·ªá ƒë·∫°t (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)

    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '00695C')  # Xanh ng·ªçc ƒë·∫≠m
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    # H√†m t√≠nh ƒëi·ªÉm BSC cho C1.5
    def tinh_diem_C15(tyle):
        """T√≠nh ƒëi·ªÉm C1.5 - Thi·∫øt l·∫≠p d·ªãch v·ª•"""
        if pd.isna(tyle): return 5
        if tyle >= 0.995: return 5
        elif tyle > 0.895: return 1 + 4 * (tyle - 0.895) / 0.10
        else: return 1

    # X·ª≠ l√Ω t·ª´ng ƒë∆°n v·ªã
    for idx, don_vi in enumerate(team_order, 1):
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)

        # L·∫•y d·ªØ li·ªáu C1.5
        c15_row = df_c15[df_c15['ƒê∆°n v·ªã'] == don_vi]
        if c15_row.empty:
            phieu_dat = 0
            tong_phieu = 0
            tyle_dat = 0
        else:
            c15_row = c15_row.iloc[0]
            phieu_dat = c15_row.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
            tong_phieu = c15_row.get('T·ªïng phi·∫øu (Sau GT)', 0)
            tyle_dat = c15_row.get('T·ª∑ l·ªá ƒë·∫°t % (Sau GT)', 0)
            if pd.notna(tyle_dat) and tyle_dat > 1:
                tyle_dat = tyle_dat / 100

        phieu_ko_dat = tong_phieu - phieu_dat if pd.notna(tong_phieu) and pd.notna(phieu_dat) else 0
        diem_bsc = tinh_diem_C15(tyle_dat)

        data = [
            short_name,
            str(int(phieu_dat)) if pd.notna(phieu_dat) else '0',
            str(int(phieu_ko_dat)) if pd.notna(phieu_ko_dat) else '0',
            str(int(tong_phieu)) if pd.notna(tong_phieu) else '0',
            format_number(tyle_dat * 100 if pd.notna(tyle_dat) else 0),
            format_number(diem_bsc)
        ]

        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'B2DFDB')

    # Th√™m d√≤ng t·ªïng (TTVT S∆°n T√¢y)
    cells = table.add_row().cells

    c15_tong_row = df_c15[df_c15['ƒê∆°n v·ªã'] == 'TTVT S∆°n T√¢y']
    if c15_tong_row.empty:
        phieu_dat_tong = 0
        tong_phieu_tong = 0
        tyle_dat_tong = 0
    else:
        c15_tong_row = c15_tong_row.iloc[0]
        phieu_dat_tong = c15_tong_row.get('Phi·∫øu ƒë·∫°t (Sau GT)', 0)
        tong_phieu_tong = c15_tong_row.get('T·ªïng phi·∫øu (Sau GT)', 0)
        tyle_dat_tong = c15_tong_row.get('T·ª∑ l·ªá ƒë·∫°t % (Sau GT)', 0)
        if pd.notna(tyle_dat_tong) and tyle_dat_tong > 1:
            tyle_dat_tong = tyle_dat_tong / 100

    phieu_ko_dat_tong = tong_phieu_tong - phieu_dat_tong if pd.notna(tong_phieu_tong) and pd.notna(phieu_dat_tong) else 0
    diem_bsc_tong = tinh_diem_C15(tyle_dat_tong)

    data_tong = [
        'TTVT S∆°n T√¢y',
        str(int(phieu_dat_tong)) if pd.notna(phieu_dat_tong) else '0',
        str(int(phieu_ko_dat_tong)) if pd.notna(phieu_ko_dat_tong) else '0',
        str(int(tong_phieu_tong)) if pd.notna(tong_phieu_tong) else '0',
        format_number(tyle_dat_tong * 100 if pd.notna(tyle_dat_tong) else 0),
        format_number(diem_bsc_tong)
    ]

    for i, value in enumerate(data_tong):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_shading(cells[i], '80CBC4')

    doc.add_paragraph()


def add_c11_exclusion_table(doc, comparison_data):
    """
    Th√™m b·∫£ng C1.1 sau gi·∫£m tr·ª´ chi ti·∫øt theo NVKT (ri√™ng bi·ªát)
    """
    if not comparison_data:
        return

    has_data = False
    
    # C1.1 SM4 - S·ª≠a ch·ªØa b√°o h·ªèng
    if 'c11_sm4' in comparison_data:
        has_data = True
        doc.add_heading('C1.1 - SAU GI·∫¢M TR·ª™ (SM4 - S·ª≠a ch·ªØa b√°o h·ªèng)', level=4)
        df = comparison_data['c11_sm4']['chi_tiet']
        
        headers = ['NVKT', 'T·ªïng phi·∫øu (Th√¥)', 'T·ªïng phi·∫øu (Sau GT)', 
                   'S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 'S·ªë phi·∫øu ƒë·∫°t (Sau GT)',
                   'T·ª∑ l·ªá % (Th√¥)', 'T·ª∑ l·ªá % (Sau GT)', 'Ch√™nh l·ªách %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'C62828')  # ƒê·ªè ƒë·∫≠m
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Ch√™nh l·ªách %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('T·ªïng phi·∫øu (Th√¥)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Th√¥)')) else '0',
                str(int(row.get('T·ªïng phi·∫øu (Sau GT)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Sau GT)')) else '0',
                str(int(row.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 0))) if pd.notna(row.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)')) else '0',
                str(int(row.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)', 0))) if pd.notna(row.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)')) else '0',
                format_number(row.get('T·ª∑ l·ªá % (Th√¥)', 0)),
                format_number(row.get('T·ª∑ l·ªá % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:  # C·ªôt ch√™nh l·ªách
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FFEBEE')
        
        # Th√™m t·ªïng h·ª£p
        if 'tong_hop' in comparison_data['c11_sm4']:
            df_th = comparison_data['c11_sm4']['tong_hop']
            if len(df_th) > 0:
                row_th = df_th.iloc[0]
                p = doc.add_paragraph()
                p.add_run('üìä T·ªïng h·ª£p C1.1 SM4: ').bold = True
                tyle_tho = row_th.get('T·ª∑ l·ªá % (Th√¥)', 0)
                tyle_sau = row_th.get('T·ª∑ l·ªá % (Sau GT)', 0)
                thay_doi = row_th.get('Thay ƒë·ªïi %', 0)
                if pd.isna(thay_doi): thay_doi = 0
                p.add_run(f'Tr∆∞·ªõc: {tyle_tho:.2f}% ‚Üí Sau: {tyle_sau:.2f}% (Œî: {thay_doi:+.2f}%)')
        
        doc.add_paragraph()
    
    # C1.1 SM2 - S·ª≠a ch·ªØa ch·ªß ƒë·ªông
    if 'c11_sm2' in comparison_data:
        has_data = True
        doc.add_heading('C1.1 - SAU GI·∫¢M TR·ª™ (SM2 - S·ª≠a ch·ªØa ch·ªß ƒë·ªông)', level=4)
        df = comparison_data['c11_sm2']['chi_tiet']
        
        headers = ['NVKT', 'T·ªïng phi·∫øu (Th√¥)', 'T·ªïng phi·∫øu (Sau GT)', 
                   'S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 'S·ªë phi·∫øu ƒë·∫°t (Sau GT)',
                   'T·ª∑ l·ªá % (Th√¥)', 'T·ª∑ l·ªá % (Sau GT)', 'Ch√™nh l·ªách %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'AD1457')  # H·ªìng ƒë·∫≠m
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Ch√™nh l·ªách %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('T·ªïng phi·∫øu (Th√¥)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Th√¥)')) else '0',
                str(int(row.get('T·ªïng phi·∫øu (Sau GT)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu (Sau GT)')) else '0',
                str(int(row.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)', 0))) if pd.notna(row.get('S·ªë phi·∫øu ƒë·∫°t (Th√¥)')) else '0',
                str(int(row.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)', 0))) if pd.notna(row.get('S·ªë phi·∫øu ƒë·∫°t (Sau GT)')) else '0',
                format_number(row.get('T·ª∑ l·ªá % (Th√¥)', 0)),
                format_number(row.get('T·ª∑ l·ªá % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FCE4EC')
        
        doc.add_paragraph()


def add_c12_exclusion_table(doc, comparison_data):
    """
    Th√™m b·∫£ng C1.2 sau gi·∫£m tr·ª´ (ri√™ng bi·ªát)
    """
    if not comparison_data:
        return
    
    # C1.2 SM1 - H·ªèng l·∫∑p l·∫°i
    if 'c12_sm1' in comparison_data:
        doc.add_heading('C1.2 - SAU GI·∫¢M TR·ª™ (SM1 - H·ªèng l·∫∑p l·∫°i)', level=4)
        df = comparison_data['c12_sm1']['chi_tiet']
        
        headers = ['NVKT', 'Phi·∫øu HLL (Th√¥)', 'Phi·∫øu HLL (Sau GT)', 
                   'Phi·∫øu BH (Th√¥)', 'Phi·∫øu BH (Sau GT)',
                   'T·ª∑ l·ªá HLL % (Th√¥)', 'T·ª∑ l·ªá HLL % (Sau GT)', 'Ch√™nh l·ªách %']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '0D47A1')  # Xanh d∆∞∆°ng ƒë·∫≠m
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            chenh_lech = row.get('Ch√™nh l·ªách %', 0)
            if pd.isna(chenh_lech):
                chenh_lech = 0
            data = [
                str(row.get('NVKT', '')),
                str(int(row.get('S·ªë phi·∫øu HLL (Th√¥)', 0))) if pd.notna(row.get('S·ªë phi·∫øu HLL (Th√¥)')) else '0',
                str(int(row.get('S·ªë phi·∫øu HLL (Sau GT)', 0))) if pd.notna(row.get('S·ªë phi·∫øu HLL (Sau GT)')) else '0',
                str(int(row.get('S·ªë phi·∫øu b√°o h·ªèng (Th√¥)', 0))) if pd.notna(row.get('S·ªë phi·∫øu b√°o h·ªèng (Th√¥)')) else '0',
                str(int(row.get('S·ªë phi·∫øu b√°o h·ªèng (Sau GT)', 0))) if pd.notna(row.get('S·ªë phi·∫øu b√°o h·ªèng (Sau GT)')) else '0',
                format_number(row.get('T·ª∑ l·ªá HLL % (Th√¥)', 0)),
                format_number(row.get('T·ª∑ l·ªá HLL % (Sau GT)', 0)),
                f"{chenh_lech:+.2f}%"
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if i == 7:
                    if chenh_lech > 0:
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif chenh_lech < 0:
                        run.font.color.rgb = RGBColor(200, 0, 0)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'E3F2FD')
        
        doc.add_paragraph()


def add_c14_exclusion_table(doc, comparison_data):
    """
    Th√™m b·∫£ng C1.4 sau gi·∫£m tr·ª´ (ri√™ng bi·ªát)
    """
    if not comparison_data or 'c14' not in comparison_data:
        return
    
    doc.add_heading('C1.4 - SAU GI·∫¢M TR·ª™ (ƒê·ªô h√†i l√≤ng kh√°ch h√†ng)', level=4)
    df = comparison_data['c14']['chi_tiet']
    
    headers = ['NVKT', 'T·ªïng KS (Th√¥)', 'KHL (Th√¥)', 'T·ª∑ l·ªá HL % (Th√¥)',
               'T·ªïng KS (Sau GT)', 'KHL (Sau GT)', 'T·ª∑ l·ªá HL % (Sau GT)', 'Ch√™nh l·ªách %']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'E65100')  # Cam ƒë·∫≠m
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        chenh_lech = row.get('Ch√™nh l·ªách %', 0)
        if pd.isna(chenh_lech):
            chenh_lech = 0
        data = [
            str(row.get('NVKT', '')),
            str(int(row.get('T·ªïng phi·∫øu KS (Th√¥)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu KS (Th√¥)')) else '0',
            str(int(row.get('S·ªë phi·∫øu KHL (Th√¥)', 0))) if pd.notna(row.get('S·ªë phi·∫øu KHL (Th√¥)')) else '0',
            format_number(row.get('T·ª∑ l·ªá HL (%) (Th√¥)', 0)),
            str(int(row.get('T·ªïng phi·∫øu KS (Sau GT)', 0))) if pd.notna(row.get('T·ªïng phi·∫øu KS (Sau GT)')) else '0',
            str(int(row.get('S·ªë phi·∫øu KHL (Sau GT)', 0))) if pd.notna(row.get('S·ªë phi·∫øu KHL (Sau GT)')) else '0',
            format_number(row.get('T·ª∑ l·ªá HL (%) (Sau GT)', 0)),
            f"{chenh_lech:+.2f}%"
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            if i == 7:
                if chenh_lech > 0:
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif chenh_lech < 0:
                    run.font.color.rgb = RGBColor(200, 0, 0)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')
    
    doc.add_paragraph()


def create_exclusion_bar_chart(comparison_data, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì bar ri√™ng cho d·ªØ li·ªáu sau gi·∫£m tr·ª´
    """
    if not comparison_data or 'tong_hop' not in comparison_data:
        return None
    
    df = comparison_data['tong_hop']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(df))
    
    tyle_sau = df['T·ª∑ l·ªá % (Sau GT)'].fillna(0).values
    chi_tieu = df['Ch·ªâ ti√™u'].values
    
    # M√†u s·∫Øc theo m·ª©c ƒë·ªô t·ªët/x·∫•u
    colors = []
    for val in tyle_sau:
        if val >= 95:
            colors.append('#4CAF50')  # Xanh l√° - t·ªët
        elif val >= 90:
            colors.append('#FFC107')  # V√†ng - trung b√¨nh
        else:
            colors.append('#F44336')  # ƒê·ªè - c·∫ßn c·∫£i thi·ªán
    
    bars = ax.bar(x, tyle_sau, color=colors, alpha=0.8, edgecolor='black', linewidth=0.5)
    
    # Th√™m gi√° tr·ªã l√™n c·ªôt
    for bar, val in zip(bars, tyle_sau):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
               f'{val:.1f}%', ha='center', va='bottom', fontsize=10, fontweight='bold')
    
    ax.set_xlabel('Ch·ªâ ti√™u', fontsize=12)
    ax.set_ylabel('T·ª∑ l·ªá (%)', fontsize=12)
    ax.set_title('T·ª∂ L·ªÜ C√ÅC CH·ªà TI√äU SAU GI·∫¢M TR·ª™', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(chi_tieu, rotation=15, ha='right', fontsize=10)
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    ax.set_ylim(0, max(tyle_sau) * 1.15 if len(tyle_sau) > 0 else 100)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_c1x_overview_table(doc, c1x_reports, comparison_data=None, unit_data=None, exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    Th√™m b·∫£ng t·ªïng quan chi ti·∫øt t·ª´ c√°c b√°o c√°o C1.x v√†o document
    N·∫øu c√≥ comparison_data, s·∫Ω th√™m b·∫£ng s·ªë li·ªáu sau gi·∫£m tr·ª´ ngay sau b·∫£ng th√¥

    Args:
        doc: Document Word
        c1x_reports: Dictionary ch·ª©a c√°c DataFrame t·ª´ load_c1x_reports()
        comparison_data: Dictionary ch·ª©a d·ªØ li·ªáu so s√°nh t·ª´ load_exclusion_comparison_data()
        unit_data: Dictionary ch·ª©a d·ªØ li·ªáu th·ªëng k√™ theo ƒë∆°n v·ªã t·ª´ load_unit_level_exclusion_data()
        exclusion_folder: Th∆∞ m·ª•c ch·ª©a d·ªØ li·ªáu gi·∫£m tr·ª´
    """
    doc.add_heading('1.3. S·ªë li·ªáu chi ti·∫øt c√°c ch·ªâ ti√™u BSC theo ƒê·ªôi/TTVT', level=2)

    # =========================================================================
    # B·∫£ng C1.1 - T·ª∑ l·ªá s·ª≠a ch·ªØa
    # =========================================================================
    # D√ôNG D·ªÆ LI·ªÜU T·ª™ So_sanh_C11_SM2.xlsx (TP1) v√† So_sanh_C11_SM4.xlsx (TP2)
    doc.add_heading('C1.1 - T·ª∑ l·ªá s·ª≠a ch·ªØa phi·∫øu ch·∫•t l∆∞·ª£ng & b√°o h·ªèng', level=3)
    
    headers = ['ƒê∆°n v·ªã', 'SC Ch·ªß ƒë·ªông (SM1)', 'ƒê·∫°t (SM2)', 'TL SC Cƒê (%)', 
               'B√°o h·ªèng (SM3)', 'ƒê·∫°t ƒêH (SM4)', 'TL SCBH (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # ƒê·ªçc d·ªØ li·ªáu t·ª´ file so s√°nh
    teams_order = ['Ph√∫c Th·ªç', 'Qu·∫£ng Oai', 'Su·ªëi Hai', 'S∆°n T√¢y', 'TTVT S∆°n T√¢y']
    team_name_map = {
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'TTVT S∆°n T√¢y': 'TTVT S∆°n T√¢y',
    }
    
    def get_short_name_c11(don_vi):
        if not don_vi: return None
        for orig, short in team_name_map.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # L·∫•y d·ªØ li·ªáu TP1 (SC Ch·ªß ƒë·ªông) t·ª´ So_sanh_C11_SM2.xlsx
    tp1_data = {}  # short_name -> {sm1, sm2, tyle, diem}
    if unit_data and 'c11_sm2' in unit_data:
        for _, row in unit_data['c11_sm2'].iterrows():
            short_name = get_short_name_c11(row.get('ƒê∆°n v·ªã', ''))
            if short_name:
                tp1_data[short_name] = {
                    'sm1': row.get('T·ªïng phi·∫øu (Th√¥)', 0),
                    'sm2': row.get('Phi·∫øu ƒë·∫°t (Th√¥)', 0),
                    'tyle': row.get('T·ª∑ l·ªá % (Th√¥)', 0),
                }
    
    # L·∫•y d·ªØ li·ªáu TP2 (B√°o h·ªèng) t·ª´ So_sanh_C11_SM4.xlsx
    tp2_data = {}  # short_name -> {sm3, sm4, tyle, diem}
    if unit_data and 'c11_sm4' in unit_data:
        for _, row in unit_data['c11_sm4'].iterrows():
            short_name = get_short_name_c11(row.get('ƒê∆°n v·ªã', ''))
            if short_name:
                tp2_data[short_name] = {
                    'sm3': row.get('T·ªïng phi·∫øu (Th√¥)', 0),   # B√°o h·ªèng = T·ªïng phi·∫øu
                    'sm4': row.get('Phi·∫øu ƒë·∫°t (Th√¥)', 0),     # ƒê·∫°t ƒêH = Phi·∫øu ƒë·∫°t
                    'tyle': row.get('T·ª∑ l·ªá % (Th√¥)', 0),
                    'diem': row.get('ƒêi·ªÉm BSC (Th√¥)', 0),
                }
    
    # L·∫•y ƒëi·ªÉm BSC t·ªïng h·ª£p t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx
    bsc_scores_c11 = {}
    bsc_data_c11 = load_bsc_unit_scores_from_comparison(exclusion_folder)
    if bsc_data_c11 and bsc_data_c11.get('units') is not None:
        for _, row in bsc_data_c11['units'].iterrows():
            short_name = get_short_name_c11(row.get('don_vi', ''))
            if short_name:
                bsc_scores_c11[short_name] = row.get('Diem_C1.1 (Tr∆∞·ªõc)', 0)
    
    # T·∫°o d·ªØ li·ªáu b·∫£ng
    for idx, team in enumerate(teams_order, 1):
        cells = table.add_row().cells
        tp1 = tp1_data.get(team, {})
        tp2 = tp2_data.get(team, {})
        bsc_score = bsc_scores_c11.get(team, 0)
        
        data = [
            team,
            str(int(tp1.get('sm1', 0))),
            str(int(tp1.get('sm2', 0))),
            format_number(tp1.get('tyle', 0)),
            str(int(tp2.get('sm3', 0))),
            str(int(tp2.get('sm4', 0))),
            format_number(tp2.get('tyle', 0)),
            format_number(bsc_score)
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')
            # T√¥ ƒë·∫≠m d√≤ng t·ªïng
            if team == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'C8E6C9')
    
    doc.add_paragraph()

    # Th√™m b·∫£ng C1.1 t·ªïng h·ª£p theo ƒë∆°n v·ªã (t·ªï) sau gi·∫£m tr·ª´ n·∫øu c√≥
    if unit_data:
        add_c11_unit_level_exclusion_table(doc, unit_data, c1x_reports)

    # Th√™m b·∫£ng C1.1 chi ti·∫øt theo NVKT sau gi·∫£m tr·ª´ n·∫øu c√≥
    # COMMENT: B·ªè b·∫£ng chi ti·∫øt t·ª´ng NVKT n√†y v√¨ ƒë√£ c√≥ trong PH·∫¶N 2
    # if comparison_data:
    #     add_c11_exclusion_table(doc, comparison_data)
    # =========================================================================
    # B·∫£ng C1.2 - T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i & s·ª± c·ªë
    # =========================================================================
    # D√ôNG D·ªÆ LI·ªÜU T·ª™ So_sanh_C12_SM1.xlsx (TP1) v√† SM4-C12-ti-le-su-co-dv-brcd.xlsx (TP2)
    doc.add_heading('C1.2 - T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i & T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª•', level=3)
    
    headers = ['ƒê∆°n v·ªã', 'HLL (SM1)', 'BH (SM2)', 'TL HLL (%)', 
               'BH SC (SM3)', 'TB (SM4)', 'TL SC (%)', 'ƒêi·ªÉm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # ƒê·ªçc d·ªØ li·ªáu t·ª´ file so s√°nh
    teams_order_c12 = ['Ph√∫c Th·ªç', 'Qu·∫£ng Oai', 'Su·ªëi Hai', 'S∆°n T√¢y', 'TTVT S∆°n T√¢y']
    team_name_map_c12 = {
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'TTVT S∆°n T√¢y': 'TTVT S∆°n T√¢y',
    }
    
    def get_short_name_c12(don_vi):
        if not don_vi: return None
        for orig, short in team_name_map_c12.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # L·∫•y d·ªØ li·ªáu TP1 (HLL) t·ª´ So_sanh_C12_SM1.xlsx
    tp1_c12_data = {}  # short_name -> {sm1, sm2, tyle}
    if unit_data and 'c12_sm1' in unit_data:
        for _, row in unit_data['c12_sm1'].iterrows():
            short_name = get_short_name_c12(row.get('ƒê∆°n v·ªã', ''))
            if short_name:
                tp1_c12_data[short_name] = {
                    'sm1': row.get('Phi·∫øu HLL (Th√¥)', 0),           # HLL (SM1)
                    'sm2': row.get('Phi·∫øu b√°o h·ªèng (Th√¥)', 0),      # BH (SM2)
                    'tyle': row.get('T·ª∑ l·ªá HLL % (Th√¥)', 0),        # TL HLL (%)
                }
    
    # L·∫•y d·ªØ li·ªáu TP2 (S·ª± c·ªë) t·ª´ SM4-C12-ti-le-su-co-dv-brcd.xlsx
    tp2_c12_data = {}  # short_name -> {sm3, sm4, tyle}
    if unit_data and 'c12_sm4' in unit_data:
        for _, row in unit_data['c12_sm4'].iterrows():
            short_name = get_short_name_c12(row.get('ƒê∆°n v·ªã', ''))
            if short_name:
                tp2_c12_data[short_name] = {
                    'sm3': row.get('Phi·∫øu b√°o h·ªèng (Th√¥)', 0),        # BH SC (SM3)
                    'sm4': row.get('T·ªïng TB (Th√¥)', 0),               # TB (SM4)
                    'tyle': row.get('T·ª∑ l·ªá b√°o h·ªèng % (Th√¥)', 0),     # TL SC (%)
                }
    
    # L·∫•y ƒëi·ªÉm BSC t·ªïng h·ª£p t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx
    bsc_scores_c12 = {}
    bsc_data_c12 = load_bsc_unit_scores_from_comparison(exclusion_folder)
    if bsc_data_c12 and bsc_data_c12.get('units') is not None:
        for _, row in bsc_data_c12['units'].iterrows():
            short_name = get_short_name_c12(row.get('don_vi', ''))
            if short_name:
                bsc_scores_c12[short_name] = row.get('Diem_C1.2 (Tr∆∞·ªõc)', 0)
    
    # T·∫°o d·ªØ li·ªáu b·∫£ng
    for idx, team in enumerate(teams_order_c12, 1):
        cells = table.add_row().cells
        tp1 = tp1_c12_data.get(team, {})
        tp2 = tp2_c12_data.get(team, {})
        bsc_score = bsc_scores_c12.get(team, 0)
        
        data = [
            team,
            str(int(tp1.get('sm1', 0) or 0)),
            str(int(tp1.get('sm2', 0) or 0)),
            format_number(tp1.get('tyle', 0) or 0),
            str(int(tp2.get('sm3', 0) or 0)),
            str(int(tp2.get('sm4', 0) or 0)),
            format_number(tp2.get('tyle', 0) or 0),
            format_number(bsc_score or 0)
        ]
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')
            if team == 'TTVT S∆°n T√¢y':
                run.font.bold = True
                set_cell_shading(cells[i], 'BBDEFB')
    
    doc.add_paragraph()

    # Th√™m b·∫£ng C1.2 t·ªïng h·ª£p theo ƒë∆°n v·ªã (t·ªï) sau gi·∫£m tr·ª´ n·∫øu c√≥
    if unit_data:
        add_c12_unit_level_exclusion_table(doc, unit_data, c1x_reports)

    # Th√™m b·∫£ng C1.2 chi ti·∫øt theo NVKT sau gi·∫£m tr·ª´ n·∫øu c√≥
    # COMMENT: B·ªè b·∫£ng chi ti·∫øt t·ª´ng NVKT n√†y v√¨ ƒë√£ c√≥ trong PH·∫¶N 2
    # if comparison_data:
    #     add_c12_exclusion_table(doc, comparison_data)
    
    # =========================================================================
    # B·∫£ng C1.3 - K√™nh TSL
    # =========================================================================
    if 'c13' in c1x_reports:
        doc.add_heading('C1.3 - Ch·ªâ ti√™u k√™nh thu√™ leased line (TSL)', level=3)
        df = c1x_reports['c13']
        
        headers = ['ƒê∆°n v·ªã', 'SC TSL (SM1)', 'ƒê·∫°t (SM2)', 'TL SC (%)', 
                   'HLL (SM3)', 'BH (SM4)', 'TL HLL (%)', 'S·ªë TB (SM6)', 'TL SC (%)', 'ƒêi·ªÉm BSC']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '6A1B9A')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            short_name = TEAM_SHORT_NAMES.get(row['ƒê∆°n v·ªã'], row['ƒê∆°n v·ªã'])
            if short_name == 'T·ªïng':
                short_name = 'TTVT S∆°n T√¢y'
            data = [
                short_name,
                str(int(row.get('SM1', 0))),
                str(int(row.get('SM2', 0))),
                format_number(row.get('T·ª∑ l·ªá s·ª≠a ch·ªØa d·ªãch v·ª• k√™nh TSL ho√†n th√†nh ƒë√∫ng th·ªùi gian quy ƒë·ªãnh', 0)),
                str(int(row.get('SM3', 0))),
                str(int(row.get('SM4', 0))),
                format_number(row.get('T·ª∑ l·ªá thu√™ bao b√°o h·ªèng d·ªãch v·ª• k√™nh TSL l·∫∑p l·∫°i', 0)),
                str(int(row.get('SM6', 0))),
                format_number(row.get('T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª• k√™nh TSL', 0)),
                format_number(row.get('Ch·ªâ ti√™u BSC', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'F3E5F5')
                if short_name == 'TTVT S∆°n T√¢y':
                    run.font.bold = True
                    set_cell_shading(cells[i], 'E1BEE7')
        
        doc.add_paragraph()
    
    # =========================================================================
    # B·∫£ng C1.4 - H√†i l√≤ng kh√°ch h√†ng
    # =========================================================================
    if 'c14' in c1x_reports:
        doc.add_heading('C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng sau s·ª≠a ch·ªØa', level=3)
        df = c1x_reports['c14']
        
        headers = ['ƒê∆°n v·ªã', 'T·ªïng phi·∫øu', 'ƒê√£ KS', 'KS TC', 'KH HL', 
                   'KHL KT PV', 'TL HL PV (%)', 'TL KH HL (%)', 'ƒêi·ªÉm BSC']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'F57C00')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            short_name = TEAM_SHORT_NAMES.get(row['ƒê∆°n v·ªã'], row['ƒê∆°n v·ªã'])
            if short_name == 'T·ªïng':
                short_name = 'TTVT S∆°n T√¢y'
            data = [
                short_name,
                str(int(row.get('T·ªïng phi·∫øu', 0))),
                str(int(row.get('SL ƒë√£ KS', 0))),
                str(int(row.get('SL KS th√†nh c√¥ng', 0))),
                str(int(row.get('SL KH h√†i l√≤ng', 0))),
                str(int(row.get('Kh√¥ng HL KT ph·ª•c v·ª•', 0))),
                format_number(row.get('T·ª∑ l·ªá HL KT ph·ª•c v·ª•', 0)),
                format_number(row.get('T·ª∑ l·ªá KH h√†i l√≤ng', 0)),
                format_number(row.get('ƒêi·ªÉm BSC', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'FFF3E0')
                if short_name == 'TTVT S∆°n T√¢y':
                    run.font.bold = True
                    set_cell_shading(cells[i], 'FFE0B2')
        
        doc.add_paragraph()

    # Th√™m b·∫£ng C1.4 t·ªïng h·ª£p theo ƒë∆°n v·ªã (t·ªï) sau gi·∫£m tr·ª´ n·∫øu c√≥
    if unit_data:
        add_c14_unit_level_exclusion_table(doc, unit_data)

    # Th√™m b·∫£ng C1.4 chi ti·∫øt theo NVKT sau gi·∫£m tr·ª´ n·∫øu c√≥
    # COMMENT: B·ªè b·∫£ng chi ti·∫øt t·ª´ng NVKT n√†y v√¨ ƒë√£ c√≥ trong PH·∫¶N 2
    # if comparison_data:
    #     add_c14_exclusion_table(doc, comparison_data)
    
    # =========================================================================
    # B·∫£ng C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• ƒë·∫°t
    # =========================================================================
    if 'c15_ttvtst' in c1x_reports:
        doc.add_heading('C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• ƒë·∫°t th·ªùi gian quy ƒë·ªãnh', level=3)
        df = c1x_reports['c15_ttvtst']
        
        headers = ['ƒê∆°n v·ªã', 'Phi·∫øu ƒë·∫°t', 'Phi·∫øu kh√¥ng ƒë·∫°t', 'T·ªïng HC', 'T·ªâ l·ªá ƒë·∫°t (%)']
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        set_table_border(table)
        
        for i, cell in enumerate(table.rows[0].cells):
            cell.text = headers[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '00796B')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        
        for idx, (_, row) in enumerate(df.iterrows(), 1):
            cells = table.add_row().cells
            don_vi = row.get('DOIVT', '')
            short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
            data = [
                short_name,
                str(int(row.get('Phi·∫øu ƒë·∫°t', 0))),
                str(int(row.get('Phi·∫øu kh√¥ng ƒë·∫°t', 0))),
                str(int(row.get('T·ªïng Ho√†n c√¥ng', 0))),
                format_number(row.get('T·ªâ l·ªá ƒë·∫°t (%)', 0))
            ]
            for i, value in enumerate(data):
                cells[i].text = value
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(10)
                if idx % 2 == 0:
                    set_cell_shading(cells[i], 'E0F2F1')
                # T√¥ ƒë·∫≠m d√≤ng TTVT S∆°n T√¢y 
                if 'TTVT' in don_vi:
                    run.font.bold = True
                    set_cell_shading(cells[i], 'B2DFDB')

        doc.add_paragraph()

    # Th√™m b·∫£ng C1.5 t·ªïng h·ª£p theo ƒë∆°n v·ªã (t·ªï) sau gi·∫£m tr·ª´ n·∫øu c√≥
    if unit_data:
        add_c15_unit_level_exclusion_table(doc, unit_data)


# =============================================================================
# H√ÄM T·∫†O BI·ªÇU ƒê·ªí
# =============================================================================
def create_team_comparison_chart(c1x_reports, output_path=None, bsc_data=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì so s√°nh ƒëi·ªÉm BSC th·ª±c t·∫ø gi·ªØa 4 t·ªï
    
    Args:
        c1x_reports: Dictionary ch·ª©a c√°c DataFrame t·ª´ load_c1x_reports()
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file ·∫£nh (None = tr·∫£ v·ªÅ bytes)
        bsc_data: Dictionary t·ª´ load_bsc_unit_scores_from_comparison() (∆∞u ti√™n s·ª≠ d·ª•ng)
    
    Returns:
        bytes ho·∫∑c str: D·ªØ li·ªáu ·∫£nh ho·∫∑c ƒë∆∞·ªùng d·∫´n file
    """
    # Chu·∫©n b·ªã d·ªØ li·ªáu t·ª´ c√°c b√°o c√°o C1.x
    teams_order = ['Ph√∫c Th·ªç', 'Qu·∫£ng Oai', 'Su·ªëi Hai', 'S∆°n T√¢y']
    
    # Kh·ªüi t·∫°o dict ch·ª©a ƒëi·ªÉm BSC (bao g·ªìm C1.5)
    bsc_scores = {team: {'C1.1': 0, 'C1.2': 0, 'C1.3': 0, 'C1.4': 0, 'C1.5': 0} for team in teams_order}
    
    # Map t√™n ƒë∆°n v·ªã
    team_name_map = {
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
    }
    
    def get_short_name(don_vi):
        if not don_vi: return None
        for orig, short in team_name_map.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return TEAM_SHORT_NAMES.get(don_vi, don_vi)
    
    # ∆ØU TI√äN: L·∫•y ƒëi·ªÉm t·ª´ bsc_data (Tong_hop_Diem_BSC_Don_Vi.xlsx) - c·ªôt (Tr∆∞·ªõc)
    if bsc_data and bsc_data.get('units') is not None and not bsc_data['units'].empty:
        print("  üìä Bi·ªÉu ƒë·ªì: S·ª≠ d·ª•ng ƒëi·ªÉm t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx (c·ªôt Tr∆∞·ªõc)")
        for _, row in bsc_data['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.1'] = row.get('Diem_C1.1 (Tr∆∞·ªõc)', 0) or 0
                bsc_scores[short_name]['C1.2'] = row.get('Diem_C1.2 (Tr∆∞·ªõc)', 0) or 0
                bsc_scores[short_name]['C1.4'] = row.get('Diem_C1.4 (Tr∆∞·ªõc)', 0) or 0
                bsc_scores[short_name]['C1.5'] = row.get('Diem_C1.5 (Tr∆∞·ªõc)', 0) or 0
        # C1.3 v·∫´n l·∫•y t·ª´ c1x_reports (kh√¥ng c√≥ trong comparison)
        if c1x_reports and 'c13' in c1x_reports:
            for _, row in c1x_reports['c13'].iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0)
    else:
        # FALLBACK: L·∫•y ƒëi·ªÉm BSC t·ª´ c1x_reports (s·ªë li·ªáu g·ªëc)
        if 'c11' in c1x_reports:
            df = c1x_reports['c11']
            for _, row in df.iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.1'] = row.get('Ch·ªâ ti√™u BSC', 0)
        
        if 'c12' in c1x_reports:
            df = c1x_reports['c12']
            for _, row in df.iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.2'] = row.get('Ch·ªâ ti√™u BSC', 0)
        
        if 'c13' in c1x_reports:
            df = c1x_reports['c13']
            for _, row in df.iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0)
        
        if 'c14' in c1x_reports:
            df = c1x_reports['c14']
            for _, row in df.iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.4'] = row.get('ƒêi·ªÉm BSC', 0)
        
        if 'c15_ttvtst' in c1x_reports:
            df = c1x_reports['c15_ttvtst']
            for _, row in df.iterrows():
                don_vi = row.get('DOIVT', '')
                short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
                if short_name in teams_order:
                    ty_le = row.get('T·ªâ l·ªá ƒë·∫°t (%)', 0)
                    if ty_le >= 99.5:
                        diem_bsc = 5.0
                    elif ty_le <= 89.5:
                        diem_bsc = 1.0
                    else:
                        diem_bsc = 1 + 4 * (ty_le - 89.5) / 10
                    bsc_scores[short_name]['C1.5'] = round(diem_bsc, 2)
    
    # T·∫°o DataFrame t·ª´ d·ªØ li·ªáu
    chart_data = pd.DataFrame(bsc_scores).T
    chart_data = chart_data.reindex(teams_order)  # ƒê·∫£m b·∫£o th·ª© t·ª±
    
    # T·∫°o bi·ªÉu ƒë·ªì
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(teams_order))
    width = 0.15  # Thu h·∫πp ƒë·ªÉ c√≥ ch·ªó cho 5 c·ªôt
    
    metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    
    for i, metric in enumerate(metrics):
        values = chart_data[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
        # Th√™m gi√° tr·ªã l√™n c·ªôt
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                       f'{val:.2f}', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('T·ªï K·ªπ thu·∫≠t', fontsize=12)
    ax.set_ylabel('ƒêi·ªÉm BSC', fontsize=12)
    ax.set_title('SO S√ÅNH ƒêI·ªÇM BSC TH·ª∞C T·∫æ GI·ªÆA C√ÅC T·ªî', fontsize=14, fontweight='bold')
    ax.set_xticks(x + width * 2)  # ƒêi·ªÅu ch·ªânh v·ªã tr√≠ label
    ax.set_xticklabels(teams_order, fontsize=11)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    # L∆∞u ho·∫∑c tr·∫£ v·ªÅ bytes
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_team_bsc_after_exclusion_chart(unit_data, c1x_reports=None, output_path=None, bsc_data=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì so s√°nh ƒëi·ªÉm BSC SAU GI·∫¢M TR·ª™ gi·ªØa 4 t·ªï
    
    Args:
        unit_data: D·ªØ li·ªáu unit_level t·ª´ load_unit_level_exclusion_data()
        c1x_reports: D·ªØ li·ªáu c1x_reports (fallback)
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file
        bsc_data: Dictionary t·ª´ load_bsc_unit_scores_from_comparison() (∆∞u ti√™n s·ª≠ d·ª•ng)
    """
    teams_order = ['Ph√∫c Th·ªç', 'Qu·∫£ng Oai', 'Su·ªëi Hai', 'S∆°n T√¢y']
    
    # Kh·ªüi t·∫°o dict ch·ª©a ƒëi·ªÉm BSC
    bsc_scores = {team: {'C1.1': 0, 'C1.2': 0, 'C1.3': 0, 'C1.4': 0, 'C1.5': 0} for team in teams_order}
    
    # Mapping t√™n ƒë·ªôi trong Excel -> t√™n ng·∫Øn
    team_name_map = {
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
    }
    
    def get_short_name(don_vi):
        if not don_vi: return None
        for orig, short in team_name_map.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return None
    
    # ================================================================
    # ∆ØU TI√äN: L·∫•y ƒëi·ªÉm t·ª´ bsc_data (Tong_hop_Diem_BSC_Don_Vi.xlsx) - c·ªôt (Sau)
    # ================================================================
    if bsc_data and bsc_data.get('units') is not None and not bsc_data['units'].empty:
        print("  üìä Bi·ªÉu ƒë·ªì sau GT: S·ª≠ d·ª•ng ƒëi·ªÉm t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx (c·ªôt Sau)")
        for _, row in bsc_data['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.1'] = row.get('Diem_C1.1 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.2'] = row.get('Diem_C1.2 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.4'] = row.get('Diem_C1.4 (Sau)', 0) or 0
                bsc_scores[short_name]['C1.5'] = row.get('Diem_C1.5 (Sau)', 0) or 0
        # C1.3 v·∫´n l·∫•y t·ª´ c1x_reports (kh√¥ng c√≥ gi·∫£m tr·ª´)
        if c1x_reports and 'c13' in c1x_reports:
            for _, row in c1x_reports['c13'].iterrows():
                don_vi = row.get('ƒê∆°n v·ªã', '')
                short_name = get_short_name(don_vi)
                if short_name in teams_order:
                    bsc_scores[short_name]['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0)
        
        # T·∫°o DataFrame v√† bi·ªÉu ƒë·ªì r·ªìi return (kh√¥ng ch·∫°y ti·∫øp ph·∫ßn t√≠nh t·ª´ unit_data)
        chart_data = pd.DataFrame(bsc_scores).T
        chart_data = chart_data.reindex(teams_order)
        
        fig, ax = plt.subplots(figsize=(14, 6))
        x = np.arange(len(teams_order))
        width = 0.15
        metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
        
        for i, metric in enumerate(metrics):
            values = chart_data[metric].fillna(0).values
            bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
            for bar, val in zip(bars, values):
                if val > 0:
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                           f'{val:.2f}', ha='center', va='bottom', fontsize=8)
        
        ax.set_xlabel('T·ªï K·ªπ thu·∫≠t', fontsize=12)
        ax.set_ylabel('ƒêi·ªÉm BSC', fontsize=12)
        ax.set_title('ƒêI·ªÇM BSC SAU GI·∫¢M TR·ª™ GI·ªÆA C√ÅC T·ªî', fontsize=14, fontweight='bold')
        ax.set_xticks(x + width * 2)
        ax.set_xticklabels(teams_order, fontsize=11)
        ax.set_ylim(0, 6)
        ax.legend(loc='upper right')
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        plt.tight_layout()
        
        if output_path:
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()
            return output_path
        else:
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            plt.close()
            buf.seek(0)
            return buf
    
    # FALLBACK: T√≠nh t·ª´ unit_data (c√°ch c≈©) - ch·ªâ ch·∫°y n·∫øu kh√¥ng c√≥ bsc_data
    # ================================================================
    # C√°c h√†m t√≠nh ƒëi·ªÉm BSC (theo ƒë√∫ng kpi_calculator.py)
    # ================================================================
    def tinh_diem_C11_TP1(kq):
        """C1.1 TP1 (30%): T·ª∑ l·ªá s·ª≠a ch·ªØa ch·ªß ƒë·ªông - kq l√† th·∫≠p ph√¢n"""
        if pd.isna(kq) or kq is None: return 5
        if kq >= 0.99: return 5
        elif kq > 0.96: return 1 + 4 * (kq - 0.96) / 0.03
        else: return 1
    
    def tinh_diem_C11_TP2(kq):
        """C1.1 TP2 (70%): T·ª∑ l·ªá s·ª≠a ch·ªØa b√°o h·ªèng ƒë√∫ng quy ƒë·ªãnh - kq l√† th·∫≠p ph√¢n"""
        if pd.isna(kq) or kq is None: return 5
        if kq >= 0.85: return 5
        elif kq >= 0.82: return 4 + (kq - 0.82) / 0.03
        elif kq >= 0.79: return 3 + (kq - 0.79) / 0.03
        elif kq >= 0.76: return 2
        else: return 1
    
    def tinh_diem_C12_TP1(kq):
        """C1.2 TP1 (50%): T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i - kq l√† th·∫≠p ph√¢n, c√†ng th·∫•p c√†ng t·ªët"""
        if pd.isna(kq) or kq is None: return 5
        if kq <= 0.025: return 5
        elif kq < 0.04: return 5 - 4 * (kq - 0.025) / 0.015
        else: return 1
    
    def tinh_diem_C12_TP2(kq):
        """C1.2 TP2 (50%): T·ª∑ l·ªá s·ª± c·ªë d·ªãch v·ª• BRCƒê - kq l√† th·∫≠p ph√¢n, c√†ng th·∫•p c√†ng t·ªët"""
        if pd.isna(kq) or kq is None: return 5
        if kq <= 0.02: return 5
        elif kq < 0.03: return 5 - 4 * (kq - 0.02) / 0.01
        else: return 1
    
    def tinh_diem_C14(kq):
        """C1.4: ƒê·ªô h√†i l√≤ng kh√°ch h√†ng - kq l√† th·∫≠p ph√¢n"""
        if pd.isna(kq) or kq is None: return 5
        if kq >= 0.995: return 5
        elif kq > 0.95: return 1 + 4 * (kq - 0.95) / 0.045
        else: return 1
    
    def tinh_diem_C15(kq):
        """C1.5: T·ªâ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• ƒë·∫°t - kq l√† th·∫≠p ph√¢n"""
        if pd.isna(kq) or kq is None: return 5
        if kq >= 0.995: return 5
        elif kq > 0.895: return 1 + 4 * (kq - 0.895) / 0.10
        else: return 1
    
    # ================================================================
    # T√≠nh C1.1 = 0.30*TP1 + 0.70*TP2
    # ================================================================
    # TP1 t·ª´ c11_sm2 (S·ª≠a ch·ªØa ch·ªß ƒë·ªông), TP2 t·ª´ c11_sm4 (S·ª≠a ch·ªØa BH)
    c11_tp1 = {}  # team -> t·ª∑ l·ªá th·∫≠p ph√¢n
    c11_tp2 = {}  # team -> t·ª∑ l·ªá th·∫≠p ph√¢n
    
    if unit_data and 'c11_sm2' in unit_data:
        df = unit_data['c11_sm2']
        for _, row in df.iterrows():
            short = get_short_name(row.get('ƒê∆°n v·ªã', ''))
            if short and short in teams_order:
                tyle = row.get('T·ª∑ l·ªá % (Sau GT)', 0) or 0
                c11_tp1[short] = tyle / 100 if tyle > 1 else tyle  # Chuy·ªÉn v·ªÅ th·∫≠p ph√¢n
    
    if unit_data and 'c11_sm4' in unit_data:
        df = unit_data['c11_sm4']
        for _, row in df.iterrows():
            short = get_short_name(row.get('ƒê∆°n v·ªã', ''))
            if short and short in teams_order:
                tyle = row.get('T·ª∑ l·ªá % (Sau GT)', 0) or 0
                c11_tp2[short] = tyle / 100 if tyle > 1 else tyle
    
    for team in teams_order:
        tp1 = c11_tp1.get(team)
        tp2 = c11_tp2.get(team)
        diem_tp1 = tinh_diem_C11_TP1(tp1)
        diem_tp2 = tinh_diem_C11_TP2(tp2)
        bsc_scores[team]['C1.1'] = round(0.30 * diem_tp1 + 0.70 * diem_tp2, 2)
    
    # ================================================================
    # T√≠nh C1.2 = 0.50*TP1 + 0.50*TP2
    # ================================================================
    # TP1 t·ª´ c12_sm1 (HLL), TP2 t·ª´ b√°o c√°o g·ªëc (kh√¥ng c√≥ trong unit_data ch∆∞a)
    c12_tp1 = {}  # team -> t·ª∑ l·ªá HLL th·∫≠p ph√¢n
    
    if unit_data and 'c12_sm1' in unit_data:
        df = unit_data['c12_sm1']
        for _, row in df.iterrows():
            short = get_short_name(row.get('ƒê∆°n v·ªã', ''))
            if short and short in teams_order:
                tyle_col = [c for c in df.columns if 'T·ª∑ l·ªá' in c and 'Sau GT' in c]
                if tyle_col:
                    tyle = row.get(tyle_col[0], 0) or 0
                    c12_tp1[short] = tyle / 100 if tyle > 1 else tyle
    
    
    # C1.2 TP2 - T·ª∑ l·ªá s·ª± c·ªë BRCƒê t·ª´ unit_data['c12_sm4'] (SAU GI·∫¢M TR·ª™)
    c12_tp2 = {}
    
    # H√†m t√≠nh ƒëi·ªÉm TP2 t·ª´ t·ª∑ l·ªá s·ª± c·ªë (‚Ä∞)
    def tinh_diem_tp2_from_percentage(tyle_percent):
        """T√≠nh ƒëi·ªÉm TP2 t·ª´ t·ª∑ l·ªá % - chuy·ªÉn sang ‚Ä∞ tr∆∞·ªõc"""
        if pd.isna(tyle_percent) or tyle_percent is None:
            return 5
        # Chuy·ªÉn % sang ‚Ä∞: 1.76% = 17.6‚Ä∞
        tyle_permil = tyle_percent * 10
        if tyle_permil <= 15:  # ‚â§1.5%
            return 5
        elif tyle_permil < 25:  # <2.5%
            return 5 - 4 * (tyle_permil - 15) / 10
        else:
            return 1
    
    if unit_data and 'c12_sm4' in unit_data:
        # ƒê·ªçc t·ª´ file SM4-C12-ti-le-su-co-dv-brcd.xlsx
        df = unit_data['c12_sm4']
        for _, row in df.iterrows():
            short = get_short_name(row.get('TEN_DOI', ''))
            if short and short in teams_order:
                tyle = row.get('T·ª∑ l·ªá b√°o h·ªèng (%) (Sau GT)', 0) or 0
                # T√≠nh ƒëi·ªÉm t·ª´ t·ª∑ l·ªá
                diem_tp2 = tinh_diem_tp2_from_percentage(tyle)
                c12_tp2[short] = diem_tp2
    elif c1x_reports and 'c12' in c1x_reports:
        # Fallback: d√πng d·ªØ li·ªáu g·ªëc n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu sau gi·∫£m tr·ª´
        df = c1x_reports['c12']
        for _, row in df.iterrows():
            don_vi = row.get('ƒê∆°n v·ªã', '')
            short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
            if short_name in teams_order:
                diem_tp2 = row.get('ƒêi·ªÉm C1.2 TP2', 5)
                c12_tp2[short_name] = diem_tp2
    
    for team in teams_order:
        tp1 = c12_tp1.get(team)
        diem_tp1 = tinh_diem_C12_TP1(tp1)
        diem_tp2 = c12_tp2.get(team, 5)  # Default 5 n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu
        bsc_scores[team]['C1.2'] = round(0.50 * diem_tp1 + 0.50 * diem_tp2, 2)
    
    # ================================================================
    # C1.3 - gi·ªØ nguy√™n t·ª´ c1x_reports (kh√¥ng c√≥ gi·∫£m tr·ª´)
    # ================================================================
    if c1x_reports and 'c13' in c1x_reports:
        df = c1x_reports['c13']
        for _, row in df.iterrows():
            don_vi = row.get('ƒê∆°n v·ªã', '')
            short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
            if short_name in teams_order:
                bsc_scores[short_name]['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0) or 0
    
    # ================================================================
    # C1.4 - sau gi·∫£m tr·ª´
    # ================================================================
    if unit_data and 'c14' in unit_data:
        df = unit_data['c14']
        for _, row in df.iterrows():
            short = get_short_name(row.get('ƒê∆°n v·ªã', ''))
            if short and short in teams_order:
                tyle_col = [c for c in df.columns if 'T·ª∑ l·ªá HL' in c and 'Sau GT' in c]
                if tyle_col:
                    tyle = row.get(tyle_col[0], 0) or 0
                    tyle_dec = tyle / 100 if tyle > 1 else tyle
                    bsc_scores[short]['C1.4'] = round(tinh_diem_C14(tyle_dec), 2)
    
    # ================================================================
    # C1.5 - S·ª¨ D·ª§NG D·ªÆ LI·ªÜU SAU GI·∫¢M TR·ª™ t·ª´ unit_data
    # ================================================================
    if unit_data and 'c15' in unit_data:
        # S·ª≠ d·ª•ng d·ªØ li·ªáu SAU GI·∫¢M TR·ª™ t·ª´ file So_sanh_C15.xlsx
        df = unit_data['c15']
        for _, row in df.iterrows():
            don_vi = row.get('ƒê∆°n v·ªã', '')
            short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
            if short_name in teams_order:
                ty_le = row.get('T·ª∑ l·ªá ƒë·∫°t % (Sau GT)', 0) or 0
                ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                bsc_scores[short_name]['C1.5'] = round(tinh_diem_C15(ty_le_dec), 2)
    elif c1x_reports and 'c15_ttvtst' in c1x_reports:
        # Fallback: n·∫øu kh√¥ng c√≥ d·ªØ li·ªáu sau gi·∫£m tr·ª´, d√πng d·ªØ li·ªáu g·ªëc
        df = c1x_reports['c15_ttvtst']
        for _, row in df.iterrows():
            don_vi = row.get('DOIVT', '')
            short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
            if short_name in teams_order:
                ty_le = row.get('T·ªâ l·ªá ƒë·∫°t (%)', 0) or 0
                ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                bsc_scores[short_name]['C1.5'] = round(tinh_diem_C15(ty_le_dec), 2)
    
    # T·∫°o DataFrame t·ª´ d·ªØ li·ªáu
    chart_data = pd.DataFrame(bsc_scores).T
    chart_data = chart_data.reindex(teams_order)
    
    # T·∫°o bi·ªÉu ƒë·ªì
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(teams_order))
    width = 0.15
    
    metrics = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    
    for i, metric in enumerate(metrics):
        values = chart_data[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=metric, color=BAR_COLORS[i])
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.05,
                       f'{val:.2f}', ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('T·ªï K·ªπ thu·∫≠t', fontsize=12)
    ax.set_ylabel('ƒêi·ªÉm BSC', fontsize=12)
    ax.set_title('ƒêI·ªÇM BSC SAU GI·∫¢M TR·ª™ GI·ªÆA C√ÅC T·ªî', fontsize=14, fontweight='bold')
    ax.set_xticks(x + width * 2)
    ax.set_xticklabels(teams_order, fontsize=11)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf

def create_nvkt_bar_chart(df_summary, team_name, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì c·ªôt so s√°nh ƒëi·ªÉm KPI theo NVKT trong 1 t·ªï
    
    Args:
        df_summary: DataFrame ch·ª©a d·ªØ li·ªáu KPI
        team_name: T√™n t·ªï c·∫ßn t·∫°o bi·ªÉu ƒë·ªì
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file (None = tr·∫£ v·ªÅ bytes)
    
    Returns:
        bytes ho·∫∑c str
    """
    # L·ªçc d·ªØ li·ªáu theo t·ªï
    df_team = df_summary[df_summary['don_vi'] == team_name].copy()
    
    if df_team.empty:
        return None
    
    # S·∫Øp x·∫øp theo t√™n NVKT
    df_team = df_team.sort_values('nvkt')
    
    # T·∫°o bi·ªÉu ƒë·ªì
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(df_team))
    width = 0.2
    
    metrics = ['Diem_C1.1', 'Diem_C1.2', 'Diem_C1.4', 'Diem_C1.5']
    labels = ['C1.1', 'C1.2', 'C1.4', 'C1.5']
    
    for i, (metric, label) in enumerate(zip(metrics, labels)):
        values = df_team[metric].fillna(0).values
        bars = ax.bar(x + i*width, values, width, label=label, color=BAR_COLORS[i])
        # Th√™m gi√° tr·ªã l√™n c·ªôt
        for bar, val in zip(bars, values):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.03,
                       f'{val:.1f}', ha='center', va='bottom', fontsize=7, rotation=90)
    
    short_name = TEAM_SHORT_NAMES.get(team_name, team_name)
    ax.set_xlabel('NVKT', fontsize=11)
    ax.set_ylabel('ƒêi·ªÉm KPI', fontsize=11)
    ax.set_title(f'ƒêI·ªÇM KPI THEO NVKT - {short_name.upper()}', fontsize=13, fontweight='bold')
    ax.set_xticks(x + width * 1.5)
    ax.set_xticklabels(df_team['nvkt'].values, fontsize=8, rotation=45, ha='right')
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right')
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


# =============================================================================
# H√ÄM T·∫†O B·∫¢NG TRONG WORD
# =============================================================================
def add_kpi_summary_table(doc, df_summary, team_name=None):
    """
    Th√™m b·∫£ng t·ªïng h·ª£p KPI v√†o document
    
    Args:
        doc: Document Word
        df_summary: DataFrame d·ªØ li·ªáu
        team_name: L·ªçc theo t·ªï (None = t·∫•t c·∫£)
    """
    if team_name:
        df = df_summary[df_summary['don_vi'] == team_name].copy()
    else:
        df = df_summary.copy()
    
    # S·∫Øp x·∫øp
    df = df.sort_values(['don_vi', 'nvkt'])
    
    # T·∫°o b·∫£ng
    headers = ['STT', 'ƒê∆°n v·ªã', 'NVKT', 'C1.1', 'C1.2', 'C1.4', 'C1.5']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], '1F4E79')
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        
        short_name = TEAM_SHORT_NAMES.get(row['don_vi'], row['don_vi'])
        
        data = [
            str(idx),
            short_name,
            row['nvkt'],
            format_number(row.get('Diem_C1.1', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            # T√¥ m√†u d√≤ng xen k·∫Ω
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F4FD')


def add_c11_detail_table(doc, df_detail, team_name=None):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.1 (Th√†nh ph·∫ßn 1 + Th√†nh ph·∫ßn 2)
    """
    if team_name:
        df = df_detail[df_detail['don_vi'] == team_name].copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    # Ti√™u ƒë·ªÅ
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.1 - Ch·∫•t l∆∞·ª£ng s·ª≠a ch·ªØa thu√™ bao BRCƒê', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('TP1 = S·ª≠a ch·ªØa ch·ªß ƒë·ªông (SCCD ‚â§72h) | TP2 = S·ª≠a ch·ªØa theo b√°o h·ªèng (SC BH)')
    
    headers = ['STT', 'NVKT', 'T·ªïng SCCD', 'ƒê·∫°t ‚â§72h', 'TL(%)', 'ƒêi·ªÉm TP1',
               'T·ªïng SC BH', 'ƒê√∫ng h·∫°n', 'TL(%)', 'ƒêi·ªÉm TP2', 'ƒêi·ªÉm C1.1']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '2E7D32')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c11_tp1_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp1_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp1', np.nan)),
            format_number(row.get('c11_tp2_tong_phieu', np.nan), 0),
            format_number(row.get('c11_tp2_phieu_dat', np.nan), 0),
            format_number(row.get('c11_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c11_tp2', np.nan)),
            format_number(row.get('Diem_C1.1', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F5E9')


def add_c12_detail_table(doc, df_detail, team_name=None):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.2 (Th√†nh ph·∫ßn 1 + Th√†nh ph·∫ßn 2)
    """
    if team_name:
        df = df_detail[df_detail['don_vi'] == team_name].copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.2 - T·ª∑ l·ªá thu√™ bao b√°o h·ªèng', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('TP1 = H·ªèng l·∫∑p (‚â•2 l·∫ßn/7 ng√†y) | TP2 = T·ª∑ l·ªá BH/TB qu·∫£n l√Ω | BH = B√°o h·ªèng | TB QL = Thu√™ bao qu·∫£n l√Ω')
    
    headers = ['STT', 'NVKT', 'H·ªèng l·∫∑p', 'T·ªïng BH', 'TL(%)', 'ƒêi·ªÉm TP1',
               'Phi·∫øu BH', 'TB QL', 'TL(‚Ä∞)', 'ƒêi·ªÉm TP2', 'ƒêi·ªÉm C1.2']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '1565C0')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(8)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c12_tp1_phieu_hll', np.nan), 0),
            format_number(row.get('c12_tp1_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp1_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp1', np.nan)),
            format_number(row.get('c12_tp2_phieu_bh', np.nan), 0),
            format_number(row.get('c12_tp2_tong_tb', np.nan), 0),
            format_number(row.get('c12_tp2_ty_le', np.nan)),
            format_number(row.get('diem_c12_tp2', np.nan)),
            format_number(row.get('Diem_C1.2', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(8)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E3F2FD')


def add_c14_detail_table(doc, df_detail, team_name=None):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng
    """
    if team_name:
        df = df_detail[df_detail['don_vi'] == team_name].copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('KS = Kh·∫£o s√°t | Kh√¥ng HL = Kh√¥ng h√†i l√≤ng | HL = H√†i l√≤ng')
    
    headers = ['STT', 'NVKT', 'T·ªïng KS', 'Kh√¥ng HL', 'T·ª∑ l·ªá HL (%)', 'ƒêi·ªÉm C1.4']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, 'F57C00')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c14_phieu_ks', np.nan), 0),
            format_number(row.get('c14_phieu_khl', np.nan), 0),
            format_number(row.get('c14_ty_le', np.nan)),
            format_number(row.get('Diem_C1.4', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'FFF3E0')


def add_c15_detail_table(doc, df_detail, team_name=None):
    """
    Th√™m b·∫£ng chi ti·∫øt C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª•
    """
    if team_name:
        df = df_detail[df_detail['don_vi'] == team_name].copy()
    else:
        df = df_detail.copy()
    
    df = df.sort_values(['don_vi', 'nvkt'])
    
    doc.add_heading('Chi ti·∫øt ch·ªâ ti√™u C1.5 - Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê ƒë·∫°t th·ªùi gian quy ƒë·ªãnh', level=3)
    
    # Ch√∫ th√≠ch
    p = doc.add_paragraph()
    p.add_run('üìã Ch√∫ th√≠ch: ').bold = True
    p.add_run('ƒê·∫°t TG = Ho√†n th√†nh ƒë√∫ng th·ªùi gian (ngo√†i CCCO: ‚â§24h, trong CCCO: theo quy ƒë·ªãnh)')
    
    headers = ['STT', 'NVKT', 'ƒê·∫°t TG', 'Kh√¥ng ƒë·∫°t', 'T·ªïng phi·∫øu', 'T·ª∑ l·ªá (%)', 'ƒêi·ªÉm C1.5']
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, header in enumerate(table.rows[0].cells):
        header.text = headers[i]
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header, '7B1FA2')
        run = header.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # D·ªØ li·ªáu
    for idx, (_, row) in enumerate(df.iterrows(), 1):
        cells = table.add_row().cells
        data = [
            str(idx),
            row['nvkt'],
            format_number(row.get('c15_phieu_dat', np.nan), 0),
            format_number(row.get('c15_phieu_khong_dat', np.nan), 0),
            format_number(row.get('c15_tong_phieu', np.nan), 0),
            format_number(row.get('c15_ty_le', np.nan)),
            format_number(row.get('Diem_C1.5', np.nan))
        ]
        
        for i, value in enumerate(data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'F3E5F5')


# =============================================================================
# H√ÄM SHC CHO B√ÅO C√ÅO T·ªîNG H·ª¢P
# =============================================================================

def load_shc_summary_by_unit(data_folder="downloads/baocao_hanoi"):
    """
    ƒê·ªçc d·ªØ li·ªáu SHC t·ªïng h·ª£p theo ƒë∆°n v·ªã t·ª´ sheet Xu_huong_theo_don_vi
    
    Returns:
        dict: {'units': {...}, 'dates': [...]} ho·∫∑c None
    """
    import glob
    
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    latest_file = max(files, key=os.path.getmtime)
    
    try:
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_don_vi')
        
        # C·ªôt ƒë·∫ßu ti√™n l√† ƒê∆°n v·ªã, c√°c c·ªôt c√≤n l·∫°i l√† ng√†y
        date_columns = [col for col in df.columns if col != 'ƒê∆°n v·ªã']
        
        result = {'units': {}, 'dates': date_columns}
        
        for _, row in df.iterrows():
            unit_name = row['ƒê∆°n v·ªã']
            values = [int(row[col]) if pd.notna(row[col]) else 0 for col in date_columns]
            result['units'][unit_name] = values
        
        return result
    except Exception as e:
        print(f"   ‚ö†Ô∏è L·ªói ƒë·ªçc SHC summary: {e}")
        return None


def load_shc_by_nvkt_for_unit(unit_name, data_folder="downloads/baocao_hanoi"):
    """
    ƒê·ªçc d·ªØ li·ªáu SHC theo t·ª´ng NVKT cho 1 ƒë∆°n v·ªã t·ª´ sheet Xu_huong_theo_NVKT
    
    Returns:
        dict: {'nvkt_list': [...], 'dates': [...], 'data': {nvkt: [values]}} ho·∫∑c None
    """
    import glob
    
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    latest_file = max(files, key=os.path.getmtime)
    
    try:
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_NVKT')
        
        # L·ªçc theo ƒë∆°n v·ªã
        df_unit = df[df['ƒê∆°n v·ªã'] == unit_name]
        
        if df_unit.empty:
            return None
        
        date_columns = [col for col in df.columns if col not in ['ƒê∆°n v·ªã', 'NVKT']]
        
        result = {
            'nvkt_list': [],
            'dates': [str(d) for d in date_columns],
            'data': {}
        }
        
        for _, row in df_unit.iterrows():
            nvkt = row['NVKT']
            values = [int(row[col]) if pd.notna(row[col]) else 0 for col in date_columns]
            result['nvkt_list'].append(nvkt)
            result['data'][nvkt] = values
        
        return result
    except Exception as e:
        print(f"   ‚ö†Ô∏è L·ªói ƒë·ªçc SHC by NVKT: {e}")
        return None


def create_nvkt_shc_grouped_chart(nvkt_data, unit_name, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì nh√≥m c·ªôt SHC theo NVKT, m·ªói ng√†y 1 m√†u kh√°c nhau
    """
    if not nvkt_data or not nvkt_data['data']:
        return None
    
    nvkt_list = nvkt_data['nvkt_list']
    dates = nvkt_data['dates']
    data = nvkt_data['data']
    
    # S·ª≠ d·ª•ng h·ªç t√™n ƒë·∫ßy ƒë·ªß
    nvkt_labels = nvkt_list
    
    # Setup figure
    fig, ax = plt.subplots(figsize=(14, 6))
    
    x = np.arange(len(nvkt_list))
    n_dates = len(dates)
    width = 0.8 / n_dates  # Chi·ªÅu r·ªông m·ªói c·ªôt
    
    # M√†u s·∫Øc cho t·ª´ng ng√†y
    colors = plt.cm.tab10(np.linspace(0, 1, n_dates))
    
    # V·∫Ω t·ª´ng ng√†y
    for i, date in enumerate(dates):
        values = [data[nvkt][i] for nvkt in nvkt_list]
        offset = (i - n_dates/2 + 0.5) * width
        bars = ax.bar(x + offset, values, width, label=date, color=colors[i])
        
        # Th√™m gi√° tr·ªã l√™n c·ªôt (ch·ªâ n·∫øu > 0)
        for bar, val in zip(bars, values):
            if val > 0:
                ax.annotate(f'{val}',
                           xy=(bar.get_x() + bar.get_width() / 2, bar.get_height()),
                           xytext=(0, 1),
                           textcoords="offset points",
                           ha='center', va='bottom',
                           fontsize=7, fontweight='bold')
    
    short_name = TEAM_SHORT_NAMES.get(unit_name, unit_name)
    ax.set_xlabel('NVKT', fontsize=11)
    ax.set_ylabel('S·ªë TB suy hao cao', fontsize=11)
    ax.set_title(f'K·∫æT QU·∫¢ X·ª¨ L√ù SUY HAO CAO - {short_name}', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x)
    ax.set_xticklabels(nvkt_labels, rotation=45, ha='right', fontsize=9)
    ax.legend(title='Ng√†y', loc='upper right', fontsize=8, ncol=2)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def create_shc_overview_chart(shc_data, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì t·ªïng h·ª£p SHC theo ng√†y cho t·∫•t c·∫£ ƒë∆°n v·ªã (stacked bar)
    """
    if not shc_data:
        return None
    
    dates = [str(d) for d in shc_data['dates']]
    units = shc_data['units']
    
    fig, ax = plt.subplots(figsize=(12, 6))
    
    x = np.arange(len(dates))
    width = 0.2
    colors = ['#2196F3', '#4CAF50', '#FF9800', '#E91E63']
    
    unit_names = list(units.keys())
    for i, unit_name in enumerate(unit_names):
        short_name = TEAM_SHORT_NAMES.get(unit_name, unit_name)
        values = units[unit_name]
        bars = ax.bar(x + i * width, values, width, label=short_name, color=colors[i % len(colors)])
        
        # Th√™m gi√° tr·ªã l√™n c·ªôt
        for bar, val in zip(bars, values):
            if val > 0:
                ax.annotate(f'{val}', xy=(bar.get_x() + bar.get_width() / 2, bar.get_height()),
                           xytext=(0, 2), textcoords="offset points",
                           ha='center', va='bottom', fontsize=8)
    
    ax.set_xlabel('Ng√†y', fontsize=11)
    ax.set_ylabel('S·ªë TB suy hao cao', fontsize=11)
    ax.set_title('XU H∆Ø·ªöNG SUY HAO CAO THEO ƒê∆†N V·ªä', fontsize=14, fontweight='bold', pad=15)
    ax.set_xticks(x + width * (len(unit_names) - 1) / 2)
    ax.set_xticklabels(dates, rotation=45, ha='right', fontsize=9)
    ax.legend(loc='upper right')
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_shc_overview_section(doc, data_folder="downloads/baocao_hanoi"):
    """
    Th√™m ph·∫ßn t·ªïng quan SHC v√†o b√°o c√°o (PH·∫¶N 1)
    """
    shc_data = load_shc_summary_by_unit(data_folder)
    
    if not shc_data:
        return
    
    doc.add_heading('1.4. T·ªïng quan Suy Hao Cao', level=2)
    
    dates = shc_data['dates']
    units = shc_data['units']
    
    # B·∫£ng t·ªïng h·ª£p
    table = doc.add_table(rows=1, cols=len(dates) + 2)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    headers = ['ƒê∆°n v·ªã'] + [str(d) for d in dates] + ['T·ªïng']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Data rows
    grand_total = 0
    for unit_name, values in units.items():
        cells = table.add_row().cells
        short_name = TEAM_SHORT_NAMES.get(unit_name, unit_name)
        cells[0].text = short_name
        
        total = sum(values)
        grand_total += total
        
        for j, val in enumerate(values):
            cells[j + 1].text = str(val)
            cells[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cells[j + 1].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if val > 5:
                run.font.bold = True
                run.font.color.rgb = RGBColor(200, 0, 0)
        
        cells[-1].text = str(total)
        cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[-1].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
    
    # D√≤ng t·ªïng
    cells = table.add_row().cells
    cells[0].text = 'T·ªîNG C·ªòNG'
    cells[0].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(cells[0], 'E3F2FD')
    
    totals_by_date = [sum(units[u][i] for u in units) for i in range(len(dates))]
    for j, total in enumerate(totals_by_date):
        cells[j + 1].text = str(total)
        cells[j + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[j + 1].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cells[j + 1], 'E3F2FD')
    
    cells[-1].text = str(grand_total)
    cells[-1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[-1].paragraphs[0].runs[0].font.bold = True
    set_cell_shading(cells[-1], 'E3F2FD')
    
    doc.add_paragraph()
    
    # Bi·ªÉu ƒë·ªì
    try:
        chart = create_shc_overview_chart(shc_data)
        if chart:
            doc.add_picture(chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì SHC: {e}")
    
    doc.add_paragraph()


def add_shc_unit_section(doc, unit_name, data_folder="downloads/baocao_hanoi"):
    """
    Th√™m ph·∫ßn SHC chi ti·∫øt cho 1 ƒë∆°n v·ªã (trong PH·∫¶N 2)
    """
    shc_data = load_shc_summary_by_unit(data_folder)
    
    if not shc_data or unit_name not in shc_data['units']:
        return
    
    dates = shc_data['dates']
    values = shc_data['units'][unit_name]
    
    short_name = TEAM_SHORT_NAMES.get(unit_name, unit_name)
    doc.add_heading(f'S·ªë li·ªáu Suy Hao Cao - {short_name}', level=3)
    
    # B·∫£ng d·ªØ li·ªáu - c·∫£i thi·ªán format
    table = doc.add_table(rows=2, cols=len(dates) + 1)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header row - bao g·ªìm c·ªôt Ch·ªâ ti√™u
    headers = ['Ng√†y'] + [str(d) for d in dates]
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)
    
    # Data row
    table.rows[1].cells[0].text = 'S·ªë TB SHC'
    table.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_shading(table.rows[1].cells[0], 'E3F2FD')
    
    for j, val in enumerate(values):
        cell = table.rows[1].cells[j + 1]
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
        run.font.bold = True
        
        # T√¥ m√†u theo m·ª©c ƒë·ªô (d·ª±a tr√™n ng∆∞·ª°ng ph√π h·ª£p cho ƒë∆°n v·ªã)
        avg_val = sum(values) / len(values) if values else 0
        if val <= avg_val * 0.5:
            set_cell_shading(cell, 'C8E6C9')  # Xanh nh·∫°t - t·ªët
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif val <= avg_val * 1.2:
            set_cell_shading(cell, 'FFF9C4')  # V√†ng nh·∫°t - trung b√¨nh
        else:
            set_cell_shading(cell, 'FFCDD2')  # ƒê·ªè nh·∫°t - cao
            run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    
    # Th·ªëng k√™ chi ti·∫øt
    total = sum(values)
    avg = total / len(values) if values else 0
    max_val = max(values) if values else 0
    min_val = min(values) if values else 0
    max_day = dates[values.index(max_val)] if values else 'N/A'
    min_day = dates[values.index(min_val)] if values else 'N/A'
    
    p = doc.add_paragraph()
    p.add_run(f'üìä T·ªïng: {total} | Trung b√¨nh: {avg:.1f}/ng√†y | ').bold = True
    p.add_run(f'Cao nh·∫•t: {max_val} ({max_day}) | Th·∫•p nh·∫•t: {min_val} ({min_day})')
    
    doc.add_paragraph()
    
    # Bi·ªÉu ƒë·ªì nh√≥m theo NVKT (nh∆∞ h√¨nh m·∫´u)
    try:
        nvkt_data = load_shc_by_nvkt_for_unit(unit_name, data_folder)
        if nvkt_data:
            chart = create_nvkt_shc_grouped_chart(nvkt_data, unit_name)
            if chart:
                doc.add_picture(chart, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì SHC cho {short_name}: {e}")
    
    doc.add_paragraph()


# =============================================================================
# H√ÄM T·∫†O B√ÅO C√ÅO WORD HO√ÄN CH·ªàNH
# =============================================================================

def generate_kpi_report(kpi_folder=DEFAULT_KPI_FOLDER, output_folder=DEFAULT_OUTPUT_FOLDER, 
                         report_month=None, report_title=None, include_exclusion=True,
                         exclusion_folder="downloads/kq_sau_giam_tru"):
    """
    T·∫°o b√°o c√°o Word ho√†n ch·ªânh v·ªõi b·∫£ng bi·ªÉu v√† bi·ªÉu ƒë·ªì KPI
    
    Args:
        kpi_folder: Th∆∞ m·ª•c ch·ª©a file KPI Excel
        output_folder: Th∆∞ m·ª•c xu·∫•t b√°o c√°o Word
        report_month: Th√°ng b√°o c√°o (vd: "01/2026"), m·∫∑c ƒë·ªãnh l√† th√°ng hi·ªán t·∫°i
        report_title: Ti√™u ƒë·ªÅ t√πy ch·ªânh
        include_exclusion: Bao g·ªìm d·ªØ li·ªáu sau gi·∫£m tr·ª´ (m·∫∑c ƒë·ªãnh True)
        exclusion_folder: Th∆∞ m·ª•c ch·ª©a d·ªØ li·ªáu sau gi·∫£m tr·ª´
        
    Returns:
        str: ƒê∆∞·ªùng d·∫´n file Word ƒë√£ t·∫°o
    """
    print("="*60)
    print("üìù B·∫ÆT ƒê·∫¶U T·∫†O B√ÅO C√ÅO WORD KPI")
    print("="*60)
    
    # X√°c ƒë·ªãnh th√°ng b√°o c√°o
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # T·∫°o th∆∞ m·ª•c output n·∫øu ch∆∞a c√≥
    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # ƒê·ªçc d·ªØ li·ªáu
    print("üìä ƒê·ªçc d·ªØ li·ªáu KPI...")
    df_summary, df_detail = load_kpi_data(kpi_folder)
    
    # L·∫•y danh s√°ch c√°c t·ªï
    teams = df_summary['don_vi'].unique()
    print(f"   T√¨m th·∫•y {len(teams)} t·ªï k·ªπ thu·∫≠t")
    
    # T·∫°o document
    doc = Document()
    
    # Thi·∫øt l·∫≠p style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # =========================================================================
    # TI√äU ƒê·ªÄ
    # =========================================================================
    created_time = datetime.now().strftime('%d/%m/%Y %H:%M')
    title = doc.add_heading(level=0)
    title_run = title.add_run(report_title or f'B√ÅO C√ÅO K·∫æT QU·∫¢ BSC/KPI\nTH√ÅNG {report_month}')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Ng√†y t·∫°o: {created_time}")
    doc.add_paragraph()
    
    # =========================================================================
    # PH·∫¶N 1: T·ªîNG QUAN
    # =========================================================================
    print("üìà T·∫°o ph·∫ßn T·ªïng quan...")
    doc.add_heading('PH·∫¶N 1: T·ªîNG QUAN', level=1)
    
    # ƒê·ªçc d·ªØ li·ªáu chi ti·∫øt t·ª´ c√°c b√°o c√°o C1.x (c·∫ßn cho bi·ªÉu ƒë·ªì BSC)
    print("üìä ƒê·ªçc d·ªØ li·ªáu chi ti·∫øt t·ª´ c√°c b√°o c√°o C1.x...")
    c1x_reports = load_c1x_reports()
    
    # ƒê·ªçc d·ªØ li·ªáu gi·∫£m tr·ª´ n·∫øu ƒë∆∞·ª£c y√™u c·∫ßu
    comparison_data = None
    if include_exclusion:
        print("üìä ƒê·ªçc d·ªØ li·ªáu so s√°nh tr∆∞·ªõc/sau gi·∫£m tr·ª´...")
        comparison_data = load_exclusion_comparison_data(exclusion_folder)
    
    # 1.1 Bi·ªÉu ƒë·ªì so s√°nh ƒëi·ªÉm BSC th·ª±c t·∫ø 4 t·ªï
    doc.add_heading('1.1. So s√°nh ƒëi·ªÉm BSC th·ª±c t·∫ø gi·ªØa c√°c t·ªï', level=2)
    if c1x_reports:
        # Load ƒëi·ªÉm BSC t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx n·∫øu c√≥
        bsc_scores_for_chart = load_bsc_unit_scores_from_comparison(exclusion_folder) if include_exclusion else None
        team_chart = create_team_comparison_chart(c1x_reports, bsc_data=bsc_scores_for_chart)
        doc.add_picture(team_chart, width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("(Kh√¥ng c√≥ d·ªØ li·ªáu C1.x ƒë·ªÉ t·∫°o bi·ªÉu ƒë·ªì)")
    doc.add_paragraph()
    
    # 1.1.b Bi·ªÉu ƒë·ªì BSC SAU GI·∫¢M TR·ª™ (ngay sau bi·ªÉu ƒë·ªì th√¥)
    if include_exclusion and comparison_data:
        doc.add_heading('1.1.b. So s√°nh ƒëi·ªÉm BSC sau gi·∫£m tr·ª´ gi·ªØa c√°c t·ªï', level=2)
        try:
            unit_data = load_unit_level_exclusion_data(exclusion_folder)
            if unit_data:
                # S·ª≠ d·ª•ng bsc_scores_for_chart ƒë√£ load ·ªü tr√™n
                bsc_after_chart = create_team_bsc_after_exclusion_chart(unit_data, c1x_reports, bsc_data=bsc_scores_for_chart)
                if bsc_after_chart:
                    doc.add_picture(bsc_after_chart, width=Inches(6.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì BSC sau gi·∫£m tr·ª´: {e}")
        doc.add_paragraph()
    
    # 1.2 Th√™m b·∫£ng th·ªëng k√™ t·ªïng h·ª£p theo t·ªï - s·ª≠ d·ª•ng ƒëi·ªÉm BSC th·ª±c t·∫ø
    doc.add_heading('1.2. Th·ªëng k√™ ƒëi·ªÉm BSC theo ƒë∆°n v·ªã', level=2)
    
    # Ch√∫ th√≠ch gi·∫£i th√≠ch c√°c ch·ªâ ti√™u
    legend = doc.add_paragraph()
    legend.add_run('üìã CH√ö TH√çCH C√ÅC CH·ªà TI√äU BSC - VI·ªÑN C·∫¢NH KH√ÅCH H√ÄNG (C)').bold = True
    
    # C1.1
    p11 = doc.add_paragraph()
    run_title = p11.add_run('C1.1 - Ch·∫•t l∆∞·ª£ng s·ª≠a ch·ªØa thu√™ bao BRCƒê: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p11.add_run('G·ªìm 2 th√†nh ph·∫ßn:\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_tp1 = p11.add_run('   ‚Ä¢ TP1 (30%): S·ª≠a ch·ªØa ch·ªß ƒë·ªông - T·ª∑ l·ªá phi·∫øu SCCD ho√†n th√†nh ‚â§72h\n')
    run_tp1.italic = True
    run_tp1.font.size = Pt(10)
    
    run_tp2 = p11.add_run('   ‚Ä¢ TP2 (70%): S·ª≠a ch·ªØa theo b√°o h·ªèng - T·ª∑ l·ªá phi·∫øu BH ho√†n th√†nh ƒë√∫ng h·∫°n')
    run_tp2.italic = True
    run_tp2.font.size = Pt(10)
    
    # C1.2
    p12 = doc.add_paragraph()
    run_title = p12.add_run('C1.2 - T·ª∑ l·ªá thu√™ bao b√°o h·ªèng: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p12.add_run('G·ªìm 2 th√†nh ph·∫ßn:\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_tp1 = p12.add_run('   ‚Ä¢ TP1 (50%): H·ªèng l·∫∑p l·∫°i - T·ª∑ l·ªá TB b√°o h·ªèng ‚â•2 l·∫ßn/7 ng√†y\n')
    run_tp1.italic = True
    run_tp1.font.size = Pt(10)
    
    run_tp2 = p12.add_run('   ‚Ä¢ TP2 (50%): T·ª∑ l·ªá s·ª± c·ªë - T·ª∑ l·ªá phi·∫øu BH / T·ªïng TB qu·∫£n l√Ω (‚Ä∞)')
    run_tp2.italic = True
    run_tp2.font.size = Pt(10)
    
    # C1.3
    p13 = doc.add_paragraph()
    run_title = p13.add_run('C1.3 - Ch·∫•t l∆∞·ª£ng s·ª≠a ch·ªØa k√™nh TSL (Leased Line): ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p13.add_run('√Åp d·ª•ng cho c√°c d·ªãch v·ª• Internet tr·ª±c ti·∫øp, k√™nh thu√™ ri√™ng, MegaWan, Metronet, Siptrunking')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    # C1.4
    p14 = doc.add_paragraph()
    run_title = p14.add_run('C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p14.add_run('T·ª∑ l·ªá kh√°ch h√†ng h√†i l√≤ng sau khi ƒë∆∞·ª£c s·ª≠a ch·ªØa (qua kh·∫£o s√°t)')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    # C1.5
    p15 = doc.add_paragraph()
    run_title = p15.add_run('C1.5 - Thi·∫øt l·∫≠p d·ªãch v·ª• BRCƒê ƒë·∫°t th·ªùi gian quy ƒë·ªãnh: ')
    run_title.bold = True
    run_title.italic = True
    run_title.font.size = Pt(10)
    
    run_desc = p15.add_run('T·ª∑ l·ªá phi·∫øu l·∫Øp ƒë·∫∑t ho√†n th√†nh ƒë√∫ng h·∫°n\n')
    run_desc.italic = True
    run_desc.font.size = Pt(10)
    
    run_detail = p15.add_run('   ‚Ä¢ Ngo√†i CCCO: ‚â§24h | Trong CCCO: Phi·∫øu tr∆∞·ªõc 17h xong trong ng√†y')
    run_detail.italic = True
    run_detail.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # =========================================================================
    # S·ª¨ D·ª§NG D·ªÆ LI·ªÜU T·ª™ Tong_hop_Diem_BSC_Don_Vi.xlsx (c√≥ c·∫£ Tr∆∞·ªõc v√† Sau GT)
    # =========================================================================
    teams_order = ['Ph√∫c Th·ªç', 'Qu·∫£ng Oai', 'Su·ªëi Hai', 'S∆°n T√¢y']
    team_name_map = {
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Ph√∫c Th·ªç': 'Ph√∫c Th·ªç',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Qu·∫£ng Oai': 'Qu·∫£ng Oai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n Su·ªëi hai': 'Su·ªëi Hai',
        'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'T·ªï K·ªπ thu·∫≠t ƒë·ªãa b√†n S∆°n T√¢y': 'S∆°n T√¢y',
        'TTVT S∆°n T√¢y': 'TTVT S∆°n T√¢y'
    }
    
    def get_short_name(don_vi):
        if not don_vi: return None
        for orig, short in team_name_map.items():
            if orig in str(don_vi) or short == don_vi:
                return short
        return don_vi
    
    # ƒê·ªçc d·ªØ li·ªáu t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx
    print("üìä ƒê·ªçc ƒëi·ªÉm BSC t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx...")
    bsc_scores = load_bsc_unit_scores_from_comparison(exclusion_folder)
    
    # Kh·ªüi t·∫°o d·ªØ li·ªáu BSC m·∫∑c ƒë·ªãnh
    bsc_data = {team: {
        'C1.1_truoc': 0, 'C1.1_sau': 0,
        'C1.2_truoc': 0, 'C1.2_sau': 0,
        'C1.3': 0,  # C1.3 gi·ªØ nguy√™n, kh√¥ng c√≥ gi·∫£m tr·ª´
        'C1.4_truoc': 0, 'C1.4_sau': 0,
        'C1.5_truoc': 0, 'C1.5_sau': 0
    } for team in teams_order + ['TTVT S∆°n T√¢y']}
    
    # ƒê·ªçc ƒëi·ªÉm t·ª´ Tong_hop_Diem_BSC_Don_Vi.xlsx
    if bsc_scores['units'] is not None and not bsc_scores['units'].empty:
        for _, row in bsc_scores['units'].iterrows():
            don_vi = row.get('don_vi', '')
            short_name = get_short_name(don_vi)
            if short_name in bsc_data:
                bsc_data[short_name]['C1.1_truoc'] = row.get('Diem_C1.1 (Tr∆∞·ªõc)', 0) or 0
                bsc_data[short_name]['C1.1_sau'] = row.get('Diem_C1.1 (Sau)', 0) or 0
                bsc_data[short_name]['C1.2_truoc'] = row.get('Diem_C1.2 (Tr∆∞·ªõc)', 0) or 0
                bsc_data[short_name]['C1.2_sau'] = row.get('Diem_C1.2 (Sau)', 0) or 0
                bsc_data[short_name]['C1.4_truoc'] = row.get('Diem_C1.4 (Tr∆∞·ªõc)', 0) or 0
                bsc_data[short_name]['C1.4_sau'] = row.get('Diem_C1.4 (Sau)', 0) or 0
                bsc_data[short_name]['C1.5_truoc'] = row.get('Diem_C1.5 (Tr∆∞·ªõc)', 0) or 0
                bsc_data[short_name]['C1.5_sau'] = row.get('Diem_C1.5 (Sau)', 0) or 0
    
    # ƒê·ªçc C1.3 t·ª´ c1x_reports (gi·ªØ nguy√™n v√¨ kh√¥ng c√≥ gi·∫£m tr·ª´)
    if c1x_reports and 'c13' in c1x_reports:
        for _, row in c1x_reports['c13'].iterrows():
            don_vi = row.get('ƒê∆°n v·ªã', '')
            short_name = get_short_name(don_vi)
            if short_name in bsc_data:
                bsc_data[short_name]['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0) or 0
            elif don_vi == 'T·ªïng':
                bsc_data['TTVT S∆°n T√¢y']['C1.3'] = row.get('Ch·ªâ ti√™u BSC', 0) or 0
    
    # T·∫°o b·∫£ng v·ªõi c·∫•u tr√∫c: ƒê∆°n v·ªã | C1.1 (Tr∆∞·ªõc/Sau) | C1.2 (Tr∆∞·ªõc/Sau) | C1.3 | C1.4 (Tr∆∞·ªõc/Sau) | C1.5 (Tr∆∞·ªõc/Sau)
    headers = ['ƒê∆°n v·ªã', 'C1.1', '', 'C1.2', '', 'C1.3', 'C1.4', '', 'C1.5', '']
    sub_headers = ['', 'Tr∆∞·ªõc', 'Sau', 'Tr∆∞·ªõc', 'Sau', '', 'Tr∆∞·ªõc', 'Sau', 'Tr∆∞·ªõc', 'Sau']
    
    table = doc.add_table(rows=2, cols=10)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header row 1 - Merge c√°c √¥
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ƒê∆°n v·ªã'
    header_cells[0].merge(table.rows[1].cells[0])  # Merge v·ªõi d√≤ng d∆∞·ªõi
    
    # C1.1 header v·ªõi merge
    header_cells[1].text = 'C1.1'
    header_cells[1].merge(header_cells[2])
    
    # C1.2 header v·ªõi merge
    header_cells[3].text = 'C1.2'
    header_cells[3].merge(header_cells[4])
    
    # C1.3 header (kh√¥ng c√≥ tr∆∞·ªõc/sau)
    header_cells[5].text = 'C1.3'
    header_cells[5].merge(table.rows[1].cells[5])
    
    # C1.4 header v·ªõi merge
    header_cells[6].text = 'C1.4'
    header_cells[6].merge(header_cells[7])
    
    # C1.5 header v·ªõi merge
    header_cells[8].text = 'C1.5'
    header_cells[8].merge(header_cells[9])
    
    # Format header row 1
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
    
    # Header row 2 (sub headers: Tr∆∞·ªõc/Sau)
    sub_header_cells = table.rows[1].cells
    sub_header_values = ['', 'Tr∆∞·ªõc', 'Sau', 'Tr∆∞·ªõc', 'Sau', '', 'Tr∆∞·ªõc', 'Sau', 'Tr∆∞·ªõc', 'Sau']
    for i, val in enumerate(sub_header_values):
        if val:
            sub_header_cells[i].text = val
            sub_header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(sub_header_cells[i], '2E75B6')
            if sub_header_cells[i].paragraphs[0].runs:
                run = sub_header_cells[i].paragraphs[0].runs[0]
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    
    # Data rows
    for idx, team in enumerate(teams_order, 1):
        cells = table.add_row().cells
        scores = bsc_data[team]
        row_data = [
            team,
            format_number(scores['C1.1_truoc']),
            format_number(scores['C1.1_sau']),
            format_number(scores['C1.2_truoc']),
            format_number(scores['C1.2_sau']),
            format_number(scores['C1.3']),
            format_number(scores['C1.4_truoc']),
            format_number(scores['C1.4_sau']),
            format_number(scores['C1.5_truoc']),
            format_number(scores['C1.5_sau'])
        ]
        for i, value in enumerate(row_data):
            cells[i].text = value
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cells[i].paragraphs[0].runs:
                run = cells[i].paragraphs[0].runs[0]
                run.font.size = Pt(9)
            if idx % 2 == 0:
                set_cell_shading(cells[i], 'E8F4FD')
    
    # D√≤ng TTVT S∆°n T√¢y
    cells = table.add_row().cells
    ttvt_scores = bsc_data['TTVT S∆°n T√¢y']
    ttvt_data = [
        'TTVT S∆°n T√¢y',
        format_number(ttvt_scores['C1.1_truoc']),
        format_number(ttvt_scores['C1.1_sau']),
        format_number(ttvt_scores['C1.2_truoc']),
        format_number(ttvt_scores['C1.2_sau']),
        format_number(ttvt_scores['C1.3']),
        format_number(ttvt_scores['C1.4_truoc']),
        format_number(ttvt_scores['C1.4_sau']),
        format_number(ttvt_scores['C1.5_truoc']),
        format_number(ttvt_scores['C1.5_sau'])
    ]
    for i, value in enumerate(ttvt_data):
        cells[i].text = value
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cells[i].paragraphs[0].runs:
            run = cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(9)
            run.font.bold = True
        set_cell_shading(cells[i], 'B2DFDB')
    
    doc.add_paragraph()
    
    # 1.4 S·ªë li·ªáu chi ti·∫øt c√°c ch·ªâ ti√™u BSC theo ƒê·ªôi/TTVT (s·ª≠ d·ª•ng c1x_reports ƒë√£ load)
    # N·∫øu c√≥ comparison_data, s·∫Ω th√™m b·∫£ng sau gi·∫£m tr·ª´ ngay sau m·ªói b·∫£ng th√¥
    if c1x_reports:
        # Load unit_data n·∫øu c·∫ßn cho b·∫£ng t·ªïng h·ª£p sau gi·∫£m tr·ª´
        unit_data = None
        if include_exclusion:
            unit_data = load_unit_level_exclusion_data(exclusion_folder)
        add_c1x_overview_table(doc, c1x_reports, comparison_data, unit_data, exclusion_folder)
    
    # 1.5 T·ªïng quan Suy Hao Cao
    print("üìä Th√™m ph·∫ßn Suy Hao Cao...")
    add_shc_overview_section(doc, data_folder="downloads/baocao_hanoi")
    
    # 1.5 S·ªê LI·ªÜU SAU GI·∫¢M TR·ª™ - T·ªîNG H·ª¢P (n·∫øu c√≥)
    if include_exclusion and comparison_data:
        print("üìä Th√™m ph·∫ßn t·ªïng h·ª£p s·ªë li·ªáu sau gi·∫£m tr·ª´...")
        doc.add_heading('1.5. T·ªïng h·ª£p s·ªë li·ªáu sau gi·∫£m tr·ª´', level=2)
        
        # Ch√∫ th√≠ch
        p_note = doc.add_paragraph()
        p_note.add_run('üìã GHI CH√ö: ').bold = True
        p_note.add_run('S·ªë li·ªáu sau gi·∫£m tr·ª´ ƒë∆∞·ª£c t√≠nh sau khi lo·∫°i b·ªè c√°c phi·∫øu b√°o h·ªèng thu·ªôc di·ªán lo·∫°i tr·ª´. B·∫£ng chi ti·∫øt ƒë√£ ƒë∆∞·ª£c hi·ªÉn th·ªã ngay sau m·ªói b·∫£ng ch·ªâ ti√™u ·ªü ph·∫ßn 1.3.')
        doc.add_paragraph()
        
        # B·∫£ng t·ªïng h·ª£p so s√°nh
        add_exclusion_summary_table(doc, comparison_data)
        
        # Bi·ªÉu ƒë·ªì ri√™ng cho d·ªØ li·ªáu sau gi·∫£m tr·ª´
        doc.add_heading('Bi·ªÉu ƒë·ªì t·ª∑ l·ªá sau gi·∫£m tr·ª´', level=3)
        try:
            exclusion_chart = create_exclusion_bar_chart(comparison_data)
            if exclusion_chart:
                doc.add_picture(exclusion_chart, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì sau gi·∫£m tr·ª´: {e}")
        
        doc.add_paragraph()
        
        # Ph·∫ßn th·ªëng k√™ theo ƒë∆°n v·ªã (T·ªï v√† TTVT)
        print("üìä Th√™m ph·∫ßn th·ªëng k√™ theo ƒë∆°n v·ªã...")
        unit_data = load_unit_level_exclusion_data(exclusion_folder)
        if unit_data:
            add_unit_level_exclusion_section(doc, unit_data, c1x_reports)
    
    doc.add_page_break()
    
    # =========================================================================
    # PH·∫¶N 2: CHI TI·∫æT THEO T·ªî
    # =========================================================================
    print("üìã T·∫°o ph·∫ßn Chi ti·∫øt theo t·ªï...")
    doc.add_heading('PH·∫¶N 2: CHI TI·∫æT THEO T·ª™NG T·ªî', level=1)
    
    # ƒê·ªçc d·ªØ li·ªáu KPI sau gi·∫£m tr·ª´ theo NVKT
    df_exclusion_nvkt = None
    df_exclusion_detail = None
    df_raw_detail = None  # D·ªØ li·ªáu Th√¥ (tr∆∞·ªõc gi·∫£m tr·ª´) t·ª´ So_sanh_*.xlsx
    if include_exclusion:
        df_exclusion_nvkt = load_nvkt_exclusion_summary(exclusion_folder)
        df_exclusion_detail = load_nvkt_exclusion_detail(exclusion_folder)
        df_raw_detail = load_nvkt_raw_detail(exclusion_folder)  # Th√¥ t·ª´ c√πng file
    
    for team_idx, team_name in enumerate(teams, 1):
        short_name = TEAM_SHORT_NAMES.get(team_name, team_name)
        print(f"   üìÅ T·ªï {team_idx}: {short_name}")
        
        doc.add_heading(f'2.{team_idx}. {short_name}', level=2)
        
        # B·∫£ng KPI t·ªïng h·ª£p c·ªßa t·ªï
        doc.add_heading(f'B·∫£ng ƒëi·ªÉm KPI t·ªïng h·ª£p', level=3)
        add_kpi_summary_table(doc, df_summary, team_name)
        doc.add_paragraph()
        
        # B·∫£ng KPI sau gi·∫£m tr·ª´ (n·∫øu c√≥)
        if df_exclusion_nvkt is not None:
            doc.add_heading('B·∫£ng ƒëi·ªÉm KPI t·ªïng h·ª£p (sau gi·∫£m tr·ª´)', level=3)
            add_kpi_summary_table_after_exclusion(doc, df_exclusion_nvkt, team_name)
            doc.add_paragraph()
        
        # Bi·ªÉu ƒë·ªì c·ªôt so s√°nh NVKT
        doc.add_heading(f'Bi·ªÉu ƒë·ªì so s√°nh ƒëi·ªÉm KPI theo NVKT', level=3)
        nvkt_chart = create_nvkt_bar_chart(df_summary, team_name)
        if nvkt_chart:
            doc.add_picture(nvkt_chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        
        # Bi·ªÉu ƒë·ªì sau gi·∫£m tr·ª´ (n·∫øu c√≥)
        if df_exclusion_nvkt is not None:
            doc.add_heading('Bi·ªÉu ƒë·ªì so s√°nh ƒëi·ªÉm KPI theo NVKT (sau gi·∫£m tr·ª´)', level=3)
            nvkt_chart_gt = create_nvkt_bar_chart_after_exclusion(df_exclusion_nvkt, team_name)
            if nvkt_chart_gt:
                doc.add_picture(nvkt_chart_gt, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        # B·∫£ng chi ti·∫øt t·ª´ng ch·ªâ ti√™u (Th√¥ - tr∆∞·ªõc gi·∫£m tr·ª´)
        if df_raw_detail is not None:
            add_c11_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c11_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.1 chi ti·∫øt sau gi·∫£m tr·ª´
        if df_exclusion_detail is not None:
            add_c11_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c12_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c12_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.2 chi ti·∫øt sau gi·∫£m tr·ª´
        if df_exclusion_detail is not None:
            add_c12_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c14_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c14_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()
        
        # C1.4 chi ti·∫øt sau gi·∫£m tr·ª´
        if df_exclusion_detail is not None:
            add_c14_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()
        
        if df_raw_detail is not None:
            add_c15_detail_table(doc, df_raw_detail, team_name)
        else:
            add_c15_detail_table(doc, df_detail, team_name)
        doc.add_paragraph()

        # C1.5 chi ti·∫øt sau gi·∫£m tr·ª´
        if df_exclusion_detail is not None:
            add_c15_detail_table_after_exclusion(doc, df_exclusion_detail, team_name)
            doc.add_paragraph()

        # S·ªë li·ªáu Suy Hao Cao cho t·ªï
        add_shc_unit_section(doc, team_name, data_folder="downloads/baocao_hanoi")
        
        # Th√™m page break sau m·ªói t·ªï (tr·ª´ t·ªï cu·ªëi)
        if team_idx < len(teams):
            doc.add_page_break()
    
    # =========================================================================
    # PH·∫¶N 3: K·∫æT LU·∫¨N
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('PH·∫¶N 3: K·∫æT LU·∫¨N V√Ä KI·∫æN NGH·ªä', level=1)
    
    # Th·ªëng k√™ t·ªïng quan
    total_nvkt = len(df_summary)
    avg_c11 = df_summary['Diem_C1.1'].mean()
    avg_c12 = df_summary['Diem_C1.2'].mean()
    avg_c14 = df_summary['Diem_C1.4'].mean()
    avg_c15 = df_summary['Diem_C1.5'].mean()
    
    doc.add_heading('3.1. Th·ªëng k√™ t·ªïng quan', level=2)
    doc.add_paragraph(f"‚Ä¢ T·ªïng s·ªë NVKT: {total_nvkt}")
    doc.add_paragraph(f"‚Ä¢ ƒêi·ªÉm trung b√¨nh C1.1: {format_number(avg_c11)}")
    doc.add_paragraph(f"‚Ä¢ ƒêi·ªÉm trung b√¨nh C1.2: {format_number(avg_c12)}")
    doc.add_paragraph(f"‚Ä¢ ƒêi·ªÉm trung b√¨nh C1.4: {format_number(avg_c14)}")
    doc.add_paragraph(f"‚Ä¢ ƒêi·ªÉm trung b√¨nh C1.5: {format_number(avg_c15)}")
    
    # ƒê√°nh gi√°
    doc.add_heading('3.2. ƒê√°nh gi√° chung', level=2)
    doc.add_paragraph("(Ph·∫ßn n√†y c·∫ßn b·ªï sung n·ªôi dung ƒë√°nh gi√° theo th·ª±c t·∫ø)")
    
    doc.add_heading('3.3. Ki·∫øn ngh·ªã', level=2)
    doc.add_paragraph("(Ph·∫ßn n√†y c·∫ßn b·ªï sung n·ªôi dung ki·∫øn ngh·ªã theo th·ª±c t·∫ø)")
    
    # =========================================================================
    # PH·ª§ L·ª§C: B·∫¢NG ƒêI·ªÇM KPI CHI TI·∫æT THEO NVKT
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('PH·ª§ L·ª§C: B·∫¢NG ƒêI·ªÇM KPI CHI TI·∫æT THEO NVKT', level=1)
    
    # B·∫£ng KPI t·ªïng h·ª£p theo NVKT
    doc.add_heading('B·∫£ng ƒëi·ªÉm KPI t·ªïng h·ª£p theo NVKT', level=2)
    add_kpi_summary_table(doc, df_summary)
    doc.add_paragraph()
    
    # T·∫°o bi·ªÉu ƒë·ªì bar cho t·ª´ng t·ªï
    doc.add_heading('Bi·ªÉu ƒë·ªì ƒëi·ªÉm KPI theo NVKT (theo t·ª´ng t·ªï)', level=2)
    
    for team_name in teams:
        short_name = TEAM_SHORT_NAMES.get(team_name, team_name)
        doc.add_heading(f'{short_name}', level=3)
        
        nvkt_chart = create_nvkt_bar_chart(df_summary, team_name)
        if nvkt_chart:
            doc.add_picture(nvkt_chart, width=Inches(6.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    # =========================================================================
    # L∆ØU FILE
    # =========================================================================
    output_file = output_path / f"Bao_cao_KPI_NVKT_{report_month.replace('/', '_')}.docx"
    doc.save(output_file)
    
    print("="*60)
    print(f"‚úÖ ƒê√É T·∫†O B√ÅO C√ÅO WORD TH√ÄNH C√îNG!")
    print(f"   üìÑ File: {output_file}")
    print("="*60)
    
    return str(output_file)


# =============================================================================
# H√ÄM T·∫†O B√ÅO C√ÅO C√Å NH√ÇN CHO T·ª™NG NVKT
# =============================================================================

def sanitize_filename(name):
    """
    Chu·∫©n h√≥a t√™n file - lo·∫°i b·ªè k√Ω t·ª± ƒë·∫∑c bi·ªát v√† thay kho·∫£ng tr·∫Øng b·∫±ng _
    """
    import re
    # Thay kho·∫£ng tr·∫Øng b·∫±ng _
    name = name.replace(' ', '_')
    # Lo·∫°i b·ªè k√Ω t·ª± ƒë·∫∑c bi·ªát (gi·ªØ l·∫°i ch·ªØ c√°i Unicode, s·ªë v√† _)
    name = re.sub(r'[^\w\s-]', '', name, flags=re.UNICODE)
    return name


def create_individual_radar_chart(nvkt_data, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì radar so s√°nh ƒëi·ªÉm KPI c·ªßa 1 NVKT
    
    Args:
        nvkt_data: Dictionary ch·ª©a ƒëi·ªÉm KPI c·ªßa NVKT
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file (None = tr·∫£ v·ªÅ bytes)
    
    Returns:
        bytes ho·∫∑c str
    """
    # L·∫•y ƒëi·ªÉm c√°c ch·ªâ ti√™u
    metrics = ['Diem_C1.1', 'Diem_C1.2', 'Diem_C1.4', 'Diem_C1.5']
    labels = ['C1.1\nS·ª≠a ch·ªØa', 'C1.2\nB√°o h·ªèng', 'C1.4\nH√†i l√≤ng', 'C1.5\nThi·∫øt l·∫≠p DV']
    
    values = []
    for m in metrics:
        val = nvkt_data.get(m, 0)
        if pd.isna(val):
            val = 0
        values.append(val)
    
    # S·ªë l∆∞·ª£ng bi·∫øn
    num_vars = len(labels)
    
    # T√≠nh g√≥c cho m·ªói tr·ª•c
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    
    # ƒê√≥ng v√≤ng radar
    values += values[:1]
    angles += angles[:1]
    
    # T·∫°o figure
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    
    # V·∫Ω radar
    ax.fill(angles, values, color='#2E86AB', alpha=0.25)
    ax.plot(angles, values, color='#2E86AB', linewidth=2, marker='o', markersize=8)
    
    # V·∫Ω ƒë∆∞·ªùng chu·∫©n 5 ƒëi·ªÉm
    target_values = [5] * (num_vars + 1)
    ax.plot(angles, target_values, color='#C73E1D', linewidth=1.5, linestyle='--', alpha=0.7, label='M·ª•c ti√™u (5 ƒëi·ªÉm)')
    
    # Thi·∫øt l·∫≠p c√°c tr·ª•c
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, fontsize=11)
    ax.set_ylim(0, 5.5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(['1', '2', '3', '4', '5'], fontsize=9)
    
    # Th√™m gi√° tr·ªã l√™n ƒëi·ªÉm
    for angle, value, label in zip(angles[:-1], values[:-1], labels):
        ax.annotate(f'{value:.2f}', 
                   xy=(angle, value), 
                   xytext=(angle, value + 0.3),
                   ha='center', va='bottom', fontsize=10, fontweight='bold')
    
    ax.legend(loc='upper right', bbox_to_anchor=(1.2, 1.1))
    ax.set_title('BI·ªÇU ƒê·ªí ƒêI·ªÇM KPI', fontsize=14, fontweight='bold', pad=20)
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_individual_summary_table(doc, nvkt_data):
    """
    Th√™m b·∫£ng t·ªïng h·ª£p ƒëi·ªÉm KPI cho 1 NVKT v√†o document
    
    Args:
        doc: Document Word
        nvkt_data: Dictionary ho·∫∑c Series ch·ª©a d·ªØ li·ªáu NVKT
    """
    # T·∫°o b·∫£ng 2 c·ªôt: Ch·ªâ ti√™u - ƒêi·ªÉm
    headers = ['Ch·ªâ ti√™u', 'ƒêi·ªÉm']
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    set_table_border(table)
    
    # Header
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1F4E79')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
    
    # D·ªØ li·ªáu c√°c ch·ªâ ti√™u
    kpi_info = [
        ('C1.1 - T·ª∑ l·ªá s·ª≠a ch·ªØa phi·∫øu ch·∫•t l∆∞·ª£ng & b√°o h·ªèng', 'Diem_C1.1'),
        ('C1.2 - T·ª∑ l·ªá b√°o h·ªèng l·∫∑p l·∫°i & s·ª± c·ªë d·ªãch v·ª•', 'Diem_C1.2'),
        ('C1.4 - ƒê·ªô h√†i l√≤ng kh√°ch h√†ng', 'Diem_C1.4'),
        ('C1.5 - T·ª∑ l·ªá thi·∫øt l·∫≠p d·ªãch v·ª• ƒë·∫°t th·ªùi gian quy ƒë·ªãnh', 'Diem_C1.5'),
    ]
    
    for idx, (label, col) in enumerate(kpi_info, 1):
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        value = nvkt_data.get(col, np.nan)
        cells[1].text = format_number(value)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, cell in enumerate(cells):
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(11)
            if idx % 2 == 0:
                set_cell_shading(cell, 'E8F4FD')
            
            # T√¥ m√†u ƒëi·ªÉm theo m·ª©c
            if i == 1 and not pd.isna(value):
                if value >= 4.5:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Xanh l√°
                    run.font.bold = True
                elif value < 3:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # ƒê·ªè
                    run.font.bold = True


def add_individual_c11_detail(doc, nvkt_data):
    """
    Th√™m chi ti·∫øt ch·ªâ ti√™u C1.1 cho 1 NVKT
    """
    doc.add_heading('2. CHI TI·∫æT CH·ªà TI√äU C1.1 - CH·∫§T L∆Ø·ª¢NG S·ª¨A CH·ªÆA BRCƒê', level=2)
    
    # Th√†nh ph·∫ßn 1
    doc.add_heading('2.1. Th√†nh ph·∫ßn 1: S·ª≠a ch·ªØa ch·ªß ƒë·ªông (SCCD ‚â§72h) - 30%', level=3)
    
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'
    set_table_border(table1)
    
    headers1 = ['T·ªïng SCCD', 'ƒê·∫°t ‚â§72h', 'T·ª∑ l·ªá (%)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table1.rows[0].cells):
        cell.text = headers1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '2E7D32')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data1 = [
        format_number(nvkt_data.get('c11_tp1_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c11_tp1_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c11_tp1_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c11_tp1', np.nan))
    ]
    for i, cell in enumerate(table1.rows[1].cells):
        cell.text = data1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Th√†nh ph·∫ßn 2
    doc.add_heading('2.2. Th√†nh ph·∫ßn 2: S·ª≠a ch·ªØa theo b√°o h·ªèng (ƒë√∫ng h·∫°n) - 70%', level=3)
    
    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    set_table_border(table2)
    
    headers2 = ['T·ªïng SC BH', 'ƒê√∫ng h·∫°n', 'T·ª∑ l·ªá (%)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table2.rows[0].cells):
        cell.text = headers2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '388E3C')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data2 = [
        format_number(nvkt_data.get('c11_tp2_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c11_tp2_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c11_tp2_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c11_tp2', np.nan))
    ]
    for i, cell in enumerate(table2.rows[1].cells):
        cell.text = data2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ƒêi·ªÉm t·ªïng h·ª£p
    p = doc.add_paragraph()
    p.add_run('‚û§ ƒêI·ªÇM T·ªîNG H·ª¢P C1.1: ').bold = True
    diem_c11 = nvkt_data.get('Diem_C1.1', np.nan)
    run = p.add_run(format_number(diem_c11))
    run.bold = True
    run.font.size = Pt(12)
    if not pd.isna(diem_c11):
        if diem_c11 >= 4.5:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif diem_c11 < 3:
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    p.add_run(' (= TP1 √ó 30% + TP2 √ó 70%)')


def add_individual_c12_detail(doc, nvkt_data):
    """
    Th√™m chi ti·∫øt ch·ªâ ti√™u C1.2 cho 1 NVKT
    """
    doc.add_heading('3. CHI TI·∫æT CH·ªà TI√äU C1.2 - T·ª∂ L·ªÜ THU√ä BAO B√ÅO H·ªéNG', level=2)
    
    # Th√†nh ph·∫ßn 1
    doc.add_heading('3.1. Th√†nh ph·∫ßn 1: H·ªèng l·∫∑p (‚â•2 l·∫ßn/7 ng√†y) - 50%', level=3)
    
    table1 = doc.add_table(rows=2, cols=4)
    table1.style = 'Table Grid'
    set_table_border(table1)
    
    headers1 = ['TB h·ªèng l·∫∑p', 'T·ªïng BH', 'T·ª∑ l·ªá (%)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table1.rows[0].cells):
        cell.text = headers1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1565C0')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data1 = [
        format_number(nvkt_data.get('c12_tp1_phieu_hll', np.nan), 0),
        format_number(nvkt_data.get('c12_tp1_phieu_bh', np.nan), 0),
        format_number(nvkt_data.get('c12_tp1_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c12_tp1', np.nan))
    ]
    for i, cell in enumerate(table1.rows[1].cells):
        cell.text = data1[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Th√†nh ph·∫ßn 2
    doc.add_heading('3.2. Th√†nh ph·∫ßn 2: T·ª∑ l·ªá BH/TB qu·∫£n l√Ω (‚Ä∞) - 50%', level=3)
    
    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    set_table_border(table2)
    
    headers2 = ['Phi·∫øu BH', 'TB qu·∫£n l√Ω', 'T·ª∑ l·ªá (‚Ä∞)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table2.rows[0].cells):
        cell.text = headers2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '1976D2')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data2 = [
        format_number(nvkt_data.get('c12_tp2_phieu_bh', np.nan), 0),
        format_number(nvkt_data.get('c12_tp2_tong_tb', np.nan), 0),
        format_number(nvkt_data.get('c12_tp2_ty_le', np.nan)),
        format_number(nvkt_data.get('diem_c12_tp2', np.nan))
    ]
    for i, cell in enumerate(table2.rows[1].cells):
        cell.text = data2[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ƒêi·ªÉm t·ªïng h·ª£p
    p = doc.add_paragraph()
    p.add_run('‚û§ ƒêI·ªÇM T·ªîNG H·ª¢P C1.2: ').bold = True
    diem_c12 = nvkt_data.get('Diem_C1.2', np.nan)
    run = p.add_run(format_number(diem_c12))
    run.bold = True
    run.font.size = Pt(12)
    if not pd.isna(diem_c12):
        if diem_c12 >= 4.5:
            run.font.color.rgb = RGBColor(0, 128, 0)
        elif diem_c12 < 3:
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    p.add_run(' (= TP1 √ó 50% + TP2 √ó 50%)')


def add_individual_c14_detail(doc, nvkt_data):
    """
    Th√™m chi ti·∫øt ch·ªâ ti√™u C1.4 cho 1 NVKT
    """
    doc.add_heading('4. CHI TI·∫æT CH·ªà TI√äU C1.4 - ƒê·ªò H√ÄI L√íNG KH√ÅCH H√ÄNG', level=2)
    
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    set_table_border(table)
    
    headers = ['Phi·∫øu KS th√†nh c√¥ng', 'Phi·∫øu KH kh√¥ng HL', 'T·ª∑ l·ªá HL (%)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, 'F57C00')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data = [
        format_number(nvkt_data.get('c14_phieu_ks', np.nan), 0),
        format_number(nvkt_data.get('c14_phieu_khl', np.nan), 0),
        format_number(nvkt_data.get('c14_ty_le', np.nan)),
        format_number(nvkt_data.get('Diem_C1.4', np.nan))
    ]
    for i, cell in enumerate(table.rows[1].cells):
        cell.text = data[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # C√¥ng th·ª©c t√≠nh ƒëi·ªÉm
    p = doc.add_paragraph()
    p.add_run('üìå C√¥ng th·ª©c t√≠nh ƒëi·ªÉm: ').bold = True
    p.add_run('‚â• 99.5% = 5 ƒëi·ªÉm, > 95% = n·ªôi suy, ‚â§ 95% = 1 ƒëi·ªÉm')


def add_individual_c15_detail(doc, nvkt_data):
    """
    Th√™m chi ti·∫øt ch·ªâ ti√™u C1.5 cho 1 NVKT
    """
    doc.add_heading('5. CHI TI·∫æT CH·ªà TI√äU C1.5 - T·ª∂ L·ªÜ THI·∫æT L·∫¨P D·ªäCH V·ª§ ƒê·∫†T', level=2)
    
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    set_table_border(table)
    
    headers = ['Phi·∫øu ƒë·∫°t', 'Phi·∫øu kh√¥ng ƒë·∫°t', 'T·ªïng phi·∫øu', 'T·ª∑ l·ªá ƒë·∫°t (%)', 'ƒêi·ªÉm']
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '7B1FA2')
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
    
    data = [
        format_number(nvkt_data.get('c15_phieu_dat', np.nan), 0),
        format_number(nvkt_data.get('c15_phieu_khong_dat', np.nan), 0),
        format_number(nvkt_data.get('c15_tong_phieu', np.nan), 0),
        format_number(nvkt_data.get('c15_ty_le', np.nan)),
        format_number(nvkt_data.get('Diem_C1.5', np.nan))
    ]
    for i, cell in enumerate(table.rows[1].cells):
        cell.text = data[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # C√¥ng th·ª©c t√≠nh ƒëi·ªÉm
    p = doc.add_paragraph()
    p.add_run('üìå C√¥ng th·ª©c t√≠nh ƒëi·ªÉm: ').bold = True
    p.add_run('‚â• 99.5% = 5 ƒëi·ªÉm, 89.5% < KQ < 99.5% = n·ªôi suy, ‚â§ 89.5% = 1 ƒëi·ªÉm')


def load_shc_trend_data(nvkt_name, data_folder="downloads/baocao_hanoi"):
    """
    ƒê·ªçc d·ªØ li·ªáu xu h∆∞·ªõng SHC cho 1 NVKT t·ª´ file Bao_cao_xu_huong_*.xlsx
    
    Args:
        nvkt_name: T√™n NVKT c·∫ßn t√¨m
        data_folder: Th∆∞ m·ª•c ch·ª©a file b√°o c√°o
    
    Returns:
        dict: {'dates': [...], 'values': [...], 'don_vi': '...'} ho·∫∑c None
    """
    import glob
    
    # T√¨m file Bao_cao_xu_huong m·ªõi nh·∫•t
    pattern = os.path.join(data_folder, "Bao_cao_xu_huong_SHC_*.xlsx")
    files = glob.glob(pattern)
    
    if not files:
        return None
    
    # S·∫Øp x·∫øp theo th·ªùi gian ch·ªânh s·ª≠a file (m·ªõi nh·∫•t cu·ªëi c√πng)
    latest_file = max(files, key=os.path.getmtime)
    print(f"   üìä S·ª≠ d·ª•ng file SHC: {os.path.basename(latest_file)}")
    
    try:
        # ƒê·ªçc sheet Xu_huong_theo_NVKT
        df = pd.read_excel(latest_file, sheet_name='Xu_huong_theo_NVKT')
        
        # T√¨m NVKT trong c·ªôt 'NVKT'
        nvkt_row = df[df['NVKT'] == nvkt_name]
        
        if nvkt_row.empty:
            return None
        
        row = nvkt_row.iloc[0]
        don_vi = row.get('ƒê∆°n v·ªã', '')
        
        # L·∫•y c√°c c·ªôt ng√†y (kh√¥ng ph·∫£i 'ƒê∆°n v·ªã' v√† 'NVKT')
        date_columns = [col for col in df.columns if col not in ['ƒê∆°n v·ªã', 'NVKT']]
        
        dates = []
        values = []
        for col in date_columns:
            dates.append(str(col))
            val = row[col]
            values.append(int(val) if pd.notna(val) else 0)
        
        return {
            'dates': dates,
            'values': values,
            'don_vi': don_vi
        }
    except Exception as e:
        print(f"   ‚ö†Ô∏è L·ªói ƒë·ªçc d·ªØ li·ªáu SHC: {e}")
        return None


def create_shc_trend_bar_chart(shc_data, nvkt_name, output_path=None):
    """
    T·∫°o bi·ªÉu ƒë·ªì c·ªôt th·ªÉ hi·ªán xu h∆∞·ªõng s·ªë TB suy hao cao theo ng√†y
    
    Args:
        shc_data: dict v·ªõi keys 'dates' v√† 'values'
        nvkt_name: T√™n NVKT
        output_path: ƒê∆∞·ªùng d·∫´n l∆∞u file (None = tr·∫£ v·ªÅ bytes)
    
    Returns:
        bytes ho·∫∑c str: Chart image
    """
    if not shc_data or not shc_data.get('dates') or not shc_data.get('values'):
        return None
    
    dates = shc_data['dates']
    values = shc_data['values']
    
    # T·∫°o figure
    fig, ax = plt.subplots(figsize=(10, 5))
    
    # V·∫Ω bi·ªÉu ƒë·ªì c·ªôt
    x_pos = range(len(dates))
    bars = ax.bar(x_pos, values, color='#2196F3', edgecolor='#1565C0', linewidth=1)
    
    # Th√™m gi√° tr·ªã l√™n c·ªôt
    for bar, val in zip(bars, values):
        height = bar.get_height()
        ax.annotate(f'{val}',
                   xy=(bar.get_x() + bar.get_width() / 2, height),
                   xytext=(0, 3),
                   textcoords="offset points",
                   ha='center', va='bottom',
                   fontsize=9, fontweight='bold')
    
    # Thi·∫øt l·∫≠p tr·ª•c
    ax.set_xticks(x_pos)
    ax.set_xticklabels(dates, rotation=45, ha='right', fontsize=9)
    ax.set_ylabel('S·ªë TB suy hao cao', fontsize=11)
    ax.set_xlabel('Ng√†y', fontsize=11)
    ax.set_title(f'XU H∆Ø·ªöNG S·ªê TB SUY HAO CAO - {nvkt_name}', fontsize=12, fontweight='bold', pad=15)
    
    # Grid v√† layout
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    # ƒê∆∞·ªùng xu h∆∞·ªõng trung b√¨nh
    if len(values) > 1:
        avg = sum(values) / len(values)
        ax.axhline(y=avg, color='#E91E63', linestyle='--', linewidth=1.5, alpha=0.7, 
                   label=f'TB: {avg:.1f}')
        ax.legend(loc='upper right')
    
    plt.tight_layout()
    
    if output_path:
        plt.savefig(output_path, dpi=150, bbox_inches='tight')
        plt.close()
        return output_path
    else:
        buf = io.BytesIO()
        plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
        plt.close()
        buf.seek(0)
        return buf


def add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi"):
    """
    Th√™m ph·∫ßn S·ªë li·ªáu Suy Hao Cao v√†o b√°o c√°o c√° nh√¢n
    Bao g·ªìm: B·∫£ng d·ªØ li·ªáu + Bi·ªÉu ƒë·ªì bar
    
    Args:
        doc: Document Word
        nvkt_name: T√™n NVKT
        data_folder: Th∆∞ m·ª•c ch·ª©a file b√°o c√°o
    """
    # Load d·ªØ li·ªáu SHC
    shc_data = load_shc_trend_data(nvkt_name, data_folder)
    
    if not shc_data:
        return  # Kh√¥ng c√≥ d·ªØ li·ªáu SHC
    
    doc.add_page_break()
    doc.add_heading('6. S·ªê LI·ªÜU SUY HAO CAO', level=2)
    
    dates = shc_data['dates']
    values = shc_data['values']
    
    # M√¥ t·∫£
    p = doc.add_paragraph()
    p.add_run('üìä Xu h∆∞·ªõng s·ªë thu√™ bao suy hao cao theo ng√†y:').bold = True
    
    doc.add_paragraph()
    
    # T·∫°o b·∫£ng d·ªØ li·ªáu (chia th√†nh c√°c nh√≥m n·∫øu nhi·ªÅu ng√†y)
    max_cols = 10  # S·ªë c·ªôt t·ªëi ƒëa m·ªói b·∫£ng
    
    for i in range(0, len(dates), max_cols):
        chunk_dates = dates[i:i+max_cols]
        chunk_values = values[i:i+max_cols]
        
        table = doc.add_table(rows=2, cols=len(chunk_dates))
        table.style = 'Table Grid'
        set_table_border(table)
        
        # Header row - Ng√†y
        for j, date in enumerate(chunk_dates):
            cell = table.rows[0].cells[j]
            cell.text = str(date)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, '1E88E5')
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        
        # Data row - S·ªë l∆∞·ª£ng
        for j, val in enumerate(chunk_values):
            cell = table.rows[1].cells[j]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            run.font.bold = True
            
            # T√¥ m√†u theo m·ª©c ƒë·ªô
            if val == 0:
                set_cell_shading(cell, 'C8E6C9')  # Xanh l√° nh·∫°t
                run.font.color.rgb = RGBColor(0, 128, 0)
            elif val <= 3:
                set_cell_shading(cell, 'FFF9C4')  # V√†ng nh·∫°t
            else:
                set_cell_shading(cell, 'FFCDD2')  # ƒê·ªè nh·∫°t
                run.font.color.rgb = RGBColor(200, 0, 0)
        
        doc.add_paragraph()
    
    # Th·ªëng k√™ t·ªïng quan
    total = sum(values)
    avg = total / len(values) if values else 0
    max_val = max(values) if values else 0
    min_val = min(values) if values else 0
    
    p = doc.add_paragraph()
    p.add_run(f'üìà T·ªïng s·ªë TB SHC trong k·ª≥: ').bold = True
    p.add_run(f'{total}')
    
    p = doc.add_paragraph()
    p.add_run(f'üìä Trung b√¨nh/ng√†y: ').bold = True
    p.add_run(f'{avg:.1f}')
    
    p = doc.add_paragraph()
    p.add_run(f'‚¨ÜÔ∏è Cao nh·∫•t: ').bold = True
    p.add_run(f'{max_val}')
    p.add_run(f'  |  ')
    p.add_run(f'‚¨áÔ∏è Th·∫•p nh·∫•t: ').bold = True
    p.add_run(f'{min_val}')
    
    doc.add_paragraph()
    
    # Bi·ªÉu ƒë·ªì bar
    try:
        chart = create_shc_trend_bar_chart(shc_data, nvkt_name)
        if chart:
            doc.add_picture(chart, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì SHC: {e}")


def generate_individual_kpi_report(nvkt_name, don_vi, kpi_folder=DEFAULT_KPI_FOLDER, 
                                    output_folder=DEFAULT_OUTPUT_FOLDER, report_month=None):
    """
    T·∫°o b√°o c√°o Word cho 1 NVKT c·ª• th·ªÉ
    
    Args:
        nvkt_name: T√™n NVKT (vd: "B√πi VƒÉn Du·∫©n")
        don_vi: T√™n ƒë∆°n v·ªã/t·ªï (vd: "T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç")
        kpi_folder: Th∆∞ m·ª•c ch·ª©a file KPI
        output_folder: Th∆∞ m·ª•c xu·∫•t b√°o c√°o
        report_month: Th√°ng b√°o c√°o (vd: "01/2026")
    
    Returns:
        str: ƒê∆∞·ªùng d·∫´n file Word ƒë√£ t·∫°o
    """
    # X√°c ƒë·ªãnh th√°ng b√°o c√°o
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # ƒê·ªçc d·ªØ li·ªáu KPI chi ti·∫øt
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    # L·ªçc d·ªØ li·ªáu cho NVKT c·ª• th·ªÉ
    nvkt_df = df_detail[(df_detail['nvkt'] == nvkt_name) & (df_detail['don_vi'] == don_vi)]
    
    if nvkt_df.empty:
        print(f"‚ö†Ô∏è Kh√¥ng t√¨m th·∫•y d·ªØ li·ªáu cho NVKT: {nvkt_name} - {don_vi}")
        return None
    
    nvkt_data = nvkt_df.iloc[0].to_dict()
    
    # T·∫°o th∆∞ m·ª•c output theo t·ªï
    short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
    folder_name = sanitize_filename(short_name)
    output_path = Path(output_folder) / "individual_reports" / folder_name
    output_path.mkdir(parents=True, exist_ok=True)
    
    # T·∫°o document
    doc = Document()
    
    # Thi·∫øt l·∫≠p style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # =========================================================================
    # TI√äU ƒê·ªÄ
    # =========================================================================
    created_time = datetime.now().strftime('%d/%m/%Y %H:%M')
    
    title = doc.add_heading(level=0)
    title_run = title.add_run('B√ÅO C√ÅO K·∫æT QU·∫¢ BSC/KPI C√Å NH√ÇN')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(level=1)
    subtitle_run = subtitle.add_run(f'TH√ÅNG {report_month}')
    subtitle_run.font.size = Pt(16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Th√¥ng tin c√° nh√¢n
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('H·ªç v√† t√™n:', nvkt_name),
        ('ƒê∆°n v·ªã:', short_name),
        ('Ng√†y t·∫°o b√°o c√°o:', created_time)
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # PH·∫¶N 1: T·ªîNG QUAN
    # =========================================================================
    doc.add_heading('1. T·ªîNG QUAN ƒêI·ªÇM KPI', level=2)
    
    # B·∫£ng t·ªïng h·ª£p
    add_individual_summary_table(doc, nvkt_data)
    doc.add_paragraph()
    
    # Bi·ªÉu ƒë·ªì radar
    try:
        radar_chart = create_individual_radar_chart(nvkt_data)
        doc.add_picture(radar_chart, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"   ‚ö†Ô∏è Kh√¥ng th·ªÉ t·∫°o bi·ªÉu ƒë·ªì radar: {e}")
    
    doc.add_page_break()
    
    # =========================================================================
    # CHI TI·∫æT T·ª™NG CH·ªà TI√äU
    # =========================================================================
    add_individual_c11_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c12_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c14_detail(doc, nvkt_data)
    doc.add_paragraph()
    
    add_individual_c15_detail(doc, nvkt_data)
    
    # =========================================================================
    # PH·∫¶N 6: S·ªê LI·ªÜU SUY HAO CAO
    # =========================================================================
    add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi")
    
    # =========================================================================
    # L∆ØU FILE
    # =========================================================================
    safe_name = sanitize_filename(nvkt_name)
    output_file = output_path / f"Bao_cao_KPI_{safe_name}_{report_month.replace('/', '_')}.docx"
    doc.save(output_file)
    
    return str(output_file)



def generate_all_individual_reports(kpi_folder=DEFAULT_KPI_FOLDER, output_folder=DEFAULT_OUTPUT_FOLDER, 
                                     report_month=None):
    """
    T·∫°o b√°o c√°o cho T·∫§T C·∫¢ NVKT
    
    Args:
        kpi_folder: Th∆∞ m·ª•c ch·ª©a file KPI
        output_folder: Th∆∞ m·ª•c xu·∫•t b√°o c√°o
        report_month: Th√°ng b√°o c√°o
    
    Returns:
        list: Danh s√°ch ƒë∆∞·ªùng d·∫´n c√°c file ƒë√£ t·∫°o
    """
    print("="*60)
    print("üìù B·∫ÆT ƒê·∫¶U T·∫†O B√ÅO C√ÅO KPI C√Å NH√ÇN CHO T·∫§T C·∫¢ NVKT")
    print("="*60)
    
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
    
    # ƒê·ªçc d·ªØ li·ªáu KPI
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_ChiTiet.xlsx"
    df_detail = pd.read_excel(detail_file)
    
    # L·∫•y danh s√°ch NVKT
    nvkt_list = df_detail[['don_vi', 'nvkt']].drop_duplicates()
    total = len(nvkt_list)
    
    print(f"üìä T√¨m th·∫•y {total} NVKT")
    print()
    
    success_files = []
    failed_count = 0
    
    for idx, row in nvkt_list.iterrows():
        don_vi = row['don_vi']
        nvkt_name = row['nvkt']
        short_name = TEAM_SHORT_NAMES.get(don_vi, don_vi)
        
        current = len(success_files) + failed_count + 1
        print(f"   [{current}/{total}] {nvkt_name} ({short_name})...", end=" ")
        
        try:
            result = generate_individual_kpi_report(
                nvkt_name=nvkt_name,
                don_vi=don_vi,
                kpi_folder=kpi_folder,
                output_folder=output_folder,
                report_month=report_month
            )
            if result:
                success_files.append(result)
                print("‚úÖ")
            else:
                failed_count += 1
                print("‚ùå (kh√¥ng c√≥ d·ªØ li·ªáu)")
        except Exception as e:
            failed_count += 1
            print(f"‚ùå ({str(e)[:30]})")
    
    print()
    print("="*60)
    print(f"‚úÖ HO√ÄN TH√ÄNH!")
    print(f"   üìÑ ƒê√£ t·∫°o: {len(success_files)} b√°o c√°o")
    print(f"   ‚ùå Th·∫•t b·∫°i: {failed_count}")
    print(f"   üìÅ Th∆∞ m·ª•c: {Path(output_folder) / 'individual_reports'}")
    print("="*60)
    
    return success_files


def generate_all_individual_reports_after_exclusion(kpi_folder, output_root, report_month=None):
    """
    T·∫°o b√°o c√°o c√° nh√¢n sau gi·∫£m tr·ª´, ph√¢n lo·∫°i theo th∆∞ m·ª•c T·ªï k·ªπ thu·∫≠t
    L∆∞u t·∫°i: {output_root}/ca_nhan/{t√™n t·ªï k·ªπ thu·∫≠t}/
    """
    print("="*60)
    print("üìù B·∫ÆT ƒê·∫¶U T·∫†O B√ÅO C√ÅO KPI C√Å NH√ÇN SAU GI·∫¢M TR·ª™")
    print("="*60)
    
    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")
        
    kpi_path = Path(kpi_folder)
    detail_file = kpi_path / "KPI_NVKT_SauGiamTru_ChiTiet.xlsx"
    summary_file = kpi_path / "KPI_NVKT_SauGiamTru_TomTat.xlsx"
    
    if not detail_file.exists():
        print(f"‚ùå Kh√¥ng t√¨m th·∫•y file: {detail_file}")
        return 0
        
    df_detail = pd.read_excel(detail_file)
    nvkt_list = df_detail[['don_vi', 'nvkt']].drop_duplicates()
    total = len(nvkt_list)
    
    print(f"üìä T√¨m th·∫•y {total} NVKT sau gi·∫£m tr·ª´")
    
    success_count = 0
    for idx, row in nvkt_list.iterrows():
        don_vi = row['don_vi']
        nvkt_name = row['nvkt']
        
        # ƒê·∫£m b·∫£o don_vi l√† chu·ªói
        don_vi_str = str(don_vi) if pd.notna(don_vi) else "Unknown"
        
        # T·∫°o th∆∞ m·ª•c cho t·ª´ng ƒê·ªôi (T·ªï)
        team_folder_name = sanitize_filename(don_vi_str)
        team_output_path = Path(output_root) / "ca_nhan" / team_folder_name
        team_output_path.mkdir(parents=True, exist_ok=True)
        
        current = success_count + 1
        print(f"   [{current}/{total}] {nvkt_name} ({don_vi})...", end=" ")
        
        try:
            # T·∫°o document m·ªõi
            doc = Document()
            
            # Thi·∫øt l·∫≠p style m·∫∑c ƒë·ªãnh cho doc
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # L·∫•y data NVKT
            # L·∫•y data NVKT - s·ª≠ d·ª•ng logic l·ªçc an to√†n v·ªõi NaN
            if pd.isna(don_vi):
                mask = (df_detail['nvkt'] == nvkt_name) & (df_detail['don_vi'].isna())
            else:
                mask = (df_detail['nvkt'] == nvkt_name) & (df_detail['don_vi'] == don_vi)
            
            nvkt_df_match = df_detail[mask]
            if nvkt_df_match.empty:
                print(f"‚ùå (Kh√¥ng t√¨m th·∫•y data)")
                continue
                
            nvkt_data = nvkt_df_match.iloc[0].to_dict()
            
            short_name = TEAM_SHORT_NAMES.get(don_vi_str, don_vi_str)
            created_time = datetime.now().strftime("%d/%m/%Y %H:%M")
            
            # Header
            header = doc.sections[0].header
            p = header.paragraphs[0]
            p.text = f"B√ÅO C√ÅO K·∫æT QU·∫¢ KPI C√Å NH√ÇN - TH√ÅNG {report_month} (SAU GI·∫¢M TR·ª™)"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Title
            title = doc.add_heading(level=0)
            title_run = title.add_run('B√ÅO C√ÅO K·∫æT QU·∫¢ BSC/KPI C√Å NH√ÇN (SAU GI·∫¢M TR·ª™)')
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(level=1)
            subtitle_run = subtitle.add_run(f'TH√ÅNG {report_month}')
            subtitle_run.font.size = Pt(16)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Th√¥ng tin c√° nh√¢n
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            info_data = [
                ('H·ªç v√† t√™n:', nvkt_name),
                ('ƒê∆°n v·ªã:', short_name),
                ('Ng√†y t·∫°o b√°o c√°o:', created_time)
            ]
            for i, (label, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = label
                info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                info_table.rows[i].cells[1].text = value
            
            doc.add_paragraph()
            
            # Ph·∫ßn 1: T·ªïng quan
            doc.add_heading('1. T·ªîNG QUAN ƒêI·ªÇM KPI', level=2)
            add_individual_summary_table(doc, nvkt_data)
            
            # Bi·ªÉu ƒë·ªì radar
            try:
                radar_chart = create_individual_radar_chart(nvkt_data)
                doc.add_picture(radar_chart, width=Inches(5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
                
            doc.add_page_break()
            
            # Chi ti·∫øt t·ª´ng ch·ªâ ti√™u
            add_individual_c11_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c12_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c14_detail(doc, nvkt_data)
            doc.add_paragraph()
            add_individual_c15_detail(doc, nvkt_data)
            
            # Suy hao cao - truy·ªÅn data_folder m·∫∑c ƒë·ªãnh
            add_individual_shc_section(doc, nvkt_name, data_folder="downloads/baocao_hanoi")
            
            # L∆∞u file
            safe_name = sanitize_filename(nvkt_name)
            current_date = datetime.now().strftime("%d_%m_%Y")
            filename = f"{safe_name}_Bao_cao_KPI_{current_date}.docx"
            output_file = team_output_path / filename
            doc.save(output_file)
            
            success_count += 1
            print("‚úÖ")
        except Exception as e:
            print(f"‚ùå (L·ªói: {e})")
            
    print(f"\n‚úÖ Ho√†n th√†nh: ƒê√£ t·∫°o {success_count}/{total} b√°o c√°o c√° nh√¢n sau gi·∫£m tr·ª´.")
    print(f"üìÅ Th∆∞ m·ª•c xu·∫•t: {output_root}/ca_nhan/")
    
    return success_count


# =============================================================================
# MAIN - Ch·∫°y tr·ª±c ti·∫øp module
# =============================================================================
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description='T·∫°o b√°o c√°o KPI/BSC cho NVKT')
    parser.add_argument('--individual', action='store_true', 
                       help='T·∫°o b√°o c√°o c√° nh√¢n cho t·ª´ng NVKT')
    parser.add_argument('--all', action='store_true',
                       help='T·∫°o b√°o c√°o cho t·∫•t c·∫£ NVKT (d√πng v·ªõi --individual)')
    parser.add_argument('--nvkt', type=str, default=None,
                       help='T√™n NVKT c·ª• th·ªÉ (d√πng v·ªõi --individual)')
    parser.add_argument('--donvi', type=str, default=None,
                       help='T√™n ƒë∆°n v·ªã c·ªßa NVKT (d√πng v·ªõi --individual --nvkt)')
    parser.add_argument('--month', type=str, default="01/2026",
                       help='Th√°ng b√°o c√°o (vd: 01/2026)')
    parser.add_argument('--kpi-folder', type=str, default="downloads/KPI",
                       help='Th∆∞ m·ª•c ch·ª©a file KPI')
    parser.add_argument('--output-folder', type=str, default="downloads/reports",
                       help='Th∆∞ m·ª•c xu·∫•t b√°o c√°o')
    
    args = parser.parse_args()
    
    if args.individual:
        if args.all:
            # T·∫°o b√°o c√°o cho t·∫•t c·∫£ NVKT
            generate_all_individual_reports(
                kpi_folder=args.kpi_folder,
                output_folder=args.output_folder,
                report_month=args.month
            )
        elif args.nvkt and args.donvi:
            # T·∫°o b√°o c√°o cho 1 NVKT c·ª• th·ªÉ
            result = generate_individual_kpi_report(
                nvkt_name=args.nvkt,
                don_vi=args.donvi,
                kpi_folder=args.kpi_folder,
                output_folder=args.output_folder,
                report_month=args.month
            )
            if result:
                print(f"‚úÖ ƒê√£ t·∫°o b√°o c√°o: {result}")
            else:
                print("‚ùå Kh√¥ng th·ªÉ t·∫°o b√°o c√°o")
        else:
            print("‚ùå L·ªói: C·∫ßn ch·ªâ ƒë·ªãnh --all ho·∫∑c --nvkt v√† --donvi")
            print("   V√≠ d·ª•: python report_generator.py --individual --all")
            print("   Ho·∫∑c:  python report_generator.py --individual --nvkt 'B√πi VƒÉn Du·∫©n' --donvi 'T·ªï K·ªπ thu·∫≠t ƒê·ªãa b√†n Ph√∫c Th·ªç'")
    else:
        # T·∫°o c·∫£ b√°o c√°o t·ªïng h·ª£p v√† b√°o c√°o c√° nh√¢n (m·∫∑c ƒë·ªãnh)
        print("=" * 60)
        print("üìä T·∫†O B√ÅO C√ÅO T·ªîNG H·ª¢P")
        print("=" * 60)
        report_path = generate_kpi_report(
            kpi_folder=args.kpi_folder,
            output_folder=args.output_folder,
            report_month=args.month
        )
        print(f"\nüìÅ ƒê∆∞·ªùng d·∫´n b√°o c√°o t·ªïng h·ª£p: {report_path}")
        
        print("\n")
        print("=" * 60)
        print("üìù T·∫†O B√ÅO C√ÅO C√Å NH√ÇN CHO T·∫§T C·∫¢ NVKT")
        print("=" * 60)
        generate_all_individual_reports(
            kpi_folder=args.kpi_folder,
            output_folder=args.output_folder,
            report_month=args.month
        )

