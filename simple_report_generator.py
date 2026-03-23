# -*- coding: utf-8 -*-
"""
Module tạo báo cáo Word đơn giản - Dữ liệu thô cấp Tổ & Trung tâm
Đọc trực tiếp từ các file C1.x report, không qua giảm trừ.

Author: Auto-generated
Date: 2026-03-11
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import os

# Thiết lập matplotlib
matplotlib.rcParams['font.family'] = 'DejaVu Sans'
matplotlib.use('Agg')

# Cấu hình
DEFAULT_DATA_FOLDER = "downloads/baocao_hanoi"
DEFAULT_OUTPUT_FOLDER = "downloads/reports"

TEAM_SHORT_NAMES = {
    "Tổ Kỹ thuật Địa bàn Phúc Thọ": "Phúc Thọ",
    "Tổ Kỹ thuật Địa bàn Quảng Oai": "Quảng Oai",
    "Tổ Kỹ thuật Địa bàn Suối hai": "Suối Hai",
    "Tổ Kỹ thuật Địa bàn Sơn Tây": "Sơn Tây",
    "TTVT Sơn Tây": "TTVT Sơn Tây",
    "Tổng": "TTVT Sơn Tây",
}

CHART_COLORS = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6B5B95']
BAR_COLORS = ['#4CAF50', '#2196F3', '#FF9800', '#E91E63', '#9C27B0']


# =============================================================================
# HÀM TIỆN ÍCH
# =============================================================================
def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def format_number(value, decimal_places=2):
    if pd.isna(value):
        return "N/A"
    return f"{value:.{decimal_places}f}"


def get_short_name(name):
    return TEAM_SHORT_NAMES.get(name, name)


def set_cell_width(cell, width_cm):
    """Đặt chiều rộng cho ô bảng"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1cm ≈ 567 twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# =============================================================================
# ĐỌC DỮ LIỆU
# =============================================================================
def load_raw_data(data_folder=DEFAULT_DATA_FOLDER):
    """
    Đọc dữ liệu thô từ 5 file báo cáo C1.x (sheet tổng hợp cấp Tổ/TT)
    
    Returns:
        dict: {'c11': df, 'c12': df, 'c13': df, 'c14': df, 'c15': df}
    """
    data_path = Path(data_folder)
    reports = {}

    file_map = {
        'c11': ('c1.1 report.xlsx', 'TH_C1.1'),
        'c12': ('c1.2 report.xlsx', 'TH_C1.2'),
        'c13': ('c1.3 report.xlsx', 'TH_C1.3'),
        'c14': ('c1.4 report.xlsx', 'TH_C1.4'),
        'c15': ('c1.5_chitiet_report.xlsx', 'TH_TTVTST'),
    }

    for key, (filename, sheet) in file_map.items():
        try:
            fpath = data_path / filename
            if fpath.exists():
                reports[key] = pd.read_excel(fpath, sheet_name=sheet)
                print(f"   ✅ Đọc {filename} ({sheet}) thành công - {len(reports[key])} dòng")
            else:
                print(f"   ⚠️ Không tìm thấy file: {filename}")
        except Exception as e:
            print(f"   ⚠️ Lỗi đọc {filename}: {e}")

    return reports


def load_shc_data(data_folder=DEFAULT_DATA_FOLDER):
    """
    Tìm và đọc file Bao_cao_xu_huong_SHC_ mới nhất, sheet Xu_huong_theo_don_vi
    
    Returns:
        DataFrame hoặc None
    """
    data_path = Path(data_folder)
    shc_files = sorted(data_path.glob('Bao_cao_xu_huong_SHC_*.xlsx'), reverse=True)
    
    if not shc_files:
        print("   ⚠️ Không tìm thấy file Bao_cao_xu_huong_SHC_")
        return None
    
    latest = shc_files[0]
    try:
        df = pd.read_excel(latest, sheet_name='Xu_huong_theo_don_vi')
        print(f"   ✅ Đọc SHC từ {latest.name} thành công - {len(df)} đơn vị, {len(df.columns)-1} ngày")
        return df
    except Exception as e:
        print(f"   ⚠️ Lỗi đọc SHC: {e}")
        return None


# =============================================================================
# TẠO BẢNG
# =============================================================================
def _add_header_row(table, headers, color='1565C0'):
    """Thêm dòng header cho bảng"""
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, color)
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)


def _add_data_row(table, data, idx, bold_last=False, bg_even='E3F2FD', bg_total='BBDEFB'):
    """Thêm dòng dữ liệu vào bảng"""
    cells = table.add_row().cells
    is_total = bold_last and idx == -1
    for i, value in enumerate(data):
        cells[i].text = str(value)
        cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(10)
        if is_total:
            run.font.bold = True
            set_cell_shading(cells[i], bg_total)
        elif idx % 2 == 0:
            set_cell_shading(cells[i], bg_even)
    return cells


def add_summary_table(doc, reports):
    """
    Bảng tổng hợp điểm BSC tất cả chỉ tiêu theo Tổ
    """
    doc.add_heading('Bảng tổng hợp điểm BSC theo đơn vị (Dữ liệu thô)', level=2)

    headers = ['STT', 'Đơn vị', 'C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, '0D47A1')

    # Lấy tên đơn vị từ C1.1 (hoặc file nào có)
    don_vi_col = 'Đơn vị'
    c15_don_vi_col = 'DOIVT'

    # Xây dựng dữ liệu theo đơn vị
    teams = []
    if 'c11' in reports:
        for _, row in reports['c11'].iterrows():
            teams.append(row[don_vi_col])

    for idx, team in enumerate(teams):
        short = get_short_name(team)
        is_total = (team == 'Tổng')

        # C1.1
        c11_val = 'N/A'
        if 'c11' in reports:
            r = reports['c11'][reports['c11'][don_vi_col] == team]
            if not r.empty:
                c11_val = format_number(r.iloc[0].get('Chỉ tiêu BSC', np.nan))

        # C1.2
        c12_val = 'N/A'
        if 'c12' in reports:
            r = reports['c12'][reports['c12'][don_vi_col] == team]
            if not r.empty:
                c12_val = format_number(r.iloc[0].get('Chỉ tiêu BSC', np.nan))

        # C1.3
        c13_val = 'N/A'
        if 'c13' in reports:
            r = reports['c13'][reports['c13'][don_vi_col] == team]
            if not r.empty:
                c13_val = format_number(r.iloc[0].get('Chỉ tiêu BSC', np.nan))

        # C1.4
        c14_val = 'N/A'
        if 'c14' in reports:
            r = reports['c14'][reports['c14'][don_vi_col] == team]
            if not r.empty:
                c14_val = format_number(r.iloc[0].get('Điểm BSC', np.nan))

        # C1.5
        c15_val = 'N/A'
        if 'c15' in reports:
            # C1.5 dùng cột DOIVT, ghép theo tên ngắn
            for _, r15 in reports['c15'].iterrows():
                s15 = get_short_name(r15[c15_don_vi_col])
                if s15 == short or (is_total and s15 == 'TTVT Sơn Tây'):
                    ty_le = r15.get('Tỉ lệ đạt (%)', np.nan)
                    if not pd.isna(ty_le):
                        # Tính điểm BSC C1.5 từ tỷ lệ
                        ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                        if ty_le_dec >= 0.995:
                            diem = 5
                        elif ty_le_dec > 0.895:
                            diem = 1 + 4 * (ty_le_dec - 0.895) / 0.10
                        else:
                            diem = 1
                        c15_val = format_number(diem)
                    break

        data = [str(idx + 1), short, c11_val, c12_val, c13_val, c14_val, c15_val]
        row_idx = -1 if is_total else idx
        _add_data_row(table, data, row_idx, bold_last=True)

    doc.add_paragraph()


def add_c11_detail_table(doc, df):
    """Bảng chi tiết C1.1"""
    doc.add_heading('Chi tiết C1.1 - Chất lượng sửa chữa thuê bao BRCĐ', level=3)

    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('SM1/SM2 = SC Chủ động (số phiếu/đạt) | SM3/SM4 = SC Báo hỏng (số phiếu/đạt)')

    col_tl_cd = 'Tỷ lệ sửa chữa phiếu chất lượng chủ động dịch vụ FiberVNN, MyTV đạt yêu cầu'
    col_tl_bh = 'Tỷ lệ phiếu sửa chữa báo hỏng dịch vụ BRCD đúng quy định không tính hẹn'

    headers = ['STT', 'Đơn vị', 'SC CĐ', 'Đạt', 'TL SC CĐ(%)', 'SC BH', 'Đạt', 'TL SC BH(%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, '2E7D32')

    for idx, (_, row) in enumerate(df.iterrows()):
        short = get_short_name(row['Đơn vị'])
        is_total = (row['Đơn vị'] == 'Tổng')
        data = [
            str(idx + 1), short,
            str(int(row.get('SM1', 0))), str(int(row.get('SM2', 0))),
            format_number(row.get(col_tl_cd, np.nan)),
            str(int(row.get('SM3', 0))), str(int(row.get('SM4', 0))),
            format_number(row.get(col_tl_bh, np.nan)),
            format_number(row.get('Chỉ tiêu BSC', np.nan))
        ]
        r_idx = -1 if is_total else idx
        _add_data_row(table, data, r_idx, bold_last=True, bg_even='E8F5E9', bg_total='C8E6C9')

    doc.add_paragraph()


def add_c12_detail_table(doc, df):
    """Bảng chi tiết C1.2"""
    doc.add_heading('Chi tiết C1.2 - Tỷ lệ thuê bao báo hỏng', level=3)

    p = doc.add_paragraph()
    p.add_run('📋 Chú thích: ').bold = True
    p.add_run('SM1/SM2 = Hỏng lặp lại | SM3/SM4 = Sự cố dịch vụ BRCĐ')

    col_tl_hll = 'Tỷ lệ thuê bao báo hỏng dịch vụ BRCĐ lặp lại'
    col_tl_sc = 'Tỷ lệ sự cố dịch vụ BRCĐ'

    headers = ['STT', 'Đơn vị', 'HLL', 'Tổng BH', 'TL HLL(%)', 'TB BH', 'Tổng TB', 'TL SC(%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, '1565C0')

    for idx, (_, row) in enumerate(df.iterrows()):
        short = get_short_name(row['Đơn vị'])
        is_total = (row['Đơn vị'] == 'Tổng')
        data = [
            str(idx + 1), short,
            str(int(row.get('SM1', 0))), str(int(row.get('SM2', 0))),
            format_number(row.get(col_tl_hll, np.nan)),
            str(int(row.get('SM3', 0))), str(int(row.get('SM4', 0))),
            format_number(row.get(col_tl_sc, np.nan)),
            format_number(row.get('Chỉ tiêu BSC', np.nan))
        ]
        r_idx = -1 if is_total else idx
        _add_data_row(table, data, r_idx, bold_last=True, bg_even='E3F2FD', bg_total='BBDEFB')

    doc.add_paragraph()


def add_c13_detail_table(doc, df):
    """Bảng chi tiết C1.3"""
    doc.add_heading('Chi tiết C1.3 - Chất lượng sửa chữa kênh TSL', level=3)

    col_tl_sc = 'Tỷ lệ sửa chữa dịch vụ kênh TSL hoàn thành đúng thời gian quy định'
    col_tl_hll = 'Tỷ lệ thuê bao báo hỏng dịch vụ kênh TSL lặp lại'
    col_tl_sucoTSL = 'Tỷ lệ sự cố dịch vụ kênh TSL'

    headers = ['STT', 'Đơn vị', 'SC đúng hạn(%)', 'HLL(%)', 'Sự cố(%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, 'E65100')

    for idx, (_, row) in enumerate(df.iterrows()):
        short = get_short_name(row['Đơn vị'])
        is_total = (row['Đơn vị'] == 'Tổng')
        data = [
            str(idx + 1), short,
            format_number(row.get(col_tl_sc, np.nan)),
            format_number(row.get(col_tl_hll, np.nan)),
            format_number(row.get(col_tl_sucoTSL, np.nan)),
            format_number(row.get('Chỉ tiêu BSC', np.nan))
        ]
        r_idx = -1 if is_total else idx
        _add_data_row(table, data, r_idx, bold_last=True, bg_even='FFF3E0', bg_total='FFE0B2')

    doc.add_paragraph()


def add_c14_detail_table(doc, df):
    """Bảng chi tiết C1.4"""
    doc.add_heading('Chi tiết C1.4 - Độ hài lòng khách hàng', level=3)

    headers = ['STT', 'Đơn vị', 'Tổng phiếu', 'Đã KS', 'KS TC', 'Hài lòng',
               'TL HL(%)', 'Điểm BSC']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, 'AD1457')

    for idx, (_, row) in enumerate(df.iterrows()):
        short = get_short_name(row['Đơn vị'])
        is_total = (row['Đơn vị'] == 'Tổng')
        data = [
            str(idx + 1), short,
            str(int(row.get('Tổng phiếu', 0))),
            str(int(row.get('SL đã KS', 0))),
            str(int(row.get('SL KS thành công', 0))),
            str(int(row.get('SL KH hài lòng', 0))),
            format_number(row.get('Tỷ lệ KH hài lòng', np.nan)),
            format_number(row.get('Điểm BSC', np.nan))
        ]
        r_idx = -1 if is_total else idx
        _add_data_row(table, data, r_idx, bold_last=True, bg_even='FCE4EC', bg_total='F8BBD0')

    doc.add_paragraph()


def add_c15_detail_table(doc, df):
    """Bảng chi tiết C1.5"""
    doc.add_heading('Chi tiết C1.5 - Thiết lập dịch vụ BRCĐ đạt thời gian quy định', level=3)

    headers = ['STT', 'Đơn vị', 'Phiếu đạt', 'Không đạt', 'Tổng HC', 'TL đạt(%)']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, '4A148C')

    for idx, (_, row) in enumerate(df.iterrows()):
        short = get_short_name(row['DOIVT'])
        is_total = (row['DOIVT'] == 'TTVT Sơn Tây')
        data = [
            str(idx + 1), short,
            str(int(row.get('Phiếu đạt', 0))),
            str(int(row.get('Phiếu không đạt', 0))),
            str(int(row.get('Tổng Hoàn công', 0))),
            format_number(row.get('Tỉ lệ đạt (%)', np.nan))
        ]
        r_idx = -1 if is_total else idx
        _add_data_row(table, data, r_idx, bold_last=True, bg_even='EDE7F6', bg_total='D1C4E9')

    doc.add_paragraph()


# =============================================================================
# TẠO BIỂU ĐỒ
# =============================================================================
def create_bsc_bar_chart(reports):
    """Biểu đồ cột nhóm: Điểm BSC 5 chỉ tiêu theo đơn vị"""
    teams_data = {}
    don_vi_col = 'Đơn vị'

    # Lấy danh sách tổ (không bao gồm Tổng)
    if 'c11' in reports:
        for _, row in reports['c11'].iterrows():
            name = row[don_vi_col]
            if name != 'Tổng':
                teams_data[get_short_name(name)] = {}

    if not teams_data:
        return None

    teams = list(teams_data.keys())

    # Lấy điểm BSC từng chỉ tiêu
    for team_full, short in TEAM_SHORT_NAMES.items():
        if short not in teams_data:
            continue
        if 'c11' in reports:
            r = reports['c11'][reports['c11'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.1'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c12' in reports:
            r = reports['c12'][reports['c12'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.2'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c13' in reports:
            r = reports['c13'][reports['c13'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.3'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c14' in reports:
            r = reports['c14'][reports['c14'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.4'] = r.iloc[0].get('Điểm BSC', 0) or 0

    # C1.5 - tên cột khác
    if 'c15' in reports:
        for _, row in reports['c15'].iterrows():
            s = get_short_name(row['DOIVT'])
            if s in teams_data:
                ty_le = row.get('Tỉ lệ đạt (%)', 0)
                if not pd.isna(ty_le):
                    ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                    if ty_le_dec >= 0.995:
                        diem = 5
                    elif ty_le_dec > 0.895:
                        diem = 1 + 4 * (ty_le_dec - 0.895) / 0.10
                    else:
                        diem = 1
                    teams_data[s]['C1.5'] = round(diem, 2)

    # Vẽ biểu đồ
    chi_tieus = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']
    fig, ax = plt.subplots(figsize=(12, 6))

    x = np.arange(len(teams))
    width = 0.15
    offsets = [-2, -1, 0, 1, 2]

    for i, ct in enumerate(chi_tieus):
        vals = [teams_data[t].get(ct, 0) for t in teams]
        bars = ax.bar(x + offsets[i] * width, vals, width, label=ct, color=BAR_COLORS[i], alpha=0.85)
        for bar in bars:
            h = bar.get_height()
            if h > 0:
                ax.annotate(f'{h:.1f}', xy=(bar.get_x() + bar.get_width() / 2, h),
                           xytext=(0, 3), textcoords="offset points",
                           ha='center', va='bottom', fontsize=8)

    ax.set_xlabel('Đơn vị', fontsize=12)
    ax.set_ylabel('Điểm BSC', fontsize=12)
    ax.set_title('SO SÁNH ĐIỂM BSC GIỮA CÁC TỔ (DỮ LIỆU THÔ)', fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(teams, fontsize=11)
    ax.set_ylim(0, 6)
    ax.legend(loc='upper right', fontsize=10)
    ax.grid(axis='y', linestyle='--', alpha=0.5)
    ax.axhline(y=5, color='green', linestyle=':', alpha=0.5, label='Mục tiêu')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def create_radar_chart(reports):
    """Biểu đồ radar tổng hợp 5 chỉ tiêu BSC"""
    teams_data = {}
    don_vi_col = 'Đơn vị'

    if 'c11' not in reports:
        return None

    for _, row in reports['c11'].iterrows():
        name = row[don_vi_col]
        if name != 'Tổng':
            teams_data[get_short_name(name)] = {}

    teams = list(teams_data.keys())
    if not teams:
        return None

    chi_tieus = ['C1.1', 'C1.2', 'C1.3', 'C1.4', 'C1.5']

    # Thu thập điểm (tương tự create_bsc_bar_chart)
    for team_full, short in TEAM_SHORT_NAMES.items():
        if short not in teams_data:
            continue
        if 'c11' in reports:
            r = reports['c11'][reports['c11'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.1'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c12' in reports:
            r = reports['c12'][reports['c12'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.2'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c13' in reports:
            r = reports['c13'][reports['c13'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.3'] = r.iloc[0].get('Chỉ tiêu BSC', 0) or 0
        if 'c14' in reports:
            r = reports['c14'][reports['c14'][don_vi_col] == team_full]
            if not r.empty:
                teams_data[short]['C1.4'] = r.iloc[0].get('Điểm BSC', 0) or 0

    if 'c15' in reports:
        for _, row in reports['c15'].iterrows():
            s = get_short_name(row['DOIVT'])
            if s in teams_data:
                ty_le = row.get('Tỉ lệ đạt (%)', 0)
                if not pd.isna(ty_le):
                    ty_le_dec = ty_le / 100 if ty_le > 1 else ty_le
                    if ty_le_dec >= 0.995:
                        diem = 5
                    elif ty_le_dec > 0.895:
                        diem = 1 + 4 * (ty_le_dec - 0.895) / 0.10
                    else:
                        diem = 1
                    teams_data[s]['C1.5'] = round(diem, 2)

    # Vẽ radar
    angles = np.linspace(0, 2 * np.pi, len(chi_tieus), endpoint=False).tolist()
    angles += angles[:1]  # Đóng vòng

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))

    colors = CHART_COLORS[:len(teams)]
    for i, team in enumerate(teams):
        values = [teams_data[team].get(ct, 0) for ct in chi_tieus]
        values += values[:1]
        ax.plot(angles, values, 'o-', linewidth=2, label=team, color=colors[i])
        ax.fill(angles, values, alpha=0.1, color=colors[i])

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(chi_tieus, fontsize=12)
    ax.set_ylim(0, 5.5)
    ax.set_title('BIỂU ĐỒ RADAR - BSC CÁC TỔ (DỮ LIỆU THÔ)', fontsize=14, fontweight='bold', pad=20)
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=10)
    ax.grid(True)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def create_single_indicator_chart(reports, chi_tieu_key, title, value_col, don_vi_col='Đơn vị', color='#2196F3'):
    """Biểu đồ cột cho 1 chỉ tiêu cụ thể"""
    if chi_tieu_key not in reports:
        return None

    df = reports[chi_tieu_key]
    dv_col = don_vi_col

    teams = []
    values = []
    for _, row in df.iterrows():
        name = row[dv_col]
        short = get_short_name(name)
        val = row.get(value_col, 0)
        if pd.isna(val):
            val = 0
        teams.append(short)
        values.append(val)

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(teams, values, color=color, alpha=0.85, edgecolor='white', linewidth=1.5)

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
               f'{val:.2f}', ha='center', va='bottom', fontsize=11, fontweight='bold')

    ax.set_ylabel('Giá trị', fontsize=12)
    ax.set_title(title, fontsize=13, fontweight='bold')
    ax.set_ylim(0, max(values) * 1.2 + 1 if values else 6)
    ax.grid(axis='y', linestyle='--', alpha=0.5)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def add_shc_detail_table(doc, df_shc):
    """Bảng chi tiết SHC theo đơn vị theo từng ngày"""
    doc.add_heading('Suy hao cao theo đơn vị', level=3)

    date_cols = [c for c in df_shc.columns if c != 'Đơn vị']
    # Rút gọn tiêu đề ngày: dd/mm/yyyy -> dd/mm
    short_dates = []
    for d in date_cols:
        s = str(d)
        if '/' in s:
            parts = s.split('/')
            short_dates.append(f"{parts[0]}/{parts[1]}")
        else:
            short_dates.append(s[:5])

    headers = ['Đơn vị'] + short_dates + ['Tổng']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    set_table_border(table)
    _add_header_row(table, headers, 'B71C1C')

    total_row_vals = [0] * len(date_cols)
    for idx, (_, row) in enumerate(df_shc.iterrows()):
        short = get_short_name(row['Đơn vị'])
        vals = [int(row.get(d, 0)) for d in date_cols]
        total = sum(vals)
        for j, v in enumerate(vals):
            total_row_vals[j] += v
        data = [short] + [str(v) for v in vals] + [str(total)]
        _add_data_row(table, data, idx, bg_even='FFEBEE', bg_total='FFCDD2')

    # Dòng tổng
    grand_total = sum(total_row_vals)
    total_data = ['TTVT Sơn Tây'] + [str(v) for v in total_row_vals] + [str(grand_total)]
    _add_data_row(table, total_data, -1, bold_last=True, bg_total='EF9A9A')

    doc.add_paragraph()


def create_shc_bar_chart(df_shc):
    """Biểu đồ cột: Số SHC ngày gần nhất theo đơn vị"""
    if df_shc is None or df_shc.empty:
        return None

    date_cols = [c for c in df_shc.columns if c != 'Đơn vị']
    if not date_cols:
        return None

    last_date = date_cols[-1]
    teams = [get_short_name(r['Đơn vị']) for _, r in df_shc.iterrows()]
    values = [int(r.get(last_date, 0)) for _, r in df_shc.iterrows()]

    fig, ax = plt.subplots(figsize=(10, 5))
    colors_bar = ['#EF5350', '#42A5F5', '#66BB6A', '#FFA726']
    bars = ax.bar(teams, values, color=colors_bar[:len(teams)], alpha=0.85, edgecolor='white', linewidth=1.5)

    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3,
               str(val), ha='center', va='bottom', fontsize=12, fontweight='bold')

    ax.set_ylabel('Số thuê bao SHC', fontsize=12)
    ax.set_title(f'SỐ THUÊ BAO SUY HAO CAO THEO ĐƠN VỊ (Ngày {last_date})', fontsize=13, fontweight='bold')
    ax.set_ylim(0, max(values) * 1.3 + 1 if values else 10)
    ax.grid(axis='y', linestyle='--', alpha=0.5)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


def create_shc_trend_chart(df_shc):
    """Biểu đồ đường: Xu hướng SHC theo ngày cho từng đơn vị"""
    if df_shc is None or df_shc.empty:
        return None

    date_cols = [c for c in df_shc.columns if c != 'Đơn vị']
    if len(date_cols) < 2:
        return None

    # Rút gọn nhãn ngày
    short_dates = []
    for d in date_cols:
        s = str(d)
        if '/' in s:
            parts = s.split('/')
            short_dates.append(f"{parts[0]}/{parts[1]}")
        else:
            short_dates.append(s[:5])

    fig, ax = plt.subplots(figsize=(12, 6))
    colors_line = ['#EF5350', '#42A5F5', '#66BB6A', '#FFA726']
    markers = ['o', 's', '^', 'D']

    for idx, (_, row) in enumerate(df_shc.iterrows()):
        team = get_short_name(row['Đơn vị'])
        vals = [int(row.get(d, 0)) for d in date_cols]
        ax.plot(short_dates, vals, marker=markers[idx % len(markers)],
               linewidth=2.5, label=team, color=colors_line[idx % len(colors_line)], markersize=6)
        # Giá trị đầu và cuối
        ax.annotate(str(vals[0]), (short_dates[0], vals[0]), textcoords='offset points',
                   xytext=(0, 8), ha='center', fontsize=8)
        ax.annotate(str(vals[-1]), (short_dates[-1], vals[-1]), textcoords='offset points',
                   xytext=(0, 8), ha='center', fontsize=8, fontweight='bold')

    ax.set_xlabel('Ngày', fontsize=12)
    ax.set_ylabel('Số thuê bao SHC', fontsize=12)
    ax.set_title('XU HƯỚNG SUY HAO CAO THEO ĐƠN VỊ', fontsize=14, fontweight='bold')
    ax.legend(loc='upper right', fontsize=10)
    ax.grid(True, linestyle='--', alpha=0.5)
    plt.xticks(rotation=45, ha='right')

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf


# =============================================================================
# HÀM CHÍNH - TẠO BÁO CÁO WORD
# =============================================================================
def generate_simple_report(data_folder=DEFAULT_DATA_FOLDER, output_folder=DEFAULT_OUTPUT_FOLDER,
                           report_month=None):
    """
    Tạo báo cáo Word đơn giản với dữ liệu thô cấp Tổ & Trung tâm

    Args:
        data_folder: Thư mục chứa file Excel
        output_folder: Thư mục xuất file Word
        report_month: Tháng báo cáo (vd: "03/2026")

    Returns:
        str: Đường dẫn file Word đã tạo
    """
    print("=" * 60)
    print("📝 TẠO BÁO CÁO ĐƠN GIẢN - DỮ LIỆU THÔ CẤP TỔ/TT")
    print("=" * 60)

    if report_month is None:
        report_month = datetime.now().strftime("%m/%Y")

    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)

    # 1. Đọc dữ liệu
    print("\n📊 Đọc dữ liệu thô...")
    reports = load_raw_data(data_folder)

    if not reports:
        print("❌ Không có dữ liệu nào để tạo báo cáo!")
        return None

    # 2. Tạo document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # TIÊU ĐỀ
    created_time = datetime.now().strftime('%d/%m/%Y %H:%M')
    title = doc.add_heading(level=0)
    title_run = title.add_run(f'BÁO CÁO TỔNG HỢP BSC/KPI\nDỮ LIỆU THÔ - THÁNG {report_month}')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'TTVT Sơn Tây - VNPT Hà Nội').bold = True
    doc.add_paragraph(f"Ngày tạo: {created_time}")
    doc.add_paragraph()

    # =========================================================================
    # PHẦN 1: TỔNG QUAN
    # =========================================================================
    print("\n📈 Tạo phần Tổng quan...")
    doc.add_heading('PHẦN 1: TỔNG QUAN', level=1)

    # Bảng tổng hợp
    add_summary_table(doc, reports)

    # Biểu đồ cột nhóm
    print("   📊 Tạo biểu đồ BSC tổng hợp...")
    try:
        bar_chart = create_bsc_bar_chart(reports)
        if bar_chart:
            doc.add_picture(bar_chart, width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    except Exception as e:
        print(f"   ⚠️ Lỗi tạo biểu đồ cột: {e}")

    # Biểu đồ radar
    print("   📊 Tạo biểu đồ radar...")
    try:
        radar = create_radar_chart(reports)
        if radar:
            doc.add_picture(radar, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
    except Exception as e:
        print(f"   ⚠️ Lỗi tạo biểu đồ radar: {e}")

    # =========================================================================
    # PHẦN 2: CHI TIẾT TỪNG CHỈ TIÊU
    # =========================================================================
    doc.add_heading('PHẦN 2: CHI TIẾT TỪNG CHỈ TIÊU', level=1)

    # C1.1
    if 'c11' in reports:
        print("   📊 Tạo chi tiết C1.1...")
        add_c11_detail_table(doc, reports['c11'])
        try:
            col_tl_bh = 'Tỷ lệ phiếu sửa chữa báo hỏng dịch vụ BRCD đúng quy định không tính hẹn'
            chart = create_single_indicator_chart(reports, 'c11', 
                'C1.1 - TỶ LỆ SỬA CHỮA BÁO HỎNG ĐÚNG QUY ĐỊNH (%)',
                col_tl_bh, color='#4CAF50')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ C1.1: {e}")
        doc.add_paragraph()

    # C1.2
    if 'c12' in reports:
        print("   📊 Tạo chi tiết C1.2...")
        add_c12_detail_table(doc, reports['c12'])
        try:
            chart = create_single_indicator_chart(reports, 'c12',
                'C1.2 - TỶ LỆ HỎNG LẶP LẠI (%)',
                'Tỷ lệ thuê bao báo hỏng dịch vụ BRCĐ lặp lại', color='#2196F3')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ C1.2: {e}")
        doc.add_paragraph()

    # C1.3
    if 'c13' in reports:
        print("   📊 Tạo chi tiết C1.3...")
        add_c13_detail_table(doc, reports['c13'])
        doc.add_paragraph()

    # C1.4
    if 'c14' in reports:
        print("   📊 Tạo chi tiết C1.4...")
        add_c14_detail_table(doc, reports['c14'])
        try:
            chart = create_single_indicator_chart(reports, 'c14',
                'C1.4 - TỶ LỆ KHÁCH HÀNG HÀI LÒNG (%)',
                'Tỷ lệ KH hài lòng', color='#E91E63')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ C1.4: {e}")
        doc.add_paragraph()

    # C1.5
    if 'c15' in reports:
        print("   📊 Tạo chi tiết C1.5...")
        add_c15_detail_table(doc, reports['c15'])
        try:
            chart = create_single_indicator_chart(reports, 'c15',
                'C1.5 - TỶ LỆ THIẾT LẬP DỊCH VỤ ĐẠT THỜI GIAN (%)',
                'Tỉ lệ đạt (%)', don_vi_col='DOIVT', color='#9C27B0')
            if chart:
                doc.add_picture(chart, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ C1.5: {e}")
        doc.add_paragraph()

    # =========================================================================
    # PHẦN 3: SUY HAO CAO
    # =========================================================================
    print("\n📊 Đọc dữ liệu Suy hao cao...")
    df_shc = load_shc_data(data_folder)
    if df_shc is not None:
        doc.add_heading('PHẦN 3: SUY HAO CAO', level=1)
        add_shc_detail_table(doc, df_shc)

        # Biểu đồ cột ngày mới nhất
        print("   📊 Tạo biểu đồ SHC...")
        try:
            bar = create_shc_bar_chart(df_shc)
            if bar:
                doc.add_picture(bar, width=Inches(6))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ SHC bar: {e}")

        # Biểu đồ xu hướng
        try:
            trend = create_shc_trend_chart(df_shc)
            if trend:
                doc.add_picture(trend, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        except Exception as e:
            print(f"   ⚠️ Lỗi biểu đồ SHC trend: {e}")

    # =========================================================================
    # LƯU FILE
    # =========================================================================
    month_str = report_month.replace('/', '_')
    output_file = output_path / f"Bao_cao_don_gian_BSC_{month_str}.docx"
    doc.save(str(output_file))

    file_size = os.path.getsize(output_file)
    print(f"\n✅ Đã tạo báo cáo: {output_file}")
    print(f"   Dung lượng: {file_size / 1024:.1f} KB")

    return str(output_file)


# =============================================================================
# CHẠY TRỰC TIẾP
# =============================================================================
if __name__ == "__main__":
    output = generate_simple_report(
        data_folder="downloads/baocao_hanoi",
        output_folder="downloads/reports",
        report_month="03/2026"
    )
    if output:
        print(f"\n🎉 Hoàn thành! File: {output}")
